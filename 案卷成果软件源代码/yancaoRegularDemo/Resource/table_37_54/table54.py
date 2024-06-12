# -*- coding:utf-8 -*-
# @Author: tanweijia
from os.path import dirname, abspath

import docx
import win32com
from win32com.client import Dispatch, constants
import yancaoRegularDemo.Resource.tools.tangyuhao_function as tyh
from yancaoRegularDemo.Resource.ReadFile import DocxData
import re
from yancaoRegularDemo.Resource.table_1_18.Table_Father import table_father
import os
from yancaoRegularDemo.Resource.tools.utils import *
from yancaoRegularDemo.Resource.tools.simple_content import Simple_Content
from docx import Document

function_description_dict = {
    'check_Description': '“本卷情况说明”：应填写“本卷材料齐全”、 “假冒伪劣卷烟待公开统一销毁”。 不能在卷内备考表中填写“无”或“/”',
    'check_SetTime': '“立卷时间”：与“结案时间”相同或者在其后',
}

# 责令改正通知书
class Table_54(table_father):
    def __init__(self, my_prefix, source_prefix):
        table_father.__init__(self)
        self.my_prefix = my_prefix
        self.source_prefix = source_prefix  # 2021-08-07版本新增
        self.mw = win32com.client.Dispatch("Word.Application")
        # self.doc = self.mw.Documents.Open(self.my_prefix + "卷内备考表_.docx")
        self.contract_text = None
        self.contract_tables_content = None
        self.file_name_real = ""

        self.all_to_check = [
            "self.check_Description()",
            # "self.check_Filer_and_Inspector()",
            "self.check_SetTime()"
        ]
    # 检查  本卷情况说明
    def check_Description(self):
        if not tyh.file_exists(self.source_prefix, "卷内备考表"):
            table_father.display(self, "文书缺失：《卷内备考表》.docx不存在", "red")
        else:
            doc = docx.Document(self.source_prefix + "卷内备考表_.docx")
            tb = doc.tables[0]
            # content是正文内容
            content = tb.cell(0, 0).text
            #print(content)
            temp = re.search("(.*?)\s*立卷人：", content)
            to_check = temp.group(1).strip()
            if to_check in ["无", "/"]:
                table_father.display(self,"本卷情况说明：不能在卷内备考表中填写“无”或“/“ ！", "red")
            else:
                # 检查首行
                temp = re.search("(.*?)\n", content)
                result = temp.group(0).strip()
                if not temp or result == "":
                    table_father.display(self,"本卷情况说明：首行应注明'本卷情况说明'，具体应填写“本卷材料齐全”或“假冒伪劣卷烟待公开统一销毁”", "red")
                    tyh.addRemarkInDoc(self.mw, self.doc, "卷内备考表", "没有在首行填写'本卷情况说明'，应填写“本卷材料齐全”或“假冒伪劣卷烟待公开统一销毁”")
                else:
                    if not result in ["本卷材料齐全", "假冒伪劣卷烟待公开统一销毁"]:
                        table_father.display(self, "本卷情况说明：'本卷情况说明'填写内容不合规，应填写“本卷材料齐全”或“假冒伪劣卷烟待公开统一销毁”", "red")
                        tyh.addRemarkInDoc(self.mw, self.doc, result, "'本卷情况说明'填写内容不合规，应填写“本卷材料齐全”或“假冒伪劣卷烟待公开统一销毁”")

    # 检查 “立卷人”：为两人  检查人”：与“立卷人”不一致；    # 电子签名不做
    # def check_Filer_and_Inspector(self):
    #     if not tyh.file_exists(self.source_prefix, "卷内备考表"):
    #         table_father.display(self, "× 《卷内备考表》.docx不存在", "red")
    #     else:
    #         doc = docx.Document(self.source_prefix + "卷内备考表_.docx")
    #         tb = doc.tables[0]
    #         # content是正文内容
    #         content = tb.cell(0, 0).text
    #         # 立卷人
    #         temp = re.search("立卷人：(.*?)\n", content)
    #         Filer_signature = temp.group(1).strip()
    #         #print(Filer_signature)
    #         Filer_list = Filer_signature.split()
    #         #print(Filer_list)
    #         if len(Filer_list) != 2:
    #             table_father.display(self,"× “立卷人” 应该为两人！", "red")
    #             tyh.addRemarkInDoc(self.mw, self.doc, Filer_signature, "立卷人” 应该为两人！")
    #         # 检查人
    #         temp = re.search("检查人：(.*?)\n", content)
    #         Inspector_signature = temp.group(1).strip()
    #         #print(Inspector_signature)
    #         if Inspector_signature in Filer_list:
    #             table_father.display(self,"× “检查人“ 不应与 “立卷人” 相同！", "red")
    #             tyh.addRemarkInDoc(self.mw, self.doc, Inspector_signature, "“检查人“ 不应与 “立卷人” 相同！")

    # 检查 “立卷时间”：与“结案时间”相同或者在其后
    def check_SetTime(self):
        if not tyh.file_exists(self.source_prefix, "卷内备考表"):
            pass
            # table_father.display(self, "× 《卷内备考表》.docx不存在", "red")
        else:
            # 要从 卷宗封面 获取结案日期？
            # if not tyh.file_exists(self.source_prefix, "卷宗封面"):
            #     table_father.display(self, "× 《卷宗封面》.docx不存在，无法从该文件获取结案日期", "red")
            # else:
                # 先获取《卷内备考表》的立案日期
            file_path = self.source_prefix + "//" + self.file_name_real
            # print(file_path)
            doc = docx.Document(file_path)
            tb = doc.tables[0]
            content = tb.cell(0, 0).text
            temp = re.search("立卷时间：(.*?)$", content)
            result = temp.group(1).strip()
            date0 = tyh.get_strtime(result)
            #print(date0)
            # 再获取《结案报告表》的结案日期
            if not tyh.file_exists(self.source_prefix, "结案报告表"):
                table_father.display(self, "立卷时间：《结案报告表》.docx不存在，无法从该文件获取结案日期", "red")
            else:
                data1 = tyh.file_exists_open(self.source_prefix, "结案报告表", DocxData)
                raw = data1.tabels_content['负责人意见'].strip()
                #print(raw)
                temp = re.search("日期[:：](.*?)$", raw)
                if temp:
                    origin_time = temp.group(1)
                    # print(origin_time)
                    end_time = tyh.get_strtime(origin_time)
                    # print(date0)
                    # print(end_time)
                    if tyh.time_differ(date0, end_time) < 0:
                        table_father.display(self, "立卷时间：“立卷时间” 应与“结案时间”（"+origin_time+"）相同或者在其后！此处有错！", "red")
                        tyh.addRemarkInDoc(self.mw, self.doc, result, "“立卷时间” 应与“结案时间”（"+origin_time+"）相同或者在其后！此处有错！")
                else:
                    table_father.display(self, "立卷时间：《结案报告表》中的负责人意见中不存在结案时间，无法获取结案日期来进行日期对比", "red")
                #print(end_time)
            #print(date1

    def check(self, contract_file_path, file_name_real):
        self.file_name_real = file_name_real
        print("正在审查" + file_name_real + "，审查结果如下：")
        self.mw = win32com.client.Dispatch("Word.Application")
        self.doc = self.mw.Documents.Open(self.my_prefix + file_name_real)
        data = DocxData(file_path=contract_file_path)
        self.contract_text = data.text
        self.contract_tables_content = data.tabels_content
        for func in self.all_to_check:
            try:
                eval(func)
            except Exception as e:
                table_father.display(self,
                                     "文档格式有误，请主观审查下列功能：" + function_description_dict[str(func)[5:-2]],
                                     "red")
                table_father.display(self, "文档存在格式错误，函数失效：" + func + ' 遇到错误:' + str(e.args))
        self.doc.Close()
        # self.mw.Quit()
        print(file_name_real + "审查完毕\n")
        info_list_result = table_father.get_info_list(self)
        return info_list_result


if __name__ == '__main__':

    source_prefix = "C:\\Users\\twj\Desktop\\test\\"
    # print(my_prefix)
    ioc = Table_54(source_prefix, source_prefix)
    contract_file_path = source_prefix + "卷内备考表_.docx"
    # print(contract_file_path)
    ioc.check(contract_file_path, "卷内备考表_.docx")