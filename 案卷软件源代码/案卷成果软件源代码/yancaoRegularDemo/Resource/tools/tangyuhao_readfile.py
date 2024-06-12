import os
import win32com
import win32com.client
import docx
from win32com import client
import re


class file_1():
    def __init__(self, file_path):
        self.file_path = file_path
        self.text = ""
        self.tabels_content = {}
        self._read()

    def remake_file(self, srcFile):
        dstFile = srcFile.replace(".docx", ".doc")
        if os.path.exists(dstFile):
            os.remove(dstFile)
        os.rename(srcFile, dstFile)
        path = dstFile
        topath = srcFile

        def readWordFile(path, toPath):
            mw = win32com.client.Dispatch("Word.Application")
            doc = mw.Documents.Open(path)
            # 将word的数据保存到另一个文件
            doc.SaveAs(toPath, 12)
            doc.Close()
            mw.Quit()

        readWordFile(path, topath)  # 读文件
        os.remove(dstFile)

    def update_dic(self, key, value):
        if key in self.tabels_content:
            if isinstance(self.tabels_content[key], list):
                self.tabels_content[key].append(value)
            else:
                self.tabels_content[key] = [self.tabels_content[key]] + [value]
        else:
            self.tabels_content[key] = value

    def _read(self):
        # self.remake_file(self.file_path)
        doc = docx.Document(self.file_path)
        for p in doc.paragraphs:
            self.text += p.text + "\n"
        tables = doc.tables

        for t in tables:

            if t.rows[0].cells[0].text == "品种":
                # 单独处理“移送财物清”单表
                prefix = "移送财物清单-"
                keys = []
                values = []
                for cell in t.rows[0].cells:
                    keys.append(prefix + cell.text)
                    values.append([])

                for row in t.rows[1:]:
                    i = 0
                    kill = 0
                    for cell in row.cells:
                        cell.text = cell.text.replace("\n", "").replace("\t", "").replace("\r", "").replace(" ", "")
                        if i == 0 and cell.text == "":
                            kill = 1
                            break
                        values[i].append(cell.text)
                        i += 1
                    if kill:
                        break
                for i in range(len(keys)):
                    self.update_dic(keys[i], values[i])

            elif "证据粘贴处" in t.rows[0].cells[0].text:
                # 证据复制(提取)单 暂时略过
                pass

            elif t.rows[0].cells[0].text == "序号":
                # 单独处理 涉案烟草专卖品核价表
                prefix = "涉案烟草专卖品核价表-"
                keys = []
                values = []
                for cell in t.rows[0].cells:
                    keys.append(prefix + cell.text)
                    values.append([])

                for row in t.rows[1:]:
                    i = 0
                    for cell in row.cells:
                        cell.text = cell.text.replace("\n", "").replace("\t", "").replace("\r", "").replace(" ", "")
                        if i == 0 and cell.text == "":
                            break
                        if i == 0 and cell.text == "合计":
                            self.update_dic(prefix + "全部数量合计", row.cells[2].text)
                            self.update_dic(prefix + "全部金额合计", row.cells[4].text)
                            break
                        values[i].append(cell.text)
                        i += 1
                for i in range(len(keys)):
                    self.update_dic(keys[i], values[i])

            elif len(t.rows) > 2 and t.rows[2].cells[0].text == "当事人" and t.rows[1].cells[0].text == "立案编号":
                # 单独处理 案件处理审批表
                normal = [0, 1, 6, 7, 8, 9, 10, 11, 12]
                index = 0
                for row in t.rows:
                    cells = []
                    for cell in row.cells:
                        cells.append(cell.text.replace("\n", "").replace("\t", "").replace("\r", "").replace(" ", ""))
                    cells_new = list(set(cells))
                    cells_new.sort(key=cells.index)
                    if index in normal:
                        i = 0
                        for cell in cells_new:
                            if i % 2 == 0:
                                key = cell
                                i += 1
                            else:
                                value = cell
                                if key != value and key != "":
                                    self.update_dic(key, value)
                                i += 1
                    else:
                        if index == 2:
                            if len(cells_new) == 5:
                                self.update_dic("当事人", cells_new[2])
                                self.update_dic(cells_new[3], cells_new[4])
                        if index == 3:
                            if len(cells_new) == 5:
                                self.update_dic(cells_new[1], cells_new[2])
                                self.update_dic(cells_new[3], cells_new[4])

                        if index == 4 and len(cells_new) == 9:
                            self.update_dic("当事人", cells_new[2])
                            self.update_dic(cells_new[3], cells_new[4])
                            self.update_dic(cells_new[5], cells_new[6])
                            self.update_dic(cells_new[7], cells_new[8])
                        if index == 5 and len(cells_new) == 5:
                            self.update_dic(cells_new[1], cells_new[2])
                            self.update_dic(cells_new[3], cells_new[4])

                    index += 1
            else:
                for row in t.rows:
                    key = None
                    value = None
                    i = 0
                    for cell in row.cells:
                        if i % 2 == 0:
                            key = cell.text.replace("\n", "").replace("\t", "").replace("\r", "").replace(" ", "")
                            i += 1
                        else:
                            value = cell.text.replace("\n", "").replace("\t", "").replace("\r", "").replace(" ", "")
                            if key != value and key != "":
                                self.update_dic(key, value)
                            i += 1


if __name__ == '__main__':
    # my_prefix = "D:\烟草\\tobacco\yancaoRegularDemo\data\\"
    # contract_file_path = my_prefix + "延长立案期限告知书_.docx"
    # data = file_1(contract_file_path)
    # text = data.text
    # # text= "  Aa     ："
    # pattern = r"(.*)：[\s\S]*你（单位）涉嫌一案"
    # text=re.findall(pattern, text)
    # text = text[0]
    text = ""
    doc = docx.Document(r'D:\烟草\副本\证据先行登记保存批准书_.docx')
    for p in doc.paragraphs:
        text += p.text + "\n"
    all_cell=[]
    for t in doc.tables:
        for row in t.rows:
            cells = []
            for cell in row.cells:
                cells.append(cell.text.replace("\n", "").replace("\t", "").replace("\r", "").replace(" ", ""))
            all_cell.append(cells)
            print(cells)
    print(all_cell)



    # for d in enumerate(data.tabels_content):
    #      print(d[1]+":"+data.tabels_content[d[1]])
