# -*- coding:utf-8 -*-
# @ModuleName: ReadFile
# @Function: 
# @Author: huhonghui、chenfumin、leitianyi
# @email: 1241328737@qq.com
# @Time: 2021/5/11 19:26
import os

from yancaoRegularDemo.Resource.tools.utils import get_root_dir
import docx


class DocxData:
    '''
        1.举报记录表                    ['举报时间', '举报形式', '举报人有关情况', '举报内容', '接待人意见', '承办部门负责人意见', '备注']
        2.立案报告表                    ['案由', '案件来源', '案发时间', '案发地点', '当事人', '证件类型及号码', '地址', '案情摘要', '承办人意见', '承办部门意见', '负责人意见']
        3.涉嫌犯罪案件移送呈报审批表       ['当事人', '性别', '移送案由', '有效证件及号码', '查获时间', '工作单位或住(地)址', '查获地点', '同案人', '主要违法事实', '承办部门意见', '局领导批示']
        4.移送财物清单                  ['移送财物清单-品种', '移送财物清单-规格', '移送财物清单-数量（条）', '移送财物清单-备注']
        5.涉案烟草专卖品核价表           ['涉案烟草专卖品核价表-全部数量合计', '涉案烟草专卖品核价表-全部金额合计', '涉案烟草专卖品核价表-序号', '涉案烟草专卖品核价表-品种规格', '涉案烟草专卖品核价表-数量(条)', '涉案烟草专卖品核价表-单价（元）', '涉案烟草专卖品核价表-合计（元）', '涉案烟草专卖品核价表-备注']
        6.案件调查终结报告              ['案由', '立案日期', '调查人', '当事人', '证件类型及号码', '调查事实', '案件性质', '处罚依据', '处理意见', '备注']
        7.案件处理审批表               ['案由', '立案编号', '立案日期', '当事人', '性别', '民族', '证件类型及号码', '住址', '联系电话', '同案人', '案件事实', '处罚依据', '承办人意见', '承办部门意见', '法制部门意见', '负责人意见']
        8.送 达 回 证                 ['送达文书名称', '送达文书文号', '受送达人', '送达地点', '送达方式', '收件人签名或盖章', '签收日期', '代收人注明代收理由', '见证人签名或盖章', '送达人签名', '备注']
        9.结案报告表                    ['案由', '立案日期', '调查人', '当事人', '地址', '案情摘要', '处理决定', '执行情况', '承办人结案理由', '承办部门意见', '负责人意见', '备注']
        10.卷    宗                   ['案由', '当事人', '承办人', '处理结果', '立案日期', '结案日期', '归档日期', '保存期限', '审批人']
        11.涉案物品返还清单             ['涉案物品返还清单-品种','涉案物品返还清单-规格','涉案物品返还清单-数量（条）','涉案物品返还清单-备注','涉案物品返还清单-损耗费','涉案物品返还清单-接收单位','涉案物品返还清单-返还单位','涉案物品返还清单-接收人','涉案物品返还清单-返还人','涉案物品返还清单-接收时间','涉案物品返还清单-返还时间']
        12.证据先行登记保存批准书
        13.证据先行登记保存通知书

    '''

    def __init__(self, file_path):
        self.file_path = file_path
        self.text = ""
        self.tabels_content = {}
        self._read()

    def _read(self):

        doc = docx.Document(self.file_path)
        for p in doc.paragraphs:
            self.text += p.text + "\n"
        tables = doc.tables

        for t in tables:

            if "移送财物清单" in self.text and t.rows[0].cells[0].text == "品种":
                # 单独处理“移送财物清单”表
                prefix = "移送财物清单-"
                keys = []
                values = []
                for cell in t.rows[0].cells:
                    keys.append(prefix + cell.text)
                    values.append([])

                for row in t.rows[1:]:
                    i = 0
                    kill = 0
                    for cell in row.cells:
                        cell.text = cell.text.replace("\n", "").replace("\t", "").replace("\r", "").replace(" ", "")
                        if i == 0 and cell.text == "":
                            kill = 1
                            break
                        values[i].append(cell.text)
                        i += 1
                    if kill:
                        break
                for i in range(len(keys)):
                    self.update_dic(keys[i], values[i])

            # elif "此卷共计" in self.text:
            elif "卷宗" in t.rows[0].cells[0].text.replace("\n", "").replace("\t", "").replace("\r", "").replace(" ", ""):
                # 单独处理 卷宗封面（卷宗）
                for row in t.rows[2:5]:
                    key = None
                    value = None
                    i = 0
                    for cell in row.cells:
                        if i % 2 == 0:
                            key = cell.text.replace("\n", "").replace("\t", "").replace("\r", "").replace(" ", "")
                            i += 1
                        else:
                            value = cell.text.replace("\n", "").replace("\t", "").replace("\r", "").replace(" ", "")
                            if key != value and key != "":
                                self.update_dic(key, value)
                            i += 1
                for row in t.rows[5:6]:  # 处理结果 一栏
                    key = None
                    value = None
                    i = 0
                    for cell in row.cells:
                        if i == 2:
                            break
                        elif i % 2 == 0:
                            key = cell.text.replace("\n", "").replace("\t", "").replace("\r", "").replace(" ", "")
                            i += 1
                        else:
                            value = cell.text.replace("\t", "").replace("\r", "").replace(" ", "")
                            if key != value and key != "":
                                self.update_dic(key, value)
                            i += 1
                for row in t.rows[6:]:
                    key = None
                    value = None
                    i = 0
                    for cell in row.cells:
                        if i % 2 == 0:
                            key = cell.text.replace("\n", "").replace("\t", "").replace("\r", "").replace(" ", "")
                            i += 1
                        else:
                            value = cell.text.replace("\n", "").replace("\t", "").replace("\r", "").replace(" ", "")
                            if key != value and key != "":
                                self.update_dic(key, value)
                            i += 1

            elif "卷内文件目录" in self.text:
                # 先提取文件名 和 页码
                c = [3, 5]
                for i in c:
                    col = t.columns[i]
                    key = col.cells[0].text.replace("\n", "").replace("\t", "").replace("\r", "").replace(" ", "")
                    value = []
                    for cell in col.cells[1:]:
                        value.append(cell.text.replace("\n", "").replace("\t", "").replace("\r", "").replace(" ", ""))
                    self.update_dic(key, value)


            elif "违法物品销毁记录表" in self.text:
                # 结果字典的样例
                # {'当事人': '无主', '案件编号': '筠烟立[2021]第4号', '销毁日期': '20220315', '销毁地点': '筠连县筠连镇水塘垃圾填埋场',
                # '品名': {'红塔山（硬经典）': ['84mm', '条', '12.00'], '云烟（紫）': ['84mm', '条', '13.00']}}
                # 先提取前两行
                j = 0
                for row in t.rows:
                    key = None
                    value = None
                    i = 0
                    k = 1
                    while not i == 6:
                        key = row.cells[i].text.replace("\n", "").replace("\t", "").replace("\r", "").replace(" ", "")
                        value = row.cells[k].text.replace("\n", "").replace("\t", "").replace("\r", "").replace(" ", "")
                        if key != value and key != "":
                            self.update_dic(key, value)
                        i += 3
                        k += 3
                    j += 1
                    if j == 2:
                        break

                # 在提取下半部分竖着的表格
                # 从第4行开始
                value = []
                temp = {}
                for row in t.rows[3:]:
                    if row.cells[0].text.strip() == "":
                        break
                    else:
                        key = row.cells[0].text.replace("\n", "").replace("\t", "").replace("\r", "").replace(" ", "")
                        value.append(
                            row.cells[1].text.replace("\n", "").replace("\t", "").replace("\r", "").replace(" ", ""))
                        value.append(
                            row.cells[2].text.replace("\n", "").replace("\t", "").replace("\r", "").replace(" ", ""))
                        value.append(
                            row.cells[3].text.replace("\n", "").replace("\t", "").replace("\r", "").replace(" ", ""))
                        temp[key] = value
                        value = []
                self.update_dic("品名", temp)
                # 读取其他项目
                # 单独读取 销毁理由
                key = t.rows[-4].cells[0].text.replace("\n", "").replace("\t", "").replace("\r", "").replace(" ", "")
                value = t.rows[-4].cells[1].text.replace("\n", "").replace("\t", "").replace("\r", "").replace(" ", "")
                self.update_dic(key, value)
                for row in t.rows[-3:-1]:
                    i = 0
                    k = 1
                    while not i == 6:
                        key = row.cells[i].text.replace("\n", "").replace("\t", "").replace("\r", "").replace(" ", "")
                        value = row.cells[k].text.replace("\n", "").replace("\t", "").replace("\r", "").replace(" ", "")
                        if key != value and key != "":
                            self.update_dic(key, value)
                        i += 3
                        k += 3

            elif "卷烟鉴别检验样品留样、损耗费用审批表" in self.text:
                text = ""
                list0 = []
                cell_set = []

                for p in doc.paragraphs:
                    text += p.text + "\n"

                t = doc.tables[0]
                for row in [t.rows[0], t.rows[2], t.rows[-1], t.rows[-2], t.rows[-3], t.rows[-4], t.rows[-5]]:
                    # print(k)
                    for cell in row.cells:
                        if cell not in cell_set:
                            cell_set.append(cell)
                            list0.append(
                                cell.text.replace("\n", "").replace("\t", "").replace("\r", "").replace(" ", ""))

                i = 0
                key = ""
                for x in list0:
                    if i % 2 == 0:
                        key = x.replace("\n", "").replace("\t", "").replace("\r", "").replace(" ", "")
                        i += 1
                    else:
                        value = x.replace("\n", "").replace("\t", "").replace("\r", "").replace(" ", "")
                        if key != value and key != "":
                            self.update_dic(key, value)
                        i += 1

                rows1 = t.rows[5:-5]
                all_list0 = []
                for row in rows1:
                    list0 = []
                    for cell in row.cells:
                        if cell not in cell_set:
                            cell_set.append(cell)
                            list0.append(
                                cell.text.replace("\n", "").replace("\t", "").replace("\r", "").replace(" ", ""))
                    list0 = list0[1:]
                    all_list0.append(list0)
                self.update_dic("卷烟质检样品损耗明细", all_list0)

                self.text = text


            elif "证据复制(提取)单" in self.text:
                text = ""
                for p in doc.paragraphs:
                    text += p.text + "\n"
                all_cell = []
                for t in doc.tables:
                    for row in t.rows:
                        cells = []
                        for cell in row.cells:
                            cells.append(
                                cell.text.replace("\n", "").replace("\t", "").replace("\r", "").replace(" ", ""))
                        all_cell.append(cells)
                self.tabels_content = all_cell
                self.text = text

            elif t.rows[0].cells[0].text == "序号":
                # 单独处理 涉案烟草专卖品核价表
                prefix = "涉案烟草专卖品核价表-"
                keys = []
                values = []
                for cell in t.rows[0].cells:
                    keys.append(prefix + cell.text)
                    values.append([])

                for row in t.rows[1:]:
                    i = 0
                    for cell in row.cells:
                        cell.text = cell.text.replace("\n", "").replace("\t", "").replace("\r", "").replace(" ", "")
                        if i == 0 and cell.text == "":
                            break
                        if i == 0 and cell.text == "合计":
                            self.update_dic(prefix + "全部数量合计", row.cells[2].text)
                            self.update_dic(prefix + "全部金额合计", row.cells[4].text)
                            break
                        values[i].append(cell.text)
                        i += 1
                for i in range(len(keys)):
                    self.update_dic(keys[i], values[i])

            elif "涉案物品返还清单" in self.text and t.rows[0].cells[0].text == "品种":
                # 单独处理 涉案物品返还清单
                # 分为三部分处理，以【鉴别检验样品损耗费合计】为中间部分，划分上中下三部分
                prefix = "涉案物品返还清单-"
                keys = []
                values = []
                line_num = 1

                # 开始处理第一部分
                for cell in t.rows[0].cells:
                    keys.append(prefix + cell.text)
                    values.append([])

                for row in t.rows[1:]:
                    i = 0
                    if i == 0 and row.cells[0].text == "合计":
                        self.update_dic(prefix + "全部数量合计", row.cells[2].text)
                        line_num += 1
                        break
                    else:
                        for cell in row.cells:
                            cell.text = cell.text.replace("/", "").replace("\n", "").replace("\t", "").replace("\r",
                                                                                                               "").replace(
                                " ",
                                "")
                            if i == 0 and cell.text == "":
                                break
                            values[i].append(cell.text)
                            i += 1
                        line_num += 1
                for i in range(len(keys)):
                    self.update_dic(keys[i], values[i])

                # 开始处理第二部分
                part_two_value = t.rows[line_num].cells[0].text.replace("\n", "").replace("\t", "").replace("\r",
                                                                                                            "").replace(
                    " ",
                    "")
                self.update_dic(prefix + "损耗费", part_two_value)
                # print("test_part2:"+part_two_value)
                line_num += 1

                # 开始处理第三部分
                part_three_value1 = t.rows[line_num].cells[0].text.replace("\n", "").replace("\t", "").replace("\r",
                                                                                                               "").replace(
                    " ",
                    "")
                part_three_value2 = t.rows[line_num].cells[2].text.replace("\n", "").replace("\t", "").replace("\r",
                                                                                                               "").replace(
                    " ",
                    "")
                self.update_dic(prefix + "接收单位", part_three_value1)
                self.update_dic(prefix + "返还单位", part_three_value2)

                # 注意，该表第三部分两个中间间隔线为透明，需再将line_num加一
                # 此处为了辨别接收人号码，不忽略/n
                line_num += 1
                part_three_value3 = t.rows[line_num].cells[0].text.replace("\t", "").replace("\r", "").replace(" ", "")
                part_three_value4 = t.rows[line_num].cells[2].text.replace("\t", "").replace("\r", "").replace(" ", "")
                self.update_dic(prefix + "接收人", part_three_value3)
                self.update_dic(prefix + "返还人", part_three_value4)

                line_num += 1
                part_three_value5 = t.rows[line_num].cells[0].text.replace("\n", "").replace("\t", "").replace("\r",
                                                                                                               "").replace(
                    " ",
                    "")
                part_three_value6 = t.rows[line_num].cells[2].text.replace("\n", "").replace("\t", "").replace("\r",
                                                                                                               "").replace(
                    " ",
                    "")
                self.update_dic(prefix + "接收时间", part_three_value5)
                self.update_dic(prefix + "返还时间", part_three_value6)

            elif len(t.rows) > 2 and t.rows[2].cells[0].text == "当事人" and t.rows[1].cells[0].text == "立案编号":
                # 单独处理 案件处理审批表
                # self.show_table_element(t)
                normal = [0, 1, 6, 7, 8, 9, 10, 11, 12]
                index = 0
                for row in t.rows:
                    cells = []
                    for cell in row.cells:
                        cells.append(cell.text.replace("\n", "").replace("\t", "").replace("\r", "").replace(" ", ""))
                    cells_new = self.get_unique_list(cells)
                    # cells_new = list(set(cells))
                    # cells_new.sort(key=cells.index)
                    if index in normal:
                        i = 0
                        for cell in cells_new:
                            if i % 2 == 0:
                                key = cell
                                i += 1
                            else:
                                value = cell
                                if key != value and key != "":
                                    self.update_dic(key, value)
                                i += 1
                    else:
                        if index == 2:
                            if len(cells_new) == 5:
                                self.update_dic("当事人", cells_new[2])
                                self.update_dic(cells_new[3], cells_new[4])
                        if index == 3:
                            if len(cells_new) == 5:
                                self.update_dic(cells_new[1], cells_new[2])
                                self.update_dic(cells_new[3], cells_new[4])

                        if index == 4 and len(cells_new) == 9:
                            self.update_dic("当事人", cells_new[2])
                            self.update_dic(cells_new[3], cells_new[4])
                            self.update_dic(cells_new[5], cells_new[6])
                            self.update_dic(cells_new[7], cells_new[8])
                        if index == 5 and len(cells_new) == 5:
                            self.update_dic(cells_new[1], cells_new[2])
                            self.update_dic(cells_new[3], cells_new[4])

                    index += 1

            elif ("证据先行登记保存批准书" in self.text or "证据先行登记保存通知书" in self.text) and t.rows[0].cells[0].text == "品种规格":
                text = ""
                for p in doc.paragraphs:
                    text += p.text + "\n"
                all_cell = []
                for t in doc.tables:
                    for row in t.rows:
                        cells = []
                        for cell in row.cells:
                            cells.append(
                                cell.text.replace("\n", "").replace("\t", "").replace("\r", "").replace(" ", ""))
                        all_cell.append(cells)
                self.tabels_content = all_cell
                self.text = text

            else:
                for row in t.rows:
                    key = None
                    value = None
                    i = 0
                    for cell in row.cells:
                        if i % 2 == 0:
                            key = cell.text.replace("\n", "").replace("\t", "").replace("\r", "").replace(" ",
                                                                                                          "").replace(
                                "\u3000", "")
                            i += 1
                        else:
                            value = cell.text.replace("\n", "").replace("\t", "").replace("\r", "").replace(" ",
                                                                                                            "").replace(
                                "\u3000", "")
                            if key != value and key != "":
                                self.update_dic(key, value)
                            i += 1

    def update_dic(self, key, value):
        if key in self.tabels_content:
            if isinstance(self.tabels_content[key], list):
                self.tabels_content[key].append(value)
            else:
                self.tabels_content[key] = [self.tabels_content[key]] + [value]
        else:
            self.tabels_content[key] = value

    def show_table_element(self, t):
        # 逐行查看一个table中的元素
        line_num = 0
        for row in t.rows[0:]:
            column_num = 0
            for cell in row.cells[0:]:
                print('第' + str(line_num) + '行第' + str(column_num) + '列的元素为：' + str(cell.text))
                column_num += 1
            line_num += 1

    def get_unique_list(self,input_list):
        # 对连续重复的元素，只保留一个
        # input:aabbaa output:aba
        reserved_element = input_list[0]
        output_list = [input_list[0]]
        for item in input_list:
            if item == reserved_element:
                continue
            else:
                reserved_element = item
                output_list.append(reserved_element)
        return output_list




if __name__ == '__main__':
    my_prefix = r"C:\\Users\\twj\Desktop\\test\\"
    contract_file_path = my_prefix + "卷宗封面_.docx"
    data = DocxData(contract_file_path, "卷宗封面")
    text = data.text
    tabels_content = data.tabels_content
    print(tabels_content)
    print(text)

# 打批住调用后面函数就行addRemarkInDoc
# f是打批注的地方

#
# processed_file_sava_dir = r'D:\烟草\tobacco\yancaoRegularDemo\data\副本'
# filePath = r'D:\烟草\tobacco\yancaoRegularDemo\data\船山区烟草专卖局.docx'
#
# copy_path = processed_file_sava_dir + "\\" + filePath.split("\\")[-1]
# filePath = str_insert(copy_path, copy_path.index(".docx"), "(已审查)")
# print(filePath)
# addRemarkInDoc(word,document,"举报时间","atttaaa")
# document.SaveAs(filePath)
# document.Close()
