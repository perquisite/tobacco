#!/usr/bin/env python
# -*- encoding: UTF-8 -*-
"""
@Project: ChinaTobaccoContract_V2_new_version_riskAdd
@File: check_Tender.py
@Author: Mr.Blonde
@Date: 2022/4/6 19:13
"""
from win32com.client import Dispatch
from win32com import client as wc
import re
from docx import Document


# 招投标文件与合同对比
# 找出七项招标风险

def doc_to_docx(file):
    if file[-3:] == "doc":
        word = wc.Dispatch("Word.Application")
        print(file)
        doc = word.Documents.Open(file)
        doc.SaveAs(file + "x", 12)
        doc.Close()
        file_path = file + "x"
        return file_path

    else:
        file_path = file
        return False


def check_Tender_file(file_path):
    file_path = doc_to_docx(file_path)

    try:
        document = Document(file_path)
    except:
        print("请确定\"" + file_path + "\"不是空文档！")  ##################################
        # return "docx_blank"
    paragraghs = document.paragraphs

    text = ""
    for p in paragraghs:
        if p.text != "":
            # 把半角全角符号一律转全角 add by qy
            text = text.replace(':', '：')
            text = text.replace('(', '（')
            text = text.replace(')', '）')
            text += p.text + "\n"

    check_dict = {}
    miss_list = []

    try:
        match = '第六章  合同主要条款([\s\S]*?)第七章'
        fileContent = re.findall(match, text)[0]
        # print("以下是匹配到的第六章内容：##################\n", fileContent)
        try:
            keyWord = '合同标的([\s\S]*?)合同金额及支付方式'
            factor = re.findall(keyWord, fileContent)[0].replace(" ", '')
            if factor == "":
                miss_list.append('合同标的')
            check_dict['合同标的'] = factor
        except:

            miss_list.append('合同标的')
            pass
        try:
            keyWord = '合同金额及支付方式([\s\S]*?)双方权利责任'
            factor = re.findall(keyWord, fileContent)[0].replace(" ", '')
            if factor == "":
                miss_list.append('付款方式')
            check_dict['付款方式'] = factor
        except:

            miss_list.append('合同付款方式')
            pass

        try:
            keyWord = '合同金额及支付方式([\s\S]*?)双方权利责任'
            factor = re.findall(keyWord, fileContent)[0].replace(" ", '')
            if factor == "":
                miss_list.append('合同价款')
            check_dict['合同价款'] = factor
        except:

            miss_list.append('合同价款')
            pass

        try:
            keyWord = '验收条款([\s\S]*?)违约责任'
            factor = re.findall(keyWord, fileContent)[0].replace(" ", '')
            if factor == "":
                miss_list.append('验收条款')
            check_dict['验收条款'] = factor
        except:

            miss_list.append('验收条款')
        try:
            keyWord = '违约责任([\s\S]*?)其他事项'
            factor = re.findall(keyWord, fileContent)[0].replace(" ", '')
            if factor == "":
                miss_list.append('违约责任')
            check_dict['违约责任'] = factor
        except:
            miss_list.append('违约责任')
        try:
            keyWord = '其他事项([\s\S]*?)代表'
            factor = re.findall(keyWord, fileContent)[0].replace(" ", '')
            if factor == "":
                miss_list.append('争议解决')
            check_dict['争议解决'] = factor
        except:

            miss_list.append('争议解决')
        try:
            keyWord = '其他事项([\s\S]*?)代表'
            factor = re.findall(keyWord, fileContent)[0].replace(" ", '')
            if factor == "":
                miss_list.append('合同生效时间')
            check_dict['合同生效时间'] = factor
        except:

            miss_list.append('合同生效时间')
    except:
        miss_list.append("招标文件第六章 合同主要条款不存在")
        pass

    miss_list_precise = []
    check_dict_precise = {}
    # check_dict是按章节分割的七个要素
    for key, item in check_dict.items():
        item = item.replace(' ', '')
        check_dict[key] = item

    print("###################################\n", check_dict)
    try:
        keyWord = '研究项目名称为《(.*?)》'
        factor = re.findall(keyWord, check_dict['合同标的'])[0].replace(" ", '')
        if factor == "":
            miss_list_precise.append('合同标的')

        check_dict_precise['合同标的'] = factor
    except:

        miss_list_precise.append('合同标的')
        pass

    try:
        keyWord = '经费按照(.*?)支付给乙方'
        factor = re.findall(keyWord, check_dict['付款方式'])[0].replace(" ", '')
        if factor == "":
            miss_list_precise.append('付款方式')
        check_dict_precise['付款方式'] = factor
    except:
        print("付款方式未找到")
        miss_list_precise.append('付款方式')
        pass

    try:
        keyWord = '总金额为人民币(.*?)元'
        factor = re.findall(keyWord, check_dict['合同价款'])[0].replace(" ", '')
        if factor == "":
            miss_list_precise.append('合同价款')
        check_dict_precise['合同价款'] = factor
    except:
        print("合同价款未找到")
        miss_list_precise.append('合同价款')
        pass

    check_dict_precise['验收条款'] = check_dict['验收条款']
    if '验收条款' in miss_list:
        miss_list_precise.append('验收条款')

    check_dict_precise['违约责任'] = check_dict['违约责任']
    if '违约责任' in miss_list:
        miss_list_precise.append('违约责任')

    check_dict_precise['争议解决'] = check_dict['争议解决']
    if '争议解决' in miss_list:
        miss_list_precise.append('争议解决')

    check_dict_precise['合同生效时间'] = check_dict['合同生效时间']
    if '合同生效时间' in miss_list:
        miss_list_precise.append('合同生效时间')

    print(check_dict_precise)
    # document.Close()

    return check_dict_precise, miss_list_precise


def check_Tender_file_contract(file_path):
    file_path = doc_to_docx(file_path)

    try:
        document = Document(file_path)
    except:
        print("请确定\"" + file_path + "\"不是空文档！")  ##################################
        # return "docx_blank"
    paragraghs = document.paragraphs

    text = ""
    for p in paragraghs:
        if p.text != "":
            # 把半角全角符号一律转全角 add by qy
            text = text.replace(':', '：')
            text = text.replace('(', '（')
            text = text.replace(')', '）')
            text += p.text + "\n"

    print(text)
    check_dict = {}
    miss_list = []
    # text = text.replace(" ", "")
    print(type(text))
    fileContent = ''
    fileContent = text

    # try:
    #     match = '第六章  合同主要条款(.*?)'
    #     fileContent = re.findall(match, text)[0].replace(" ", "")  # 找到主体名并去掉空格
    # fileContent=text

    try:
        keyWord = '合同标的([\s\S]*?)合同金额及支付方式'
        factor = re.findall(keyWord, fileContent)[0].replace(" ", '')
        if factor == "":
            miss_list.append('合同标的')
        check_dict['合同标的'] = factor
    except:
        print("合同标的未找到")
        miss_list.append('合同标的')
        pass
    try:
        keyWord = '合同金额及支付方式([\s\S]*?)双方权利责任'
        factor = re.findall(keyWord, fileContent)[0].replace(" ", '')
        if factor == "":
            miss_list.append('付款方式')
        check_dict['付款方式'] = factor
    except:
        print("合同付款方式未找到")
        miss_list.append('合同付款方式')
        pass
    try:
        keyWord = '合同金额及支付方式([\s\S]*?)双方权利责任'
        factor = re.findall(keyWord, fileContent)[0].replace(" ", '')
        if factor == "":
            miss_list.append('合同价款')
        check_dict['合同价款'] = factor
    except:
        # print("合同价款未找到")
        miss_list.append('合同价款')
        pass
    try:
        keyWord = '验收条款([\s\S]*?)违约责任'
        factor = re.findall(keyWord, fileContent)[0].replace(" ", '')
        if factor == "":
            miss_list.append('验收条款')
        check_dict['验收条款'] = factor
    except:
        # print("验收条款未找到")
        miss_list.append('验收条款')
    try:
        keyWord = '违约责任([\s\S]*?)其他事项'
        factor = re.findall(keyWord, fileContent)[0].replace(" ", '')
        if factor == "":
            miss_list.append('违约责任')
        check_dict['违约责任'] = factor
    except:
        # print("违约责任未找到")
        miss_list.append('违约责任')
    try:
        keyWord = '其他事项([\s\S]*?)代表'
        factor = re.findall(keyWord, fileContent)[0].replace(" ", '')
        if factor == "":
            miss_list.append('争议解决')
        check_dict['争议解决'] = factor
    except:
        # print("争议解决未找到")
        miss_list.append('争议解决')
    try:
        keyWord = '其他事项([\s\S]*?)代表'
        factor = re.findall(keyWord, fileContent)[0].replace(" ", '')
        if factor == "":
            miss_list.append('合同生效时间')
        check_dict['合同生效时间'] = factor
    except:
        # print("合同生效时间")
        miss_list.append('合同生效时间')

    # print(check_dict)
    # print(miss_list)
    miss_list_precise = []
    check_dict_precise = {}
    # check_dict是按章节分割的七个要素
    for key, item in check_dict.items():
        item = item.replace(' ', '')
        check_dict[key] = item

    print("###################################\n", check_dict)
    try:
        keyWord = '研究项目名称为《(.*?)》'
        factor = re.findall(keyWord, check_dict['合同标的'])[0].replace(" ", '')
        if factor == "":
            miss_list_precise.append('合同标的')
        # print(factor)
        check_dict_precise['合同标的'] = factor
    except:
        # print("合同标的未找到1")
        miss_list_precise.append('合同标的')
        pass

    try:
        keyWord = '经费按照(.*?)支付给乙方'
        factor = re.findall(keyWord, check_dict['付款方式'])[0].replace(" ", '')
        if factor == "":
            miss_list_precise.append('付款方式')
        check_dict_precise['付款方式'] = factor
    except:
        # print("付款方式未找到")
        miss_list_precise.append('付款方式')
        pass

    try:
        keyWord = '总金额为人民币(.*?)元'
        factor = re.findall(keyWord, check_dict['合同价款'])[0].replace(" ", '')
        if factor == "":
            miss_list_precise.append('合同价款')
        check_dict_precise['合同价款'] = factor
    except:
        # print("合同价款未找到")
        miss_list_precise.append('合同价款')
        pass

    check_dict_precise['验收条款'] = check_dict['验收条款']
    if '验收条款' in miss_list:
        miss_list_precise.append('验收条款')

    check_dict_precise['违约责任'] = check_dict['违约责任']
    if '违约责任' in miss_list:
        miss_list_precise.append('违约责任')

    check_dict_precise['争议解决'] = check_dict['争议解决']
    if '争议解决' in miss_list:
        miss_list_precise.append('争议解决')

    check_dict_precise['合同生效时间'] = check_dict['合同生效时间']
    if '合同生效时间' in miss_list:
        miss_list_precise.append('合同生效时间')

    print(check_dict_precise)
    # document.Close()

    return check_dict_precise, miss_list_precise


# def check_Tender_file_contract(file_path):
#     file_path = doc_to_docx(file_path)
#
#     try:
#         document=Document(file_path)
#     except:
#         # print("请确定\"" + file_path + "\"不是空文档！")  ##################################
#         pass
#     finally:
#         document = Document(file_path)
#     # document = Document(file_path)
#     paragraghs = document.paragraphs
#
#     text = ""
#     for p in paragraghs:
#         if p.text != "":
#             # 把半角全角符号一律转全角 add by qy
#             text = text.replace(':', '：')
#             text = text.replace('(', '（')
#             text = text.replace(')', '）')
#             text += p.text + "\n"
#
#     print(text)
#     check_dict = {}
#     miss_list = []
#     # text = text.replace(" ", "")
#     print(type(text))
#
#     # try:
#     #     match = '第六章  合同主要条款(.*?)'
#     #     fileContent = re.findall(match, text)[0].replace(" ", "")  # 找到主体名并去掉空格
#     # fileContent=text
#
#     try:
#         keyWord = '合同标的([\s\S]*?)合同金额及支付方式'
#         factor = re.findall(keyWord, fileContent)[0].replace("", ' ')
#         if factor == "":
#             miss_list.append('合同标的')
#         check_dict['合同标的'] = factor
#     except:
#         print("合同标的未找到")
#         miss_list.append('合同标的')
#         pass
#     try:
#         keyWord = '合同金额及支付方式([\s\S]*?)双方权利责任'
#         factor = re.findall(keyWord, fileContent)[0].replace("", ' ')
#         if factor == "":
#             miss_list.append('付款方式')
#         check_dict['付款方式'] = factor
#     except:
#         print("合同付款方式未找到")
#         miss_list.append('合同付款方式')
#         pass
#     try:
#         keyWord = '合同金额及支付方式([\s\S]*?)双方权利责任'
#         factor = re.findall(keyWord, fileContent)[0].replace("", ' ')
#         if factor == "":
#             miss_list.append('合同价款')
#         check_dict['合同价款'] = factor
#     except:
#         print("合同价款未找到")
#         miss_list.append('合同价款')
#         pass
#     try:
#         keyWord = '验收条款([\s\S]*?)违约责任'
#         factor = re.findall(keyWord, fileContent)[0].replace("", ' ')
#         if factor == "":
#             miss_list.append('验收条款')
#         check_dict['验收条款'] = factor
#     except:
#         print("验收条款未找到")
#         miss_list.append('验收条款')
#     try:
#         keyWord = '违约责任([\s\S]*?)其他事项'
#         factor = re.findall(keyWord, fileContent)[0].replace("", ' ')
#         if factor == "":
#             miss_list.append('违约责任')
#         check_dict['违约责任'] = factor
#     except:
#         print("违约责任未找到")
#         miss_list.append('违约责任')
#     try:
#         keyWord = '其他事项([\s\S]*?)代表'
#         factor = re.findall(keyWord, fileContent)[0].replace("", ' ')
#         if factor == "":
#             miss_list.append('争议解决')
#         check_dict['争议解决'] = factor
#     except:
#         print("争议解决未找到")
#         miss_list.append('争议解决')
#     try:
#         keyWord = '其他事项([\s\S]*?)代表'
#         factor = re.findall(keyWord, fileContent)[0].replace("", ' ')
#         if factor == "":
#             miss_list.append('合同生效时间')
#         check_dict['合同生效时间'] = factor
#     except:
#         print("合同生效时间")
#         miss_list.append('合同生效时间')
#
#
#     print(check_dict)
#     print(miss_list)
#     miss_list_precise = []
#     check_dict_precise = {}
#     # check_dict是按章节分割的七个要素
#     for key, item in check_dict.items():
#         item = item.replace(' ', '')
#         print(key)
#         print(item)
#         check_dict[key]=item
#     print("###################################\n",check_dict)
#     try:
#         keyWord = '研究项目名称为《(.*?)》'
#         factor = re.findall(keyWord, check_dict['合同标的'])[0].replace("", ' ')
#         if factor == "":
#             miss_list_precise.append('合同标的')
#         check_dict_precise['合同标的'] = factor
#     except:
#         print("合同标的未找到")
#         miss_list_precise.append('合同标的')
#         pass
#
#     try:
#         keyWord = '经费按照(.*?)支付给乙方'
#         factor = re.findall(keyWord, check_dict['付款方式'])[0].replace("", ' ')
#         if factor == "":
#             miss_list_precise.append('付款方式')
#         check_dict_precise['付款方式'] = factor
#     except:
#         print("付款方式未找到")
#         miss_list_precise.append('付款方式')
#         pass
#
#     try:
#         keyWord = '总金额为人民币(.*?)万元'
#         factor = re.findall(keyWord, check_dict['合同价款'])[0].replace("", ' ')
#         if factor == "":
#             miss_list_precise.append('合同价款')
#         check_dict_precise['合同价款'] = factor
#     except:
#         print("合同价款未找到")
#         miss_list_precise.append('合同价款')
#         pass
#
#     check_dict_precise['验收条款'] = check_dict['验收条款']
#     if '验收条款' in miss_list:
#         miss_list_precise.append('验收条款')
#
#     check_dict_precise['违约责任'] = check_dict['违约责任']
#     if '违约责任' in miss_list:
#         miss_list_precise.append('违约责任')
#
#     check_dict_precise['争议解决'] = check_dict['争议解决']
#     if '争议解决' in miss_list:
#         miss_list_precise.append('争议解决')
#
#     check_dict_precise['合同生效时间'] = check_dict['合同生效时间']
#     if '合同生效时间' in miss_list:
#         miss_list_precise.append('合同生效时间')
#
#     return check_dict_precise, miss_list_precise


# 不同条款的描述方式
check_statement_dict = {
    '合同标的': '',
    '付款方式': '',
    '合同价款': '',
    '验收条款': '',
    '违约'
    ''
    ''
    '责任': '',
    '争议解E:\projects\pycharm\hetong\范例\3\合同文本 (3).doc决': '',
    '合同生效时间': '',
}

missList = []
checkDict = {}
missList_contract = []
checkDict_contract = {}

checkDict, missList = check_Tender_file('E:\projects\pycharm\hetong\范例\\2\招标文件客户满意度调查2021----修改版0811.doc')
checkDict_contract, missList_contract = check_Tender_file_contract('E:\projects\pycharm\hetong\范例\\2\合同文本 (3).doc')
# print("#################checkDict####################\n", checkDict)
print("\033[1;31m#################missList 招标缺失####################\n", missList)
print("\033[1;33m#################checkDict 招标查找####################")
for i in checkDict.keys():
    print(f"{i}----{checkDict[i]}")
print("\033[1;31m#################missList_contract 合同缺失####################\n", missList_contract)
print("\033[1;33m#################checkDict_contract 合同查找####################")
for i in checkDict_contract.keys():
    print(f"\033[1;33m{i}----{checkDict_contract[i]}\033[0m")

def compare_results(dict1, dict2):
    # differ=set(dict1.items())^set(dict2.items())
    # print(differ)
    differ = {}
    for key, item in dict1.items():
        if item != '':
            if dict2[key] != item:
                differ[key] = item
            else:
                pass
        else:
            pass
    return differ


differ = compare_results(checkDict, checkDict_contract)

for i in differ.keys():
    print(f"{i}----{differ[i]}\n")
