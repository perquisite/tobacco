# coding:utf-8
import os
import re

from architecture_zhaobiao_read import A_DocReader
from utils import UnifiedSocialCreditIdentifier, checkIdCard, isTelPhoneNumber, isRightDate, checkQQ, checkEmail, \
    str_insert, addRemarkInDoc, digital_to_Upper, is_contain_dot, isEmail, checkEntersAndSpace, is_youbian, is_bankcard, \
    table_ok, check_money1, check_money, check_money_split
from win32com.client import Dispatch
import helpful as hp
import datetime


# add by tyh
def get_strtime(s):
    if "年" not in s or "月" not in s or "日" not in s:
        return False
    out = ""
    s = s.replace("年", "-").replace("月", "-").replace("日", " ").replace("/", "-").strip().replace(" ", "").replace("；",
                                                                                                                   "").replace(
        "。", "").replace("；", "").replace("。", "")
    for i in s:
        if i in ['1', '2', '3', '4', '5', '6', '7', '8', '9', '0', '-']:
            out += i
    out = re.sub("\s+", " ", out)
    t = ""
    t = re.findall('(\d{4}-\d{1,2}-\d{1,2})', out)
    if t:
        t = t[0]
        return t
    else:
        return False


def processFunc3(tables, text, filePath, processed_file_sava_dir, filePath_zhaobiao):
    try:
        word = Dispatch('Word.Application')
        word.Documents.close()
        word.Quit()
    except:
        pass
    finally:
        word.Visible = 0
        document = word.Documents.Open(FileName=filePath)

    factors = {}
    factors_ok = []
    factors_error = {}
    factors_to_inform = {}
    zhaobiao_exist = 0
    # print(text, tables)

    # 缺失的要素
    missObject = ""

    try:
        dir=os.path.split(filePath)[0]
        filePath_zhaobiao=dir+'/招标文件.docx'
        reader = A_DocReader(filePath_zhaobiao)
        text_zhaobiao = reader.get_text()
        architecture_dict = reader.get_info(text_zhaobiao)
        zhaobiao_exist = 1
        # addRemarkInDoc(word, document, " ", filePath_zhaobiao)
        print('招标文件提取：'+filePath_zhaobiao)
        for k, v in architecture_dict.items():
            print(str(k) + ':' + str(v))
            # addRemarkInDoc(word, document, " ", str(k) + ':' + str(v))
    except:
        zhaobiao_exist = 0
        print('招标文件提取失败')
        addRemarkInDoc(word, document, "", f"招标文件提取失败")

    # 第一二部分
    if 1 == True:
        # 第一部分 合同协议书_甲方乙方
        fabaoren = None
        chengbaoren = None
        try:
            match = '甲方（发包人）：(.*?)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "").replace("\t", "")
            factors["主体&甲方（发包人）："] = factor
            if factor != "" and factor != "；" and factor != "。":
                fabaoren = factor
                factors_ok.append("主体&甲方（发包人）：")

            else:
                factors_error["主体&甲方（发包人）："] = "第一部分开头：甲方（发包人）：未填写完整"
                addRemarkInDoc(word, document, "甲方（发包人）：", f"第一部分开头：要素填写错误：甲方（发包人）：未填写完整")
        except:
            missObject += "第一部分开头：要素“甲方（发包人）：”缺失\n"
        addRemarkInDoc(word, document, "甲方（发包人）：", f"第一部分开头：请审核是否与交易相对方提供的身份证明文件一致")
        factors_to_inform["主体&甲方（发包人）："] = "第一部分开头：请审核是否与交易相对方提供的身份证明文件一致"

        try:
            match = '住所地：(.*?)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["主体&甲方（发包人）住所地"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("主体&甲方（发包人）住所地")
            else:
                factors_error["主体&甲方（发包人）住所地"] = "第一部分开头：住所地未填写完整"
                addRemarkInDoc(word, document, "甲方（发包人）", f"第一部分开头：甲方（发包人）住所地未填写完整")
        except:
            missObject += "第一部分开头：要素“甲方（发包人）住所地”缺失\n"

        fabaoren_fading = None
        try:
            match = '法定代表人/负责人：(.*?)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["主体&甲方（发包人）法定代表人/负责人"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("主体&甲方（发包人）法定代表人/负责人")
                fabaoren_fading = factor
            else:
                factors_error["主体&甲方（发包人）法定代表人/负责人"] = "第一部分开头：法定代表人/负责人未填写完整"
                addRemarkInDoc(word, document, "甲方（发包人）", f"第一部分开头：要素填写错误：甲方（发包人）法定代表人/负责人未填写完整")
        except:
            missObject += "第一部分开头：要素“甲方（发包人）法定代表人/负责人”缺失\n"

        try:
            match = '统一社会信用代码：(.*?)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["主体&甲方（发包人）统一社会信用代码/身份证号码"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if UnifiedSocialCreditIdentifier().check_code(factor, 'sc'):
                    factors_ok.append("主体&甲方（发包人）统一社会信用代码/身份证号码")
                else:
                    # factors_error["主体&甲方（发包人）统一社会信用代码/身份证号码"] = "主体&甲方（发包人）统一社会信用代码/身份证号码未填写正确"
                    if checkIdCard(factor) == 'ok':
                        factors_ok.append("主体&甲方（发包人）统一社会信用代码/身份证号码")
                    else:
                        rs = checkIdCard(factor)
                        factors_error["主体&甲方（发包人）统一社会信用代码/身份证号码"] = "第一部分开头：统一社会信用代码未填写正确或" + rs
                        addRemarkInDoc(word, document, "甲方（发包人）", f"第一部分开头：请核对并完善统一社会信用代码/身份证号码" + rs)
            else:
                factors_error["主体&甲方（发包人）统一社会信用代码/身份证号码"] = "第一部分开头：统一社会信用代码/身份证号码未填写完整"
                addRemarkInDoc(word, document, "甲方（发包人）", f"第一部分开头：请核对并完善统一社会信用代码/身份证号码")
        except:
            missObject += "第一部分开头：要素“甲方（发包人）统一社会信用代码/身份证号码”缺失\n"

        try:
            match = '联系电话：(.*?)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["主体&甲方（发包人）联系电话"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if isTelPhoneNumber(factor) != "Error":
                    factors_ok.append("主体&甲方（发包人）联系电话")
                else:
                    factors_error["主体&甲方（发包人）联系电话"] = "第一部分开头：联系电话填写有误"

                    addRemarkInDoc(word, document, "甲方（发包人）", f"第一部分开头：甲方（发包人）联系电话填写错误")
            else:
                factors_error["主体&甲方（发包人）联系电话"] = "第一部分开头：联系电话未填写完整"

                addRemarkInDoc(word, document, "甲方（发包人）", f"第一部分开头：甲方（发包人）联系电话未填写完整")
        except:
            missObject += "第一部分开头：要素“甲方（发包人）联系电话”缺失\n"

        try:
            match = '乙方（承包人）：(.*?)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "").replace("\t", "")
            factors["主体&乙方（承包人）："] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("主体&乙方（承包人）：")
                chengbaoren = factor
            else:
                factors_error["主体&乙方（承包人）："] = "第一部分开头：乙方（承包人）：未填写完整"
                addRemarkInDoc(word, document, "乙方（承包人）：", f"第一部分开头：要素填写错误：乙方（承包人）：未填写完整")
        except:
            missObject += "第一部分开头：要素“乙方（承包人）：”缺失\n"
        addRemarkInDoc(word, document, "乙方（承包人）：", f"第一部分开头：请审核是否与交易相对方提供的身份证明文件一致")
        factors_to_inform["主体&乙方（承包人）："] = "第一部分开头：请审核是否与交易相对方提供的身份证明文件一致"

        try:
            match = '住所地：(.*?)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["主体&乙方（承包人）住所地"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("主体&乙方（承包人）住所地")
                addRemarkInDoc(word, document, "乙方（承包人）：", f"第一部分开头：请审核是否与交易相对方提供的身份证明文件一致")
            else:
                factors_error["主体&乙方（承包人）住所地"] = "第一部分开头：住所地未填写完整"

                addRemarkInDoc(word, document, "乙方（承包人）", f"第一部分开头：乙方（承包人）住所地未填写完整")
        except:
            missObject += "第一部分开头：要素“乙方（承包人）住所地”缺失\n"

        chengbaoren_fading = None
        try:
            match = '法定代表人/负责人：(.*?)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["主体&乙方（承包人）法定代表人/负责人"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("主体&乙方（承包人）法定代表人/负责人")
                chengbaoren_fading = factor
            else:
                factors_error["主体&乙方（承包人）法定代表人/负责人"] = "第一部分开头：第一部分开头：法定代表人/负责人未填写完整"

                addRemarkInDoc(word, document, "乙方（承包人）", f"第一部分开头：要素填写错误：乙方（承包人）法定代表人/负责人未填写完整")
        except:
            missObject += "第一部分开头：要素“乙方（承包人）法定代表人/负责人”缺失\n"

        try:
            match = '统一社会信用代码：(.*?)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["主体&乙方（承包人）统一社会信用代码/身份证号码"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if UnifiedSocialCreditIdentifier().check_code(factor, 'sc'):
                    factors_ok.append("主体&乙方（承包人）统一社会信用代码/身份证号码")
                else:
                    # factors_error["主体&乙方（承包人）统一社会信用代码/身份证号码"] = "主体&乙方（承包人）统一社会信用代码/身份证号码未填写正确"
                    if checkIdCard(factor) == 'ok':
                        factors_ok.append("主体&乙方（承包人）统一社会信用代码/身份证号码")
                    else:
                        rs = checkIdCard(factor)
                        factors_error["主体&乙方（承包人）统一社会信用代码/身份证号码"] = "第一部分开头：统一社会信用代码未填写正确或" + rs

                        addRemarkInDoc(word, document, "乙方（承包人）", f"第一部分开头：请核对并完善统一社会信用代码/身份证号码" + rs)
            else:
                factors_error["主体&乙方（承包人）统一社会信用代码/身份证号码"] = "统一社会信用代码/身份证号码未填写完整"

                addRemarkInDoc(word, document, "乙方（承包人）", f"第一部分开头：请核对并完善统一社会信用代码/身份证号码")
        except:
            missObject += "第一部分开头：要素“乙方（承包人）统一社会信用代码/身份证号码”缺失\n"

        try:
            match = '联系电话：(.*?)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["主体&乙方（承包人）联系电话"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if isTelPhoneNumber(factor) != "Error":
                    factors_ok.append("主体&乙方（承包人）联系电话")
                else:
                    factors_error["主体&乙方（承包人）联系电话"] = "第一部分开头：联系电话填写有误"

                    addRemarkInDoc(word, document, "乙方（承包人）", f"第一部分开头：乙方（承包人）联系电话填写错误")
            else:
                factors_error["主体&乙方（承包人）联系电话"] = "第一部分开头：联系电话未填写完整"

                addRemarkInDoc(word, document, "乙方（承包人）", f"第一部分开头：乙方（承包人）联系电话未填写完整")
        except:
            missObject += "第一部分开头：要素“乙方（承包人）联系电话”缺失\n"

        # 第一部分_工程概况
        factors_to_inform["工程概况"] = "请比对该部分内容是否与招标文件的“项目概况”内容一致"
        addRemarkInDoc(word, document, "一、工程概况", f"第一部分_工程概况：请比对该部分内容是否与招标文件的“项目概况”内容一致")
        gongchengmingcheng = None
        try:
            match = '工程名称：(.*?)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["工程名称"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("工程名称")
                gongchengmingcheng = factor
            else:
                factors_error["工程名称"] = "第一部分_工程概况：工程名称未填写完整"
                addRemarkInDoc(word, document, "工程名称", f"第一部分_工程概况：工程名称未填写完整")

            try:
                if zhaobiao_exist == 1:
                    if factor != architecture_dict['工程名称']:
                        factors_error["工程名称"] = "第一部分_工程概况：工程名称与招标文件不一致"
                        addRemarkInDoc(word, document, "工程名称", f"第一部分_工程概况：工程名称与招标文件不一致")
            except:
                addRemarkInDoc(word, document, "工程名称", f"第一部分_工程概况：招标文件工程名称提取错误")
        except:
            missObject += "第一部分_工程概况：要素“工程名称”缺失\n"
        gongchengdidian = None
        try:
            match = '工程地点：(.*?)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["工程地点"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("工程地点")
                gongchengdidian = factor
            else:
                factors_error["工程地点"] = "第一部分_工程概况：工程地点未填写完整"

                addRemarkInDoc(word, document, "工程地点", f"第一部分_工程概况：工程地点未填写完整")
            try:
                if zhaobiao_exist == 1:
                    if factor != architecture_dict['工程地点']:
                        factors_error["工程地点"] = "第一部分_工程概况：工程地点与招标文件不一致"
                        addRemarkInDoc(word, document, "工程地点", f"第一部分_工程概况：工程地点与招标文件不一致")
            except:
                addRemarkInDoc(word, document, "工程地点", f"第一部分_工程概况：招标文件工程地点提取错误")
        except:
            missObject += "第一部分_工程概况：要素“工程地点”缺失\n"

        try:
            match = '工程立项批准文号：(.*?)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["工程立项批准文号"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("工程立项批准文号")
            else:
                factors_error["工程立项批准文号"] = "第一部分_工程概况：工程立项批准文号未填写完整"

                addRemarkInDoc(word, document, "工程立项批准文号", f"第一部分_工程概况：工程立项批准文号未填写完整")
            try:
                if zhaobiao_exist == 1:
                    if factor != architecture_dict['工程立项批准文号']:
                        factors_error["工程立项批准文号"] = "第一部分_工程概况：工程立项批准文号与招标文件不一致"
                        addRemarkInDoc(word, document, "工程立项批准文号", f"第一部分_工程概况：工程立项批准文号与招标文件不一致")
            except:
                addRemarkInDoc(word, document, "工程立项批准文号", f"第一部分_工程概况：招标文件工程立项批准文号提取错误")
        except:
            missObject += "第一部分_工程概况：要素“工程立项批准文号”缺失\n"

        try:
            match = '资金来源：(.*?)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["资金来源"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("资金来源")
            else:
                factors_error["资金来源"] = "第一部分_工程概况：资金来源未填写完整"

                addRemarkInDoc(word, document, "资金来源", f"第一部分_工程概况：资金来源未填写完整")

            try:
                if zhaobiao_exist == 1:
                    if factor != architecture_dict['资金来源']:
                        factors_error["资金来源"] = "第一部分_工程概况：资金来源与招标文件不一致"
                        addRemarkInDoc(word, document, "资金来源", f"第一部分_工程概况：资金来源与招标文件不一致")
            except:
                addRemarkInDoc(word, document, "资金来源", f"第一部分_工程概况：招标文件资金来源提取错误")
        except:
            missObject += "第一部分_工程概况：要素“资金来源”缺失\n"

        try:
            match = '工程内容（应与招标文件的“项目概况”内容一致）：(.*?)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["工程内容"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("工程内容")
            else:
                factors_error["工程内容"] = "第一部分_工程概况：工程内容未填写完整"

                addRemarkInDoc(word, document, "工程内容", f"第一部分_工程概况：工程内容未填写完整")
            try:
                if zhaobiao_exist == 1:
                    if factor != architecture_dict['工程内容']:
                        factors_error["工程内容"] = "第一部分_工程概况：工程内容与招标文件不一致"
                        addRemarkInDoc(word, document, "工程内容", f"第一部分_工程概况：工程内容与招标文件不一致")
            except:
                addRemarkInDoc(word, document, "工程内容", f"第一部分_工程概况：招标文件提取错误")
        except:
            missObject += "第一部分_工程概况：要素“工程内容”缺失\n"

        try:
            if "附件1：承包人承揽工程项目一览表" in text:
                factors_ok.append("详见《承包人承揽工程项目一览表》（附件1）")
            else:
                factors_error["详见《承包人承揽工程项目一览表》（附件1）"] = "第一部分_工程概况：未附有附件1或者附件1名称与该条表述不一致"

                addRemarkInDoc(word, document, "详见《承包人承揽工程项目一览表》（附件1）", f"第一部分_工程概况：未附有附件1或者附件1名称与该条表述不一致")
        except:
            missObject += "第一部分_工程概况：要素“详见《承包人承揽工程项目一览表》（附件1）”缺失\n"

        # 第一部分_合同工期

        factors_to_inform["二、合同工期"] = "请比对该部分内容是否与招标文件的“工期”内容一致"
        addRemarkInDoc(word, document, "二、合同工期", f"第一部分_合同工期：请比对该部分内容是否与招标文件的“工期”内容一致")
        start_time = None
        try:
            match = '计划开工日期：(.*?)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "").replace("。", "")
            factors["计划开工日期"] = factor
            if get_strtime(factor):
                factors_ok.append("计划开工日期")
                start_time = get_strtime(factor)
                year = start_time.split('-')[0]
                month = start_time.split('-')[1]
                day = start_time.split('-')[2]
                start_time = datetime.date(int(year), int(month), int(day))

            else:
                factors_error["计划开工日期"] = "第一部分_合同工期：计划开工日期未填写完整"

                addRemarkInDoc(word, document, "计划开工日期", f"第一部分_合同工期：计划开工日期未填写完整")
        except:
            missObject += "第一部分_合同工期：要素“计划开工日期”缺失\n"

        end_time = None
        try:
            match = '计划竣工日期：(.*?)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "").replace("。", "")
            factors["计划竣工日期"] = factor
            if get_strtime(factor):
                factors_ok.append("计划竣工日期")
                end_time = get_strtime(factor)
                year = end_time.split('-')[0]
                month = end_time.split('-')[1]
                day = end_time.split('-')[2]
                end_time = datetime.date(int(year), int(month), int(day))

            else:
                factors_error["计划竣工日期"] = "第一部分_合同工期：计划竣工日期未填写完整"

                addRemarkInDoc(word, document, "计划竣工日期", f"第一部分_合同工期：计划竣工日期未填写完整")
        except:
            missObject += "第一部分_合同工期：要素“计划竣工日期”缺失\n"

        try:
            match = '工期总日历天数：(.*?)天.*'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "").replace("\ue5e5", "")
            factors["工期总日历天数"] = factor
            if factor != "" and factor != "；" and factor != "。":

                if start_time is not None and end_time is not None:
                    days_sum = abs((start_time - end_time).days)
                    if int(factor) != int(days_sum) + 1:
                        factors_error["工期总日历天数"] = "第一部分_合同工期：工期总日历天数可能错误"
                        addRemarkInDoc(word, document, "工期总日历天数", f"第一部分_合同工期：工期总日历天数可能错误")
                else:
                    factors_ok.append("工期总日历天数")
            else:
                factors_error["工期总日历天数"] = "第一部分_合同工期：工期总日历天数未填写完整"

                addRemarkInDoc(word, document, "工期总日历天数", f"第一部分_合同工期：工期总日历天数未填写完整")
            try:
                if zhaobiao_exist == 1:
                    if float(factor) != float(architecture_dict['计划工期']):
                        factors_error["工期总日历天数："] = "第一部分_合同工期：计划工期与招标文件不一致"
                        addRemarkInDoc(word, document, "工期总日历天数：", f"第一部分_合同工期：计划工期与招标文件不一致")
            except:
                addRemarkInDoc(word, document, "工期总日历天数：", f"第一部分_合同工期：招标文件计划工期提取错误")
        except:
            missObject += "第一部分_合同工期：要素“工期总日历天数”缺失\n"

        # 第一部分_签约合同价与合同价格形式
        try:
            addRemarkInDoc(word, document, "签约合同价为：", f"第一部分_签约合同价与合同价格形式：检查该部分与招标文件是否一致")
            factors_to_inform["签约合同价"] = "检查该部分与招标文件是否一致"
            match = '签约合同价为：\n.*人民币(.*)；\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "").replace("\ue5e5", "")
            if factor == '':
                factors_error["签约合同价"] = "第一部分_签约合同价与合同价格形式：签约合同价无法提取"
                addRemarkInDoc(word, document, "签约合同价：", f"第一部分_签约合同价与合同价格形式：签约合同价无法提取")
            else:
                factor = factor.replace(" ", "").replace("；", "").replace("。", "").replace("；", "").replace("。", "")
                factors["签约合同价"] = factor
                if check_money(factor) != False:
                    factors_ok.append("签约合同价")
                else:
                    factors_error["签约合同价"] = "第一部分_签约合同价与合同价格形式：签约合同价未填写完整或大小写不一致"
                    addRemarkInDoc(word, document, "签约合同价为：", f"第一部分_签约合同价与合同价格形式：签约合同价未填写完整或大小写不一致")
                try:
                    if zhaobiao_exist == 1:
                        if float(check_money_split(factor)[1]) != float(architecture_dict['签约合同价']):
                            factors_error["签约合同价"] = "第一部分_签约合同价与合同价格形式：签约合同价与招标文件不一致"
                            addRemarkInDoc(word, document, "签约合同价为：", f"第一部分_签约合同价与合同价格形式：签约合同价与招标文件不一致")
                except:
                    addRemarkInDoc(word, document, "签约合同价为：", f"第一部分_签约合同价与合同价格形式：招标文件签约合同价提取错误")
        except:
            missObject += "第一部分_签约合同价与合同价格形式：要素“签约合同价”缺失\n"

        try:
            addRemarkInDoc(word, document, "安全文明施工费：", f"第一部分_签约合同价与合同价格形式：检查与招标文件的“价款”是否一致")
            match = '安全文明施工费：\n人民币(.*)；\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "").replace("\ue5e5", "")
            if factor == '':
                factors_error["安全文明施工费"] = "第一部分_签约合同价与合同价格形式：安全文明施工费无法提取"

                addRemarkInDoc(word, document, "安全文明施工费：", f"第一部分_签约合同价与合同价格形式：安全文明施工费无法提取")
            else:
                factor = factor.replace(" ", "").replace("；", "").replace("。", "").replace("；", "").replace("。", "")
                factors["安全文明施工费"] = factor
                if check_money(factor) != False:
                    factors_ok.append("安全文明施工费")
                else:
                    factors_error["安全文明施工费"] = "第一部分_签约合同价与合同价格形式：安全文明施工费未填写完整或大小写不一致"
                    addRemarkInDoc(word, document, "安全文明施工费：", f"第一部分_签约合同价与合同价格形式：安全文明施工费未填写完整或大小写不一致")
                try:
                    if zhaobiao_exist == 1:
                        if float(check_money_split(factor)[1]) != float(architecture_dict['安全文明施工费']):
                            factors_error["安全文明施工费"] = "第一部分_签约合同价与合同价格形式：安全文明施工费与招标文件不一致"
                            addRemarkInDoc(word, document, "安全文明施工费：", f"第一部分_签约合同价与合同价格形式：安全文明施工费与招标文件不一致")
                except:
                    addRemarkInDoc(word, document, "安全文明施工费：", f"第一部分_签约合同价与合同价格形式：招标文件安全文明施工费提取错误")
        except:
            missObject += "第一部分_签约合同价与合同价格形式：要素“安全文明施工费”缺失\n"

        try:
            addRemarkInDoc(word, document, "材料和工程设备暂估价金额：", f"第一部分_签约合同价与合同价格形式：检查与招标文件的“价款”是否一致")
            match = '.*材料和工程设备暂估价金额：\n.*人民币(.*?)；\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "").replace("\ue5e5", "")
            if factor == '':
                factors_error["材料和工程设备暂估价金额"] = "第一部分_签约合同价与合同价格形式：材料和工程设备暂估价金额无法提取"

                addRemarkInDoc(word, document, "材料和工程设备暂估价金额：", f"第一部分_签约合同价与合同价格形式：材料和工程设备暂估价金额无法提取")
            else:
                factor = factor.replace(" ", "").replace("；", "").replace("。", "").replace("；", "").replace("。", "")
                factors["材料和工程设备暂估价金额"] = factor
                if check_money(factor) != False:
                    factors_ok.append("材料和工程设备暂估价金额")
                else:
                    factors_error["材料和工程设备暂估价金额"] = "第一部分_签约合同价与合同价格形式：材料和工程设备暂估价金额未填写完整或大小写不一致"
                    addRemarkInDoc(word, document, "材料和工程设备暂估价金额：", f"第一部分_签约合同价与合同价格形式：材料和工程设备暂估价金额未填写完整或大小写不一致")
                try:
                    if zhaobiao_exist == 1:
                        if float(check_money_split(factor)[1]) != float(architecture_dict['材料和工程设备暂估价金额']):
                            factors_error["材料和工程设备暂估价金额"] = "第一部分_签约合同价与合同价格形式：材料和工程设备暂估价金额与招标文件不一致"
                            addRemarkInDoc(word, document, "材料和工程设备暂估价金额：", f"第一部分_签约合同价与合同价格形式：材料和工程设备暂估价金额与招标文件不一致")
                except:
                    addRemarkInDoc(word, document, "材料和工程设备暂估价金额：", f"第一部分_签约合同价与合同价格形式：招标文件材料和工程设备暂估价金额提取错误")
        except:
            missObject += "第一部分_签约合同价与合同价格形式：要素“材料和工程设备暂估价金额”缺失\n"

        try:
            addRemarkInDoc(word, document, "专业工程暂估价金额：", f"第一部分_签约合同价与合同价格形式：检查与招标文件的“价款”是否一致")
            match = '.*专业工程暂估价金额：\n.*人民币(.*?)；\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "").replace("\ue5e5", "")
            if factor == '':
                factors_error["专业工程暂估价金额"] = "第一部分_签约合同价与合同价格形式：专业工程暂估价金额无法提取"

                addRemarkInDoc(word, document, "专业工程暂估价金额：", f"第一部分_签约合同价与合同价格形式：专业工程暂估价金额无法提取")
            else:
                factor = factor.replace(" ", "").replace("；", "").replace("。", "").replace("；", "").replace("。", "")
                factors["专业工程暂估价金额"] = factor
                if check_money(factor) != False:
                    factors_ok.append("专业工程暂估价金额")
                else:
                    factors_error["专业工程暂估价金额"] = "第一部分_签约合同价与合同价格形式：专业工程暂估价金额未填写完整或大小写不一致"
                    addRemarkInDoc(word, document, "专业工程暂估价金额：", f"第一部分_签约合同价与合同价格形式：专业工程暂估价金额未填写完整或大小写不一致")
                try:
                    if zhaobiao_exist == 1:
                        if float(check_money_split(factor)[1]) != float(architecture_dict['专业工程暂估价金额']):
                            factors_error["专业工程暂估价金额"] = "第一部分_签约合同价与合同价格形式：专业工程暂估价金额与招标文件不一致"
                            addRemarkInDoc(word, document, "专业工程暂估价金额：", f"第一部分_签约合同价与合同价格形式：专业工程暂估价金额与招标文件不一致")
                except:
                    addRemarkInDoc(word, document, "专业工程暂估价金额：", f"第一部分_签约合同价与合同价格形式：招标文件专业工程暂估价金额提取错误")
        except:
            missObject += "第一部分_签约合同价与合同价格形式：要素“专业工程暂估价金额”缺失\n"

        try:
            addRemarkInDoc(word, document, "暂列金额：", f"第一部分_签约合同价与合同价格形式：检查与招标文件的“价款”是否一致")
            match = '.*暂列金额：\n.*人民币(.*?)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "").replace("\ue5e5", "")
            if factor == '':
                factors_error["暂列金额"] = "第一部分_签约合同价与合同价格形式：暂列金额无法提取"
                addRemarkInDoc(word, document, "暂列金额：", f"第一部分_签约合同价与合同价格形式：暂列金额无法提取")
            else:
                factor = factor.replace(" ", "").replace("；", "").replace("。", "").replace("；", "").replace("。", "")
                factors["暂列金额"] = factor
                if check_money(factor) != False:
                    factors_ok.append("暂列金额")
                else:
                    factors_error["暂列金额"] = "第一部分_签约合同价与合同价格形式：暂列金额未填写完整或大小写不一致"
                    addRemarkInDoc(word, document, "暂列金额：", f"第一部分_签约合同价与合同价格形式：暂列金额未填写完整或大小写不一致")
                try:
                    if zhaobiao_exist == 1:
                        if float(check_money_split(factor)[1]) != float(architecture_dict['暂列金额']):
                            factors_error["专业工程暂估价金额"] = "第一部分_签约合同价与合同价格形式：暂列金额与招标文件不一致"
                            addRemarkInDoc(word, document, "暂列金额：", f"第一部分_签约合同价与合同价格形式：暂列金额与招标文件不一致")
                except:
                    addRemarkInDoc(word, document, "暂列金额：", f"第一部分_签约合同价与合同价格形式：招标文件暂列金额提取错误")
        except:
            missObject += "第一部分_签约合同价与合同价格形式：要素“暂列金额”缺失\n"

        try:
            match = '合同价格形式：(.*?)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["合同价格形式"] = factor
            if factor != "" and factor != "；" and factor != "。" and factor in ['单价合同', '总价合同']:
                factors_ok.append("合同价格形式")
            else:
                factors_error["合同价格形式"] = "第一部分_签约合同价与合同价格形式：合同价格形式未填写正确"

                addRemarkInDoc(word, document, "2.合同价格形式：", f"第一部分_签约合同价与合同价格形式：合同价格形式未填写正确")
        except:
            missObject += "第一部分_签约合同价与合同价格形式：要素“合同价格形式”缺失\n"

        # 第一部分_项目经理
        try:
            match = '承包人项目经理（按中标文件载明的项目经理填写）(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["承包人项目经理"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("承包人项目经理")
                factors_to_inform["承包人项目经理"] = "第一部分_项目经理：请比对与招标文件中的“项目经理”是否一致"
                addRemarkInDoc(word, document, "承包人项目经理", f"第一部分_项目经理：请比对与招标文件中的“项目经理”是否一致")
            else:
                factors_error["承包人项目经理"] = "第一部分_项目经理：承包人项目经理未填写完整"
                addRemarkInDoc(word, document, "承包人项目经理", f"第一部分_项目经理：承包人项目经理未填写完整")
            try:
                if zhaobiao_exist == 1:
                    if check_money_split(factor)[1] != architecture_dict['承包人项目经理']:
                        factors_error["承包人项目经理"] = "第一部分_项目经理：承包人项目经理与招标文件不一致"
                        addRemarkInDoc(word, document, "承包人项目经理", f"第一部分_签约合同价与合同价格形式：承包人项目经理与招标文件不一致")
            except:
                addRemarkInDoc(word, document, "承包人项目经理", f"第一部分_签约合同价与合同价格形式：招标文件承包人项目经理提取错误")
        except:
            missObject += "第一部分_项目经理：要素“承包人项目经理”缺失\n"

        # 第一部分_签订时间
        sign_date = None
        try:
            match = '九、签订时间\n.*本合同于(.*)签订。\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["签订时间"] = factor
            if get_strtime(factor) != False:
                sign_date = get_strtime(factor)
                year = sign_date.split('-')[0]
                month = sign_date.split('-')[1]
                day = sign_date.split('-')[2]
                sign_date = datetime.date(int(year), int(month), int(day))
                if start_time and (start_time - sign_date).days > 0:
                    factors_error["签订时间"] = "第一部分_签订时间：签订时间是否在计划开工时间之前"
                    addRemarkInDoc(word, document, "九、签订时间", f"第一部分_签订时间：签订时间没有在计划开工时间之前")
                else:
                    factors_ok.append("签订时间")
                    factors_to_inform["签订时间"] = "第一部分_签订时间：请核实合同签订时间是否在中标通知书发出之日30日内签订"
                    addRemarkInDoc(word, document, "九、签订时间", f"第一部分_签订时间：请核实合同签订时间是否在中标通知书发出之日30日内签订")
            else:
                factors_error["签订时间"] = "第一部分_签订时间：签订时间未填写完整"

                addRemarkInDoc(word, document, "九、签订时间", f"第一部分_签订时间：签订时间未填写完整")
        except:
            missObject += "第一部分_签订时间：要素“签订时间”缺失\n"

        # 第一部分_签订地点
        try:
            match = '十、签订地点\n.*本合同在(.*)签订。\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["签订地点"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("签订地点")
            else:
                factors_error["签订地点"] = "第一部分_签订地点：签订地点未填写完整"

                addRemarkInDoc(word, document, "签订地点", f"第一部分_签订地点：签订地点未填写完整")
        except:
            missObject += "第一部分_签订地点：要素“签订地点”缺失\n"

        # 第一部分_结尾签字
        addRemarkInDoc(word, document, "均具有同等法律效力",
                       f"第一部分_结尾签字：请核对印章的名称与合同主体信息中的名称是否一致，请使用企业信息查询工具等核实各项填写内容是否与公示信息一致")
        factors_to_inform["第一部分_结尾印章"] = "请核对印章的名称与合同主体信息中的名称是否一致，请使用企业信息查询工具等核实各项填写内容是否与公示信息一致"

        try:
            match = '发包人：(.*)[（]公章[）].*承包人.*\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "")
            factors["第一部分_结尾发包人"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    if factor != fabaoren:
                        factors_error["第一部分_结尾发包人"] = "第一部分_结尾签字：和合同发包人不同"
                        addRemarkInDoc(word, document, "均具有同等法律效力", f"第一部分_结尾签字：和合同发包人不同")
                    else:
                        factors_ok.append("第一部分_结尾发包人")
                except:
                    addRemarkInDoc(word, document, "均具有同等法律效力", f"第一部分_结尾签字：发包人提取失败")
            else:
                factors_error["第一部分_结尾发包人"] = "第一部分_结尾签字：第一部分_结尾发包人未填写完整"
                addRemarkInDoc(word, document, "均具有同等法律效力", f"第一部分_结尾签字：第一部分_结尾发包人未填写完整")
        except:
            missObject += "第一部分_结尾签字：要素“第一部分_结尾发包人”缺失\n"

        try:
            match = '发包人：.*承包人：(.*)[（]公章[）]\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "")
            factors["第一部分_结尾承包人"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    if factor != chengbaoren:
                        factors_error["第一部分_结尾承包人"] = "第一部分_结尾签字：和合同承包不同"
                        addRemarkInDoc(word, document, "均具有同等法律效力", f"第一部分_结尾签字：和合同承包不同")
                    else:
                        factors_ok.append("第一部分_结尾承包人")
                except:
                    addRemarkInDoc(word, document, "均具有同等法律效力", f"第一部分_结尾签字：承包人提取失败")
            else:
                factors_error["第一部分_结尾承包人"] = "第一部分_结尾签字：第一部分_结尾承包人未填写完整"
                addRemarkInDoc(word, document, "均具有同等法律效力", f"第一部分_结尾签字：第一部分_结尾承包人未填写完整")
        except:
            missObject += "第一部分_结尾签字：要素“第一部分_结尾承包人”缺失\n"

        try:
            match = '.*法定代表人或其委托代理人：(.*)法定代表人或其委托代理人：(.*)\n'
            factor = list(re.findall(match, text)[0])
            factor[0] = factor[0].replace(" ", "")
            factor[1] = factor[1].replace(" ", "")
            if factor == ['', '']:
                match = '.*法定代表人或其委托代理人：.*法定代表人或其委托代理人：.*\n(.*)\n'
                factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
                factors["法定代表人或其委托代理人签字"] = factor
                if factor != "" and factor != "；" and factor != "。" and factor.split("（签字）")[1] != '' and \
                        factor.split("（签字）")[2] != '':
                    factors_ok.append("法定代表人或其委托代理人签字")
                else:
                    factors_error["法定代表人或其委托代理人签字"] = "第一部分_结尾签字：法定代表人或其委托代理人签字未填写完整"
                    addRemarkInDoc(word, document, "均具有同等法律效力", f"第一部分_结尾签字：法定代表人或其委托代理人签字未填写完整")
            else:
                factors["法定代表人或其委托代理人签字"] = factor
                if factor != [] and factor != ['', ''] and factor[0] != '' and \
                        factor[1] != '':
                    factors_ok.append("法定代表人或其委托代理人签字")
                else:
                    factors_error["法定代表人或其委托代理人签字"] = "第一部分_结尾签字：法定代表人或其委托代理人签字未填写完整"
                    addRemarkInDoc(word, document, "均具有同等法律效力", f"第一部分_结尾签字：法定代表人或其委托代理人签字未填写完整")
        except:
            missObject += "第一部分_结尾签字：要素“法定代表人或其委托代理人签字”缺失\n"

        # 第一部分_结尾要素
        try:
            match = '.*统一社会信用代码：(.*)统一社会信用代码：(.*)\n'
            factor = re.findall(match, text)[0]
            factor = [factor[0].replace(" ", "").replace("；", "").replace("。", "").replace("；", "").replace("。", ""),
                      factor[1].replace(" ", "").replace("；", "").replace("。", "").replace("；", "").replace("。", "")]

            factors["统一社会信用代码"] = factor
            if factor != ["", ""]:
                if UnifiedSocialCreditIdentifier().check_code(factor[0],
                                                              'sc') and UnifiedSocialCreditIdentifier().check_code(
                    factor[1], 'sc'):
                    factors_ok.append("统一社会信用代码")
                else:
                    factors_error["主体&出租方统一社会信用代码/身份证号码"] = "第一部分_结尾要素：统一社会信用代码未填写正确或"

                    addRemarkInDoc(word, document, "均具有同等法律效力", f"第一部分_结尾要素：请核对并完善统一社会信用代码")
            else:
                factors_error["统一社会信用代码"] = "第一部分_结尾要素：统一社会信用代码未填写完整"

                addRemarkInDoc(word, document, "均具有同等法律效力", f"第一部分_结尾要素：统一社会信用代码未填写完整")
        except:
            missObject += "第一部分_结尾要素：要素“统一社会信用代码”缺失\n"

        dizhi = None
        try:
            match = '.*地  址：(.*)地  址：(.*)\n'
            factor = re.findall(match, text)[0]
            factor = [factor[0].replace(" ", "").replace("；", "").replace("。", "").replace("；", "").replace("。", ""),
                      factor[1].replace(" ", "").replace("；", "").replace("。", "").replace("；", "").replace("。", "")]
            factors["地  址"] = factor
            if factor != ["", ""]:
                factors_ok.append("地  址")
                dizhi = factor
            else:
                factors_error["地  址"] = "第一部分_结尾要素：地  址未填写完整"

                addRemarkInDoc(word, document, "均具有同等法律效力", f"第一部分_结尾要素：地  址未填写完整")
        except:
            missObject += "第一部分_结尾要素：要素“地  址”缺失\n"

        try:
            match = '.*邮政编码：(.*)邮政编码：(.*)\n'
            factor = re.findall(match, text)[0]
            factor = [factor[0].replace(" ", "").replace("；", "").replace("。", "").replace("；", "").replace("。", ""),
                      factor[1].replace(" ", "").replace("；", "").replace("。", "").replace("；", "").replace("。", "")]
            factors["邮政编码"] = factor
            if factor != ["", ""]:
                if is_youbian(factor[0]) and is_youbian(factor[1]):
                    factors_ok.append("邮政编码")
                else:
                    factors_error["邮政编码"] = "第一部分_结尾要素：邮政编码填写错误"

                    addRemarkInDoc(word, document, "均具有同等法律效力", f"第一部分_结尾要素：邮政编码填写错误")
            else:
                factors_error["邮政编码"] = "第一部分_结尾要素：邮政编码未填写完整"

                addRemarkInDoc(word, document, "均具有同等法律效力", f"第一部分_结尾要素：邮政编码未填写完整")
        except:
            missObject += "第一部分_结尾要素：要素“邮政编码”缺失\n"
        fabaoren_fading = None
        chengbaoren_fading = None
        fadingdaibiaoren = None
        try:
            match = '.*法定代表人：(.*)法定代表人：(.*)\n'
            factor = re.findall(match, text)[0]
            factor = [factor[0].replace(" ", "").replace("；", "").replace("。", "").replace("；", "").replace("。", ""),
                      factor[1].replace(" ", "").replace("；", "").replace("。", "").replace("；", "").replace("。", "")]
            factors["法定代表人"] = factor
            if factor != ["", ""]:
                factors_ok.append("法定代表人")
                fabaoren_fading = factor[0]
                chengbaoren_fading = factor[1]
                fadingdaibiaoren = factor
            else:
                factors_error["法定代表人"] = "第一部分_结尾要素：法定代表人未填写完整"

                addRemarkInDoc(word, document, "均具有同等法律效力", f"第一部分_结尾要素：法定代表人未填写完整")
        except:
            missObject += "第一部分_结尾要素：要素“法定代表人”缺失\n"
        fabaoren_daili = None
        chengbaoren_daili = None
        weituodailiren = None
        try:
            match = '.*法定代表人：.*\n.*委托代理人：(.*)委托代理人：(.*)\n'
            factor = re.findall(match, text)[0]
            factor = [factor[0].replace(" ", "").replace("；", "").replace("。", "").replace("；", "").replace("。", ""),
                      factor[1].replace(" ", "").replace("；", "").replace("。", "").replace("；", "").replace("。", "")]
            factors["委托代理人"] = factor
            if factor != ["", ""]:
                factors_ok.append("委托代理人")
                fabaoren_daili = factor[0]
                chengbaoren_daili = factor[1]
                weituodailiren = factor
            else:
                factors_error["委托代理人"] = "第一部分_结尾要素：委托代理人未填写完整"

                addRemarkInDoc(word, document, "均具有同等法律效力", f"第一部分_结尾要素：委托代理人未填写完整")
        except:
            missObject += "第一部分_结尾要素：要素“委托代理人”缺失\n"
        dianhua = None

        try:
            match = '.*电  话：(.*)电  话：(.*)\n'
            factor = re.findall(match, text)[0]
            factor = [factor[0].replace(" ", "").replace("；", "").replace("。", "").replace("\t", "").replace("。", ""),
                      factor[1].replace(" ", "").replace("；", "").replace("。", "").replace("\t", "").replace("。", "")]
            factors["电  话"] = factor
            if factor != ["", ""]:
                if isTelPhoneNumber(factor[0]) == True and isTelPhoneNumber(factor[1]) == True:
                    factors_ok.append("电  话")
                    dianhua = factor
                else:
                    factors_error["电  话"] = "第一部分_结尾要素：电  话未填写错误"

                    addRemarkInDoc(word, document, "均具有同等法律效力", f"第一部分_结尾要素：电  话未填写错误")
            else:
                factors_error["电  话"] = "第一部分_结尾要素：电  话未填写完整"

                addRemarkInDoc(word, document, "均具有同等法律效力", f"第一部分_结尾要素：电  话未填写完整")
        except:
            missObject += "第一部分_结尾要素：要素“电  话”缺失\n"

        try:
            match = '.*传  真：(.*)传  真：(.*)\n'
            factor = re.findall(match, text)[0]
            factor = [factor[0].replace(" ", "").replace("；", "").replace("。", "").replace("\t", "").replace("。", ""),
                      factor[1].replace(" ", "").replace("；", "").replace("。", "").replace("\t", "").replace("。", "")]
            factors["传  真"] = factor
            if factor != ["", ""]:
                factors_ok.append("传  真")
            else:
                factors_error["传  真"] = "第一部分_结尾要素：传  真未填写完整"

                addRemarkInDoc(word, document, "均具有同等法律效力", f"第一部分_结尾要素：传  真未填写完整")
        except:
            missObject += "第一部分_结尾要素：要素“传  真”缺失\n"

        try:
            match = '.*电子信箱：(.*)电子信箱：(.*)\n'
            factor = re.findall(match, text)[0]
            factor = [factor[0].replace(" ", "").replace("；", "").replace("。", "").replace("\t", "").replace("。", ""),
                      factor[1].replace(" ", "").replace("；", "").replace("。", "").replace("\t", "").replace("。", "")]
            factors["电子信箱"] = factor
            if factor != ["", ""]:
                factors_ok.append("电子信箱")
            else:
                factors_error["电子信箱"] = "第一部分_结尾要素：电子信箱未填写完整"

                addRemarkInDoc(word, document, "均具有同等法律效力", f"第一部分_结尾要素：电子信箱未填写完整")
        except:
            missObject += "第一部分_结尾要素：要素“电子信箱”缺失\n"

        try:
            match = '.*开户银行：(.*)开户银行：(.*)\n'
            factor = re.findall(match, text)[0]
            factor = [factor[0].replace(" ", "").replace("；", "").replace("。", "").replace("\t", "").replace("。", ""),
                      factor[1].replace(" ", "").replace("；", "").replace("。", "").replace("\t", "").replace("。", "")]
            factors["开户银行"] = factor
            if factor != ["", ""]:
                if is_bankcard(factor[0]) and is_bankcard(factor[1]):
                    factors_ok.append("开户银行")
                else:
                    factors_error["开户银行"] = "第一部分_结尾要素：开户银行填写错误"

                    addRemarkInDoc(word, document, "均具有同等法律效力", f"第一部分_结尾要素：开户银行填写错误")
            else:
                factors_error["开户银行"] = "第一部分_结尾要素：开户银行未填写完整"

                addRemarkInDoc(word, document, "均具有同等法律效力", f"第一部分_结尾要素：开户银行未填写完整")
        except:
            missObject += "第一部分_结尾要素：要素“开户银行”缺失\n"

        try:
            match = '.*发包人账号：(.*)承包人账号：(.*)\n'
            factor = re.findall(match, text)[0]
            factor = [factor[0].replace(" ", "").replace("；", "").replace("。", "").replace("\t", "").replace("。", ""),
                      factor[1].replace(" ", "").replace("；", "").replace("。", "").replace("\t", "").replace("。", "")]
            factors["发包人账号_承包人账号"] = factor
            if factor != ["", ""]:
                factors_ok.append("发包人账号_承包人账号")
            else:
                factors_error["发包人账号_承包人账号"] = "第一部分_结尾要素：发包人账号_承包人账号未填写完整"

                addRemarkInDoc(word, document, "均具有同等法律效力", f"第一部分_结尾要素：发包人账号_承包人账号未填写完整")
        except:
            missObject += "第一部分_结尾要素：要素“发包人账号_承包人账号”缺失\n"

    # 第三部分_1
    if 1 == True:
        # 第三部分_1
        try:
            match = '.*合同\n其他合同文件包括：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["词语定义_合同"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor.replace(" ", '') != '经双方确认进入合同的其它文件':
                    factors_error["词语定义_合同"] = "第三部分_1.1 词语定义:请详细了解该文件内容，并将该文件列入本合同附件"
                    addRemarkInDoc(word, document, "其他合同文件包括：", f"第三部分_1.1 词语定义:请详细了解该文件内容，并将该文件列入本合同附件")
                else:
                    factors_ok.append("词语定义_合同")
            else:
                factors_error["词语定义_合同"] = "词语定义_合同未填写完整"

                addRemarkInDoc(word, document, "其他合同文件包括：", f"第三部分_1.1 词语定义:词语定义_合同未填写完整")
        except:
            missObject += "第三部分_1.1 词语定义:要素“词语定义_合同”缺失\n"

        try:
            match = '.*监理人：\n名    称：(.*)\n资质类别和等级：.*\n联系电话：.*\n电子信箱：.*\n通信地址：.*\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["监理人_名称"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("监理人_名称")
            else:
                factors_error["监理人_名称"] = "第三部分_1.1 词语定义:监理人_名称未填写完整"
                addRemarkInDoc(word, document, "1.1.2.1监理人：", f"第三部分_1.1 词语定义:监理人_名称未填写完整")
        except:
            missObject += "第三部分_1.1 词语定义:要素“监理人_名称”缺失\n"

        try:
            match = '.*监理人：\n名    称：.*\n资质类别和等级：(.*)\n联系电话：.*\n电子信箱：.*\n通信地址：.*\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["监理人_资质类别和等级"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("监理人_资质类别和等级")
            else:
                factors_error["监理人_资质类别和等级"] = "第三部分_1.1 词语定义:监理人_资质类别和等级未填写完整"

                addRemarkInDoc(word, document, "1.1.2.1监理人：", f"第三部分_1.1 词语定义:监理人_资质类别和等级未填写完整")
        except:
            missObject += "第三部分_1.1 词语定义:要素“监理人_资质类别和等级”缺失\n"

        try:
            match = '.*监理人：\n名    称：.*\n资质类别和等级：.*\n联系电话：(.*)\n电子信箱：.*\n通信地址：.*\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["监理人_联系电话"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("监理人_联系电话")
            else:
                factors_error["监理人_联系电话"] = "第三部分_1.1 词语定义:监理人_联系电话未填写完整"

                addRemarkInDoc(word, document, "1.1.2.1监理人：", f"第三部分_1.1 词语定义:监理人_联系电话未填写完整")
        except:
            missObject += "第三部分_1.1 词语定义:要素“监理人_联系电话”缺失\n"

        try:
            match = '.*监理人：\n名    称：.*\n资质类别和等级：.*\n联系电话：.*\n电子信箱：(.*)\n通信地址：.*\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["监理人_电子信箱"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("监理人_电子信箱")
            else:
                factors_error["监理人_电子信箱"] = "第三部分_1.1 词语定义:监理人_电子信箱未填写完整"

                addRemarkInDoc(word, document, "1.1.2.1监理人：", f"第三部分_1.1 词语定义:监理人_电子信箱未填写完整")
        except:
            missObject += "第三部分_1.1 词语定义:要素“监理人_电子信箱”缺失\n"

        try:
            match = '.*监理人：\n名    称：.*\n资质类别和等级：.*\n联系电话：.*\n电子信箱：.*\n通信地址：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["监理人_通信地址"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("监理人_通信地址")
            else:
                factors_error["监理人_通信地址"] = "第三部分_1.1 词语定义:监理人_通信地址未填写完整"

                addRemarkInDoc(word, document, "1.1.2.1监理人：", f"第三部分_1.1 词语定义:监理人_通信地址未填写完整")
        except:
            missObject += "第三部分_1.1 词语定义:要素“监理人_通信地址”缺失\n"

        try:
            match = '.*设计人：\n名    称：(.*)\n资质类别和等级：.*\n联系电话：.*\n电子信箱：.*\n通信地址：.*\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["设计人_名称"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("设计人_名称")
            else:
                factors_error["设计人_名称"] = "第三部分_1.1 词语定义:设计人_名称未填写完整"

                addRemarkInDoc(word, document, "1.1.2.2 设计人：", f"第三部分_1.1 词语定义:设计人_名称未填写完整")
        except:
            missObject += "第三部分_1.1 词语定义:要素“设计人_名称”缺失\n"

        try:
            match = '.*设计人：\n名    称：.*\n资质类别和等级：(.*)\n联系电话：.*\n电子信箱：.*\n通信地址：.*\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["设计人_资质类别和等级"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("设计人_资质类别和等级")
            else:
                factors_error["设计人_资质类别和等级"] = "第三部分_1.1 词语定义:设计人_资质类别和等级未填写完整"

                addRemarkInDoc(word, document, "1.1.2.2 设计人：", f"第三部分_1.1 词语定义:设计人_资质类别和等级未填写完整")
        except:
            missObject += "第三部分_1.1 词语定义:要素“设计人_资质类别和等级”缺失\n"

        try:
            match = '.*设计人：\n名    称：.*\n资质类别和等级：.*\n联系电话：(.*)\n电子信箱：.*\n通信地址：.*\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["设计人_联系电话"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("设计人_联系电话")
            else:
                factors_error["设计人_联系电话"] = "第三部分_1.1 词语定义:设计人_联系电话未填写完整"
                addRemarkInDoc(word, document, "1.1.2.2 设计人：", f"第三部分_1.1 词语定义:设计人_联系电话未填写完整")
        except:
            missObject += "第三部分_1.1 词语定义:要素“设计人_联系电话”缺失\n"

        try:
            match = '.*设计人：\n名    称：.*\n资质类别和等级：.*\n联系电话：.*\n电子信箱：(.*)\n通信地址：.*\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["设计人_电子信箱"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("设计人_电子信箱")
            else:
                factors_error["设计人_电子信箱"] = "第三部分_1.1 词语定义:设计人_电子信箱未填写完整"

                addRemarkInDoc(word, document, "1.1.2.2 设计人：", f"第三部分_1.1 词语定义:设计人_电子信箱未填写完整")
        except:
            missObject += "第三部分_1.1 词语定义:要素“设计人_电子信箱”缺失\n"

        try:
            match = '.*设计人：\n名    称：.*\n资质类别和等级：.*\n联系电话：.*\n电子信箱：.*\n通信地址：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["设计人_通信地址"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("设计人_通信地址")
            else:
                factors_error["设计人_通信地址"] = "第三部分_1.1 词语定义:设计人_通信地址未填写完整"
                factors_error["设计人_通信地址"] = "第三部分_1.1 词语定义:设计人_通信地址未填写完整"
                addRemarkInDoc(word, document, "1.1.2.2 设计人：", f"第三部分_1.1 词语定义:设计人_通信地址未填写完整")
        except:
            missObject += "第三部分_1.1 词语定义:要素“设计人_通信地址”缺失\n"

        addRemarkInDoc(word, document, "1.1.2 合同当事人及其他相关方", f"第三部分_1.1 词语定义:请核实该项内容与《监理合同》、《设计合同》约定一致")

        # 1.1.3
        try:
            match = '1.1.3.1 作为施工现场组成部分的其他场所包括：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["1.1.3.1"] = factor
            if factor != "" and factor != "；" and factor != "。":

                if factor != '无':
                    factors_error["1.1.3.1"] = "第三部分_1.1 词语定义:请了解组成施工现场的其他场所的位置、是否已具备施工条件等情况，并将该等情况详细填写）"
                    addRemarkInDoc(word, document, "1.1.3.1",
                                   f"第三部分_1.1 词语定义:请了解组成施工现场的其他场所的位置、是否已具备施工条件等情况，并将该等情况详细填写）")
                else:
                    factors_ok.append("1.1.3.1")
            else:
                factors_error["1.1.3.1"] = "第三部分_1.1 词语定义:工程和设备未填写完整"

                addRemarkInDoc(word, document, "1.1.3.1", f"第三部分_1.1 词语定义:工程和设备未填写完整")
        except:
            missObject += "第三部分_1.1 词语定义:要素“1.1.3.1”缺失\n"

        try:
            match = '1.1.3.2 永久占地包括：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["1.1.3.2"] = factor
            if factor != "" and factor != "；" and factor != "。":

                if factor != '指为实施本合同工程而需要的一切永久占用的土地':
                    factors_error["1.1.3.2"] = "第三部分_1.1 词语定义:请了解永久占地的详细情况，并将该等情况详细填写"
                    addRemarkInDoc(word, document, "1.1.3.2", f"第三部分_1.1 词语定义:请了解永久占地的详细情况，并将该等情况详细填写")
                else:
                    factors_ok.append("1.1.3.2")
            else:
                factors_error["1.1.3.2"] = "第三部分_1.1 词语定义:1.1.3.2未填写完整"

                addRemarkInDoc(word, document, "1.1.3.2", f"第三部分_1.1 词语定义:1.1.3.2未填写完整")
        except:
            missObject += "第三部分_1.1 词语定义:要素“1.1.3.2”缺失\n"

        try:
            match = '1.1.3.3 临时占地包括：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["1.1.3.3"] = factor
            if factor != "" and factor != "；" and factor != "。":

                if factor != '指为实施本合同工程而需要的一切临时占用的土地，包括施工所用的临时支线、便道、便桥和现场的临时出入通道，以及生产（办公）、生活等临时设施用地等':
                    factors_error["1.1.3.3"] = "第三部分_1.1 词语定义:请了解临时占地的详细情况，并将该等情况详细填写"
                    addRemarkInDoc(word, document, "1.1.3.3", f"第三部分_1.1 词语定义:请了解临时占地的详细情况，并将该等情况详细填写")
                else:
                    factors_ok.append("1.1.3.3")
            else:
                factors_error["1.1.3.3"] = "第三部分_1.1 词语定义:1.1.3.3未填写完整"

                addRemarkInDoc(word, document, "1.1.3.3", f"第三部分_1.1 词语定义:1.1.3.3未填写完整")
        except:
            missObject += "第三部分_1.1 词语定义:第三部分_1.1 词语定义:要素“1.1.3.3”缺失\n"

        # 1.3
        try:
            match = '适用于合同的其他规范性文件：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["1.3法律"] = factor
            if factor != "" and factor != "；" and factor != "。":

                if factor != '无':
                    factors_error["1.3法律"] = "第三部分_1.3法律 :请详细了解所涉及的其他规范性文件，并填写详细名称"
                    addRemarkInDoc(word, document, "1.3法律", f"第三部分_1.3法律 :请详细了解所涉及的其他规范性文件，并填写详细名称")
                else:
                    factors_ok.append("1.3法律")
            else:
                factors_error["1.3法律"] = "第三部分_1.3法律 :适用于合同的其他规范性文件未填写完整"

                addRemarkInDoc(word, document, "1.3法律", f"第三部分_1.3法律 :适用于合同的其他规范性文件未填写完整")
        except:
            missObject += "第三部分_1.3法律 :要素“1.3法律”缺失\n"

        # 3.1.4
        try:
            match = '1.4.1适用于工程的标准规范包括：((.|\n)*)1.4.2 发包*'
            factor = re.findall(match, text)[0][0].replace(" ", "").replace("；", "").replace("。", "").replace("；",
                                                                                                              "").replace(
                "。", "").replace('\n', '')
            factors["1.4.1适用于工程的标准规范包括"] = factor
            if factor != "" and factor != "；" and factor != "。":

                if factor != '（1）本招标工程遵守设计图纸明确的技术规范；（2）执行国家现行的施工、质量检测及验收规范；（3' \
                             '）依照设计文件的要求，本招标工程项目的施工、材料、设备须达到现行中华人民共和国及省、自治区、直辖市或行业的现行工程建设标准和规范的要求。'.replace("；",
                                                                                                            "").replace(
                    "。",
                    ""):
                    factors_error["1.4.1适用于工程的标准规范包括"] = "第三部分_1.4 标准和规范 :请详细了解标准、规范的内容，并填写详细名称"
                    addRemarkInDoc(word, document, "1.4.1适用于工程的标准规范包括：", f"第三部分_1.4 标准和规范 :请详细了解标准、规范的内容，并填写详细名称")
                else:
                    factors_ok.append("1.4.1适用于工程的标准规范包括")
            else:
                factors_error["1.4.1适用于工程的标准规范包括"] = "第三部分_1.4 标准和规范 :适用于工程的标准规范未填写完整"

                addRemarkInDoc(word, document, "1.4.1适用于工程的标准规范包括：", f"第三部分_1.4 标准和规范 :适用于工程的标准规范未填写完整")
        except:
            missObject += "要素“第三部分_1.4 标准和规范 :1.4.1适用于工程的标准规范包括”缺失\n"

        try:
            match = '1.4.2 发包人提供国外标准、规范的名称：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["发包人提供国外标准、规范的名称"] = factor
            if factor != "" and factor != "；" and factor != "。":

                if factor != '无':
                    factors_error["发包人提供国外标准、规范的名称"] = "第三部分_1.4 标准和规范 :请详细了解标准、规范的内容，并填写详细名称"
                    addRemarkInDoc(word, document, "1.4.2 发包人提供国外标准、规范的名称", f"第三部分_1.4 标准和规范 :请详细了解标准、规范的内容，并填写详细名称")
                else:
                    factors_ok.append("发包人提供国外标准、规范的名称")
            else:
                factors_error["发包人提供国外标准、规范的名称"] = "第三部分_1.4 标准和规范 :发包人提供国外标准、规范的名称未填写完整"

                addRemarkInDoc(word, document, "1.4.2 发包人提供国外标准、规范的名称", f"第三部分_1.4 标准和规范 :发包人提供国外标准、规范的名称未填写完整")
        except:
            missObject += "第三部分_1.4 标准和规范 :要素“发包人提供国外标准、规范的名称”缺失\n"

        try:
            match = '.*发包人提供国外标准、规范的份数：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["发包人提供国外标准、规范的份数"] = factor
            if factor != "" and factor != "；" and factor != "。":

                if factor != '无':
                    factors_error["发包人提供国外标准、规范的份数"] = "第三部分_1.4 标准和规范 :请详细了解标准、规范的内容，并填写详细名称"
                    addRemarkInDoc(word, document, "发包人提供国外标准、规范的份数", f"第三部分_1.4 标准和规范 :请详细了解标准、规范的内容，并填写详细名称")
                else:
                    factors_ok.append("发包人提供国外标准、规范的份数")
            else:
                factors_error["发包人提供国外标准、规范的份数"] = "第三部分_1.4 标准和规范 :发包人提供国外标准、规范的份数未填写完整"

                addRemarkInDoc(word, document, "发包人提供国外标准、规范的份数", f"第三部分_1.4 标准和规范 :发包人提供国外标准、规范的份数未填写完整")
        except:
            missObject += "第三部分_1.4 标准和规范 :要素“发包人提供国外标准、规范的份数”缺失\n"

        try:
            match = '.*发包人提供国外标准、规范的名称：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["发包人提供国外标准、规范的名称"] = factor
            if factor != "" and factor != "；" and factor != "。":

                if factor != '无':
                    factors_error["发包人提供国外标准、规范的名称"] = "第三部分_1.4 标准和规范 :请详细了解标准、规范的内容，并填写详细名称"
                    addRemarkInDoc(word, document, "发包人提供国外标准、规范的名称", f"第三部分_1.4 标准和规范 :请详细了解标准、规范的内容，并填写详细名称")
                else:
                    factors_ok.append("发包人提供国外标准、规范的名称")
            else:
                factors_error["发包人提供国外标准、规范的名称"] = "第三部分_1.4 标准和规范 :发包人提供国外标准、规范的名称未填写完整"

                addRemarkInDoc(word, document, "发包人提供国外标准、规范的名称", f"第三部分_1.4 标准和规范 :发包人提供国外标准、规范的名称未填写完整")
        except:
            missObject += "第三部分_1.4 标准和规范 :要素“发包人提供国外标准、规范的名称”缺失\n"

        try:
            match = '1.4.3发包人对工程的技术标准和功能要求的特殊要求：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["发包人对工程的技术标准和功能要求的特殊要求"] = factor
            if factor != "" and factor != "；" and factor != "。":

                if factor != '无':
                    factors_error["发包人对工程的技术标准和功能要求的特殊要求"] = "第三部分_1.4 标准和规范 :请详细了解标准、规范的内容，并填写详细名称"
                    addRemarkInDoc(word, document, "发包人对工程的技术标准和功能要求的特殊要求", f"第三部分_1.4 标准和规范 :请详细了解标准、规范的内容，并填写详细名称")
                else:
                    factors_ok.append("发包人对工程的技术标准和功能要求的特殊要求")
            else:
                factors_error["发包人对工程的技术标准和功能要求的特殊要求"] = "第三部分_1.4 标准和规范 :发包人对工程的技术标准和功能要求的特殊要求未填写完整"

                addRemarkInDoc(word, document, "发包人对工程的技术标准和功能要求的特殊要求", f"第三部分_1.4 标准和规范 :发包人对工程的技术标准和功能要求的特殊要求未填写完整")
        except:
            missObject += "要素“第三部分_1.4 标准和规范 :发包人对工程的技术标准和功能要求的特殊要求”缺失\n"

        # 3.1.5
        try:
            match = '合同文件组成及优先顺序为：((.|\n)*)\n1.6 图纸和承包人文件.*'
            factor = re.findall(match, text)[0][0].replace(" ", "").replace("。", "").replace("；", "").replace("\t",
                                                                                                              "").replace(
                "。",
                "").replace(
                "\n", "")
            factors["合同文件的优先顺序"] = factor
            if factor != "" and factor != "；" and factor != "。":
                x = '(1)合同协议书；(2)中标通知书；(3)投标函及投标函附录；(4)专用合同条款及其附件；(' \
                    '5)通用合同条款（以住房城乡建设部工商总局印发的建设工程施工合同（示范文本）（GF-2017-0201）中通用合同条款为准）；(6)技术标准和要求；（7）施工图纸及其他设计文件；(' \
                    '8)已标价工程量清单；(9)其他合同文件'.replace(" ", "").replace("。", "").replace("；", "").replace("\t", "").replace(
                    "。", "").replace(
                    '(',
                    '（').replace(
                    ')', '）')
                if factor != x:
                    factors_error["合同文件的优先顺序"] = "第三部分_1.5 合同文件的优先顺序:请了解组成施工现场的其他场所的位置、是否已具备施工条件等情况，并将该等情况详细填写）"
                    addRemarkInDoc(word, document, "合同文件的优先顺序",
                                   f"第三部分_1.5 合同文件的优先顺序:请了解组成施工现场的其他场所的位置、是否已具备施工条件等情况，并将该等情况详细填写）")
                else:
                    factors_ok.append("合同文件的优先顺序")
            else:
                factors_error["合同文件的优先顺序"] = "第三部分_1.5 合同文件的优先顺序:合同文件的优先顺序未填写完整"
                addRemarkInDoc(word, document, "合同文件的优先顺序", f"第三部分_1.5 合同文件的优先顺序:合同文件的优先顺序未填写完整")
        except:
            missObject += "第三部分_1.5 合同文件的优先顺序:要素“合同文件的优先顺序”缺失\n"

        # 3.1.6
        try:
            match = '发包人向承包人提供图纸的期限：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["发包人向承包人提供图纸的期限"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '合同正式签署生效后7天内':
                    factors_error[
                        "发包人向承包人提供图纸的期限"] = "第三部分_1.6 图纸和承包人文件:请核实施工图、勘察报告、图审报告三报告的提供主体。若需发包人提供的图纸超出该三报告范围，请核实超出范围报告的费用承担主体、由发包人提供的理由并在该条中对该情况进行批注"
                    addRemarkInDoc(word, document, "发包人向承包人提供图纸的期限",
                                   f"第三部分_1.6 图纸和承包人文件:请核实施工图、勘察报告、图审报告三报告的提供主体。若需发包人提供的图纸超出该三报告范围，请核实超出范围报告的费用承担主体、由发包人提供的理由并在该条中对该情况进行批注")
                else:
                    factors_ok.append("发包人向承包人提供图纸的期限")
            else:
                factors_error["发包人向承包人提供图纸的期限"] = "第三部分_1.6 图纸和承包人文件:发包人向承包人提供图纸的期限未填写完整"
                addRemarkInDoc(word, document, "发包人向承包人提供图纸的期限", f"第三部分_1.6 图纸和承包人文件:发包人向承包人提供图纸的期限未填写完整")
        except:
            missObject += "第三部分_1.6 图纸和承包人文件:要素“发包人向承包人提供图纸的期限”缺失\n"

        try:
            match = '发包人向承包人提供图纸的数量：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["发包人向承包人提供图纸的数量"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '2份，承包人需要增加图纸套数的，由承包人自行复制并承担相应费用':
                    factors_error[
                        "发包人向承包人提供图纸的数量"] = "第三部分_1.6 图纸和承包人文件:请核实施工图、勘察报告、图审报告三报告的提供主体。若需发包人提供的图纸超出该三报告范围，请核实超出范围报告的费用承担主体、由发包人提供的理由并在该条中对该情况进行批注"
                    addRemarkInDoc(word, document, "发包人向承包人提供图纸的数量",
                                   f"第三部分_1.6 图纸和承包人文件:请核实施工图、勘察报告、图审报告三报告的提供主体。若需发包人提供的图纸超出该三报告范围，请核实超出范围报告的费用承担主体、由发包人提供的理由并在该条中对该情况进行批注")
                else:
                    factors_ok.append("发包人向承包人提供图纸的数量")
            else:
                factors_error["发包人向承包人提供图纸的数量"] = "第三部分_1.6 图纸和承包人文件:发包人向承包人提供图纸的数量未填写完整"
                addRemarkInDoc(word, document, "发包人向承包人提供图纸的数量", f"第三部分_1.6 图纸和承包人文件:发包人向承包人提供图纸的数量未填写完整")
        except:
            missObject += "第三部分_1.6 图纸和承包人文件:要素“发包人向承包人提供图纸的数量”缺失\n"

        try:
            match = '发包人向承包人提供图纸的内容：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["发包人向承包人提供图纸的内容"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '施工图、勘察报告、图审报告':
                    factors_error[
                        "发包人向承包人提供图纸的内容"] = "第三部分_1.6 图纸和承包人文件:请核实施工图、勘察报告、图审报告三报告的提供主体。若需发包人提供的图纸超出该三报告范围，请核实超出范围报告的费用承担主体、由发包人提供的理由并在该条中对该情况进行批注"
                    addRemarkInDoc(word, document, "发包人向承包人提供图纸的内容",
                                   f"第三部分_1.6 图纸和承包人文件:请核实施工图、勘察报告、图审报告三报告的提供主体。若需发包人提供的图纸超出该三报告范围，请核实超出范围报告的费用承担主体、由发包人提供的理由并在该条中对该情况进行批注")
                else:
                    factors_ok.append("发包人向承包人提供图纸的内容")
            else:
                factors_error["发包人向承包人提供图纸的内容"] = "第三部分_1.6 图纸和承包人文件:发包人向承包人提供图纸的内容未填写完整"
                addRemarkInDoc(word, document, "发包人向承包人提供图纸的内容", f"第三部分_1.6 图纸和承包人文件:发包人向承包人提供图纸的内容未填写完整")
        except:
            missObject += "第三部分_1.6 图纸和承包人文件:要素“发包人向承包人提供图纸的内容”缺失\n"

        try:
            match = '需要由承包人提供的文件，包括：\n((.|\n)*)承包人提供的文件的期限为*'
            factor = re.findall(match, text)[0][0].replace(" ", "").replace("；", "").replace("。", "").replace("；",
                                                                                                              "").replace(
                "。", "").replace("\n", "")
            factors["需要由承包人提供的文件"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '1、招标文件要求承包人提供的相关文件；2、应由承包人提供的必要的深化设计图、加工图和大样图等'.replace("；", "").replace("。", ""):
                    factors_error["需要由承包人提供的文件"] = "第三部分_1.6 图纸和承包人文件:请了解与推荐条款不一致的原因并将相关原因进行批注"
                    addRemarkInDoc(word, document, "需要由承包人提供的文件，包括：", f"第三部分_1.6 图纸和承包人文件:请了解与推荐条款不一致的原因并将相关原因进行批注")
                else:
                    factors_ok.append("需要由承包人提供的文件")
            else:
                factors_error["需要由承包人提供的文件"] = "第三部分_1.6 图纸和承包人文件:需要由承包人提供的文件未填写完整"
                addRemarkInDoc(word, document, "需要由承包人提供的文件，包括：", f"第三部分_1.6 图纸和承包人文件:需要由承包人提供的文件未填写完整")
        except:
            missObject += "第三部分_1.6 图纸和承包人文件:要素“需要由承包人提供的文件”缺失\n"

        try:
            match = '承包人提供的文件的期限为：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["承包人提供的文件的期限"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '根据项目进度按发包人和监理人要求':
                    factors_error["承包人提供的文件的期限"] = "第三部分_1.6 图纸和承包人文件:请了解与推荐条款不一致的原因并将相关原因进行批注"
                    addRemarkInDoc(word, document, "承包人提供的文件的期限为：", f"第三部分_1.6 图纸和承包人文件:请了解与推荐条款不一致的原因并将相关原因进行批注")
                else:
                    factors_ok.append("承包人提供的文件的期限")
            else:
                factors_error["承包人提供的文件的期限"] = "第三部分_1.6 图纸和承包人文件:承包人提供的文件的期限未填写完整"
                addRemarkInDoc(word, document, "承包人提供的文件的期限为：", f"第三部分_1.6 图纸和承包人文件:承包人提供的文件的期限未填写完整")
        except:
            missObject += "第三部分_1.6 图纸和承包人文件:要素“承包人提供的文件的期限”缺失\n"

        try:
            match = '承包人提供的文件的数量为：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["承包人提供的文件的数量"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '根据现场所需数量提供':
                    factors_error["承包人提供的文件的数量"] = "第三部分_1.6 图纸和承包人文件:请了解与推荐条款不一致的原因并将相关原因进行批注"
                    addRemarkInDoc(word, document, "承包人提供的文件的数量为：", f"第三部分_1.6 图纸和承包人文件:请了解与推荐条款不一致的原因并将相关原因进行批注")
                else:
                    factors_ok.append("承包人提供的文件的数量")
            else:
                factors_error["承包人提供的文件的数量"] = "第三部分_1.6 图纸和承包人文件:承包人提供的文件的数量未填写完整"
                addRemarkInDoc(word, document, "承包人提供的文件的数量为：", f"第三部分_1.6 图纸和承包人文件:承包人提供的文件的数量未填写完整")
        except:
            missObject += "第三部分_1.6 图纸和承包人文件:要素“承包人提供的文件的数量”缺失\n"

        try:
            match = '承包人提供的文件的形式为：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["承包人提供的文件的形式"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '书面形式':
                    factors_error["承包人提供的文件的形式"] = "第三部分_1.6 图纸和承包人文件:请了解与推荐条款不一致的原因并将相关原因进行批注"
                    addRemarkInDoc(word, document, "承包人提供的文件的形式为：", f"第三部分_1.6 图纸和承包人文件:请了解与推荐条款不一致的原因并将相关原因进行批注")
                else:
                    factors_ok.append("承包人提供的文件的形式")
            else:
                factors_error["承包人提供的文件的形式"] = "第三部分_1.6 图纸和承包人文件:承包人提供的文件的形式未填写完整"
                addRemarkInDoc(word, document, "承包人提供的文件的形式为：", f"第三部分_1.6 图纸和承包人文件:承包人提供的文件的形式未填写完整")
        except:
            missObject += "第三部分_1.6 图纸和承包人文件:要素“承包人提供的文件的形式”缺失\n"

        try:
            match = '发包人审批承包人文件的期限：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["发包人审批承包人文件的期限"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '收到承包人提供的文件（盖公章）后14天内':
                    factors_error["发包人审批承包人文件的期限"] = "第三部分_1.6 图纸和承包人文件:了解与推荐条款不一致的原因并将相关原因进行批注"
                    addRemarkInDoc(word, document, "发包人审批承包人文件的期限：", f"第三部分_1.6 图纸和承包人文件:了解与推荐条款不一致的原因并将相关原因进行批注")
                else:
                    factors_ok.append("发包人审批承包人文件的期限")
            else:
                factors_error["发包人审批承包人文件的期限"] = "第三部分_1.6 图纸和承包人文件:发包人审批承包人文件的期限未填写完整"
                addRemarkInDoc(word, document, "发包人审批承包人文件的期限：", f"第三部分_1.6 图纸和承包人文件:发包人审批承包人文件的期限未填写完整")
        except:
            missObject += "第三部分_1.6 图纸和承包人文件:要素“发包人审批承包人文件的期限”缺失\n"

        try:
            match = '关于现场图纸准备的约定：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["关于现场图纸准备的约定"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '承包人自行在现场准备施工图2套（至少1套为原件）':
                    factors_error["关于现场图纸准备的约定"] = "第三部分_1.6 图纸和承包人文件:承包人自行在现场准备施工图2套（至少1套为原件）"
                    addRemarkInDoc(word, document, "关于现场图纸准备的约定：", f"第三部分_1.6 图纸和承包人文件:承包人自行在现场准备施工图2套（至少1套为原件）")
                else:
                    factors_ok.append("关于现场图纸准备的约定")
            else:
                factors_error["关于现场图纸准备的约定"] = "第三部分_1.6 图纸和承包人文件:关于现场图纸准备的约定未填写完整"
                addRemarkInDoc(word, document, "关于现场图纸准备的约定：", f"第三部分_1.6 图纸和承包人文件:关于现场图纸准备的约定未填写完整")
        except:
            missObject += "第三部分_1.6 图纸和承包人文件:要素“关于现场图纸准备的约定”缺失\n"

        # 3.1.7
        try:
            match = '1.7.1发包人和承包人应当在(.*)前将与合同有关的通知.*'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["书面函件送达时间"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if get_strtime(factor) == False:
                    factors_error["书面函件送达时间"] = "第三部分_1.7 联络:书面函件送达时间错误"
                    addRemarkInDoc(word, document, "1.7.1发包人和承包人应当在", f"第三部分_1.7 联络:书面函件送达时间错误")
                else:
                    factors_ok.append("书面函件送达时间")
            else:
                factors_error["书面函件送达时间"] = "第三部分_1.7 联络:书面函件送达时间未填写完整"
                addRemarkInDoc(word, document, "书面函件送达时间：", f"第三部分_1.7 联络:书面函件送达时间未填写完整")
        except:
            missObject += "第三部分_1.7 联络:要素“书面函件送达时间”缺失\n"

        try:
            match = '1.7.2 发包人接收文件的地点：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["发包人接收文件的地点"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '发包人施工场地管理机构的办公地点':
                    factors_error["发包人接收文件的地点"] = "第三部分_1.7 联络:请了解与推荐条款不一致的原因并将相关原因进行批注"
                    addRemarkInDoc(word, document, "发包人接收文件的地点", f"第三部分_1.7 联络:请了解与推荐条款不一致的原因并将相关原因进行批注")
                else:
                    factors_ok.append("发包人接收文件的地点")
            else:
                factors_error["发包人接收文件的地点"] = "第三部分_1.7 联络:发包人接收文件的地点未填写完整"
                addRemarkInDoc(word, document, "发包人接收文件的地点", f"第三部分_1.7 联络:发包人接收文件的地点未填写完整")
        except:
            missObject += "第三部分_1.7 联络:要素“发包人接收文件的地点”缺失\n"

        try:
            match = '发包人指定的接收人为：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["发包人指定的接收人"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("发包人指定的接收人")
            else:
                factors_error["发包人指定的接收人"] = "第三部分_1.7 联络:发包人指定的接收人未填写完整"
                addRemarkInDoc(word, document, "发包人指定的接收人为：", f"第三部分_1.7 联络:发包人指定的接收人未填写完整")
        except:
            missObject += "第三部分_1.7 联络:要素“发包人指定的接收人”缺失\n"

        try:
            match = '承包人接收文件的地点：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["承包人接收文件的地点"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '承包人施工场地管理机构的办公地点'.replace(" ", "").replace("；", "").replace("。", ""):
                    factors_error["承包人接收文件的地点"] = "第三部分_1.7 联络:请了解与推荐条款不一致的原因并将相关原因进行批注"
                    addRemarkInDoc(word, document, "承包人接收文件的地点：", f"第三部分_1.7 联络:请了解与推荐条款不一致的原因并将相关原因进行批注")
                else:
                    factors_ok.append("承包人接收文件的地点")
            else:
                factors_error["承包人接收文件的地点"] = "第三部分_1.7 联络:承包人接收文件的地点未填写完整"
                addRemarkInDoc(word, document, "承包人接收文件的地点：", f"第三部分_1.7 联络:承包人接收文件的地点未填写完整")
        except:
            missObject += "第三部分_1.7 联络:要素“承包人接收文件的地点”缺失\n"

        try:
            match = '承包人指定的接收人为：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["承包人指定的接收人"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("承包人指定的接收人")
            else:
                factors_error["承包人指定的接收人"] = "第三部分_1.7 联络:承包人指定的接收人未填写完整"
                addRemarkInDoc(word, document, "承包人指定的接收人为：", f"第三部分_1.7 联络:承包人指定的接收人未填写完整")
        except:
            missObject += "第三部分_1.7 联络:要素“承包人指定的接收人”缺失\n"

        try:
            match = '监理人接收文件的地点：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["监理人接收文件的地点"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '监理人施工场地管理机构的办公地点'.replace(" ", "").replace("；", "").replace("。", ""):
                    factors_error["监理人接收文件的地点"] = "第三部分_1.7 联络:请了解与推荐条款不一致的原因并将相关原因进行批注"
                    addRemarkInDoc(word, document, "监理人接收文件的地点：", f"第三部分_1.7 联络:请了解与推荐条款不一致的原因并将相关原因进行批注")
                else:
                    factors_ok.append("监理人接收文件的地点")
            else:
                factors_error["监理人接收文件的地点"] = "第三部分_1.7 联络:监理人接收文件的地点未填写完整"
                addRemarkInDoc(word, document, "监理人接收文件的地点：", f"第三部分_1.7 联络:监理人接收文件的地点未填写完整")
        except:
            missObject += "第三部分_1.7 联络:要素“监理人接收文件的地点”缺失\n"

        try:
            match = '监理人指定的接收人为：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["监理人指定的接收人"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("监理人指定的接收人")
            else:
                factors_error["监理人指定的接收人"] = "第三部分_1.7 联络:监理人指定的接收人未填写完整"
                addRemarkInDoc(word, document, "监理人指定的接收人为：", f"第三部分_1.7 联络:监理人指定的接收人未填写完整")
        except:
            missObject += "第三部分_1.7 联络:要素“监理人指定的接收人”缺失\n"

        # 3.1.8
        try:
            match = '关于出入现场的权利的约定：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["关于出入现场的权利的约定"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if '自行办理和完善出入现场的批准手续并承担费用' in factor or '自行办理、完善出入现场的批准手续并承担费用' in factor or '自行完善、办理出入现场的批准手续并承担费用' in factor or '自行完善和办理出入现场的批准手续并承担费用' in factor:
                    factors_ok.append("关于出入现场的权利的约定")
                else:
                    factors_error["关于出入现场的权利的约定"] = "第三部分_1.8 交通运输:请核实该条权利义务，原则上由承包人自行办理、完善出入现场的批准手续并承担费用"
                    addRemarkInDoc(word, document, "关于出入现场的权利的约定",
                                   f"第三部分_1.8 交通运输:请核实该条权利义务，原则上由承包人自行办理、完善出入现场的批准手续并承担费用")
            else:
                factors_error["关于出入现场的权利的约定"] = "第三部分_1.8 交通运输:关于出入现场的权利的约定未填写完整"
                addRemarkInDoc(word, document, "关于出入现场的权利的约定", f"第三部分_1.8 交通运输:关于出入现场的权利的约定未填写完整")
        except:
            missObject += "第三部分_1.8 交通运输:要素“关于出入现场的权利的约定”缺失\n"

        try:
            match = '关于场外交通和场内交通的边界的约定：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["关于场外交通和场内交通的边界的约定"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '以业主提供项目相关标段总图为界'.replace(" ", "").replace("；", "").replace("。", ""):
                    factors_error["关于场外交通和场内交通的边界的约定"] = "第三部分_1.8 交通运输:请核实本条约定的场外与场内交通边界与业主提供的标段总图相关内容是否一致"
                    addRemarkInDoc(word, document, "关于场外交通和场内交通的边界的约定：",
                                   f"第三部分_1.8 交通运输:请核实本条约定的场外与场内交通边界与业主提供的标段总图相关内容是否一致")
                else:
                    factors_ok.append("关于场外交通和场内交通的边界的约定")
            else:
                factors_error["关于场外交通和场内交通的边界的约定"] = "第三部分_1.8 交通运输:关于场外交通和场内交通的边界的约定未填写完整"
                addRemarkInDoc(word, document, "关于场外交通和场内交通的边界的约定：", f"第三部分_1.8 交通运输:关于场外交通和场内交通的边界的约定未填写完整")
        except:
            missObject += "第三部分_1.8 交通运输:要素“关于场外交通和场内交通的边界的约定”缺失\n"

        try:
            match = '关于发包人向承包人免费提供满足工程施工需要的场内道路和交通设施的约定：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["关于发包人向承包人免费提供满足工程施工需要的场内道路和交通设施的约定"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '施工所需的场内临时道路和交通设施的修建、维护、养护和管理均由承包人负责承担，相关费用由承包人承担'.replace(" ", "").replace("；",
                                                                                                         "").replace(
                    "。", ""):
                    factors_error[
                        "关于发包人向承包人免费提供满足工程施工需要的场内道路和交通设施的约定"] = "第三部分_1.8 交通运输:核实填写内容与双方约定是否一致。原则上应约定 ‘施工所需的场内临时道路和交通设施的修建、维护、养护和管理均由承包人负责承担，相关费用由承包人承担‘"
                    addRemarkInDoc(word, document, "关于发包人向承包人免费提供满足工程施工需要的场内道路和交通设施的约定：",
                                   f"第三部分_1.8 交通运输:核实填写内容与双方约定是否一致。原则上应约定 ‘施工所需的场内临时道路和交通设施的修建、维护、养护和管理均由承包人负责承担，相关费用由承包人承担‘")
                else:
                    factors_ok.append("关于发包人向承包人免费提供满足工程施工需要的场内道路和交通设施的约定")
            else:
                factors_error[
                    "关于发包人向承包人免费提供满足工程施工需要的场内道路和交通设施的约定"] = "第三部分_1.8 交通运输:关于发包人向承包人免费提供满足工程施工需要的场内道路和交通设施的约定未填写完整"
                addRemarkInDoc(word, document, "关于发包人向承包人免费提供满足工程施工需要的场内道路和交通设施的约定：",
                               f"第三部分_1.8 交通运输:关于发包人向承包人免费提供满足工程施工需要的场内道路和交通设施的约定未填写完整")
        except:
            missObject += "要素“第三部分_1.8 交通运输:关于发包人向承包人免费提供满足工程施工需要的场内道路和交通设施的约定”缺失\n"

        try:
            match = '运输超大件或超重件所需的道路和桥梁临时加固改造费用和其他有关费用由(.*)承担。\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["超大件和超重件的运输承担人"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '承包人'.replace(" ", "").replace("；", "").replace("。", ""):
                    factors_error["超大件和超重件的运输承担人"] = "第三部分_1.8 交通运输:请核实填写内容与双方约定是否一致。原则上应约定由‘承包人‘承担"
                    addRemarkInDoc(word, document, "1.8.4超大件和超重件的运输", f"第三部分_1.8 交通运输:请核实填写内容与双方约定是否一致。原则上应约定由‘承包人‘承担")
                else:
                    factors_ok.append("超大件和超重件的运输承担人")
            else:
                factors_error["超大件和超重件的运输承担人"] = "第三部分_1.8 交通运输:超大件和超重件的运输承担人未填写完整"
                addRemarkInDoc(word, document, "1.8.4超大件和超重件的运输", f"第三部分_1.8 交通运输:超大件和超重件的运输承担人未填写完整")
        except:
            missObject += "第三部分_1.8 交通运输:要素“超大件和超重件的运输承担人”缺失\n"

        # 3.1.9
        try:
            match = '1.9.1关于发包人提供给承包人的图纸、发包人为实施工程自行编制或委托编制的技术规范以及反映发包人关于合同要求或其他类似性质的文件的著作权的归属：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["技术规范以及反映发包人关于合同要求或其他类似性质的文件的著作权的归属"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '发包人所有'.replace(" ", "").replace("；", "").replace("。", ""):
                    factors_error["技术规范以及反映发包人关于合同要求或其他类似性质的文件的著作权的归属"] = "第三部分_1.9 知识产权:请核实填写内容与双方约定是否一致。原则上应约定由发包人所有"
                    addRemarkInDoc(word, document,
                                   "1.9.1关于发包人提供给承包人的图纸、发包人为实施工程自行编制或委托编制的技术规范以及反映发包人关于合同要求或其他类似性质的文件的著作权的归属：",
                                   f"第三部分_1.9 知识产权:请核实填写内容与双方约定是否一致。原则上应约定由发包人所有")
                else:
                    factors_ok.append("技术规范以及反映发包人关于合同要求或其他类似性质的文件的著作权的归属")
            else:
                factors_error[
                    "技术规范以及反映发包人关于合同要求或其他类似性质的文件的著作权的归属"] = "第三部分_1.9 知识产权:技术规范以及反映发包人关于合同要求或其他类似性质的文件的著作权的归属未填写完整"
                addRemarkInDoc(word, document,
                               "1.9.1关于发包人提供给承包人的图纸、发包人为实施工程自行编制或委托编制的技术规范以及反映发包人关于合同要求或其他类似性质的文件的著作权的归属：",
                               f"第三部分_1.9 知识产权:技术规范以及反映发包人关于合同要求或其他类似性质的文件的著作权的归属未填写完整")
        except:
            missObject += "第三部分_1.9 知识产权:要素“技术规范以及反映发包人关于合同要求或其他类似性质的文件的著作权的归属”缺失\n"

        try:
            match = '关于发包人提供的上述文件的使用限制的要求：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["关于发包人提供的上述文件的使用限制的要求"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '按通用条款执行'.replace(" ", "").replace("；", "").replace("。", ""):
                    factors_error["关于发包人提供的上述文件的使用限制的要求"] = "第三部分_1.9 知识产权:请核实双方关于该项使用限制的约定"
                    addRemarkInDoc(word, document, "关于发包人提供的上述文件的使用限制的要求：", f"第三部分_1.9 知识产权:请核实双方关于该项使用限制的约定")
                else:
                    factors_ok.append("关于发包人提供的上述文件的使用限制的要求")
            else:
                factors_error["关于发包人提供的上述文件的使用限制的要求"] = "第三部分_1.9 知识产权:关于发包人提供的上述文件的使用限制的要求未填写完整"
                addRemarkInDoc(word, document, "关于发包人提供的上述文件的使用限制的要求：", f"第三部分_1.9 知识产权:关于发包人提供的上述文件的使用限制的要求未填写完整")
        except:
            missObject += "第三部分_1.9 知识产权:要素“关于发包人提供的上述文件的使用限制的要求”缺失\n"

        try:
            match = '1.9.2 关于承包人为实施工程所编制文件的著作权的归属：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["关于承包人为实施工程所编制文件的著作权的归属"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '除署名权以外的著作权属于发包人'.replace(" ", "").replace("；", "").replace("。", ""):
                    factors_error["关于承包人为实施工程所编制文件的著作权的归属"] = "第三部分_1.9 知识产权:请核实填写内容与双方约定是否一致。原则上应约定’除署名权以外的著作权属于发包人’"
                    addRemarkInDoc(word, document, "1.9.2 关于承包人为实施工程所编制文件的著作权的归属：",
                                   f"第三部分_1.9 知识产权:请核实填写内容与双方约定是否一致。原则上应约定’除署名权以外的著作权属于发包人’")
                else:
                    factors_ok.append("关于承包人为实施工程所编制文件的著作权的归属")
            else:
                factors_error["关于承包人为实施工程所编制文件的著作权的归属"] = "第三部分_1.9 知识产权:关于承包人为实施工程所编制文件的著作权的归属未填写完整"
                addRemarkInDoc(word, document, "1.9.2 关于承包人为实施工程所编制文件的著作权的归属：",
                               f"第三部分_1.9 知识产权:关于承包人为实施工程所编制文件的著作权的归属未填写完整")
        except:
            missObject += "第三部分_1.9 知识产权:要素“关于承包人为实施工程所编制文件的著作权的归属”缺失\n"

        try:
            match = '关于承包人提供的上述文件的使用限制的要求：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["关于承包人提供的上述文件的使用限制的要求"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '按通用条款执行'.replace(" ", "").replace("；", "").replace("。", ""):
                    factors_error["关于承包人提供的上述文件的使用限制的要求"] = "第三部分_1.9 知识产权:请核实双方关于该项使用限制的约定"
                    addRemarkInDoc(word, document, "关于承包人提供的上述文件的使用限制的要求：", f"第三部分_1.9 知识产权:请核实双方关于该项使用限制的约定")
                else:
                    factors_ok.append("关于承包人提供的上述文件的使用限制的要求")
            else:
                factors_error["关于承包人提供的上述文件的使用限制的要求"] = "第三部分_1.9 知识产权:关于承包人提供的上述文件的使用限制的要求未填写完整"
                addRemarkInDoc(word, document, "关于承包人提供的上述文件的使用限制的要求：", f"第三部分_1.9 知识产权:关于承包人提供的上述文件的使用限制的要求未填写完整")
        except:
            missObject += "第三部分_1.9 知识产权:要素“关于承包人提供的上述文件的使用限制的要求”缺失\n"

        try:
            match = '1.9.4 承包人在施工过程中所采用的专利、专有技术、技术秘密的使用费的承担方式：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["承包人在施工过程中所采用的专利、专有技术、技术秘密的使用费的承担方式"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '由承包人自行承担'.replace(" ", "").replace("；", "").replace("。", ""):
                    factors_error[
                        "承包人在施工过程中所采用的专利、专有技术、技术秘密的使用费的承担方式"] = "第三部分_1.9 知识产权:请核实填写内容与双方约定是否一致。原则上应约定由‘承包人‘自行承担。"
                    addRemarkInDoc(word, document, "1.9.4 承包人在施工过程中所采用的专利、专有技术、技术秘密的使用费的承担方式：",
                                   f"第三部分_1.9 知识产权:请核实填写内容与双方约定是否一致。原则上应约定由‘承包人‘自行承担。")
                else:
                    factors_ok.append("承包人在施工过程中所采用的专利、专有技术、技术秘密的使用费的承担方式")
            else:
                factors_error[
                    "承包人在施工过程中所采用的专利、专有技术、技术秘密的使用费的承担方式"] = "第三部分_1.9 知识产权:承包人在施工过程中所采用的专利、专有技术、技术秘密的使用费的承担方式未填写完整"
                addRemarkInDoc(word, document, "1.9.4 承包人在施工过程中所采用的专利、专有技术、技术秘密的使用费的承担方式：",
                               f"第三部分_1.9 知识产权:承包人在施工过程中所采用的专利、专有技术、技术秘密的使用费的承担方式未填写完整")
        except:
            missObject += "第三部分_1.9 知识产权:要素“承包人在施工过程中所采用的专利、专有技术、技术秘密的使用费的承担方式”缺失\n"

        try:
            match = '出现工程量清单错误时，是否调整合同价格：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["出现工程量清单错误时，是否调整合同价格"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '是'.replace(" ", "").replace("；", "").replace("。", ""):
                    factors_error["出现工程量清单错误时，是否调整合同价格"] = "第三部分_1.10工程量清单错误的修正:请核实双方对工程量偏差范围的具体约定"
                    addRemarkInDoc(word, document, "出现工程量清单错误时，是否调整合同价格：", f"第三部分_1.10工程量清单错误的修正:请核实双方对工程量偏差范围的具体约定")
                else:
                    factors_ok.append("出现工程量清单错误时，是否调整合同价格")
            else:
                factors_error["出现工程量清单错误时，是否调整合同价格"] = "第三部分_1.10工程量清单错误的修正:出现工程量清单错误时，是否调整合同价格未填写完整"
                addRemarkInDoc(word, document, "出现工程量清单错误时，是否调整合同价格：", f"第三部分_1.10工程量清单错误的修正:出现工程量清单错误时，是否调整合同价格未填写完整")
        except:
            missObject += "第三部分_1.10工程量清单错误的修正:要素“出现工程量清单错误时，是否调整合同价格”缺失\n"

        try:
            match = '允许调整合同价格的工程量偏差范围：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["允许调整合同价格的工程量偏差范围"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '按照承包人实际完成的工程量调整合同价格。工程量计算规则执行国家标准《建设工程工程量清单计价规范》（GB50500—2013）或其适用的修订版本。除合同另有约定外，承包人实际完成的工程量按约定的工程量计算规则和有合同约束力的图纸进行计量。'.replace(
                        " ", "").replace("；", "").replace("。", ""):
                    factors_error["允许调整合同价格的工程量偏差范围"] = "第三部分_1.10工程量清单错误的修正:请核实双方对工程量偏差范围的具体约定"
                    addRemarkInDoc(word, document, "允许调整合同价格的工程量偏差范围：", f"第三部分_1.10工程量清单错误的修正:请核实双方对工程量偏差范围的具体约定")
                else:
                    factors_ok.append("允许调整合同价格的工程量偏差范围")
            else:
                factors_error["允许调整合同价格的工程量偏差范围"] = "第三部分_1.10工程量清单错误的修正:允许调整合同价格的工程量偏差范围未填写完整"
                addRemarkInDoc(word, document, "允许调整合同价格的工程量偏差范围：", f"第三部分_1.10工程量清单错误的修正:允许调整合同价格的工程量偏差范围未填写完整")
        except:
            missObject += "第三部分_1.10工程量清单错误的修正:要素“允许调整合同价格的工程量偏差范围”缺失\n"

    # 第三部分_2
    if 1 == True:
        factors_to_inform["2.1 发包人代表"] = "请核实该条信息是否与发包人代表提供的身份证件所载内容是否一致"
        addRemarkInDoc(word, document, "2.1 发包人代表", f"第三部分_2. 发包人:请核实该条信息是否与发包人代表提供的身份证件所载内容是否一致")
        try:
            match = '2.2.1项目负责人\n姓    名：(.*)；\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["发包人代表_项目负责人_姓名"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("发包人代表_项目负责人_姓名")
            else:
                factors_error["发包人代表_项目负责人_姓名"] = "第三部分_2. 发包人:发包人代表_项目负责人_姓名未填写完整"
                addRemarkInDoc(word, document, "2.2.1项目负责人", f"第三部分_2. 发包人:发包人代表_项目负责人_姓名未填写完整")
        except:
            missObject += "第三部分_2. 发包人:要素“发包人代表_项目负责人_姓名”缺失\n"

        try:
            match = '2.2.1项目负责人\n姓    名：.*；\n身份证号：(.*)；\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["发包人代表_项目负责人_身份证"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if checkIdCard(factor) != 'ok':
                    factors_error["发包人代表_项目负责人_身份证"] = "第三部分_2. 发包人:发包人代表_身份证填写错误：" + checkIdCard(factor)
                    addRemarkInDoc(word, document, "2.2.1项目负责人", f"第三部分_2. 发包人:发包人代表_身份证填写错误：" + checkIdCard(factor))
                else:
                    factors_ok.append("发包人代表_项目负责人_身份证")
            else:
                factors_error["发包人代表_项目负责人_身份证"] = "第三部分_2. 发包人:发包人代表_项目负责人_身份证未填写完整"
                addRemarkInDoc(word, document, "2.2.1项目负责人", f"第三部分_2. 发包人:发包人代表_项目负责人_身份证未填写完整")
        except:
            missObject += "第三部分_2. 发包人:要素“发包人代表_项目负责人_身份证”缺失\n"

        try:
            match = '2.2.1项目负责人\n姓    名：.*；\n身份证号：.*；\n职    务：(.*)；\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["发包人代表_项目负责人_职务"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("发包人代表_项目负责人_职务")
            else:
                factors_error["发包人代表_项目负责人_职务"] = "第三部分_2. 发包人:发包人代表_项目负责人_职务未填写完整"
                addRemarkInDoc(word, document, "2.2.1项目负责人", f"第三部分_2. 发包人:发包人代表_项目负责人_职务未填写完整")
        except:
            missObject += "第三部分_2. 发包人:要素“发包人代表_项目负责人_职务”缺失\n"

        try:
            match = '2.2.1项目负责人\n姓    名：.*；\n身份证号：.*；\n职    务：.*；\n联系电话：(.*)；\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["发包人代表_项目负责人_电话"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if isTelPhoneNumber(factor) == 'Error':
                    factors_error["发包人代表_项目负责人_电话"] = "第三部分_2. 发包人:电话填写错误"
                    addRemarkInDoc(word, document, "2.2.1项目负责人", f"第三部分_2. 发包人:电话填写错误")
                else:
                    factors_ok.append("发包人代表_项目负责人_电话")
            else:
                factors_error["发包人代表_项目负责人_电话"] = "第三部分_2. 发包人:发包人代表_项目负责人_电话未填写完整"
                addRemarkInDoc(word, document, "2.2.1项目负责人", f"第三部分_2. 发包人:发包人代表_项目负责人_电话未填写完整")
        except:
            missObject += "第三部分_2. 发包人:要素“发包人代表_项目负责人_电话”缺失\n"

        try:
            match = '2.2.1项目负责人\n姓    名：.*；\n身份证号：.*；\n职    务：.*；\n联系电话：.*；\n电子信箱：(.*)；\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["发包人代表_项目负责人_电子信箱"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("发包人代表_项目负责人_电子信箱")
            else:
                factors_error["发包人代表_项目负责人_电子信箱"] = "第三部分_2. 发包人:发包人代表_项目负责人_电子信箱未填写完整"
                addRemarkInDoc(word, document, "2.2.1项目负责人", f"第三部分_2. 发包人:发包人代表_项目负责人_电子信箱未填写完整")
        except:
            missObject += "第三部分_2. 发包人:要素“发包人代表_项目负责人_电子信箱”缺失\n"

        try:
            match = '2.2.1项目负责人\n姓    名：.*；\n身份证号：.*；\n职    务：.*；\n联系电话：.*；\n电子信箱：.*；\n通信地址：(.*)。\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["发包人代表_项目负责人_通信地址"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("发包人代表_项目负责人_通信地址")
            else:
                factors_error["发包人代表_项目负责人_通信地址"] = "第三部分_2. 发包人:发包人代表_项目负责人_通信地址未填写完整"
                addRemarkInDoc(word, document, "2.2.1项目负责人", f"第三部分_2. 发包人:发包人代表_项目负责人_通信地址未填写完整")
        except:
            missObject += "第三部分_2. 发包人:要素“发包人代表_项目负责人_通信地址”缺失\n"

        try:
            match = '2.2.2现场代表\n姓    名：(.*)；\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["发包人代表_现场代表_姓名"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("发包人代表_现场代表_姓名")
            else:
                factors_error["发包人代表_现场代表_姓名"] = "第三部分_2. 发包人:发包人代表_现场代表_姓名未填写完整"
                addRemarkInDoc(word, document, "2.2.2现场代表", f"第三部分_2. 发包人:发包人代表_现场代表_姓名未填写完整")
        except:
            missObject += "第三部分_2. 发包人:要素“发包人代表_现场代表_姓名”缺失\n"

        try:
            match = '2.2.2现场代表\n姓    名：.*；\n身份证号：(.*)；\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["发包人代表_现场代表_身份证"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if checkIdCard(factor) != 'ok':
                    factors_error["发包人代表_现场代表_身份证"] = "第三部分_2. 发包人:发包人代表_身份证填写错误：" + checkIdCard(factor)
                    addRemarkInDoc(word, document, "2.2.2现场代表", f"第三部分_2. 发包人:发包人代表_身份证填写错误：" + checkIdCard(factor))
                else:
                    factors_ok.append("发包人代表_现场代表_身份证")
            else:
                factors_error["发包人代表_现场代表_身份证"] = "第三部分_2. 发包人:发包人代表_现场代表_身份证未填写完整"
                addRemarkInDoc(word, document, "2.2.2现场代表", f"第三部分_2. 发包人:发包人代表_现场代表_身份证未填写完整")
        except:
            missObject += "第三部分_2. 发包人:要素“发包人代表_现场代表_身份证”缺失\n"

        try:
            match = '2.2.2现场代表\n姓    名：.*；\n身份证号：.*；\n职    务：(.*)；\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["发包人代表_现场代表_职务"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("发包人代表_现场代表_职务")
            else:
                factors_error["发包人代表_现场代表_职务"] = "第三部分_2. 发包人:发包人代表_现场代表_职务未填写完整"
                addRemarkInDoc(word, document, "2.2.2现场代表", f"第三部分_2. 发包人:发包人代表_现场代表_职务未填写完整")
        except:
            missObject += "第三部分_2. 发包人:要素“发包人代表_现场代表_职务”缺失\n"

        try:
            match = '2.2.2现场代表\n姓    名：.*；\n身份证号：.*；\n职    务：.*；\n联系电话：(.*)；\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["发包人代表_现场代表_电话"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if isTelPhoneNumber(factor) == 'Error':
                    factors_error["发包人代表_现场代表_电话"] = "第三部分_2. 发包人:电话填写错误"
                    addRemarkInDoc(word, document, "2.2.2现场代表", f"第三部分_2. 发包人:电话填写错误")
                else:
                    factors_ok.append("发包人代表_现场代表_电话")
            else:
                factors_error["发包人代表_现场代表_电话"] = "第三部分_2. 发包人:发包人代表_现场代表_电话未填写完整"
                addRemarkInDoc(word, document, "2.2.2现场代表", f"第三部分_2. 发包人:发包人代表_现场代表_电话未填写完整")
        except:
            missObject += "第三部分_2. 发包人:要素“发包人代表_现场代表_电话”缺失\n"

        try:
            match = '2.2.2现场代表\n姓    名：.*；\n身份证号：.*；\n职    务：.*；\n联系电话：.*；\n电子信箱：(.*)；\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["发包人代表_现场代表_电子信箱"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("发包人代表_现场代表_电子信箱")
            else:
                factors_error["发包人代表_现场代表_电子信箱"] = "第三部分_2. 发包人:发包人代表_现场代表_电子信箱未填写完整"
                addRemarkInDoc(word, document, "2.2.2现场代表", f"第三部分_2. 发包人:发包人代表_现场代表_电子信箱未填写完整")
        except:
            missObject += "第三部分_2. 发包人:要素“发包人代表_现场代表_电子信箱”缺失\n"

        try:
            match = '2.2.2现场代表\n姓    名：.*；\n身份证号：.*；\n职    务：.*；\n联系电话：.*；\n电子信箱：.*；\n通信地址：(.*)。\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["发包人代表_现场代表_通信地址"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("发包人代表_现场代表_通信地址")
            else:
                factors_error["发包人代表_现场代表_通信地址"] = "第三部分_2. 发包人:发包人代表_现场代表_通信地址未填写完整"
                addRemarkInDoc(word, document, "2.2.2现场代表", f"第三部分_2. 发包人:发包人代表_现场代表_通信地址未填写完整")
        except:
            missObject += "第三部分_2. 发包人:要素“发包人代表_现场代表_通信地址”缺失\n"

        try:
            match = '2.2.3技术负责人\n姓    名：(.*)；\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["发包人代表_技术负责人_姓名"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("发包人代表_技术负责人_姓名")
            else:
                factors_error["发包人代表_技术负责人_姓名"] = "第三部分_2. 发包人:发包人代表_技术负责人_姓名未填写完整"
                addRemarkInDoc(word, document, "2.2.3技术负责人", f"第三部分_2. 发包人:发包人代表_技术负责人_姓名未填写完整")
        except:
            missObject += "第三部分_2. 发包人:要素“发包人代表_技术负责人_姓名”缺失\n"

        try:
            match = '2.2.3技术负责人\n姓    名：.*；\n身份证号：(.*)；\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["发包人代表_技术负责人_身份证"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if checkIdCard(factor) != 'ok':
                    factors_error["发包人代表_技术负责人_身份证"] = "第三部分_2. 发包人:发包人代表_身份证填写错误：" + checkIdCard(factor)
                    addRemarkInDoc(word, document, "2.2.3技术负责人", f"第三部分_2. 发包人:发包人代表_身份证填写错误：" + checkIdCard(factor))
                else:
                    factors_ok.append("发包人代表_技术负责人_身份证")
            else:
                factors_error["发包人代表_技术负责人_身份证"] = "第三部分_2. 发包人:发包人代表_技术负责人_身份证未填写完整"
                addRemarkInDoc(word, document, "2.2.3技术负责人", f"第三部分_2. 发包人:发包人代表_技术负责人_身份证未填写完整")
        except:
            missObject += "第三部分_2. 发包人:要素“发包人代表_技术负责人_身份证”缺失\n"

        try:
            match = '2.2.3技术负责人\n姓    名：.*；\n身份证号：.*；\n职    务：(.*)；\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["发包人代表_技术负责人_职务"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("发包人代表_技术负责人_职务")
            else:
                factors_error["发包人代表_技术负责人_职务"] = "第三部分_2. 发包人:发包人代表_技术负责人_职务未填写完整"
                addRemarkInDoc(word, document, "2.2.3技术负责人", f"第三部分_2. 发包人:发包人代表_技术负责人_职务未填写完整")
        except:
            missObject += "第三部分_2. 发包人:要素“发包人代表_技术负责人_职务”缺失\n"

        try:
            match = '2.2.3技术负责人\n姓    名：.*；\n身份证号：.*；\n职    务：.*；\n联系电话：(.*)；\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["发包人代表_技术负责人_电话"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if isTelPhoneNumber(factor) == 'Error':
                    factors_error["发包人代表_技术负责人_电话"] = "第三部分_2. 发包人:电话填写错误"
                    addRemarkInDoc(word, document, "2.2.3技术负责人", f"第三部分_2. 发包人:电话填写错误")
                else:
                    factors_ok.append("发包人代表_技术负责人_电话")
            else:
                factors_error["发包人代表_技术负责人_电话"] = "第三部分_2. 发包人:发包人代表_技术负责人_电话未填写完整"
                addRemarkInDoc(word, document, "2.2.3技术负责人", f"第三部分_2. 发包人:发包人代表_技术负责人_电话未填写完整")
        except:
            missObject += "第三部分_2. 发包人:要素“发包人代表_技术负责人_电话”缺失\n"

        try:
            match = '2.2.3技术负责人\n姓    名：.*；\n身份证号：.*；\n职    务：.*；\n联系电话：.*；\n电子信箱：(.*)；\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["发包人代表_技术负责人_电子信箱"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("发包人代表_技术负责人_电子信箱")
            else:
                factors_error["发包人代表_技术负责人_电子信箱"] = "第三部分_2. 发包人:发包人代表_技术负责人_电子信箱未填写完整"
                addRemarkInDoc(word, document, "2.2.3技术负责人", f"第三部分_2. 发包人:发包人代表_技术负责人_电子信箱未填写完整")
        except:
            missObject += "第三部分_2. 发包人:要素“发包人代表_技术负责人_电子信箱”缺失\n"

        try:
            match = '2.2.3技术负责人\n姓    名：.*；\n身份证号：.*；\n职    务：.*；\n联系电话：.*；\n电子信箱：.*；\n通信地址：(.*)。\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["发包人代表_技术负责人_通信地址"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("发包人代表_技术负责人_通信地址")
            else:
                factors_error["发包人代表_技术负责人_通信地址"] = "第三部分_2. 发包人:发包人代表_技术负责人_通信地址未填写完整"
                addRemarkInDoc(word, document, "2.2.3技术负责人", f"第三部分_2. 发包人:发包人代表_技术负责人_通信地址未填写完整")
        except:
            missObject += "第三部分_2. 发包人:要素“发包人代表_技术负责人_通信地址”缺失\n"

        try:
            match = '2.4.1 提供施工现场\n关于发包人移交施工现场的期限要求：在(.*)之前提供，.*'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["2.4.1时间"] = factor
            if get_strtime(factor):
                factors_ok.append("2.4.1时间")
                time_2_4_1 = get_strtime(factor)
                year = time_2_4_1.split('-')[0]
                month = time_2_4_1.split('-')[1]
                day = time_2_4_1.split('-')[2]
                time_2_4_1 = datetime.date(int(year), int(month), int(day))
                try:
                    if (time_2_4_1 - start_time).days > 0:
                        factors_ok.append("2.4.1 提供施工现场时间")
                    else:
                        factors_error["2.4.1时间"] = "第三部分_2. 发包人:2.4.1时间约定日期没有在《合同协议书》第二条约定的合同计划开工日期之前"
                        addRemarkInDoc(word, document, "2.4.1 提供施工现场",
                                       f"第三部分_2. 发包人:2.4.1时间约定日期没有在《合同协议书》第二条约定的合同计划开工日期之前")
                except:
                    factors_error["2.4.1时间"] = "第三部分_2. 发包人:合同计划开工日期提取失败"
                    addRemarkInDoc(word, document, "2.4.1 提供施工现场", f"第三部分_2. 发包人:合同计划开工日期提取失败")
            else:
                factors_error["2.4.1时间"] = "2.4.1时间未填写完整"
                addRemarkInDoc(word, document, "2.4.1 提供施工现场", f"第三部分_2. 发包人:2.4.1时间未填写完整")
        except:
            missObject += "第三部分_2. 发包人:要素“2.4.1时间”缺失\n"

    # 第三部分_3-5
    if 1 == True:
        # 3.1
        try:
            match = '（9）承包人提交的竣工资料的内容：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["（9）承包人提交的竣工资料的内容"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '根据建设行政主管部门和(或)城市建设档案管理部门的规定，完成竣工资料（含竣工图电子文档、声像档案）和备案资料的收集、整理、立卷、归档工作'.replace(" ",
                                                                                                            "").replace(
                    "；", "").replace("。", ""):
                    factors_error[
                        "（9）承包人提交的竣工资料的内容"] = "第三部分_3. 承包人:请核实该条约定的资料完整性是否符合建设行政主管部门和(或)城市建设档案管理部门的规定，原则上应由承包人根据建设行政主管部门和(或)城市建设档案管理部门的规定，完成全部竣工资料（含竣工图电子文档、声像档案）和备案资料的收集、整理、立卷、归档工作。"
                    addRemarkInDoc(word, document, "（9）承包人提交的竣工资料的内容",
                                   f"第三部分_3. 承包人:请核实该条约定的资料完整性是否符合建设行政主管部门和(或)城市建设档案管理部门的规定，原则上应由承包人根据建设行政主管部门和(或)城市建设档案管理部门的规定，完成全部竣工资料（含竣工图电子文档、声像档案）和备案资料的收集、整理、立卷、归档工作。")
                else:
                    factors_ok.append("（9）承包人提交的竣工资料的内容")
            else:
                factors_error["（9）承包人提交的竣工资料的内容"] = "第三部分_3. 承包人:（9）承包人提交的竣工资料的内容未填写完整"
                addRemarkInDoc(word, document, "（9）承包人提交的竣工资料的内容", f"第三部分_3. 承包人:（9）承包人提交的竣工资料的内容未填写完整")
        except:
            missObject += "第三部分_3. 承包人:要素“（9）承包人提交的竣工资料的内容”缺失\n"

        try:
            match = '承包人需要提交的竣工资料套数：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["承包人需要提交的竣工资料套数"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '纸质资料一式四套（并按相关要求装订成册）、电子资料一套'.replace(" ", "").replace("；", "").replace("。", ""):
                    factors_error["承包人需要提交的竣工资料套数"] = "第三部分_3. 承包人:要求承包人提供的竣工资料至少需要纸质资料一式四套（并按相关要求装订成册）、电子资料一套"
                    addRemarkInDoc(word, document, "承包人需要提交的竣工资料套数",
                                   f"第三部分_3. 承包人:要求承包人提供的竣工资料至少需要纸质资料一式四套（并按相关要求装订成册）、电子资料一套")
                else:
                    factors_ok.append("承包人需要提交的竣工资料套数")
            else:
                factors_error["承包人需要提交的竣工资料套数"] = "第三部分_3. 承包人:承包人需要提交的竣工资料套数未填写完整"
                addRemarkInDoc(word, document, "承包人需要提交的竣工资料套数", f"第三部分_3. 承包人:承包人需要提交的竣工资料套数未填写完整")
        except:
            missObject += "第三部分_3. 承包人:要素“承包人需要提交的竣工资料套数”缺失\n"

        try:
            match = '承包人提交的竣工资料的费用承担：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["承包人提交的竣工资料的费用承担"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '由承包人承担'.replace(" ", "").replace("；", "").replace("。", ""):
                    factors_error["承包人提交的竣工资料的费用承担"] = "第三部分_3. 承包人:原则上由承包人承担"
                    addRemarkInDoc(word, document, "承包人提交的竣工资料的费用承担", f"第三部分_3. 承包人:原则上由承包人承担")
                else:
                    factors_ok.append("承包人提交的竣工资料的费用承担")
            else:
                factors_error["承包人提交的竣工资料的费用承担"] = "第三部分_3. 承包人:承包人提交的竣工资料的费用承担未填写完整"
                addRemarkInDoc(word, document, "承包人提交的竣工资料的费用承担", f"第三部分_3. 承包人:承包人提交的竣工资料的费用承担未填写完整")
        except:
            missObject += "第三部分_3. 承包人:要素“承包人提交的竣工资料的费用承担”缺失\n"

        try:
            match = '承包人提交的竣工资料移交时间：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["承包人提交的竣工资料移交时间"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    match = "竣工验收通过后(.*)天内"
                    x = re.findall(match, factor)[0].replace(" ", "")
                    if int(x) < 3 or int(x) > 15:
                        factors_error["承包人提交的竣工资料移交时间"] = "第三部分_3. 承包人:请核实超出推荐条款的理由"
                        addRemarkInDoc(word, document, "承包人提交的竣工资料移交时间", f"第三部分_3. 承包人:请核实超出推荐条款的理由")
                    else:
                        factors_ok.append("承包人提交的竣工资料移交时间")
                except:
                    factors_error["承包人提交的竣工资料移交时间"] = "第三部分_3. 承包人:时间无法提取"
                    addRemarkInDoc(word, document, "承包人提交的竣工资料移交时间", f"第三部分_3. 承包人:时间无法提取")

            else:
                factors_error["承包人提交的竣工资料移交时间"] = "承包人提交的竣工资料移交时间未填写完整"
                addRemarkInDoc(word, document, "承包人提交的竣工资料移交时间", f"第三部分_3. 承包人:承包人提交的竣工资料移交时间未填写完整")
        except:
            missObject += "第三部分_3. 承包人:要素“承包人提交的竣工资料移交时间”缺失\n"

        try:
            match = '承包人提交的竣工资料形式要求：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["承包人提交的竣工资料形式要求"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '按照发包方要求提交'.replace(" ", "").replace("；", "").replace("。", ""):
                    factors_error["承包人提交的竣工资料形式要求"] = "第三部分_3. 承包人:原则上应约定按发包方要求提交"
                    addRemarkInDoc(word, document, "承包人提交的竣工资料形式要求", f"第三部分_3. 承包人:原则上应约定按发包方要求提交")
                else:
                    factors_ok.append("承包人提交的竣工资料形式要求")
            else:
                factors_error["承包人提交的竣工资料形式要求"] = "第三部分_3. 承包人:承包人提交的竣工资料形式要求未填写完整"
                addRemarkInDoc(word, document, "承包人提交的竣工资料形式要求", f"第三部分_3. 承包人:承包人提交的竣工资料形式要求未填写完整")
        except:
            missObject += "第三部分_3. 承包人:要素“承包人提交的竣工资料形式要求”缺失\n"

        # 3.2
        try:
            match = '3.2.1 项目经理：\n姓    名：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["3.2.1姓名"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("3.2.1姓名")
            else:
                factors_error["3.2.1姓名"] = "第三部分_3. 承包人:3.2.1姓名未填写完整"
                addRemarkInDoc(word, document, "3.2.1 项目经理：", f"第三部分_3. 承包人:3.2.1姓名未填写完整")
        except:
            missObject += "第三部分_3. 承包人:要素“3.2.1姓名”缺失\n"

        try:
            match = '3.2.1 项目经理：\n姓    名：.*\n身份证号：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["3.2.1身份证号"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if checkIdCard(factor) != 'ok':
                    factors_error["3.2.1身份证号"] = "第三部分_3. 承包人:3.2.1身份证号填写错误:" + checkIdCard(factor)
                    addRemarkInDoc(word, document, "3.2.1 项目经理：", f"第三部分_3. 承包人:3.2.1身份证号填写错误:" + checkIdCard(factor))
                else:
                    factors_ok.append("3.2.1身份证号")
            else:
                factors_error["3.2.1身份证号"] = "第三部分_3. 承包人:3.2.1身份证号未填写完整"
                addRemarkInDoc(word, document, "3.2.1 项目经理：", f"第三部分_3. 承包人:3.2.1身份证号未填写完整")
        except:
            missObject += "第三部分_3. 承包人:要素“3.2.1身份证号”缺失\n"

        try:
            match = '3.2.1 项目经理：\n姓    名：.*\n身份证号：.*\n建造师执业资格等级：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["3.2.1建造师执业资格等级"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("3.2.1建造师执业资格等级")
            else:
                factors_error["3.2.1建造师执业资格等级"] = "第三部分_3. 承包人:3.2.1建造师执业资格等级未填写完整"
                addRemarkInDoc(word, document, "3.2.1 项目经理：", f"第三部分_3. 承包人:3.2.1建造师执业资格等级未填写完整")
        except:
            missObject += "第三部分_3. 承包人:要素“3.2.1建造师执业资格等级”缺失\n"

        try:
            match = '3.2.1 项目经理：\n姓    名：.*\n身份证号：.*\n建造师执业资格等级：.*\n建造师注册证书号：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["3.2.1建造师注册证书号"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("3.2.1建造师注册证书号")
            else:
                factors_error["3.2.1建造师注册证书号"] = "第三部分_3. 承包人:3.2.1建造师注册证书号未填写完整"
                addRemarkInDoc(word, document, "3.2.1 项目经理：", f"第三部分_3. 承包人:3.2.1建造师注册证书号未填写完整")
        except:
            missObject += "第三部分_3. 承包人:要素“3.2.1建造师注册证书号”缺失\n"

        try:
            match = '3.2.1 项目经理：\n姓    名：.*\n身份证号：.*\n建造师执业资格等级：.*\n建造师注册证书号：.*\n建造师执业印章号：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["3.2.1建造师执业印章号"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("3.2.1建造师执业印章号")
            else:
                factors_error["3.2.1建造师执业印章号"] = "第三部分_3. 承包人:3.2.1建造师执业印章号未填写完整"
                addRemarkInDoc(word, document, "3.2.1 项目经理：", f"第三部分_3. 承包人:3.2.1建造师执业印章号未填写完整")
        except:
            missObject += "第三部分_3. 承包人:要素“3.2.1建造师执业印章号”缺失\n"

        try:
            match = '3.2.1 项目经理：\n姓    名：.*\n身份证号：.*\n建造师执业资格等级：.*\n建造师注册证书号：.*\n建造师执业印章号：.*\n安全生产考核合格证书号：(.*)'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["3.2.1安全生产考核合格证书号"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("3.2.1安全生产考核合格证书号")
            else:
                factors_error["3.2.1安全生产考核合格证书号"] = "第三部分_3. 承包人:3.2.1安全生产考核合格证书号未填写完整"
                addRemarkInDoc(word, document, "3.2.1 项目经理：", f"第三部分_3. 承包人:3.2.1安全生产考核合格证书号未填写完整")
        except:
            missObject += "第三部分_3. 承包人:要素“3.2.1安全生产考核合格证书号”缺失\n"

        try:
            match = '3.2.1 项目经理：\n姓    名：.*\n身份证号：.*\n建造师执业资格等级：.*\n建造师注册证书号：.*\n建造师执业印章号：.*\n安全生产考核合格证书号：.*\n联系电话：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["3.2.1联系电话"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if isTelPhoneNumber(factor) == 'Error':
                    factors_error["3.2.1联系电话"] = "第三部分_3. 承包人:3.2.1联系电话填写错误"
                    addRemarkInDoc(word, document, "3.2.1 项目经理：", f"第三部分_3. 承包人:3.2.1联系电话填写错误")
                else:
                    factors_ok.append("3.2.1联系电话")
            else:
                factors_error["3.2.1联系电话"] = "第三部分_3. 承包人:3.2.1联系电话未填写完整"
                addRemarkInDoc(word, document, "3.2.1 项目经理：", f"第三部分_3. 承包人:3.2.1联系电话未填写完整")
        except:
            missObject += "第三部分_3. 承包人:要素“3.2.1联系电话”缺失\n"

        try:
            match = '3.2.1 项目经理：\n姓    名：.*\n身份证号：.*\n建造师执业资格等级：.*\n建造师注册证书号：.*\n建造师执业印章号：.*\n安全生产考核合格证书号：.*\n联系电话.*\n电子信箱：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["3.2.1电子信箱"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("3.2.1电子信箱")
            else:
                factors_error["3.2.1电子信箱"] = "第三部分_3. 承包人:3.2.1电子信箱未填写完整"
                addRemarkInDoc(word, document, "3.2.1 项目经理：", f"第三部分_3. 承包人:3.2.1电子信箱未填写完整")
        except:
            missObject += "第三部分_3. 承包人:要素“3.2.1电子信箱”缺失\n"

        try:
            match = '3.2.1 项目经理：\n姓    名：.*\n身份证号：.*\n建造师执业资格等级：.*\n建造师注册证书号：.*\n建造师执业印章号：.*\n安全生产考核合格证书号：.*\n联系电话.*\n电子信箱：.*\n通信地址：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["3.2.1通信地址"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("3.2.1通信地址")
            else:
                factors_error["3.2.1通信地址"] = "第三部分_3. 承包人:3.2.1通信地址未填写完整"
                addRemarkInDoc(word, document, "3.2.1 项目经理：", f"第三部分_3. 承包人:3.2.1通信地址未填写完整")
        except:
            missObject += "要素“3.2.1通信地址”缺失\n"

        factors_to_inform["3.2.1通信地址"] = "请核对项目经理资质证书并确定上述内容的准确性"
        addRemarkInDoc(word, document, "3.2.1 项目经理：", f"第三部分_3. 承包人:请核对项目经理资质证书并确定上述内容的准确性")

        try:
            match = '承包人对项目经理的授权范围如下：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["承包人对项目经理的授权范围如下"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '受承包人委托全面负责工程安全、质量、进度、投资四大控制要素，完善合同、信息、组织管理体系，做好项目施工现场全面管理工作'.replace(" ", "").replace("；",
                                                                                                                    "").replace(
                    "。", ""):
                    factors_error[
                        "承包人对项目经理的授权范围如下"] = "第三部分_3. 承包人:项目经理权限必须包括“全面负责工程安全、质量、进度、投资四大控制要素，完善合同、信息、组织管理体系，做好项目施工现场全面管理工作"
                    addRemarkInDoc(word, document, "承包人对项目经理的授权范围如下",
                                   f"第三部分_3. 承包人:项目经理权限必须包括“全面负责工程安全、质量、进度、投资四大控制要素，完善合同、信息、组织管理体系，做好项目施工现场全面管理工作")
                else:
                    factors_ok.append("承包人对项目经理的授权范围如下")
            else:
                factors_error["承包人对项目经理的授权范围如下"] = "第三部分_3. 承包人:承包人对项目经理的授权范围如下未填写完整"
                addRemarkInDoc(word, document, "承包人对项目经理的授权范围如下", f"第三部分_3. 承包人:承包人对项目经理的授权范围如下未填写完整")
        except:
            missObject += "第三部分_3. 承包人:要素“承包人对项目经理的授权范围如下”缺失\n"

        try:
            match = '关于项目经理每月在施工现场的时间要求：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["关于项目经理每月在施工现场的时间要求"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '派驻本合同工程项目的项目经理和项目技术负责人应履行到场签到手续，签到表由发包人、监理人共同审核确认；项目经理和项目技术负责人每月应保证不低于施工时间的80%的驻工地时间。不够天数的，按照不足天数每天次支付违约金5000元/人/天'.replace(
                        " ", "").replace("；", "").replace("。", ""):
                    factors_error["关于项目经理每月在施工现场的时间要求"] = "第三部分_3. 承包人:应该约定项目经理的最低驻场时间及违反该约定的违约责任"
                    addRemarkInDoc(word, document, "关于项目经理每月在施工现场的时间要求", f"第三部分_3. 承包人:应该约定项目经理的最低驻场时间及违反该约定的违约责任")
                else:
                    factors_ok.append("关于项目经理每月在施工现场的时间要求")
            else:
                factors_error["关于项目经理每月在施工现场的时间要求"] = "第三部分_3. 承包人:关于项目经理每月在施工现场的时间要求未填写完整"
                addRemarkInDoc(word, document, "关于项目经理每月在施工现场的时间要求", f"第三部分_3. 承包人:关于项目经理每月在施工现场的时间要求未填写完整")
        except:
            missObject += "第三部分_3. 承包人:要素“关于项目经理每月在施工现场的时间要求”缺失\n"

        try:
            match = '承包人未提交劳动合同，以及没有为项目经理缴纳社会保险证明的违约责任：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["承包人未提交劳动合同，以及没有为项目经理缴纳社会保险证明的违约责任"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '按通用条款执行'.replace(" ", "").replace("；", "").replace("。", ""):
                    factors_error[
                        "承包人未提交劳动合同，以及没有为项目经理缴纳社会保险证明的违约责任"] = "第三部分_3. 承包人:承包人未提交劳动合同，若通用条款约定较本条约定更为严格，应该约定“按通用条款执行“"
                    addRemarkInDoc(word, document, "承包人未提交劳动合同，以及没有为项目经理缴纳社会保险证明的违约责任",
                                   f"第三部分_3. 承包人:承包人未提交劳动合同，若通用条款约定较本条约定更为严格，应该约定“按通用条款执行“")
                else:
                    factors_ok.append("承包人未提交劳动合同，以及没有为项目经理缴纳社会保险证明的违约责任")
            else:
                factors_error[
                    "承包人未提交劳动合同，以及没有为项目经理缴纳社会保险证明的违约责任"] = "第三部分_3. 承包人:承包人未提交劳动合同，以及没有为项目经理缴纳社会保险证明的违约责任未填写完整"
                addRemarkInDoc(word, document, "承包人未提交劳动合同，以及没有为项目经理缴纳社会保险证明的违约责任",
                               f"第三部分_3. 承包人:承包人未提交劳动合同，以及没有为项目经理缴纳社会保险证明的违约责任未填写完整")
        except:
            missObject += "第三部分_3. 承包人:要素“承包人未提交劳动合同，以及没有为项目经理缴纳社会保险证明的违约责任”缺失\n"

        try:
            match = '项目经理未经批准，擅自离开施工现场的违约责任：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["项目经理未经批准，擅自离开施工现场的违约责任"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '承包人派驻到项目上的项目经理、项目副经理（如有）、项目技术负责人离开现场需报总监理工程师批准并经发包人同意。从擅自离开施工现场次日起每天向发包人支付违约金5000元，按日累计，此违约金在合同价款中扣除。超过20天视为承包人擅自更换项目经理'.replace(
                        " ", "").replace("；", "").replace("。", ""):
                    factors_error["项目经理未经批准，擅自离开施工现场的违约责任"] = "第三部分_3. 承包人:项目经理擅自离场的违约责任应至少与其违反最低驻场时间的违约责任约定一致"
                    addRemarkInDoc(word, document, "项目经理未经批准，擅自离开施工现场的违约责任",
                                   f"第三部分_3. 承包人:项目经理擅自离场的违约责任应至少与其违反最低驻场时间的违约责任约定一致")
                else:
                    factors_ok.append("项目经理未经批准，擅自离开施工现场的违约责任")
            else:
                factors_error["项目经理未经批准，擅自离开施工现场的违约责任"] = "第三部分_3. 承包人:项目经理未经批准，擅自离开施工现场的违约责任未填写完整"
                addRemarkInDoc(word, document, "项目经理未经批准，擅自离开施工现场的违约责任", f"第三部分_3. 承包人:项目经理未经批准，擅自离开施工现场的违约责任未填写完整")
        except:
            missObject += "第三部分_3. 承包人:要素“项目经理未经批准，擅自离开施工现场的违约责任”缺失\n"

        try:
            match = '3.2.2 承包人擅自更换项目经理的违约责任：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["3.2.2 承包人擅自更换项目经理的违约责任"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '项目经理必须与承包人投标文件承诺的项目经理一致，承包人确需更换项目经理需事前经发包人书面同意；承包人擅自更换项目经理的，发包人有权处违约金5万元/人.次，经监理人发出书面警告仍不纠正的，发包人有权解除合同，因此造成的工期延误、工程质量等一切损失由承包人全部承担'.replace(
                        " ", "").replace("；", "").replace("。", ""):
                    factors_error[
                        "3.2.2 承包人擅自更换项目经理的违约责任"] = "第三部分_3. 承包人:该条违约责任不应低于：承包人擅自更换项目经理的，视为根本违约，应按5万元/人.次的标准支付违约金；因此造成的工期延误、工程质量等一切损失由承包人全部承担。发包人据此有单方合同解除权。"
                    addRemarkInDoc(word, document, "3.2.2 承包人擅自更换项目经理的违约责任",
                                   f"第三部分_3. 承包人:该条违约责任不应低于：承包人擅自更换项目经理的，视为根本违约，应按5万元/人.次的标准支付违约金；因此造成的工期延误、工程质量等一切损失由承包人全部承担。发包人据此有单方合同解除权。")
                else:
                    factors_ok.append("3.2.2 承包人擅自更换项目经理的违约责任")
            else:
                factors_error["3.2.2 承包人擅自更换项目经理的违约责任"] = "第三部分_3. 承包人:3.2.2 承包人擅自更换项目经理的违约责任未填写完整"
                addRemarkInDoc(word, document, "3.2.2 承包人擅自更换项目经理的违约责任", f"第三部分_3. 承包人:3.2.2 承包人擅自更换项目经理的违约责任未填写完整")
        except:
            missObject += "第三部分_3. 承包人:要素“3.2.2 承包人擅自更换项目经理的违约责任”缺失\n"

        try:
            match = '  3.2.3 承包人无正当理由拒绝更换项目经理的违约责任：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["  3.2.3 承包人无正当理由拒绝更换项目经理的违约责任"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '如发包人认为承包人委派的项目经理不能胜任本工程的，且承包人拒绝更换项目经理的处违约金5万元/人.次，经监理人发出书面警告仍不纠正的，发包人有权解除合同，因此造成的工期延误、工程质量等一切损失由承包人全部承担'.replace(
                        " ", "").replace("；", "").replace("。", ""):
                    factors_error[
                        "  3.2.3 承包人无正当理由拒绝更换项目经理的违约责任"] = "第三部分_3. 承包人:该条违约责任不应低于：承包人无正当理由拒绝更换项目经理的，视为根本违约，应按5万元/人.次的标准支付违约金；因此造成的工期延误、工程质量等一切损失由承包人全部承担。发包人据此有单方合同解除权"
                    addRemarkInDoc(word, document, "  3.2.3 承包人无正当理由拒绝更换项目经理的违约责任",
                                   f"第三部分_3. 承包人:该条违约责任不应低于：承包人无正当理由拒绝更换项目经理的，视为根本违约，应按5万元/人.次的标准支付违约金；因此造成的工期延误、工程质量等一切损失由承包人全部承担。发包人据此有单方合同解除权")
                else:
                    factors_ok.append("  3.2.3 承包人无正当理由拒绝更换项目经理的违约责任")
            else:
                factors_error["  3.2.3 承包人无正当理由拒绝更换项目经理的违约责任"] = "  第三部分_3. 承包人:3.2.3 承包人无正当理由拒绝更换项目经理的违约责任未填写完整"
                addRemarkInDoc(word, document, "  3.2.3 承包人无正当理由拒绝更换项目经理的违约责任",
                               f"  第三部分_3. 承包人:3.2.3 承包人无正当理由拒绝更换项目经理的违约责任未填写完整")
        except:
            missObject += "第三部分_3. 承包人:要素“  3.2.3 承包人无正当理由拒绝更换项目经理的违约责任”缺失\n"

        # 3.3
        try:
            match = '3.3.1 承包人提交项目管理机构及施工现场管理人员安排报告的期限：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["3.3.1 承包人提交项目管理机构及施工现场管理人员安排报告的期限"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '按通用条款执行'.replace(" ", "").replace("；", "").replace("。", ""):
                    factors_error["3.3.1 承包人提交项目管理机构及施工现场管理人员安排报告的期限"] = "第三部分_3. 承包人:该条期限不应多于合同签订之日起15日"
                    addRemarkInDoc(word, document, "3.3.1 承包人提交项目管理机构及施工现场管理人员安排报告的期限",
                                   f"第三部分_3. 承包人:该条期限不应多于合同签订之日起15日")
                else:
                    factors_ok.append("3.3.1 承包人提交项目管理机构及施工现场管理人员安排报告的期限")
            else:
                factors_error[
                    "3.3.1 承包人提交项目管理机构及施工现场管理人员安排报告的期限"] = "第三部分_3. 承包人:3.3.1 承包人提交项目管理机构及施工现场管理人员安排报告的期限未填写完整"
                addRemarkInDoc(word, document, "3.3.1 承包人提交项目管理机构及施工现场管理人员安排报告的期限",
                               f"第三部分_3. 承包人:3.3.1 承包人提交项目管理机构及施工现场管理人员安排报告的期限未填写完整")
        except:
            missObject += "第三部分_3. 承包人:要素“3.3.1 承包人提交项目管理机构及施工现场管理人员安排报告的期限”缺失\n"

        try:
            match = '3.3.3 承包人无正当理由拒绝撤换主要施工管理人员的违约责任：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["3.3.3 承包人无正当理由拒绝撤换主要施工管理人员的违约责任"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '处违约金1万元/人.次'.replace(" ", "").replace("；", "").replace("。", ""):
                    factors_error["3.3.3 承包人无正当理由拒绝撤换主要施工管理人员的违约责任"] = "第三部分_3. 承包人:该条违约责任不应低于“违约金1万元/人/次”的标准"
                    addRemarkInDoc(word, document, "3.3.3 承包人无正当理由拒绝撤换主要施工管理人员的违约责任",
                                   f"第三部分_3. 承包人:该条违约责任不应低于“违约金1万元/人/次”的标准")
                else:
                    factors_ok.append("3.3.3 承包人无正当理由拒绝撤换主要施工管理人员的违约责任")
            else:
                factors_error["3.3.3 承包人无正当理由拒绝撤换主要施工管理人员的违约责任"] = "第三部分_3. 承包人:3.3.3 承包人无正当理由拒绝撤换主要施工管理人员的违约责任未填写完整"
                addRemarkInDoc(word, document, "3.3.3 承包人无正当理由拒绝撤换主要施工管理人员的违约责任",
                               f"第三部分_3. 承包人:3.3.3 承包人无正当理由拒绝撤换主要施工管理人员的违约责任未填写完整")
        except:
            missObject += "第三部分_3. 承包人:要素“3.3.3 承包人无正当理由拒绝撤换主要施工管理人员的违约责任”缺失\n"

        try:
            match = '3.3.4 承包人主要施工管理人员离开施工现场的批准要求：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["3.3.4 承包人主要施工管理人员离开施工现场的批准要求"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '主要管理人员离开现场超过1天，应向发包人书面申请，经发包人审批后方能离开'.replace(" ", "").replace("；", "").replace("。", ""):
                    factors_error["3.3.4 承包人主要施工管理人员离开施工现场的批准要求"] = "第三部分_3. 承包人:必须经发包人审批"
                    addRemarkInDoc(word, document, "3.3.4 承包人主要施工管理人员离开施工现场的批准要求", f"第三部分_3. 承包人:必须经发包人审批")
                else:
                    factors_ok.append("3.3.4 承包人主要施工管理人员离开施工现场的批准要求")
            else:
                factors_error["3.3.4 承包人主要施工管理人员离开施工现场的批准要求"] = "第三部分_3. 承包人:3.3.4 承包人主要施工管理人员离开施工现场的批准要求未填写完整"
                addRemarkInDoc(word, document, "3.3.4 承包人主要施工管理人员离开施工现场的批准要求",
                               f"第三部分_3. 承包人:3.3.4 承包人主要施工管理人员离开施工现场的批准要求未填写完整")
        except:
            missObject += "第三部分_3. 承包人:要素“3.3.4 承包人主要施工管理人员离开施工现场的批准要求”缺失\n"

        try:
            match = '3.3.5承包人擅自更换主要施工管理人员的违约责任：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["3.3.5承包人擅自更换主要施工管理人员的违约责任"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '处违约金1万元/人.次'.replace(" ", "").replace("；", "").replace("。", ""):
                    factors_error["3.3.5承包人擅自更换主要施工管理人员的违约责任"] = "第三部分_3. 承包人:该条违约责任不应低于“违约金1万元/人/次”的标准"
                    addRemarkInDoc(word, document, "3.3.5承包人擅自更换主要施工管理人员的违约责任",
                                   f"第三部分_3. 承包人:该条违约责任不应低于“违约金1万元/人/次”的标准")
                else:
                    factors_ok.append("3.3.5承包人擅自更换主要施工管理人员的违约责任")
            else:
                factors_error["3.3.5承包人擅自更换主要施工管理人员的违约责任"] = "第三部分_3. 承包人:3.3.5承包人擅自更换主要施工管理人员的违约责任未填写完整"
                addRemarkInDoc(word, document, "3.3.5承包人擅自更换主要施工管理人员的违约责任",
                               f"第三部分_3. 承包人:3.3.5承包人擅自更换主要施工管理人员的违约责任未填写完整")
        except:
            missObject += "第三部分_3. 承包人:要素“3.3.5承包人擅自更换主要施工管理人员的违约责任”缺失\n"

        try:
            match = '承包人主要施工管理人员擅自离开施工现场的违约责任：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["承包人主要施工管理人员擅自离开施工现场的违约责任"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '处违约金 5000元/人.次'.replace(" ", "").replace("；", "").replace("。", ""):
                    factors_error["承包人主要施工管理人员擅自离开施工现场的违约责任"] = "第三部分_3. 承包人:该条违约责任不应低于“违约金5000元/人/次”的标准"
                    addRemarkInDoc(word, document, "承包人主要施工管理人员擅自离开施工现场的违约责任",
                                   f"第三部分_3. 承包人:该条违约责任不应低于“违约金5000元/人/次”的标准")
                else:
                    factors_ok.append("承包人主要施工管理人员擅自离开施工现场的违约责任")
            else:
                factors_error["承包人主要施工管理人员擅自离开施工现场的违约责任"] = "第三部分_3. 承包人:承包人主要施工管理人员擅自离开施工现场的违约责任未填写完整"
                addRemarkInDoc(word, document, "承包人主要施工管理人员擅自离开施工现场的违约责任", f"第三部分_3. 承包人:承包人主要施工管理人员擅自离开施工现场的违约责任未填写完整")
        except:
            missObject += "第三部分_3. 承包人:要素“承包人主要施工管理人员擅自离开施工现场的违约责任”缺失\n"

        try:
            match = '禁止分包的工程包括：(.*)'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["禁止分包的工程"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("禁止分包的工程")
                if factor != '按通用条款执行':
                    factors_to_inform["禁止分包的工程"] = "第三部分_3.5 分包:请比对通用条款的该条规定，若禁止分包的工程超过通用条款范围，请核实并要求合理解释"
                    addRemarkInDoc(word, document, "禁止分包的工程", f"第三部分_3.5 分包:请比对通用条款的该条规定，若禁止分包的工程超过通用条款范围，请核实并要求合理解释")
            else:
                factors_error["禁止分包的工程"] = "第三部分_3.5 分包:禁止分包的工程未填写完整"
                addRemarkInDoc(word, document, "禁止分包的工程", f"第三部分_3.5 分包:禁止分包的工程未填写完整")
        except:
            missObject += "第三部分_3.5 分包:要素“禁止分包的工程”缺失\n"

        try:
            match = '主体结构、关键性工作的范围：(.*)'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["主体结构、关键性工作的范围"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("主体结构、关键性工作的范围")
                if factor != '按通用条款执行':
                    factors_to_inform["主体结构、关键性工作的范围"] = "第三部分_3.5 分包:请比对通用条款的该条规定，若禁止分包的工程超过通用条款范围，请核实并要求合理解释"
                    addRemarkInDoc(word, document, "主体结构、关键性工作的范围",
                                   f"第三部分_3.5 分包:请比对通用条款的该条规定，若禁止分包的工程超过通用条款范围，请核实并要求合理解释")
            else:
                factors_error["主体结构、关键性工作的范围"] = "第三部分_3.5 分包:主体结构、关键性工作的范围未填写完整"
                addRemarkInDoc(word, document, "主体结构、关键性工作的范围", f"第三部分_3.5 分包:主体结构、关键性工作的范围未填写完整")
        except:
            missObject += "第三部分_3.5 分包:要素“主体结构、关键性工作的范围”缺失\n"

        try:
            match = '允许分包的专业工程包括：(.*?)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["允许分包的专业工程包括"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("允许分包的专业工程包括")
                if factor != '无':
                    factors_to_inform[
                        "允许分包的专业工程包括"] = "第三部分_3.5 分包:请核实允许专业分包工程的必要性，并关注专业分包工程的专业分包合同，核实本条内容是否与专业分包合同约定一致"
                    addRemarkInDoc(word, document, "允许分包的专业工程包括",
                                   f"第三部分_3.5 分包:请核实允许专业分包工程的必要性，并关注专业分包工程的专业分包合同，核实本条内容是否与专业分包合同约定一致")
            else:
                factors_error["允许分包的专业工程包括"] = "第三部分_3.5 分包:允许分包的专业工程包括未填写完整"
                addRemarkInDoc(word, document, "允许分包的专业工程包括", f"第三部分_3.5 分包:允许分包的专业工程包括未填写完整")
        except:
            missObject += "第三部分_3.5 分包:要素“允许分包的专业工程包括”缺失\n"

        try:
            match = '其他关于分包的约定：\n(.*?)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["其他关于分包的约定"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("其他关于分包的约定")
                if factor != '无':
                    factors_to_inform["其他关于分包的约定"] = "第三部分_3.5 分包:请核实允许专业分包工程的必要性，并关注专业分包工程的专业分包合同，核实本条内容是否与专业分包合同约定一致"
                    addRemarkInDoc(word, document, "其他关于分包的约定",
                                   f"第三部分_3.5 分包:请核实允许专业分包工程的必要性，并关注专业分包工程的专业分包合同，核实本条内容是否与专业分包合同约定一致")
            else:
                factors_error["其他关于分包的约定"] = "第三部分_3.5 分包:其他关于分包的约定未填写完整"
                addRemarkInDoc(word, document, "其他关于分包的约定", f"第三部分_3.5 分包:其他关于分包的约定未填写完整")
        except:
            missObject += "第三部分_3.5 分包:要素“其他关于分包的约定”缺失\n"

        try:
            match = '关于分包合同价款支付的约定：\n(.*?)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["关于分包合同价款支付的约定"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("关于分包合同价款支付的约定")
                if factor != '无':
                    factors_to_inform[
                        "关于分包合同价款支付的约定"] = "第三部分_3.5 分包:请核实允许专业分包工程的必要性，并关注专业分包工程的专业分包合同，核实本条内容是否与专业分包合同约定一致"
                    addRemarkInDoc(word, document, "关于分包合同价款支付的约定",
                                   f"第三部分_3.5 分包:请核实允许专业分包工程的必要性，并关注专业分包工程的专业分包合同，核实本条内容是否与专业分包合同约定一致")
            else:
                factors_error["关于分包合同价款支付的约定"] = "第三部分_3.5 分包:关于分包合同价款支付的约定未填写完整"
                addRemarkInDoc(word, document, "关于分包合同价款支付的约定", f"第三部分_3.5 分包:关于分包合同价款支付的约定未填写完整")
        except:
            missObject += "第三部分_3.5 分包:要素“关于分包合同价款支付的约定”缺失\n"

        try:
            match = '承包人负责照管工程及工程相关的材料、工程设备的起始时间：(.*)'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["承包人负责照管工程及工程相关的材料、工程设备的起始时间"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("承包人负责照管工程及工程相关的材料、工程设备的起始时间")
                if factor != '按通用条款执行':
                    factors_to_inform[
                        "承包人负责照管工程及工程相关的材料、工程设备的起始时间"] = "第三部分_3.6 工程照管与成品、半成品保护:承包人负责照管工程及工程相关的材料、工程设备的起始时间原则上应始于承包人正式进场施工或视为进场施工，终于工程验收合格并交付"
                    addRemarkInDoc(word, document, "承包人负责照管工程及工程相关的材料、工程设备的起始时间",
                                   f"第三部分_3.6 工程照管与成品、半成品保护:承包人负责照管工程及工程相关的材料、工程设备的起始时间原则上应始于承包人正式进场施工或视为进场施工，终于工程验收合格并交付")
            else:
                factors_error["承包人负责照管工程及工程相关的材料、工程设备的起始时间"] = "第三部分_3.6 工程照管与成品、半成品保护:承包人负责照管工程及工程相关的材料、工程设备的起始时间未填写完整"
                addRemarkInDoc(word, document, "承包人负责照管工程及工程相关的材料、工程设备的起始时间",
                               f"第三部分_3.6 工程照管与成品、半成品保护:承包人负责照管工程及工程相关的材料、工程设备的起始时间未填写完整")
        except:
            missObject += "第三部分_3.6 工程照管与成品、半成品保护:要素“承包人负责照管工程及工程相关的材料、工程设备的起始时间”缺失\n"

        try:
            match = '承包人是否提供履约担保：(.*?)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["承包人是否提供履约担保"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("承包人是否提供履约担保")
            else:
                factors_error["承包人是否提供履约担保"] = "第三部分_3.6 工程照管与成品、半成品保护:承包人是否提供履约担保未填写完整"
                addRemarkInDoc(word, document, "承包人是否提供履约担保", f"第三部分_3.6 工程照管与成品、半成品保护:承包人是否提供履约担保未填写完整")
        except:
            missObject += "第三部分_3.6 工程照管与成品、半成品保护:要素“承包人是否提供履约担保”缺失\n"

        try:
            match = '承包人提供履约担保的形式、金额及期限的：\n履约担保形式：(.*?)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["承包人提供履约担保的形式、金额及期限的"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '现金担保或银行保函担保' and factor != "现金担保" and factor != '银行保函担保':
                    factors_error["承包人提供履约担保的形式、金额及期限的履约担保形式"] = "第三部分_3.6 工程照管与成品、半成品保护:承包人提供履约担保的形式、金额及期限的履约担保形式未填写完整"
                    addRemarkInDoc(word, document, "承包人提供履约担保的形式、金额及期限的",
                                   f"第三部分_3.6 工程照管与成品、半成品保护:承包人提供履约担保的形式、金额及期限的履约担保形式未填写完整")
                else:
                    factors_ok.append("承包人提供履约担保的形式、金额及期限的履约担保形式")
            else:
                factors_error["承包人提供履约担保的形式、金额及期限的"] = "承包人提供履约担保的形式、金额及期限的履约担保形式未填写完整"
                addRemarkInDoc(word, document, "承包人提供履约担保的形式、金额及期限的",
                               f"第三部分_3.6 工程照管与成品、半成品保护:承包人提供履约担保的形式、金额及期限的履约担保形式未填写完整")
        except:
            missObject += "第三部分_3.6 工程照管与成品、半成品保护:要素“承包人提供履约担保的形式、金额及期限的履约担保形式”缺失\n"

        try:
            match = '基本履约保证金 ：中标价的(.*?)%'
            factor = re.findall(match, text)[0].replace(" ", "").replace("【", "").replace("】", "").replace("；",
                                                                                                           "").replace(
                "。", "")
            factors["基本履约保证金：中标价的"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("基本履约保证金：中标价的")
                factors_to_inform["基本履约保证金：中标价的"] = "第三部分_3.6 工程照管与成品、半成品保护:该金额是否与下列基本履约保证金计算公式计算所得结果一致"
                addRemarkInDoc(word, document, "基本履约保证金 ：", f"第三部分_3.6 工程照管与成品、半成品保护:该金额是否与下列基本履约保证金计算公式计算所得结果一致")
            else:
                factors_error["基本履约保证金：中标价的"] = "第三部分_3.6 工程照管与成品、半成品保护:基本履约保证金：中标价的未填写完整"
                addRemarkInDoc(word, document, "基本履约保证金 ：", f"第三部分_3.6 工程照管与成品、半成品保护:基本履约保证金 ：中标价的未填写完整")
        except:
            missObject += "第三部分_3.6 工程照管与成品、半成品保护:要素“基本履约保证金 ：中标价的”缺失\n"

        flag0 = 0
        flag1 = 0
        try:
            match = '基本履约保证金＝中标价（扣除暂列金及专业工程暂估价部分）(.*?)元'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["基本履约保证金＝中标价"] = factor
            if factor != "" and factor != "；" and factor != "。":
                flag0 = 1
                zhongbiaojia = float(factor)
                factors_ok.append("基本履约保证金＝中标价")
                factors_to_inform["基本履约保证金＝中标价"] = "第三部分_3.6 工程照管与成品、半成品保护:请根据中标通知书核实该条计算公式的中标价是否正确"
                addRemarkInDoc(word, document, "基本履约保证金＝中标价", f"第三部分_3.6 工程照管与成品、半成品保护:请根据中标通知书核实该条计算公式的中标价是否正确")
            else:
                factors_error["基本履约保证金＝中标价"] = "第三部分_3.6 工程照管与成品、半成品保护:基本履约保证金＝中标价未填写完整"
                addRemarkInDoc(word, document, "基本履约保证金＝中标价", f"第三部分_3.6 工程照管与成品、半成品保护:基本履约保证金＝中标价未填写完整")
        except:
            missObject += "第三部分_3.6 工程照管与成品、半成品保护:要素“基本履约保证金＝中标价”缺失\n"

        try:
            match = '基本履约保证金＝中标价（扣除暂列金及专业工程暂估价部分）.*?元(.*?)%'
            factor = re.findall(match, text)[0].replace(" ", "")
            match = '.*的(.*)'
            factor = re.findall(match, factor)[0].replace(" ", "").replace("【", "").replace("】", "").replace("；",
                                                                                                             "").replace(
                "。", "")
            factors["基本履约保证金＝中标价（扣除暂列金及专业工程暂估价部分）"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor.isdigit() and 10 <= float(factor) <= 20:
                    flag1 = 1
                    percentage = float(factor)
                    factors_ok.append("基本履约保证金＝中标价（扣除暂列金及专业工程暂估价部分）")
                    factors_to_inform[
                        "基本履约保证金＝中标价（扣除暂列金及专业工程暂估价部分）"] = "第三部分_3.6 工程照管与成品、半成品保护:请核实该计算公式中的比例是否符合双方约定，请取得财务部门确认"
                    addRemarkInDoc(word, document, "基本履约保证金＝中标价",
                                   f"第三部分_3.6 工程照管与成品、半成品保护:请核实该计算公式中的比例是否符合双方约定，请取得财务部门确认")
                else:
                    factors_error[
                        "基本履约保证金＝中标价（扣除暂列金及专业工程暂估价部分）"] = "第三部分_3.6 工程照管与成品、半成品保护:基本履约保证金＝中标价（扣除暂列金及专业工程暂估价部分）比例数值在10-20中选择"
                    addRemarkInDoc(word, document, "基本履约保证金＝中标价",
                                   f"第三部分_3.6 工程照管与成品、半成品保护:基本履约保证金＝中标价（扣除暂列金及专业工程暂估价部分）比例数值在10-20中选择")
            else:
                factors_error[
                    "基本履约保证金＝中标价（扣除暂列金及专业工程暂估价部分）"] = "第三部分_3.6 工程照管与成品、半成品保护:基本履约保证金＝中标价（扣除暂列金及专业工程暂估价部分）未填写完整"
                addRemarkInDoc(word, document, "基本履约保证金＝中标价",
                               f"第三部分_3.6 工程照管与成品、半成品保护:基本履约保证金＝中标价（扣除暂列金及专业工程暂估价部分）未填写完整")
        except:
            missObject += "第三部分_3.6 工程照管与成品、半成品保护:要素“基本履约保证金＝中标价（扣除暂列金及专业工程暂估价部分）”缺失\n"

        danbaojin = False
        try:
            if flag0 and flag1:
                danbaojin = zhongbiaojia * percentage * 0.01
        except:
            danbaojin = False

        try:
            match = '履约保证金的退还：竣工验收合格后，承包人提交退还申请后(.*?)天内'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["履约保证金的退还：竣工验收合格后，承包人提交退还申请后"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if int(factor) >= 28:
                    factors_ok.append("履约保证金的退还：竣工验收合格后，承包人提交退还申请后")
                    factors_to_inform["履约保证金的退还：竣工验收合格后，承包人提交退还申请后"] = "第三部分_3.6 工程照管与成品、半成品保护:请取得财务部门确认"
                    addRemarkInDoc(word, document, "履约保证金的退还：竣工验收合格后，承包人提交退还申请后", f"第三部分_3.6 工程照管与成品、半成品保护:请取得财务部门确认")
                else:
                    factors_error[
                        "履约保证金的退还：竣工验收合格后，承包人提交退还申请后"] = "第三部分_3.6 工程照管与成品、半成品保护:履约保证金的退还：竣工验收合格后，承包人提交退还申请后数值不少于28天"
                    addRemarkInDoc(word, document, "履约保证金的退还：竣工验收合格后，承包人提交退还申请后",
                                   f"第三部分_3.6 工程照管与成品、半成品保护:履约保证金的退还：竣工验收合格后，承包人提交退还申请后数值不少于28天")
            else:
                factors_error["履约保证金的退还：竣工验收合格后，承包人提交退还申请后"] = "第三部分_3.6 工程照管与成品、半成品保护:履约保证金的退还：竣工验收合格后，承包人提交退还申请后未填写完整"
                addRemarkInDoc(word, document, "履约保证金的退还：竣工验收合格后，承包人提交退还申请后",
                               f"第三部分_3.6 工程照管与成品、半成品保护:履约保证金的退还：竣工验收合格后，承包人提交退还申请后未填写完整")
        except:
            missObject += "第三部分_3.6 工程照管与成品、半成品保护:要素“履约保证金的退还：竣工验收合格后，承包人提交退还申请后”缺失\n"

        # 3.4
        try:
            match = '关于监理人的监理内容：(.*?)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["关于监理人的监理内容"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("关于监理人的监理内容")
            else:
                factors_error["关于监理人的监理内容"] = "第三部分_4. 监理人:关于监理人的监理内容未填写完整"
                addRemarkInDoc(word, document, "关于监理人的监理内容", f"第三部分_4. 监理人:关于监理人的监理内容未填写完整")
        except:
            missObject += "第三部分_4. 监理人:要素“关于监理人的监理内容”缺失\n"

        try:
            match = '关于监理人的监理权限：(.*?)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["关于监理人的监理权限"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("关于监理人的监理权限")
            else:
                factors_error["关于监理人的监理权限"] = "第三部分_4. 监理人:关于监理人的监理权限未填写完整"
                addRemarkInDoc(word, document, "关于监理人的监理权限", f"第三部分_4. 监理人:关于监理人的监理权限未填写完整")
        except:
            missObject += "第三部分_4. 监理人:要素“关于监理人的监理权限”缺失\n"

        try:
            match = '关于监理人在施工现场的办公场所、生活场所的提供和费用承担的约定：(.*?)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["关于监理人在施工现场的办公场所、生活场所的提供和费用承担的约定"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("关于监理人在施工现场的办公场所、生活场所的提供和费用承担的约定")

            else:
                factors_error["关于监理人在施工现场的办公场所、生活场所的提供和费用承担的约定"] = "第三部分_4. 监理人:关于监理人在施工现场的办公场所、生活场所的提供和费用承担的约定未填写完整"
                addRemarkInDoc(word, document, "关于监理人在施工现场的办公场所、生活场所的提供和费用承担的约定",
                               f"第三部分_4. 监理人:关于监理人在施工现场的办公场所、生活场所的提供和费用承担的约定未填写完整")
        except:
            missObject += "第三部分_4. 监理人:要素“关于监理人在施工现场的办公场所、生活场所的提供和费用承担的约定”缺失\n"

        factors_to_inform["4.1监理人的一般规定"] = "第三部分_4. 监理人:请核对本条约定与监理合同的相关约定是否一致"
        addRemarkInDoc(word, document, "4.1监理人的一般规定", f"第三部分_4. 监理人:请核对本条约定与监理合同的相关约定是否一致")

        try:
            match = '总监理工程师：\n姓    名：(.*?)；'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["总监理工程师姓名"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("总监理工程师姓名")
            else:
                factors_error["总监理工程师姓名"] = "第三部分_4. 监理人:总监理工程师姓名未填写完整"
                addRemarkInDoc(word, document, "4.2 监理人员", f"第三部分_4. 监理人:总监理工程师姓名未填写完整")
        except:
            missObject += "第三部分_4. 监理人:要素“总监理工程师姓名”缺失\n"

        try:
            match = '总监理工程师：\n姓    名：.*；\n职    务：(.*?)；'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["总监理工程师职务"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("总监理工程师职务")
            else:
                factors_error["总监理工程师职务"] = "第三部分_4. 监理人:总监理工程师职务未填写完整"
                addRemarkInDoc(word, document, "4.2 监理人员", f"第三部分_4. 监理人:总监理工程师职务未填写完整")
        except:
            missObject += "第三部分_4. 监理人:要素“总监理工程师职务”缺失\n"

        try:
            match = '总监理工程师：\n姓    名：.*；\n职    务：.*；\n监理工程师执业资格证书号：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["监理工程师执业资格证书号"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("监理工程师执业资格证书号")
            else:
                factors_error["监理工程师执业资格证书号"] = "第三部分_4. 监理人:监理工程师执业资格证书号未填写完整"
                addRemarkInDoc(word, document, "4.2 监理人员", f"第三部分_4. 监理人:监理工程师执业资格证书号未填写完整")
        except:
            missObject += "第三部分_4. 监理人:要素“总监理工程师监理工程师执业资格证书号”缺失\n"

        try:
            match = '总监理工程师：\n姓    名：.*；\n职    务：.*；\n监理工程师执业资格证书号：.*\n联系电话：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["总监理工程师联系电话"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("总监理工程师联系电话")
            else:
                factors_error["总监理工程师联系电话"] = "第三部分_4. 监理人:总监理工程师联系电话未填写完整"
                addRemarkInDoc(word, document, "4.2 监理人员", f"第三部分_4. 监理人:总监理工程师联系电话未填写完整")
        except:
            missObject += "第三部分_4. 监理人:要素“总监理工程师联系电话”缺失\n"

        try:
            match = '总监理工程师：\n姓    名：.*；\n职    务：.*；\n监理工程师执业资格证书号：.*\n联系电话：.*\n电子信箱：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["总监理工程师电子信箱"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("总监理工程师电子信箱")
            else:
                factors_error["总监理工程师电子信箱"] = "第三部分_4. 监理人:总监理工程师电子信箱未填写完整"
                addRemarkInDoc(word, document, "4.2 监理人员", f"第三部分_4. 监理人:总监理工程师电子信箱未填写完整")
        except:
            missObject += "第三部分_4. 监理人:要素“总监理工程师电子信箱”缺失\n"

        try:
            match = '总监理工程师：\n姓    名：.*；\n职    务：.*；\n监理工程师执业资格证书号：.*\n联系电话：.*\n电子信箱：.*\n通信地址：(.*)'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["总监理工程师通信地址"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("总监理工程师通信地址")
            else:
                factors_error["总监理工程师通信地址"] = "第三部分_4. 监理人:总监理工程师通信地址未填写完整"
                addRemarkInDoc(word, document, "4.2 监理人员", f"第三部分_4. 监理人:总监理工程师通信地址未填写完整")
        except:
            missObject += "第三部分_4. 监理人:要素“总监理工程师通信地址”缺失\n"

        try:
            match = '总监理工程师：\n姓    名：.*；\n职    务：.*；\n监理工程师执业资格证书号：.*\n联系电话：.*\n电子信箱：.*\n通信地址：.*\n关于监理人的其他约定：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["关于监理人的其他约定"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("关于监理人的其他约定")
                factors_to_inform["关于监理人的其他约定"] = "第三部分_4. 监理人:请核对本条约定与监理合同的相关约定是否一致"
                addRemarkInDoc(word, document, "4.2 监理人员", f"第三部分_4. 监理人:请核对本条约定与监理合同的相关约定是否一致")
            else:
                factors_error["关于监理人的其他约定"] = "第三部分_4. 监理人:关于监理人的其他约定未填写完整"
                addRemarkInDoc(word, document, "4.2 监理人员", f"第三部分_4. 监理人:关于监理人的其他约定未填写完整")
        except:
            missObject += "第三部分_4. 监理人:要素“关于监理人的其他约定”缺失\n"

        try:
            match = '在发包人和承包人不能通过协商达成一致意见时，发包人授权监理人对以下事项进行确定：\n(.*?)。\n'
            factor = re.findall(match, text, re.S)[0].replace(" ", "").replace("；", "").replace("。", "").replace("；",
                                                                                                                 "").replace(
                "。", "")
            factors["在发包人和承包人不能通过协商达成一致意见时，发包人授权监理人对以下事项进行确定"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("在发包人和承包人不能通过协商达成一致意见时，发包人授权监理人对以下事项进行确定")
                try:
                    match = '（1）材料、设备技术参数是否满足设计文件及招标文件要求；.*?（2）工法、技术措施是否满足设计文件及验收规范要求；.*?（3）工程半成品、成品的感观质量是否满足验收规范要求。\n'
                    factor = re.findall(match, text, re.S)[0].replace(" ", "").replace("；", "").replace("。",
                                                                                                        "").replace(
                        "；", "").replace("。", "")
                except:
                    factors_to_inform["在发包人和承包人不能通过协商达成一致意见时，发包人授权监理人对以下事项进行确定"] = "第三部分_4. 监理人:请核对本条约定与监理合同的相关约定是否一致"
                    addRemarkInDoc(word, document, "在发包人和承包人不能通过协商达成一致意见时，发包人授权监理人对以下事项进行确定",
                                   f"第三部分_4. 监理人:请核对本条约定与监理合同的相关约定是否一致")
            else:
                factors_error[
                    "在发包人和承包人不能通过协商达成一致意见时，发包人授权监理人对以下事项进行确定"] = "第三部分_4. 监理人:在发包人和承包人不能通过协商达成一致意见时，发包人授权监理人对以下事项进行确定未填写完整"
                addRemarkInDoc(word, document, "在发包人和承包人不能通过协商达成一致意见时，发包人授权监理人对以下事项进行确定",
                               f"第三部分_4. 监理人:在发包人和承包人不能通过协商达成一致意见时，发包人授权监理人对以下事项进行确定未填写完整")
        except:
            missObject += "第三部分_4. 监理人:要素“在发包人和承包人不能通过协商达成一致意见时，发包人授权监理人对以下事项进行确定”缺失\n"

        # 3.5
        try:
            match = '特殊质量标准和要求：(.*?)。\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["特殊质量标准和要求"] = factor
            if factor != "" and factor != "；" and factor != "。":

                if factor != '无'.replace(" ", "").replace("；", "").replace("。", ""):
                    factors_error["特殊质量标准和要求"] = "第三部分_5. 工程质量:请核实双方对特殊质量的标准与要求，并将相应标准作为合同附件"
                    addRemarkInDoc(word, document, "5.1.1", f"第三部分_5. 工程质量:请核实双方对特殊质量的标准与要求，并将相应标准作为合同附件")
                else:
                    factors_ok.append("特殊质量标准和要求")
            else:
                factors_error["特殊质量标准和要求"] = "第三部分_5. 工程质量:特殊质量标准和要求未填写完整"
                addRemarkInDoc(word, document, "5.1.1", f"第三部分_5. 工程质量:特殊质量标准和要求未填写完整")
        except:
            missObject += "第三部分_5. 工程质量:要素“特殊质量标准和要求”缺失\n"

        try:
            match = '关于工程奖项的约定：(.*?)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["关于工程奖项的约定"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '无'.replace(" ", "").replace("；", "").replace("。", ""):
                    factors_error["关于工程奖项的约定"] = "第三部分_5. 工程质量:请核实双方对特殊质量的标准与要求，并将相应标准作为合同附件"
                    addRemarkInDoc(word, document, "5.1.1", f"第三部分_5. 工程质量:请核实双方对特殊质量的标准与要求，并将相应标准作为合同附件")
                else:
                    factors_ok.append("关于工程奖项的约定")
            else:
                factors_error["关于工程奖项的约定"] = "第三部分_5. 工程质量:关于工程奖项的约定未填写完整"
                addRemarkInDoc(word, document, "关于工程奖项的约定", f"第三部分_5. 工程质量:关于工程奖项的约定未填写完整")
        except:
            missObject += "第三部分_5. 工程质量:要素“关于工程奖项的约定”缺失\n"

        try:
            match = '承包人提前通知监理人隐蔽工程检查的期限的约定：\n(.*)。'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["承包人提前通知监理人隐蔽工程检查的期限的约定"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("承包人提前通知监理人隐蔽工程检查的期限的约定")
                if factor != '按通用条款执行':
                    factors_to_inform[
                        "承包人提前通知监理人隐蔽工程检查的期限的约定"] = "第三部分_5. 工程质量:请认真比对通用条款的约定，原则上承包人提前通知监理人隐蔽工程检查的期限不应小于隐蔽工程封闭前72小时"
                    addRemarkInDoc(word, document, "承包人提前通知监理人隐蔽工程检查的期限的约定",
                                   f"第三部分_5. 工程质量:请认真比对通用条款的约定，原则上承包人提前通知监理人隐蔽工程检查的期限不应小于隐蔽工程封闭前72小时")
            else:
                factors_error["承包人提前通知监理人隐蔽工程检查的期限的约定"] = "第三部分_5. 工程质量:承包人提前通知监理人隐蔽工程检查的期限的约定未填写完整"
                addRemarkInDoc(word, document, "承包人提前通知监理人隐蔽工程检查的期限的约定", f"第三部分_5. 工程质量:承包人提前通知监理人隐蔽工程检查的期限的约定未填写完整")
        except:
            missObject += "第三部分_5. 工程质量:要素“承包人提前通知监理人隐蔽工程检查的期限的约定”缺失\n"

        try:
            match = '监理人不能按时进行检查时，应提前(.*?)小时'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["监理人不能按时进行检查时，应提前"] = factor
            if factor != "" and factor != "；" and factor != "。" and factor != '（24-40之间选择）':
                factor = re.search(r'[0-9]+', factor).group(0)
                if 24 <= int(factor) < 40:
                    factors_ok.append("监理人不能按时进行检查时，应提前")
                else:
                    factors_error["监理人不能按时进行检查时，应提前"] = "第三部分_5. 工程质量:24-40之间选择"
                    addRemarkInDoc(word, document, "监理人不能按时进行检查时，应提前", f"第三部分_5. 工程质量:24-40之间选择")
            else:
                factors_error["监理人不能按时进行检查时，应提前"] = "第三部分_5. 工程质量:监理人不能按时进行检查时，应提前时间未填写完整"
                addRemarkInDoc(word, document, "监理人不能按时进行检查时，应提前", f"第三部分_5. 工程质量:监理人不能按时进行检查时，应提前时间未填写完整")
        except:
            missObject += "第三部分_5. 工程质量:要素“监理人不能按时进行检查时，应提前”缺失\n"

        try:
            match = '监理人不能按时进行检查时，应提前.*小时提交书面延期要求。\n关于延期最长不得超过：(.*)小时。\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["关于延期最长不得超过"] = factor
            if factor != "" and factor != "；" and factor != "。" and factor != '（48-72之间选择）':
                factor = re.search(r'[0-9]+', factor).group(0)
                if 48 <= int(factor) < 72:
                    factors_ok.append("关于延期最长不得超过")
                else:
                    factors_error["关于延期最长不得超过"] = "第三部分_5. 工程质量:48-72之间选择"
                    addRemarkInDoc(word, document, "关于延期最长不得超过", f"第三部分_5. 工程质量:48-72之间选择")
            else:
                factors_error["关于延期最长不得超过"] = "第三部分_5. 工程质量:关于延期最长不得超过时间未填写完整"
                addRemarkInDoc(word, document, "关于延期最长不得超过", f"第三部分_5. 工程质量:关于延期最长不得超过时间未填写完整")
        except:
            missObject += "第三部分_5. 工程质量:要素“关于延期最长不得超过”缺失\n"

    # 第三部分_6-11
    if 1 == True:
        # 3.6
        try:
            match = '6.1.1 项目安全生产的达标目标及相应事项的约定：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["6.1.1 项目安全生产的达标目标及相应事项的约定"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '无'.replace(" ", "").replace("；", "").replace("。", ""):
                    factors_error["6.1.1 项目安全生产的达标目标及相应事项的约定"] = "第三部分_6. 安全文明施工与环境保护:请认真核对该条约定的具体内容"
                    addRemarkInDoc(word, document, "6.1.1 项目安全生产的达标目标及相应事项的约定", f"第三部分_6. 安全文明施工与环境保护:请认真核对该条约定的具体内容")
                else:
                    factors_ok.append("6.1.1 项目安全生产的达标目标及相应事项的约定")
            else:
                factors_error["6.1.1 项目安全生产的达标目标及相应事项的约定"] = "第三部分_6. 安全文明施工与环境保护:6.1.1 项目安全生产的达标目标及相应事项的约定未填写完整"
                addRemarkInDoc(word, document, "6.1.1 项目安全生产的达标目标及相应事项的约定",
                               f"第三部分_6. 安全文明施工与环境保护:6.1.1 项目安全生产的达标目标及相应事项的约定未填写完整")
        except:
            missObject += "第三部分_6. 安全文明施工与环境保护:要素“6.1.1 项目安全生产的达标目标及相应事项的约定”缺失\n"

        try:
            match = '6.1.2 关于治安保卫的特别约定：(.*)\n'
            factor = re.findall(match, text)[0].replace('\xa0', ' ').replace(" ", "").replace("。", "").replace("；",
                                                                                                               "").replace(
                "\t", "")
            factors["6.1.2 关于治安保卫的特别约定"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '承包人在其施工场地内，根据工程建设及安全的需要提供并负责维护施工使用的照明、围栏等相关设施，并配备看守或警卫。承包人未履行上述义务造成工程、财产的损失和人员伤害，由承包人承担全部责任并负担所发生的全部费用'.replace(
                        " ", "").replace("；", "").replace("。", ""):
                    factors_error["6.1.2 关于治安保卫的特别约定"] = "第三部分_6. 安全文明施工与环境保护:施工期间，施工场地的治安保卫应由承包人负责并承担相应费用"
                    addRemarkInDoc(word, document, "6.1.2 关于治安保卫的特别约定",
                                   f"第三部分_6. 安全文明施工与环境保护:施工期间，施工场地的治安保卫应由承包人负责并承担相应费用")
                else:
                    factors_ok.append("6.1.2 关于治安保卫的特别约定")
            else:
                factors_error["6.1.2 关于治安保卫的特别约定"] = "第三部分_6. 安全文明施工与环境保护:6.1.2 关于治安保卫的特别约定未填写完整"
                addRemarkInDoc(word, document, "6.1.2 关于治安保卫的特别约定", f"第三部分_6. 安全文明施工与环境保护:6.1.2 关于治安保卫的特别约定未填写完整")
        except:
            missObject += "第三部分_6. 安全文明施工与环境保护:要素“6.1.2 关于治安保卫的特别约定”缺失\n"

        try:
            match = '关于编制施工场地治安管理计划的约定：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["关于编制施工场地治安管理计划的约定"] = factor
            if factor != "" and factor != "；" and factor != "。":
                pattern = r'.*工程开工后(.*)天内.*'
                try:
                    d = re.findall(pattern, factor)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t",
                                                                                                                  "")
                    if int(d) < 1 or int(d) > 15:
                        factors_error[
                            "关于编制施工场地治安管理计划的约定"] = "第三部分_6. 安全文明施工与环境保护:施工场地治安管理计划应由承包人在合理期限内提交，合理期限应定于工程开工后1-15天内"
                        addRemarkInDoc(word, document, "关于编制施工场地治安管理计划的约定",
                                       f"第三部分_6. 安全文明施工与环境保护:施工场地治安管理计划应由承包人在合理期限内提交，合理期限应定于工程开工后1-15天内")
                    else:
                        factors_ok.append("关于编制施工场地治安管理计划的约定")
                except:
                    factors_error["关于编制施工场地治安管理计划的约定"] = "第三部分_6. 安全文明施工与环境保护:合理期限提取错误"
                    addRemarkInDoc(word, document, "关于编制施工场地治安管理计划的约定", f"第三部分_6. 安全文明施工与环境保护:合理期限提取错误")
            else:
                factors_error["关于编制施工场地治安管理计划的约定"] = "第三部分_6. 安全文明施工与环境保护:关于编制施工场地治安管理计划的约定未填写完整"
                addRemarkInDoc(word, document, "关于编制施工场地治安管理计划的约定", f"第三部分_6. 安全文明施工与环境保护:关于编制施工场地治安管理计划的约定未填写完整")
        except:
            missObject += "第三部分_6. 安全文明施工与环境保护:要素“关于编制施工场地治安管理计划的约定”缺失\n"

        try:
            match = '合同当事人对文明施工的要求：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["合同当事人对文明施工的要求"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '按发包人要求'.replace(" ", "").replace("；", "").replace("。", ""):
                    factors_error["合同当事人对文明施工的要求"] = "第三部分_6. 安全文明施工与环境保护:请核实发包人对文明施工的要求与本条约定是否一致"
                    addRemarkInDoc(word, document, "合同当事人对文明施工的要求", f"第三部分_6. 安全文明施工与环境保护:请核实发包人对文明施工的要求与本条约定是否一致")
                else:
                    factors_ok.append("合同当事人对文明施工的要求")
            else:
                factors_error["合同当事人对文明施工的要求"] = "第三部分_6. 安全文明施工与环境保护:合同当事人对文明施工的要求未填写完整"
                addRemarkInDoc(word, document, "合同当事人对文明施工的要求", f"第三部分_6. 安全文明施工与环境保护:合同当事人对文明施工的要求未填写完整")
        except:
            missObject += "要素“合同当事人对文明施工的要求”缺失\n"

        try:
            match = '6.1.6 关于安全文明施工费支付比例和支付期限的约定：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t",
                                                                                                           "").replace(
                '\t', "")
            factors["6.1.6 关于安全文明施工费支付比例和支付期限的约定"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '合同签订完成后，承包人提出申请后28天内发包人支付安全文明施工措施费用基本费的%（工期在一年及一年以上的为50%，一年以内的为70%）即￥元（大写人民币：）。当工程进度款累计支付达到合同金额的％时（此数值同安全文明施工费第一次支付比例一致），其余部分安全文明施工措施费用与进度款同期支付'.replace(
                        " ", "").replace("；", "").replace("。", ""):
                    factors_error["6.1.6 关于安全文明施工费支付比例和支付期限的约定"] = "第三部分_6. 安全文明施工与环境保护:请核实双方对安全文明施工费支付的约定与本条约定是否一致"
                    addRemarkInDoc(word, document, "6.1.6 关于安全文明施工费支付比例和支付期限的约定",
                                   f"第三部分_6. 安全文明施工与环境保护:请核实双方对安全文明施工费支付的约定与本条约定是否一致")
                else:
                    factors_ok.append("6.1.6 关于安全文明施工费支付比例和支付期限的约定")
            else:
                factors_error["6.1.6 关于安全文明施工费支付比例和支付期限的约定"] = "第三部分_6. 安全文明施工与环境保护:6.1.6 关于安全文明施工费支付比例和支付期限的约定未填写完整"
                addRemarkInDoc(word, document, "6.1.6 关于安全文明施工费支付比例和支付期限的约定",
                               f"第三部分_6. 安全文明施工与环境保护:6.1.6 关于安全文明施工费支付比例和支付期限的约定未填写完整")
        except:
            missObject += "第三部分_6. 安全文明施工与环境保护:要素“6.1.6 关于安全文明施工费支付比例和支付期限的约定”缺失\n"

        # 3.7
        try:
            match = '7.1.1 合同当事人约定的施工组织设计应包括的其他内容：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["7.1.1 合同当事人约定的施工组织设计应包括的其他内容"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '无'.replace(" ", "").replace("；", "").replace("。", ""):
                    factors_error["7.1.1 合同当事人约定的施工组织设计应包括的其他内容"] = "第三部分_7. 工期和进度:请核实双方约定的施工组织设计"
                    addRemarkInDoc(word, document, "7.1.1 合同当事人约定的施工组织设计应包括的其他内容", f"第三部分_7. 工期和进度:请核实双方约定的施工组织设计")
                else:
                    factors_ok.append("7.1.1 合同当事人约定的施工组织设计应包括的其他内容")
            else:
                factors_error["7.1.1 合同当事人约定的施工组织设计应包括的其他内容"] = "第三部分_7. 工期和进度:7.1.1 合同当事人约定的施工组织设计应包括的其他内容未填写完整"
                addRemarkInDoc(word, document, "7.1.1 合同当事人约定的施工组织设计应包括的其他内容",
                               f"第三部分_7. 工期和进度:7.1.1 合同当事人约定的施工组织设计应包括的其他内容未填写完整")
        except:
            missObject += "第三部分_7. 工期和进度:要素“7.1.1 合同当事人约定的施工组织设计应包括的其他内容”缺失\n"

        try:
            match = '承包人提交详细施工组织设计的期限的约定：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["承包人提交详细施工组织设计的期限的约定"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '按通用条款执行'.replace(" ", "").replace("；", "").replace("。", ""):
                    factors_error["承包人提交详细施工组织设计的期限的约定"] = "第三部分_7. 工期和进度:请认真比对通用条款的约定，原则上承包人提交详细施工组织设计的期限不应超过通用条款约定"
                    addRemarkInDoc(word, document, "承包人提交详细施工组织设计的期限的约定",
                                   f"第三部分_7. 工期和进度:请认真比对通用条款的约定，原则上承包人提交详细施工组织设计的期限不应超过通用条款约定")
                else:
                    factors_ok.append("承包人提交详细施工组织设计的期限的约定")
            else:
                factors_error["承包人提交详细施工组织设计的期限的约定"] = "第三部分_7. 工期和进度:承包人提交详细施工组织设计的期限的约定未填写完整"
                addRemarkInDoc(word, document, "承包人提交详细施工组织设计的期限的约定", f"第三部分_7. 工期和进度:承包人提交详细施工组织设计的期限的约定未填写完整")
        except:
            missObject += "第三部分_7. 工期和进度:要素“承包人提交详细施工组织设计的期限的约定”缺失\n"

        try:
            match = '发包人和监理人在收到详细的施工组织设计后确认或提出修改意见的期限：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["发包人和监理人在收到详细的施工组织设计后确认或提出修改意见的期限"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '在监理人收到施工组织设计后7天内确认或提出修改意见'.replace(" ", "").replace("；", "").replace("。", ""):
                    factors_error[
                        "发包人和监理人在收到详细的施工组织设计后确认或提出修改意见的期限"] = "第三部分_7. 工期和进度:原则上发包人与监理人提出确认或修改针对施工组织计划的意见的期限应在7-30日期限内选择"
                    addRemarkInDoc(word, document, "发包人和监理人在收到详细的施工组织设计后确认或提出修改意见的期限",
                                   f"第三部分_7. 工期和进度:原则上发包人与监理人提出确认或修改针对施工组织计划的意见的期限应在7-30日期限内选择")
                else:
                    factors_ok.append("发包人和监理人在收到详细的施工组织设计后确认或提出修改意见的期限")
            else:
                factors_error[
                    "发包人和监理人在收到详细的施工组织设计后确认或提出修改意见的期限"] = "第三部分_7. 工期和进度:发包人和监理人在收到详细的施工组织设计后确认或提出修改意见的期限未填写完整"
                addRemarkInDoc(word, document, "发包人和监理人在收到详细的施工组织设计后确认或提出修改意见的期限",
                               f"第三部分_7. 工期和进度:发包人和监理人在收到详细的施工组织设计后确认或提出修改意见的期限未填写完整")
        except:
            missObject += "第三部分_7. 工期和进度:要素“发包人和监理人在收到详细的施工组织设计后确认或提出修改意见的期限”缺失\n"

        try:
            match = '发包人和监理人在收到修订的施工进度计划后确认或提出修改意见的期限：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["发包人和监理人在收到修订的施工进度计划后确认或提出修改意见的期限"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '在收到修订的施工进度计划后7天内完成审核和批准或提出修改意见'.replace(" ", "").replace("；", "").replace("。", ""):
                    factors_error[
                        "发包人和监理人在收到修订的施工进度计划后确认或提出修改意见的期限"] = "第三部分_7. 工期和进度:原则上发包人与监理人提出确认或修改针对施工进度计划的意见的期限应在7-15日期限内选择"
                    addRemarkInDoc(word, document, "发包人和监理人在收到修订的施工进度计划后确认或提出修改意见的期限",
                                   f"第三部分_7. 工期和进度:原则上发包人与监理人提出确认或修改针对施工进度计划的意见的期限应在7-15日期限内选择")
                else:
                    factors_ok.append("发包人和监理人在收到修订的施工进度计划后确认或提出修改意见的期限")
            else:
                factors_error[
                    "发包人和监理人在收到修订的施工进度计划后确认或提出修改意见的期限"] = "第三部分_7. 工期和进度:发包人和监理人在收到修订的施工进度计划后确认或提出修改意见的期限未填写完整"
                addRemarkInDoc(word, document, "发包人和监理人在收到修订的施工进度计划后确认或提出修改意见的期限",
                               f"第三部分_7. 工期和进度:发包人和监理人在收到修订的施工进度计划后确认或提出修改意见的期限未填写完整")
        except:
            missObject += "第三部分_7. 工期和进度:要素“发包人和监理人在收到修订的施工进度计划后确认或提出修改意见的期限”缺失\n"

        try:
            match = '关于承包人提交工程开工报审表的期限：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["关于承包人提交工程开工报审表的期限"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '取得施工许可证后3天内'.replace(" ", "").replace("；", "").replace("。", ""):
                    factors_error["关于承包人提交工程开工报审表的期限"] = "第三部分_7. 工期和进度:原则上承包人提交工程开工报审表的期限应在1-7日期限内选择"
                    addRemarkInDoc(word, document, "关于承包人提交工程开工报审表的期限", f"第三部分_7. 工期和进度:原则上承包人提交工程开工报审表的期限应在1-7日期限内选择")
                else:
                    factors_ok.append("关于承包人提交工程开工报审表的期限")
            else:
                factors_error["关于承包人提交工程开工报审表的期限"] = "第三部分_7. 工期和进度:关于承包人提交工程开工报审表的期限未填写完整"
                addRemarkInDoc(word, document, "关于承包人提交工程开工报审表的期限", f"第三部分_7. 工期和进度:关于承包人提交工程开工报审表的期限未填写完整")
        except:
            missObject += "第三部分_7. 工期和进度:要素“关于承包人提交工程开工报审表的期限”缺失\n"

        try:
            match = '关于发包人应完成的其他开工准备工作及期限：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["关于发包人应完成的其他开工准备工作及期限"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '及时准备完善办理施工手续的相关资料'.replace(" ", "").replace("；", "").replace("。", ""):
                    factors_error["关于发包人应完成的其他开工准备工作及期限"] = "第三部分_7. 工期和进度:请核实该条内容与双方约定是否一致"
                    addRemarkInDoc(word, document, "关于发包人应完成的其他开工准备工作及期限", f"第三部分_7. 工期和进度:请核实该条内容与双方约定是否一致")
                else:
                    factors_ok.append("关于发包人应完成的其他开工准备工作及期限")
            else:
                factors_error["关于发包人应完成的其他开工准备工作及期限"] = "第三部分_7. 工期和进度:关于发包人应完成的其他开工准备工作及期限未填写完整"
                addRemarkInDoc(word, document, "关于发包人应完成的其他开工准备工作及期限", f"第三部分_7. 工期和进度:关于发包人应完成的其他开工准备工作及期限未填写完整")
        except:
            missObject += "第三部分_7. 工期和进度:要素“关于发包人应完成的其他开工准备工作及期限”缺失\n"

        try:
            match = '关于承包人应完成的其他开工准备工作及期限：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["关于承包人应完成的其他开工准备工作及期限"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '承包人应在合同签订后5天内明确人员、材料、设备等开工必备条件落实到位的合理时限，并按该时限要求进行准备，超出时限要求的，发包人有权追究其违约责任 '.replace(" ",
                                                                                                                 "").replace(
                    "；", "").replace("。", ""):
                    factors_error["关于承包人应完成的其他开工准备工作及期限"] = "第三部分_7. 工期和进度:原则上承包人完成其他开工准备的工作期限应在1-15日期限内选择"
                    addRemarkInDoc(word, document, "关于承包人应完成的其他开工准备工作及期限",
                                   f"第三部分_7. 工期和进度:原则上承包人完成其他开工准备的工作期限应在1-15日期限内选择")
                else:
                    factors_ok.append("关于承包人应完成的其他开工准备工作及期限")
            else:
                factors_error["关于承包人应完成的其他开工准备工作及期限"] = "第三部分_7. 工期和进度:关于承包人应完成的其他开工准备工作及期限未填写完整"
                addRemarkInDoc(word, document, "关于承包人应完成的其他开工准备工作及期限", f"第三部分_7. 工期和进度:关于承包人应完成的其他开工准备工作及期限未填写完整")
        except:
            missObject += "要素“关于承包人应完成的其他开工准备工作及期限”缺失\n"

        try:
            match = '因发包人原因造成监理人未能在计划开工日期之日起(.*)天内发出开工通知的'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["7.3.2开工通知"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '180'.replace(" ", "").replace("；", "").replace("。", ""):
                    factors_error["7.3.2开工通知"] = "第三部分_7. 工期和进度:原则该条期限不应少于180日"
                    addRemarkInDoc(word, document, "7.3.2开工通知", f"第三部分_7. 工期和进度:原则该条期限不应少于180日")
                else:
                    factors_ok.append("7.3.2开工通知")
            else:
                factors_error["7.3.2开工通知"] = "第三部分_7. 工期和进度:7.3.2开工通知未填写完整"
                addRemarkInDoc(word, document, "7.3.2开工通知", f"第三部分_7. 工期和进度:7.3.2开工通知未填写完整")
        except:
            missObject += "第三部分_7. 工期和进度:要素“7.3.2开工通知”缺失\n"

        try:
            match = '7.4.1发包人通过监理人向承包人提供测量基准点、基准线和水准点及其书面资料的期限：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["7.4.1发包人通过监理人向承包人提供测量基准点、基准线和水准点及其书面资料的期限"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '无'.replace(" ", "").replace("；", "").replace("。", ""):
                    factors_error["7.4.1发包人通过监理人向承包人提供测量基准点、基准线和水准点及其书面资料的期限"] = "第三部分_7. 工期和进度:请核实该条内容与双方约定是否一致"
                    addRemarkInDoc(word, document, "7.4.1发包人通过监理人向承包人提供测量基准点、基准线和水准点及其书面资料的期限",
                                   f"第三部分_7. 工期和进度:请核实该条内容与双方约定是否一致")
                else:
                    factors_ok.append("第三部分_7. 工期和进度:7.4.1发包人通过监理人向承包人提供测量基准点、基准线和水准点及其书面资料的期限")
            else:
                factors_error[
                    "7.4.1发包人通过监理人向承包人提供测量基准点、基准线和水准点及其书面资料的期限"] = "第三部分_7. 工期和进度:7.4.1发包人通过监理人向承包人提供测量基准点、基准线和水准点及其书面资料的期限未填写完整"
                addRemarkInDoc(word, document, "7.4.1发包人通过监理人向承包人提供测量基准点、基准线和水准点及其书面资料的期限",
                               f"第三部分_7. 工期和进度:7.4.1发包人通过监理人向承包人提供测量基准点、基准线和水准点及其书面资料的期限未填写完整")
        except:
            missObject += "第三部分_7. 工期和进度:要素“7.4.1发包人通过监理人向承包人提供测量基准点、基准线和水准点及其书面资料的期限”缺失\n"

        content_7_5_2_1 = None
        content_7_5_2_3 = None
        try:
            match = '因承包人原因造成工期延误，逾期竣工违约金的计算方法为：\n逾期天数乘以(.*)元/天支付逾期违约金，逾期超过.*天的，发包人有权解除合同，由此造成的损失全部由承包人承担。\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["因承包人原因造成工期延误，逾期竣工违约金的计算方法_价格"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    factor = int(factor)
                    if factor < 5000:
                        factors_error[
                            "因承包人原因造成工期延误，逾期竣工违约金的计算方法_价格"] = "第三部分_7. 工期和进度:承包人逾期竣工承担的违约责任，不应低于以下约定：按5000元/天标准支付逾期违约金"
                        addRemarkInDoc(word, document, "因承包人原因造成工期延误，逾期竣工违约金的计算方法",
                                       f"第三部分_7. 工期和进度:承包人逾期竣工承担的违约责任，不应低于以下约定：按5000元/天标准支付逾期违约金")
                    else:
                        content_7_5_2_1 = factor
                        factors_ok.append("因承包人原因造成工期延误，逾期竣工违约金的计算方法_价格")
                except:
                    factors_error["因承包人原因造成工期延误，逾期竣工违约金的计算方法_价格"] = "第三部分_7. 工期和进度:逾期违约金提取错误"
                    addRemarkInDoc(word, document, "因承包人原因造成工期延误，逾期竣工违约金的计算方法", f"第三部分_7. 工期和进度:逾期违约金提取错误")
            else:
                factors_error["因承包人原因造成工期延误，逾期竣工违约金的计算方法_价格"] = "第三部分_7. 工期和进度:因承包人原因造成工期延误，逾期竣工违约金的计算方法_价格未填写完整"
                addRemarkInDoc(word, document, "因承包人原因造成工期延误，逾期竣工违约金的计算方法",
                               f"第三部分_7. 工期和进度:因承包人原因造成工期延误，逾期竣工违约金的计算方法_价格未填写完整")
        except:
            missObject += "第三部分_7. 工期和进度:要素“因承包人原因造成工期延误，逾期竣工违约金的计算方法_价格”缺失\n"

        try:
            match = '因承包人原因造成工期延误，逾期竣工违约金的计算方法为：\n逾期天数乘以.*元/天支付逾期违约金，逾期超过(.*)天的，发包人有权解除合同，由此造成的损失全部由承包人承担。\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["因承包人原因造成工期延误，逾期竣工违约金的计算方法_天数"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    factor = int(factor)
                    if factor > 90:
                        factors_error["因承包人原因造成工期延误，逾期竣工违约金的计算方法_天数"] = "第三部分_7. 工期和进度:逾期天数不能超过90"
                        addRemarkInDoc(word, document, "因承包人原因造成工期延误，逾期竣工违约金的计算方法",
                                       f"第三部分_7. 工期和进度:逾期天数不能超过90")
                    else:
                        factors_ok.append("因承包人原因造成工期延误，逾期竣工违约金的计算方法_天数")
                except:
                    factors_error["因承包人原因造成工期延误，逾期竣工违约金的计算方法_天数"] = "第三部分_7. 工期和进度:逾期时间提取错误"
                    addRemarkInDoc(word, document, "因承包人原因造成工期延误，逾期竣工违约金的计算方法", f"第三部分_7. 工期和进度:逾期时间提取错误")
            else:
                factors_error["因承包人原因造成工期延误，逾期竣工违约金的计算方法_天数"] = "第三部分_7. 工期和进度:因承包人原因造成工期延误，逾期竣工违约金的计算方法_天数未填写完整"
                addRemarkInDoc(word, document, "因承包人原因造成工期延误，逾期竣工违约金的计算方法",
                               f"第三部分_7. 工期和进度:因承包人原因造成工期延误，逾期竣工违约金的计算方法_天数未填写完整")
        except:
            missObject += "第三部分_7. 工期和进度:要素“因承包人原因造成工期延误，逾期竣工违约金的计算方法_天数”缺失\n"

        try:
            match = '因承包人原因造成工期延误，逾期竣工违约金的上限：最高不超过合同总价的(.*)%。\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["因承包人原因造成工期延误，逾期竣工违约金的上限"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    factor = int(factor)
                    if factor < 5 or factor > 20:
                        factors_error["因承包人原因造成工期延误，逾期竣工违约金的上限"] = "第三部分_7. 工期和进度:逾期竣工的违约金上限应应在5%-20%之间选择"
                        addRemarkInDoc(word, document, "因承包人原因造成工期延误，逾期竣工违约金的上限",
                                       f"第三部分_7. 工期和进度:逾期竣工的违约金上限应应在5%-20%之间选择")
                    else:
                        content_7_5_2_3 = factor
                        factors_ok.append("因承包人原因造成工期延误，逾期竣工违约金的上限")
                except:
                    factors_error["因承包人原因造成工期延误，逾期竣工违约金的上限"] = "第三部分_7. 工期和进度:上限提取错误"
                    addRemarkInDoc(word, document, "因承包人原因造成工期延误，逾期竣工违约金的上限", f"第三部分_7. 工期和进度:上限提取错误")
            else:
                factors_error["因承包人原因造成工期延误，逾期竣工违约金的上限"] = "第三部分_7. 工期和进度:因承包人原因造成工期延误，逾期竣工违约金的上限未填写完整"
                addRemarkInDoc(word, document, "因承包人原因造成工期延误，逾期竣工违约金的上限", f"第三部分_7. 工期和进度:因承包人原因造成工期延误，逾期竣工违约金的上限未填写完整")
        except:
            missObject += "第三部分_7. 工期和进度:要素“因承包人原因造成工期延误，逾期竣工违约金的上限”缺失\n"

        try:
            match = '不利物质条件的其他情形和有关约定：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["不利物质条件的其他情形和有关约定"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '无'.replace(" ", "").replace("；", "").replace("。", ""):
                    factors_error["不利物质条件的其他情形和有关约定"] = "第三部分_7. 工期和进度:请核实该条内容与双方约定是否一致"
                    addRemarkInDoc(word, document, "不利物质条件的其他情形和有关约定", f"第三部分_7. 工期和进度:请核实该条内容与双方约定是否一致")
                else:
                    factors_ok.append("不利物质条件的其他情形和有关约定")
            else:
                factors_error["不利物质条件的其他情形和有关约定"] = "第三部分_7. 工期和进度:不利物质条件的其他情形和有关约定未填写完整"
                addRemarkInDoc(word, document, "不利物质条件的其他情形和有关约定", f"第三部分_7. 工期和进度:不利物质条件的其他情形和有关约定未填写完整")
        except:
            missObject += "第三部分_7. 工期和进度:要素“不利物质条件的其他情形和有关约定”缺失\n"

        try:
            match = '7.9.2提前竣工的奖励：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["7.9.2提前竣工的奖励"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '无'.replace(" ", "").replace("；", "").replace("。", ""):
                    factors_error["7.9.2提前竣工的奖励"] = "第三部分_7. 工期和进度:原则上不应有提前竣工奖励。若有特别约定，请核实提前竣工奖励的必要性与合理性"
                    addRemarkInDoc(word, document, "7.9.2提前竣工的奖励",
                                   f"第三部分_7. 工期和进度:原则上不应有提前竣工奖励。若有特别约定，请核实提前竣工奖励的必要性与合理性")
                else:
                    factors_ok.append("7.9.2提前竣工的奖励")
            else:
                factors_error["7.9.2提前竣工的奖励"] = "第三部分_7. 工期和进度:7.9.2提前竣工的奖励未填写完整"
                addRemarkInDoc(word, document, "7.9.2提前竣工的奖励", f"第三部分_7. 工期和进度:7.9.2提前竣工的奖励未填写完整")
        except:
            missObject += "第三部分_7. 工期和进度:要素“7.9.2提前竣工的奖励”缺失\n"

        try:
            match = '8.4.1发包人供应的材料设备的保管费用的承担：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["8.4.1发包人供应的材料设备的保管费用的承担"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '由承包人承担'.replace(" ", "").replace("；", "").replace("。", ""):
                    factors_error["8.4.1发包人供应的材料设备的保管费用的承担"] = "第三部分_8. 材料与设备:原则上应由承包人承担"
                    addRemarkInDoc(word, document, "8.4.1发包人供应的材料设备的保管费用的承担", f"第三部分_8. 材料与设备:原则上应由承包人承担")
                else:
                    factors_ok.append("8.4.1发包人供应的材料设备的保管费用的承担")
            else:
                factors_error["8.4.1发包人供应的材料设备的保管费用的承担"] = "第三部分_8. 材料与设备:8.4.1发包人供应的材料设备的保管费用的承担未填写完整"
                addRemarkInDoc(word, document, "8.4.1发包人供应的材料设备的保管费用的承担", f"第三部分_8. 材料与设备:8.4.1发包人供应的材料设备的保管费用的承担未填写完整")
        except:
            missObject += "第三部分_8. 材料与设备:要素“8.4.1发包人供应的材料设备的保管费用的承担”缺失\n"

        try:
            match = '需要承包人报送样品的材料或工程设备，样品的种类、名称、规格、数量要求：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["需要承包人报送样品的材料或工程设备，样品的种类、名称、规格、数量要求"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '按发包人要求'.replace(" ", "").replace("；", "").replace("。", ""):
                    factors_error["需要承包人报送样品的材料或工程设备，样品的种类、名称、规格、数量要求"] = "第三部分_8. 材料与设备:原则上应约定为 按发包人要求执行"
                    addRemarkInDoc(word, document, "需要承包人报送样品的材料或工程设备，样品的种类、名称、规格、数量要求",
                                   f"第三部分_8. 材料与设备:原则上应约定为 按发包人要求执行")
                else:
                    factors_ok.append("需要承包人报送样品的材料或工程设备，样品的种类、名称、规格、数量要求")
            else:
                factors_error[
                    "需要承包人报送样品的材料或工程设备，样品的种类、名称、规格、数量要求"] = "第三部分_8. 材料与设备:需要承包人报送样品的材料或工程设备，样品的种类、名称、规格、数量要求未填写完整"
                addRemarkInDoc(word, document, "需要承包人报送样品的材料或工程设备，样品的种类、名称、规格、数量要求",
                               f"第三部分_8. 材料与设备:需要承包人报送样品的材料或工程设备，样品的种类、名称、规格、数量要求未填写完整")
        except:
            missObject += "第三部分_8. 材料与设备:要素“需要承包人报送样品的材料或工程设备，样品的种类、名称、规格、数量要求”缺失\n"

        try:
            match = '关于修建临时设施费用承担的约定：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["关于修建临时设施费用承担的约定"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '承包人应按监理人、发包人批准的施工组织设计自行修建临时设施并自行承担临时设施的费用。需要发包人办理申请手续和承担相关费用的临时占地：承包人应自行承担修建临时设施的费用，需要临时占地的，由发包人配合承包人自行办理相关手续，承包人承担相应费用。承包人应无偿分别为发包人、跟审（若有）、项管（若有）、监理等单位在施工现场提供临时办公室'.replace(
                        " ", "").replace("；", "").replace("。", ""):
                    factors_error["关于修建临时设施费用承担的约定"] = "第三部分_8. 材料与设备:临时设施费（包括临时占地费用与设施搭建费用）已包含在工程价款中，应全部由承包人承担"
                    addRemarkInDoc(word, document, "关于修建临时设施费用承担的约定",
                                   f"第三部分_8. 材料与设备:临时设施费（包括临时占地费用与设施搭建费用）已包含在工程价款中，应全部由承包人承担")
                else:
                    factors_ok.append("关于修建临时设施费用承担的约定")
            else:
                factors_error["关于修建临时设施费用承担的约定"] = "第三部分_8. 材料与设备:关于修建临时设施费用承担的约定未填写完整"
                addRemarkInDoc(word, document, "关于修建临时设施费用承担的约定", f"第三部分_8. 材料与设备:关于修建临时设施费用承担的约定未填写完整")
        except:
            missObject += "第三部分_8. 材料与设备:要素“关于修建临时设施费用承担的约定”缺失\n"

        try:
            match = '施工现场需要配置的试验场所：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["施工现场需要配置的试验场所"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '根据规范和施工现场需求设置'.replace(" ", "").replace("；", "").replace("。", ""):
                    factors_error["施工现场需要配置的试验场所"] = "第三部分_9. 试验与检验:请核实该条内容与双方约定是否一致，原则上应根据发包人与监理人要求执行"
                    addRemarkInDoc(word, document, "施工现场需要配置的试验场所", f"第三部分_9. 试验与检验:请核实该条内容与双方约定是否一致，原则上应根据发包人与监理人要求执行")
                else:
                    factors_ok.append("施工现场需要配置的试验场所")
            else:
                factors_error["施工现场需要配置的试验场所"] = "第三部分_9. 试验与检验:施工现场需要配置的试验场所未填写完整"
                addRemarkInDoc(word, document, "施工现场需要配置的试验场所", f"第三部分_9. 试验与检验:施工现场需要配置的试验场所未填写完整")
        except:
            missObject += "第三部分_9. 试验与检验:要素“施工现场需要配置的试验场所”缺失\n"

        try:
            match = '施工现场需要配备的试验设备：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["施工现场需要配备的试验设备"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '根据规范和施工现场需求设置'.replace(" ", "").replace("；", "").replace("。", ""):
                    factors_error["施工现场需要配备的试验设备"] = "第三部分_9. 试验与检验:请核实该条内容与双方约定是否一致，原则上应根据发包人与监理人要求执行"
                    addRemarkInDoc(word, document, "施工现场需要配备的试验设备", f"第三部分_9. 试验与检验:请核实该条内容与双方约定是否一致，原则上应根据发包人与监理人要求执行")
                else:
                    factors_ok.append("施工现场需要配备的试验设备")
            else:
                factors_error["施工现场需要配备的试验设备"] = "第三部分_9. 试验与检验:施工现场需要配备的试验设备未填写完整"
                addRemarkInDoc(word, document, "施工现场需要配备的试验设备", f"第三部分_9. 试验与检验:施工现场需要配备的试验设备未填写完整")
        except:
            missObject += "第三部分_9. 试验与检验:要素“施工现场需要配备的试验设备”缺失\n"

        try:
            match = '施工现场需要具备的其他试验条件：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["施工现场需要具备的其他试验条件"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '根据规范和施工现场需求设置'.replace(" ", "").replace("；", "").replace("。", ""):
                    factors_error["施工现场需要具备的其他试验条件"] = "第三部分_9. 试验与检验:请核实该条内容与双方约定是否一致，原则上应根据发包人与监理人要求执行"
                    addRemarkInDoc(word, document, "施工现场需要具备的其他试验条件",
                                   f"第三部分_9. 试验与检验:请核实该条内容与双方约定是否一致，原则上应根据发包人与监理人要求执行")
                else:
                    factors_ok.append("施工现场需要具备的其他试验条件")
            else:
                factors_error["施工现场需要具备的其他试验条件"] = "第三部分_9. 试验与检验:施工现场需要具备的其他试验条件未填写完整"
                addRemarkInDoc(word, document, "施工现场需要具备的其他试验条件", f"第三部分_9. 试验与检验:施工现场需要具备的其他试验条件未填写完整")
        except:
            missObject += "第三部分_9. 试验与检验:要素“施工现场需要具备的其他试验条件”缺失\n"

        try:
            match = '现场工艺试验的有关约定：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["现场工艺试验的有关约定"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '按监理人和发包人要求'.replace(" ", "").replace("；", "").replace("。", ""):
                    factors_error["现场工艺试验的有关约定"] = "第三部分_9. 试验与检验:请核实该条内容与双方约定是否一致，原则上应根据发包人与监理人要求执行"
                    addRemarkInDoc(word, document, "现场工艺试验的有关约定", f"第三部分_9. 试验与检验:请核实该条内容与双方约定是否一致，原则上应根据发包人与监理人要求执行")
                else:
                    factors_ok.append("现场工艺试验的有关约定")
            else:
                factors_error["现场工艺试验的有关约定"] = "第三部分_9. 试验与检验:现场工艺试验的有关约定未填写完整"
                addRemarkInDoc(word, document, "现场工艺试验的有关约定", f"第三部分_9. 试验与检验:现场工艺试验的有关约定未填写完整")
        except:
            missObject += "第三部分_9. 试验与检验:要素“现场工艺试验的有关约定”缺失\n"

        try:
            match = '关于变更的范围的约定：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["关于变更的范围的约定"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '按通用条款执行'.replace(" ", "").replace("；", "").replace("。", ""):
                    factors_error["关于变更的范围的约定"] = "第三部分_10. 变更:请核实该条内容与双方约定是否一致"
                    addRemarkInDoc(word, document, "关于变更的范围的约定", f"第三部分_10. 变更:请核实该条内容与双方约定是否一致")
                else:
                    factors_ok.append("关于变更的范围的约定")
            else:
                factors_error["关于变更的范围的约定"] = "第三部分_10. 变更:关于变更的范围的约定未填写完整"
                addRemarkInDoc(word, document, "关于变更的范围的约定", f"第三部分_10. 变更:关于变更的范围的约定未填写完整")
        except:
            missObject += "第三部分_10. 变更:要素“关于变更的范围的约定”缺失\n"

        try:
            match = '监理人审查承包人合理化建议的期限：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["监理人审查承包人合理化建议的期限"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '收到承包人书面建议后7天内'.replace(" ", "").replace("；", "").replace("。", ""):
                    factors_error["监理人审查承包人合理化建议的期限"] = "第三部分_10. 变更:该条期限在3-15日期间选择"
                    addRemarkInDoc(word, document, "监理人审查承包人合理化建议的期限", f"第三部分_10. 变更:该条期限在3-15日期间选择")
                else:
                    factors_ok.append("监理人审查承包人合理化建议的期限")
            else:
                factors_error["监理人审查承包人合理化建议的期限"] = "第三部分_10. 变更:监理人审查承包人合理化建议的期限未填写完整"
                addRemarkInDoc(word, document, "监理人审查承包人合理化建议的期限", f"第三部分_10. 变更:监理人审查承包人合理化建议的期限未填写完整")
        except:
            missObject += "第三部分_10. 变更:要素“监理人审查承包人合理化建议的期限”缺失\n"

        try:
            match = '发包人审批承包人合理化建议的期限：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["发包人审批承包人合理化建议的期限"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '收到承包人书面建议后7天内'.replace(" ", "").replace("；", "").replace("。", ""):
                    factors_error["发包人审批承包人合理化建议的期限"] = "第三部分_10. 变更:该条期限在3-30日期间选择"
                    addRemarkInDoc(word, document, "发包人审批承包人合理化建议的期限", f"第三部分_10. 变更:该条期限在3-30日期间选择")
                else:
                    factors_ok.append("发包人审批承包人合理化建议的期限")
            else:
                factors_error["发包人审批承包人合理化建议的期限"] = "第三部分_10. 变更:发包人审批承包人合理化建议的期限未填写完整"
                addRemarkInDoc(word, document, "发包人审批承包人合理化建议的期限", f"第三部分_10. 变更:发包人审批承包人合理化建议的期限未填写完整")
        except:
            missObject += "第三部分_10. 变更:要素“发包人审批承包人合理化建议的期限”缺失\n"

        try:
            match = '承包人提出的合理化建议降低了合同价格或者提高了工程经济效益的奖励的方法和金额为：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["承包人提出的合理化建议降低了合同价格或者提高了工程经济效益的奖励的方法和金额为"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '无'.replace(" ", "").replace("；", "").replace("。", ""):
                    factors_error[
                        "承包人提出的合理化建议降低了合同价格或者提高了工程经济效益的奖励的方法和金额为"] = "第三部分_10. 变更:根据实际情况填写。原则上“无”，钢结构深化设计可结合实际考虑奖励措施"
                    addRemarkInDoc(word, document, "承包人提出的合理化建议降低了合同价格或者提高了工程经济效益的奖励的方法和金额为",
                                   f"第三部分_10. 变更:根据实际情况填写。原则上“无”，钢结构深化设计可结合实际考虑奖励措施")
                else:
                    factors_ok.append("承包人提出的合理化建议降低了合同价格或者提高了工程经济效益的奖励的方法和金额为")
            else:
                factors_error[
                    "承包人提出的合理化建议降低了合同价格或者提高了工程经济效益的奖励的方法和金额为"] = "第三部分_10. 变更:承包人提出的合理化建议降低了合同价格或者提高了工程经济效益的奖励的方法和金额为未填写完整"
                addRemarkInDoc(word, document, "承包人提出的合理化建议降低了合同价格或者提高了工程经济效益的奖励的方法和金额为",
                               f"第三部分_10. 变更:承包人提出的合理化建议降低了合同价格或者提高了工程经济效益的奖励的方法和金额为未填写完整")
        except:
            missObject += "第三部分_10. 变更:要素“承包人提出的合理化建议降低了合同价格或者提高了工程经济效益的奖励的方法和金额为”缺失\n"

        try:
            match = '对于依法必须招标的暂估价项目的确认和批准采取第(.*)种方式确定。\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["10.7.1 依法必须招标的暂估价项目"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '1'.replace(" ", "").replace("；", "").replace("。", ""):
                    factors_error["10.7.1 依法必须招标的暂估价项目"] = "第三部分_10. 变更:若采用单价合同，应使用推荐条款"
                    addRemarkInDoc(word, document, "10.7.1 依法必须招标的暂估价项目", f"第三部分_10. 变更:若采用单价合同，应使用推荐条款")
                else:
                    factors_ok.append("10.7.1 依法必须招标的暂估价项目")
            else:
                factors_error["10.7.1 依法必须招标的暂估价项目"] = "第三部分_10. 变更:10.7.1 依法必须招标的暂估价项目未填写完整"
                addRemarkInDoc(word, document, "10.7.1 依法必须招标的暂估价项目", f"第三部分_10. 变更:10.7.1 依法必须招标的暂估价项目未填写完整")
        except:
            missObject += "第三部分_10. 变更:要素“10.7.1 依法必须招标的暂估价项目”缺失\n"

        try:
            match = '对于不属于依法必须招标的暂估价项目的确认和批准采取第(.*)种方式确定。\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["10.7.2 不属于依法必须招标的暂估价项目"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '3'.replace(" ", "").replace("；", "").replace("。", ""):
                    factors_error["10.7.2 不属于依法必须招标的暂估价项目"] = "第三部分_10. 变更:若采用单价合同，应使用推荐条款"
                    addRemarkInDoc(word, document, "10.7.2 不属于依法必须招标的暂估价项目", f"第三部分_10. 变更:若采用单价合同，应使用推荐条款")
                else:
                    factors_ok.append("10.7.2 不属于依法必须招标的暂估价项目")
            else:
                factors_error["10.7.2 不属于依法必须招标的暂估价项目"] = "第三部分_10. 变更:10.7.2 不属于依法必须招标的暂估价项目未填写完整"
                addRemarkInDoc(word, document, "10.7.2 不属于依法必须招标的暂估价项目", f"第三部分_10. 变更:10.7.2 不属于依法必须招标的暂估价项目未填写完整")
        except:
            missObject += "第三部分_10. 变更:要素“10.7.2 不属于依法必须招标的暂估价项目”缺失\n"

        try:
            match = '承包人直接实施的暂估价项目的约定：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["承包人直接实施的暂估价项目的约定"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '发包人在工程量清单中给定暂估价的材料和工程设备，若发包人委托承包人采购，承包人在采购前15天，其质量、品牌需征得发包人、项目管理公司（若有）、监理人书面签字认可，否则在结算时按零价格处理。材料价格以甲方最终确认的价格为准'.replace(
                        " ", "").replace("；", "").replace("。", ""):
                    factors_error["承包人直接实施的暂估价项目的约定"] = "第三部分_10. 变更:若采用单价合同，应使用推荐条款"
                    addRemarkInDoc(word, document, "承包人直接实施的暂估价项目的约定", f"第三部分_10. 变更:若采用单价合同，应使用推荐条款")
                else:
                    factors_ok.append("承包人直接实施的暂估价项目的约定")
            else:
                factors_error["承包人直接实施的暂估价项目的约定"] = "第三部分_10. 变更:承包人直接实施的暂估价项目的约定未填写完整"
                addRemarkInDoc(word, document, "承包人直接实施的暂估价项目的约定", f"第三部分_10. 变更:承包人直接实施的暂估价项目的约定未填写完整")
        except:
            missObject += "第三部分_10. 变更:要素“承包人直接实施的暂估价项目的约定”缺失\n"

        try:
            match = '合同当事人关于暂列金额使用的约定：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["合同当事人关于暂列金额使用的约定"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '按照发包人的要求使用'.replace(" ", "").replace("；", "").replace("。", ""):
                    factors_error["合同当事人关于暂列金额使用的约定"] = "第三部分_10. 变更:请核实本条内容是否与双方约定一致，原则上应按照发包人要求使用"
                    addRemarkInDoc(word, document, "合同当事人关于暂列金额使用的约定", f"第三部分_10. 变更:请核实本条内容是否与双方约定一致，原则上应按照发包人要求使用")
                else:
                    factors_ok.append("合同当事人关于暂列金额使用的约定")
            else:
                factors_error["合同当事人关于暂列金额使用的约定"] = "第三部分_10. 变更:合同当事人关于暂列金额使用的约定未填写完整"
                addRemarkInDoc(word, document, "合同当事人关于暂列金额使用的约定", f"第三部分_10. 变更:合同当事人关于暂列金额使用的约定未填写完整")
        except:
            missObject += "第三部分_10. 变更:要素“合同当事人关于暂列金额使用的约定”缺失\n"

        try:
            match = '专用合同条款①承包人在已标价工程量清单或预算书中载明的材料单价低于基准价格的：专用合同条款合同履行期间材料单价涨幅以基准价格为基础超过(.*)%时，或材料单价跌幅以已标价工程量清单'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["专用合同条款①_1"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    factor = int(factor)
                    if factor > 5 or factor < 3:
                        factors_error["专用合同条款①_1"] = "第三部分_11. 价格调整:各项数值在3-5之间选择"
                        addRemarkInDoc(word, document, "专用合同条款①", f"第三部分_11. 价格调整:各项数值在3-5之间选择")
                    else:
                        factors_ok.append("专用合同条款①_1")
                except:
                    factors_error["专用合同条款①_1"] = "第三部分_11. 价格调整:数值提取错误"
                    addRemarkInDoc(word, document, "专用合同条款①", f"第三部分_11. 价格调整:数值提取错误")
            else:
                factors_error["专用合同条款①_1"] = "第三部分_11. 价格调整:专用合同条款①_1未填写完整"
                addRemarkInDoc(word, document, "专用合同条款①", f"第三部分_11. 价格调整:专用合同条款①_1未填写完整")
        except:
            missObject += "第三部分_11. 价格调整:要素“专用合同条款①_1”缺失\n"

        try:
            match = '或材料单价跌幅以已标价工程量清单或预算书中载明材料单价为基础超过(.*)%时，其超过部分据实'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["专用合同条款①_2"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    factor = int(factor)
                    if factor > 5 or factor < 3:
                        factors_error["专用合同条款①_2"] = "第三部分_11. 价格调整:各项数值在3-5之间选择"
                        addRemarkInDoc(word, document, "专用合同条款①", f"第三部分_11. 价格调整:各项数值在3-5之间选择")
                    else:
                        factors_ok.append("专用合同条款①_2")
                except:
                    factors_error["专用合同条款①_2"] = "第三部分_11. 价格调整:数值提取错误"
                    addRemarkInDoc(word, document, "专用合同条款①", f"第三部分_11. 价格调整:数值提取错误")
            else:
                factors_error["专用合同条款①_2"] = "第三部分_11. 价格调整:专用合同条款①_2未填写完整"
                addRemarkInDoc(word, document, "专用合同条款①", f"第三部分_11. 价格调整:专用合同条款①_2未填写完整")
        except:
            missObject += "第三部分_11. 价格调整:要素“专用合同条款①_2”缺失\n"

        try:
            match = '②承包人在已标价工程量清单或预算书中载明的材料单价高于基准价格的：专用合同条款合同履行期间材料单价跌幅以基准价格为基础超过(.*)%时，材料单价涨幅以已标'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["专用合同条款②_1"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    factor = int(factor)
                    if factor > 5 or factor < 3:
                        factors_error["专用合同条款②_1"] = "第三部分_11. 价格调整:各项数值在3-5之间选择"
                        addRemarkInDoc(word, document, "第三部分_11. 价格调整:②承包人在已标价工程量清单或预算书中载明的材料单价高于基准价格的",
                                       f"各项数值在3-5之间选择")
                    else:
                        factors_ok.append("专用合同条款②_1")
                except:
                    factors_error["专用合同条款②_1"] = "第三部分_11. 价格调整:数值提取错误"
                    addRemarkInDoc(word, document, "第三部分_11. 价格调整:②承包人在已标价工程量清单或预算书中载明的材料单价高于基准价格的", f"数值提取错误")
            else:
                factors_error["专用合同条款②_1"] = "第三部分_11. 价格调整:专用合同条款②_1未填写完整"
                addRemarkInDoc(word, document, "②承包人在已标价工程量清单或预算书中载明的材料单价高于基准价格的", f"第三部分_11. 价格调整:专用合同条款②_1未填写完整")
        except:
            missObject += "第三部分_11. 价格调整:要素“专用合同条款②_1”缺失\n"

        try:
            match = '材料单价涨幅以已标价工程量清单或预算书中载明材料单价为基础超过(.*)%时，其超过部分'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["专用合同条款②_2"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    factor = int(factor)
                    if factor > 5 or factor < 3:
                        factors_error["专用合同条款②_2"] = "第三部分_11. 价格调整:各项数值在3-5之间选择"
                        addRemarkInDoc(word, document, "②承包人在已标价工程量清单或预算书中载明的材料单价高于基准价格的",
                                       f"第三部分_11. 价格调整:各项数值在3-5之间选择")
                    else:
                        factors_ok.append("专用合同条款②_2")
                except:
                    factors_error["专用合同条款②_2"] = "第三部分_11. 价格调整:数值提取错误"
                    addRemarkInDoc(word, document, "②承包人在已标价工程量清单或预算书中载明的材料单价高于基准价格的", f"第三部分_11. 价格调整:数值提取错误")
            else:
                factors_error["专用合同条款②_2"] = "第三部分_11. 价格调整:专用合同条款②_2未填写完整"
                addRemarkInDoc(word, document, "②承包人在已标价工程量清单或预算书中载明的材料单价高于基准价格的", f"第三部分_11. 价格调整:专用合同条款②_2未填写完整")
        except:
            missObject += "第三部分_11. 价格调整:要素“专用合同条款②_2”缺失\n"

        try:
            match = '专用合同条款合同履行期间材料单价涨跌幅以基准单价为基础超过±(.*)%时，其超过部分据实调'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["专用合同条款③_1"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    factor = int(factor)
                    if factor > 5 or factor < 3:
                        factors_error["专用合同条款③_1"] = "第三部分_11. 价格调整:各项数值在3-5之间选择"
                        addRemarkInDoc(word, document, "第三部分_11. 价格调整:③承包人在已标价工程量清单或预算书中载明的材料单价等于基准单价的",
                                       f"各项数值在3-5之间选择")
                    else:
                        factors_ok.append("专用合同条款③_1")
                except:
                    factors_error["专用合同条款③_1"] = "第三部分_11. 价格调整:数值提取错误"
                    addRemarkInDoc(word, document, "③承包人在已标价工程量清单或预算书中载明的材料单价等于基准单价的", f"第三部分_11. 价格调整:数值提取错误")
            else:
                factors_error["专用合同条款③_1"] = "第三部分_11. 价格调整:专用合同条款③_1未填写完整"
                addRemarkInDoc(word, document, "③承包人在已标价工程量清单或预算书中载明的材料单价等于基准单价的", f"第三部分_11. 价格调整:专用合同条款③_1未填写完整")
        except:
            missObject += "第三部分_11. 价格调整:要素“专用合同条款③_1”缺失\n"

    # 第三部分_12-21
    if 1 == True:
        # 3.12
        try:
            match = '预付款支付比例或金额：合同金额（扣除中标价暂列金、专业工程暂估价及安全文明施工费）的(.*)%\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["预付款支付比例或金额"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    factor = int(factor)
                    if factor < 10 or factor > 20:
                        factors_error["预付款支付比例或金额"] = "第三部分_12. 合同价格、计量与支付:应该在在10-20间选择"
                        addRemarkInDoc(word, document, "12.2.1 预付款的支付", f"第三部分_12. 合同价格、计量与支付:应该在在10-20间选择")
                    else:
                        factors_ok.append("预付款支付比例或金额")
                except:
                    factors_error["预付款支付比例或金额"] = "第三部分_12. 合同价格、计量与支付:提取错误"
                    addRemarkInDoc(word, document, "12.2.1 预付款的支付", f"第三部分_12. 合同价格、计量与支付:提取错误")
            else:
                factors_error["预付款支付比例或金额"] = "第三部分_12. 合同价格、计量与支付:预付款支付比例或金额未填写完整"
                addRemarkInDoc(word, document, "12.2.1 预付款的支付", f"第三部分_12. 合同价格、计量与支付:预付款支付比例或金额未填写完整")
        except:
            missObject += "第三部分_12. 合同价格、计量与支付:要素“预付款支付比例或金额”缺失\n"

        try:
            match = '预付款支付期限：承包人完善办理相关资料并且办理施工许可证后(.*)天内。\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["预付款支付期限"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    factor = int(factor)
                    if factor < 15 or factor > 30:
                        factors_error["预付款支付期限"] = "第三部分_12. 合同价格、计量与支付:应该在在15-30间选择"
                        addRemarkInDoc(word, document, "第三部分_12. 合同价格、计量与支付:12.2.1 预付款的支付", f"应该在在15-30间选择")
                    else:
                        factors_ok.append("预付款支付期限")
                except:
                    factors_error["预付款支付期限"] = "第三部分_12. 合同价格、计量与支付:提取错误"
                    addRemarkInDoc(word, document, "12.2.1 预付款的支付", f"第三部分_12. 合同价格、计量与支付:提取错误")
            else:
                factors_error["预付款支付期限"] = "第三部分_12. 合同价格、计量与支付:预付款支付期限未填写完整"
                addRemarkInDoc(word, document, "12.2.1 预付款的支付", f"第三部分_12. 合同价格、计量与支付:预付款支付期限未填写完整")
        except:
            missObject += "第三部分_12. 合同价格、计量与支付:要素“预付款支付期限”缺失\n"

        try:
            match = '预付款扣回的方式：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["预付款扣回的方式"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("预付款扣回的方式")
            else:
                factors_error["预付款扣回的方式"] = "第三部分_12. 合同价格、计量与支付:预付款扣回的方式未填写完整"
                addRemarkInDoc(word, document, "预付款扣回的方式", f"第三部分_12. 合同价格、计量与支付:预付款扣回的方式未填写完整")
        except:
            missObject += "要素“预付款扣回的方式”缺失\n"

        factors_to_inform["预付款扣回的方式"] = '第三部分_12. 合同价格、计量与支付:可根据项目实际情况约定扣回比例和扣回时间。请要求财务部门确定该条款内容。'
        addRemarkInDoc(word, document, "预付款扣回的方式", f"第三部分_12. 合同价格、计量与支付:可根据项目实际情况约定扣回比例和扣回时间。请要求财务部门确定该条款内容。")

        try:
            match = '承包人提交预付款担保的期限：预付款支付前(.*)天。\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["承包人提交预付款担保的期限"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    factor = int(factor)
                    if factor < 1 or factor > 7:
                        factors_error["承包人提交预付款担保的期限"] = "第三部分_12. 合同价格、计量与支付:该项数值在1-7之间选择"
                        addRemarkInDoc(word, document, "承包人提交预付款担保的期限", f"第三部分_12. 合同价格、计量与支付:该项数值在1-7之间选择")
                    else:
                        factors_ok.append("承包人提交预付款担保的期限")
                except:
                    factors_error["承包人提交预付款担保的期限"] = "第三部分_12. 合同价格、计量与支付:提取错误"
                    addRemarkInDoc(word, document, "承包人提交预付款担保的期限", f"第三部分_12. 合同价格、计量与支付:提取错误")
            else:
                factors_error["承包人提交预付款担保的期限"] = "第三部分_12. 合同价格、计量与支付:承包人提交预付款担保的期限未填写完整"
                addRemarkInDoc(word, document, "承包人提交预付款担保的期限", f"第三部分_12. 合同价格、计量与支付:承包人提交预付款担保的期限未填写完整")
        except:
            missObject += "第三部分_12. 合同价格、计量与支付:要素“承包人提交预付款担保的期限”缺失\n"

        try:
            match = '预付款担保的形式为：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["预付款担保的形式为"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("预付款担保的形式为")
            else:
                factors_error["预付款担保的形式为"] = "第三部分_12. 合同价格、计量与支付:预付款担保的形式为未填写完整"
                addRemarkInDoc(word, document, "预付款担保的形式为", f"第三部分_12. 合同价格、计量与支付:预付款担保的形式为未填写完整")
        except:
            missObject += "要素“预付款担保的形式为”缺失\n"

        # 12.3
        try:
            match = '关于计量周期的约定：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["关于计量周期的约定"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '按月进行'.replace(" ", "").replace("；", "").replace("。", ""):
                    factors_error["关于计量周期的约定"] = "第三部分_12. 合同价格、计量与支付:原则上应按月进行"
                    addRemarkInDoc(word, document, "关于计量周期的约定", f"第三部分_12. 合同价格、计量与支付:原则上应按月进行")
                else:
                    factors_ok.append("关于计量周期的约定")
            else:
                factors_error["关于计量周期的约定"] = "第三部分_12. 合同价格、计量与支付:关于计量周期的约定未填写完整"
                addRemarkInDoc(word, document, "关于计量周期的约定", f"第三部分_12. 合同价格、计量与支付:关于计量周期的约定未填写完整")
        except:
            missObject += "第三部分_12. 合同价格、计量与支付:要素“关于计量周期的约定”缺失\n"

        # 12.4
        try:
            match = '关于付款周期的约定：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["关于付款周期的约定"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("关于付款周期的约定")
            else:
                factors_error["关于付款周期的约定"] = "第三部分_12. 合同价格、计量与支付:关于付款周期的约定未填写完整"
                addRemarkInDoc(word, document, "关于付款周期的约定", f"第三部分_12. 合同价格、计量与支付:关于付款周期的约定未填写完整")
        except:
            missObject += "第三部分_12. 合同价格、计量与支付:要素“关于付款周期的约定”缺失\n"

        factors_to_inform["关于付款周期的约定"] = '第三部分_12. 合同价格、计量与支付:请财务部门确定付款周期的约定内容'
        addRemarkInDoc(word, document, "关于付款周期的约定", f"第三部分_12. 合同价格、计量与支付:请财务部门确定付款周期的约定内容")

        try:
            match = '关于进度付款申请单编制的约定：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["关于进度付款申请单编制的约定"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != ' 除满足通用条款12.4.2条外，还应按发包人具体要求进行编制'.replace(" ", "").replace("；", "").replace("。", ""):
                    factors_error["关于进度付款申请单编制的约定"] = "第三部分_12. 合同价格、计量与支付:除满足通用条款12.4.2条外，还应按发包人具体要求进行编制。"
                    addRemarkInDoc(word, document, "关于进度付款申请单编制的约定",
                                   f"第三部分_12. 合同价格、计量与支付:除满足通用条款12.4.2条外，还应按发包人具体要求进行编制。")
                else:
                    factors_ok.append("关于进度付款申请单编制的约定")
            else:
                factors_error["关于进度付款申请单编制的约定"] = "第三部分_12. 合同价格、计量与支付:关于进度付款申请单编制的约定未填写完整"
                addRemarkInDoc(word, document, "关于进度付款申请单编制的约定", f"第三部分_12. 合同价格、计量与支付:关于进度付款申请单编制的约定未填写完整")
        except:
            missObject += "第三部分_12. 合同价格、计量与支付:要素“关于进度付款申请单编制的约定”缺失\n"

        try:
            match = '监理人审查并报送发包人的期限：收到承包人进度付款申请单以及相关资料后(.*)天内完成审查并报送发包人。\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["（1）监理人审查并报送发包人的期限"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    factor = int(factor)
                    if factor < 1 or factor > 7:
                        factors_error["（1）监理人审查并报送发包人的期限"] = "第三部分_12. 合同价格、计量与支付:该项数值在1-7之间选择"
                        addRemarkInDoc(word, document, "（1）监理人审查并报送发包人的期限", f"第三部分_12. 合同价格、计量与支付:该项数值在1-7之间选择")
                    else:
                        factors_ok.append("（1）监理人审查并报送发包人的期限")
                except:
                    factors_error["（1）监理人审查并报送发包人的期限"] = "第三部分_12. 合同价格、计量与支付:提取错误"
                    addRemarkInDoc(word, document, "（1）监理人审查并报送发包人的期限", f"第三部分_12. 合同价格、计量与支付:提取错误")
            else:
                factors_error["（1）监理人审查并报送发包人的期限"] = "第三部分_12. 合同价格、计量与支付:（1）监理人审查并报送发包人的期限未填写完整"
                addRemarkInDoc(word, document, "（1）监理人审查并报送发包人的期限", f"第三部分_12. 合同价格、计量与支付:（1）监理人审查并报送发包人的期限未填写完整")
        except:
            missObject += "第三部分_12. 合同价格、计量与支付:要素“（1）监理人审查并报送发包人的期限”缺失\n"

        try:
            match = '发包人完成审批并签发进度款支付证书的期限：收到监理人审查确认后的进度付款申请单以及相关资料后(.*)天内完成审批并签发；若发.*'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["发包人完成审批并签发进度款支付证书的期限_1"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    factor = int(factor)
                    if factor < 3 or factor > 30:
                        factors_error["发包人完成审批并签发进度款支付证书的期限_1"] = "第三部分_12. 合同价格、计量与支付:该项数值在3-30之间选择"
                        addRemarkInDoc(word, document, "发包人完成审批并签发进度款支付证书的期限", f"第三部分_12. 合同价格、计量与支付:该项数值在3-30之间选择")
                    else:
                        factors_ok.append("发包人完成审批并签发进度款支付证书的期限_1")
                except:
                    factors_error["发包人完成审批并签发进度款支付证书的期限_1"] = "第三部分_12. 合同价格、计量与支付:提取错误"
                    addRemarkInDoc(word, document, "发包人完成审批并签发进度款支付证书的期限", f"第三部分_12. 合同价格、计量与支付:提取错误")
            else:
                factors_error["发包人完成审批并签发进度款支付证书的期限_1"] = "第三部分_12. 合同价格、计量与支付:发包人完成审批并签发进度款支付证书的期限_1未填写完整"
                addRemarkInDoc(word, document, "发包人完成审批并签发进度款支付证书的期限",
                               f"第三部分_12. 合同价格、计量与支付:发包人完成审批并签发进度款支付证书的期限_1未填写完整")
        except:
            missObject += "第三部分_12. 合同价格、计量与支付:要素“发包人完成审批并签发进度款支付证书的期限_1”缺失\n"

        try:
            match = '在收到正式的本期跟踪审计报告后(.*)天内完成审批并签发。\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["发包人完成审批并签发进度款支付证书的期限_2"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    factor = int(factor)
                    if factor < 3 or factor > 30:
                        factors_error["发包人完成审批并签发进度款支付证书的期限_2"] = "第三部分_12. 合同价格、计量与支付:该项数值在3-30之间选择"
                        addRemarkInDoc(word, document, "发包人完成审批并签发进度款支付证书的期限", f"第三部分_12. 合同价格、计量与支付:该项数值在3-30之间选择")
                    else:
                        factors_ok.append("发包人完成审批并签发进度款支付证书的期限_2")
                except:
                    factors_error["发包人完成审批并签发进度款支付证书的期限_2"] = "第三部分_12. 合同价格、计量与支付:提取错误"
                    addRemarkInDoc(word, document, "发包人完成审批并签发进度款支付证书的期限", f"第三部分_12. 合同价格、计量与支付:提取错误")
            else:
                factors_error["发包人完成审批并签发进度款支付证书的期限_2"] = "第三部分_12. 合同价格、计量与支付:发包人完成审批并签发进度款支付证书的期限_2未填写完整"
                addRemarkInDoc(word, document, "发包人完成审批并签发进度款支付证书的期限",
                               f"第三部分_12. 合同价格、计量与支付:发包人完成审批并签发进度款支付证书的期限_2未填写完整")
        except:
            missObject += "第三部分_12. 合同价格、计量与支付:要素“发包人完成审批并签发进度款支付证书的期限_2”缺失\n"

        try:
            match = '（2）发包人支付进度款的期限：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["（2）发包人支付进度款的期限"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("（2）发包人支付进度款的期限")
            else:
                factors_error["（2）发包人支付进度款的期限"] = "第三部分_12. 合同价格、计量与支付:（2）发包人支付进度款的期限未填写完整"
                addRemarkInDoc(word, document, "（2）发包人支付进度款的期限", f"第三部分_12. 合同价格、计量与支付:（2）发包人支付进度款的期限未填写完整")
        except:
            missObject += "第三部分_12. 合同价格、计量与支付:要素“（2）发包人支付进度款的期限”缺失\n"
        factors_to_inform["（2）发包人支付进度款的期限"] = '第三部分_12. 合同价格、计量与支付:请要求财务部门确定该条款内容'
        addRemarkInDoc(word, document, "（2）发包人支付进度款的期限", f"第三部分_12. 合同价格、计量与支付:请要求财务部门确定该条款内容")

        try:
            match = '3、单价合同的总价项目支付分解表的编制与审批：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["3、单价合同的总价项目支付分解表的编制与审批"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '以计划工期按月分摊进入每期支付'.replace(" ", "").replace("；", "").replace("。", ""):
                    factors_error["3、单价合同的总价项目支付分解表的编制与审批"] = "第三部分_12. 合同价格、计量与支付:请要求财务部门确定该条款内容"
                    addRemarkInDoc(word, document, "3、单价合同的总价项目支付分解表的编制与审批", f"第三部分_12. 合同价格、计量与支付:请要求财务部门确定该条款内容")
                else:
                    factors_ok.append("3、单价合同的总价项目支付分解表的编制与审批")
            else:
                factors_error["3、单价合同的总价项目支付分解表的编制与审批"] = "第三部分_12. 合同价格、计量与支付:3、单价合同的总价项目支付分解表的编制与审批未填写完整"
                addRemarkInDoc(word, document, "3、单价合同的总价项目支付分解表的编制与审批",
                               f"第三部分_12. 合同价格、计量与支付:3、单价合同的总价项目支付分解表的编制与审批未填写完整")
        except:
            missObject += "第三部分_12. 合同价格、计量与支付:要素“3、单价合同的总价项目支付分解表的编制与审批”缺失\n"

        # 3.13
        try:
            match = '13.1.2监理人不能按时进行验收时，应提前(.*)小时提交书面延期要求'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["13.1.2监理人不能按时进行验收时"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    factor = int(factor)
                    if factor < 24 or factor > 48:
                        factors_error["13.1.2监理人不能按时进行验收时"] = "第三部分_13. 验收和工程试车:在24-48之间选择"
                        addRemarkInDoc(word, document, "13.1.2监理人不能按时进行验收时", f"第三部分_13. 验收和工程试车:在24-48之间选择")
                    else:
                        factors_ok.append("13.1.2监理人不能按时进行验收时")
                except:
                    factors_error["13.1.2监理人不能按时进行验收时"] = "第三部分_13. 验收和工程试车:提取错误"
                    addRemarkInDoc(word, document, "13.1.2监理人不能按时进行验收时", f"第三部分_13. 验收和工程试车:提取错误")
            else:
                factors_error["13.1.2监理人不能按时进行验收时"] = "第三部分_13. 验收和工程试车:13.1.2监理人不能按时进行验收时未填写完整"
                addRemarkInDoc(word, document, "13.1.2监理人不能按时进行验收时", f"第三部分_13. 验收和工程试车:13.1.2监理人不能按时进行验收时未填写完整")
        except:
            missObject += "第三部分_13. 验收和工程试车:要素“13.1.2监理人不能按时进行验收时”缺失\n"

        try:
            match = '13.1.2监理人不能按时进行验收时，应提前.*小时提交书面延期要求。\n关于延期最长不得超过：(.*)小时。\n13.2 竣工验.*'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["13.1关于延期最长不得超过"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    factor = int(factor)
                    if factor < 48 or factor > 72:
                        factors_error["13.1关于延期最长不得超过"] = "第三部分_13. 验收和工程试车:在48-72之间选择"
                        addRemarkInDoc(word, document, "13.1 分部分项工程验收", f"第三部分_13. 验收和工程试车:在48-72之间选择")
                    else:
                        factors_ok.append("13.1关于延期最长不得超过")
                except:
                    factors_error["13.1关于延期最长不得超过"] = "第三部分_13. 验收和工程试车:提取错误"
                    addRemarkInDoc(word, document, "13.1 分部分项工程验收", f"第三部分_13. 验收和工程试车:提取错误")
            else:
                factors_error["13.1关于延期最长不得超过"] = "第三部分_13. 验收和工程试车:13.1关于延期最长不得超过未填写完整"
                addRemarkInDoc(word, document, "13.1 分部分项工程验收", f"第三部分_13. 验收和工程试车:13.1关于延期最长不得超过未填写完整")
        except:
            missObject += "第三部分_13. 验收和工程试车:要素“13.1关于延期最长不得超过”缺失\n"

        try:
            match = '承包人向发包人移交工程的期限：竣工验收后(.*)天内完成实体及资料移交。\n发包人未按本合同约定接收全部或部分.*'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["13.2.5移交、接收全部与部分工程"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    factor = int(factor)
                    if factor < 15 or factor > 28:
                        factors_error["13.2.5移交、接收全部与部分工程"] = "第三部分_13. 验收和工程试车:在15-28之间选择"
                        addRemarkInDoc(word, document, "13.2.5移交、接收全部与部分工程", f"第三部分_13. 验收和工程试车:在15-28之间选择")
                    else:
                        factors_ok.append("第三部分_13. 验收和工程试车:13.2.5移交、接收全部与部分工程")
                except:
                    factors_error["13.2.5移交、接收全部与部分工程"] = "第三部分_13. 验收和工程试车:提取错误"
                    addRemarkInDoc(word, document, "13.2.5移交、接收全部与部分工程", f"第三部分_13. 验收和工程试车:提取错误")
            else:
                factors_error["13.2.5移交、接收全部与部分工程"] = "第三部分_13. 验收和工程试车:13.2.5移交、接收全部与部分工程未填写完整"
                addRemarkInDoc(word, document, "13.2.5移交、接收全部与部分工程", f"第三部分_13. 验收和工程试车:13.2.5移交、接收全部与部分工程未填写完整")
        except:
            missObject += "第三部分_13. 验收和工程试车:要素“13.2.5移交、接收全部与部分工程”缺失\n"

        # 13.3
        try:
            match = '关于投料试车相关事项的约定：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["关于投料试车相关事项的约定"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("关于投料试车相关事项的约定")
            else:
                factors_error["关于投料试车相关事项的约定"] = "第三部分_13. 验收和工程试车:关于投料试车相关事项的约定未填写完整"
                addRemarkInDoc(word, document, "关于投料试车相关事项的约定", f"第三部分_13. 验收和工程试车:关于投料试车相关事项的约定未填写完整")
        except:
            missObject += "第三部分_13. 验收和工程试车:要素“关于投料试车相关事项的约定”缺失\n"
        factors_to_inform['关于投料试车相关事项的约定'] = '第三部分_13. 验收和工程试车:核实双方约定是否与本条约定一致'
        addRemarkInDoc(word, document, "关于投料试车相关事项的约定", f"第三部分_13. 验收和工程试车:核实双方约定是否与本条约定一致")

        try:
            match = '承包人完成竣工退场的期限：颁发工程接收证书后(.*)天内。\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["承包人完成竣工退场的期限"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    factor = int(factor)
                    if factor < 1 or factor > 14:
                        factors_error["承包人完成竣工退场的期限"] = "第三部分_13. 验收和工程试车:在1-14之间选择"
                        addRemarkInDoc(word, document, "承包人完成竣工退场的期限", f"第三部分_13. 验收和工程试车:在1-14之间选择")
                    else:
                        factors_ok.append("承包人完成竣工退场的期限")
                except:
                    factors_error["承包人完成竣工退场的期限"] = "第三部分_13. 验收和工程试车:提取错误"
                    addRemarkInDoc(word, document, "承包人完成竣工退场的期限", f"第三部分_13. 验收和工程试车:提取错误")
            else:
                factors_error["承包人完成竣工退场的期限"] = "第三部分_13. 验收和工程试车:承包人完成竣工退场的期限未填写完整"
                addRemarkInDoc(word, document, "承包人完成竣工退场的期限", f"第三部分_13. 验收和工程试车:承包人完成竣工退场的期限未填写完整")
        except:
            missObject += "第三部分_13. 验收和工程试车:要素“承包人完成竣工退场的期限”缺失\n"

        # 3.14
        try:
            match = '承包人提交竣工结算申请单的期限：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["承包人提交竣工结算申请单的期限"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("承包人提交竣工结算申请单的期限")
            else:
                factors_error["承包人提交竣工结算申请单的期限"] = "第三部分_14. 竣工结算:14.1 竣工结算申请承包人提交竣工结算申请单的期限未填写完整"
                addRemarkInDoc(word, document, "承包人提交竣工结算申请单的期限", f"第三部分_14. 竣工结算:14.1 竣工结算申请承包人提交竣工结算申请单的期限未填写完整")
        except:
            missObject += "第三部分_14. 竣工结算:要素“承包人提交竣工结算申请单的期限”缺失\n"
        factors_to_inform['承包人提交竣工结算申请单的期限'] = '第三部分_14. 竣工结算:14.1 竣工结算申请请财务部门确定该条内容'
        addRemarkInDoc(word, document, "承包人提交竣工结算申请单的期限", f"第三部分_14. 竣工结算:14.1 竣工结算申请请财务部门确定该条内容")

        try:
            match = '发包人审批竣工付款申请单的期限：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["发包人审批竣工付款申请单的期限"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("发包人审批竣工付款申请单的期限")
            else:
                factors_error["发包人审批竣工付款申请单的期限"] = "第三部分_14. 竣工结算:14.2 竣工结算审核发包人审批竣工付款申请单的期限未填写完整"
                addRemarkInDoc(word, document, "发包人审批竣工付款申请单的期限", f"第三部分_14. 竣工结算:14.2 竣工结算审核发包人审批竣工付款申请单的期限未填写完整")
        except:
            missObject += "第三部分_14. 竣工结算:要素“发包人审批竣工付款申请单的期限”缺失\n"
        factors_to_inform['发包人审批竣工付款申请单的期限'] = '第三部分_14. 竣工结算:14.2 竣工结算审核请财务部门确定该条内容'
        addRemarkInDoc(word, document, "发包人审批竣工付款申请单的期限", f"第三部分_14. 竣工结算:14.2 竣工结算审核请财务部门确定该条内容")

        try:
            match = '发包人完成竣工付款的期限：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["发包人完成竣工付款的期限"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("发包人完成竣工付款的期限")
            else:
                factors_error["发包人完成竣工付款的期限"] = "第三部分_14. 竣工结算:发包人完成竣工付款的期限未填写完整"
                addRemarkInDoc(word, document, "发包人完成竣工付款的期限", f"第三部分_14. 竣工结算:14.2 竣工结算审核发包人完成竣工付款的期限未填写完整")
        except:
            missObject += "第三部分_14. 竣工结算:14.2 竣工结算审核要素“发包人完成竣工付款的期限”缺失\n"
        factors_to_inform['发包人完成竣工付款的期限'] = '第三部分_14. 竣工结算:14.2 竣工结算审核请财务部门确定该条内容'
        addRemarkInDoc(word, document, "发包人完成竣工付款的期限", f"第三部分_14. 竣工结算:14.2 竣工结算审核请财务部门确定该条内容")

        try:
            match = '承包人提交最终结清申请单的份数：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["承包人提交最终结清申请单的份数"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("承包人提交最终结清申请单的份数")
            else:
                factors_error["承包人提交最终结清申请单的份数"] = "第三部分_14. 竣工结算:14.4.1 最终结清申请单承包人提交最终结清申请单的份数未填写完整"
                addRemarkInDoc(word, document, "承包人提交最终结清申请单的份数", f"第三部分_14. 竣工结算:14.4.1 最终结清申请单承包人提交最终结清申请单的份数未填写完整")
        except:
            missObject += "第三部分_14. 竣工结算:要素“承包人提交最终结清申请单的份数”缺失\n"
        factors_to_inform['承包人提交最终结清申请单的份数'] = '第三部分_14. 竣工结算:14.4.1 最终结清申请单不应少于3份'
        addRemarkInDoc(word, document, "承包人提交最终结清申请单的份数", f"第三部分_14. 竣工结算:14.4.1 最终结清申请单不应少于3份")

        try:
            match = '承包人提交最终结算申请单的期限：在缺陷责任期终止证书颁发后(.*)天内。\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["承包人提交最终结算申请单的期限"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    factor = int(factor)
                    if factor < 1 or factor > 14:
                        factors_error["承包人提交最终结算申请单的期限"] = "第三部分_14. 竣工结算:14.4.1 最终结清 数值在1-14之间选择"
                        addRemarkInDoc(word, document, "承包人提交最终结算申请单的期限", f"第三部分_14. 竣工结算:14.4.1 最终结清 数值在1-14之间选择")
                    else:
                        factors_ok.append("承包人提交最终结算申请单的期限")
                except:
                    factors_error["承包人提交最终结算申请单的期限"] = "第三部分_14. 竣工结算:14.4.1 最终结清 无法提取"
                    addRemarkInDoc(word, document, "承包人提交最终结算申请单的期限", f"第三部分_14. 竣工结算:14.4.1 最终结清 无法提取")
            else:
                factors_error["承包人提交最终结算申请单的期限"] = "第三部分_14. 竣工结算:14.4.1 最终结清 承包人提交最终结算申请单的期限未填写完整"
                addRemarkInDoc(word, document, "承包人提交最终结算申请单的期限", f"第三部分_14. 竣工结算:14.4.1 最终结清 承包人提交最终结算申请单的期限未填写完整")
        except:
            missObject += "第三部分_14. 竣工结算:14.4.1 最终结清 要素“承包人提交最终结算申请单的期限”缺失\n"

        try:
            match = '发包人完成最终结清申请单的审批并颁发最终结清证书的期限：承包人提交最终结算申请单后(.*)天内。\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["发包人完成最终结清申请单的审批并颁发最终结清证书的期限"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    factor = int(factor)
                    if factor > 14 and factor < 1:
                        factors_error["发包人完成最终结清申请单的审批并颁发最终结清证书的期限"] = "第三部分_14. 竣工结算:14.4.2 最终结清证书和支付 数值在1-14之间选择"
                        addRemarkInDoc(word, document, "发包人完成最终结清申请单的审批并颁发最终结清证书的期限",
                                       f"第三部分_14. 竣工结算:14.4.2 最终结清证书和支付 数值在1-14之间选择")
                    else:
                        factors_ok.append("发包人完成最终结清申请单的审批并颁发最终结清证书的期限")
                except:
                    factors_error["发包人完成最终结清申请单的审批并颁发最终结清证书的期限"] = "第三部分_14. 竣工结算:14.4.2 最终结清证书和支付 提取错误"
                    addRemarkInDoc(word, document, "发包人完成最终结清申请单的审批并颁发最终结清证书的期限",
                                   f"第三部分_14. 竣工结算:14.4.2 最终结清证书和支付 提取错误")
            else:
                factors_error[
                    "发包人完成最终结清申请单的审批并颁发最终结清证书的期限"] = "第三部分_14. 竣工结算:14.4.2 最终结清证书和支付 发包人完成最终结清申请单的审批并颁发最终结清证书的期限未填写完整"
                addRemarkInDoc(word, document, "发包人完成最终结清申请单的审批并颁发最终结清证书的期限",
                               f"第三部分_14. 竣工结算:14.4.2 最终结清证书和支付 发包人完成最终结清申请单的审批并颁发最终结清证书的期限未填写完整")
        except:
            missObject += "第三部分_14. 竣工结算:14.4.2 最终结清证书和支付 要素“发包人完成最终结清申请单的审批并颁发最终结清证书的期限”缺失\n"

        try:
            match = '发包人完成支付的期限：发包人完成最终结清申请单的审批并颁发最终结清证书，承包人提交书面申请后(.*)天内。\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["发包人完成支付的期限"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    factor = int(factor)
                    if factor > 28 and factor < 15:
                        factors_error["发包人完成支付的期限"] = "第三部分_14. 竣工结算:14.4.2 最终结清证书和支付 数值在15-28之间选择"
                        addRemarkInDoc(word, document, "发包人完成支付的期限", f"第三部分_14. 竣工结算:14.4.2 最终结清证书和支付 数值在15-28之间选择")
                    else:
                        factors_ok.append("发包人完成支付的期限")
                except:
                    factors_error["发包人完成支付的期限"] = "第三部分_14. 竣工结算:14.4.2 最终结清证书和支付 提取错误"
                    addRemarkInDoc(word, document, "发包人完成支付的期限", f"第三部分_14. 竣工结算:14.4.2 最终结清证书和支付 提取错误")
            else:
                factors_error["发包人完成支付的期限"] = "第三部分_14. 竣工结算:14.4.2 最终结清证书和支付 发包人完成支付的期限未填写完整"
                addRemarkInDoc(word, document, "发包人完成支付的期限", f"第三部分_14. 竣工结算:14.4.2 最终结清证书和支付 发包人完成支付的期限未填写完整")
        except:
            missObject += "第三部分_14. 竣工结算:14.4.2 最终结清证书和支付 要素“发包人完成支付的期限”缺失\n"

        # 3.15
        quexianzerenqi = None
        try:
            match = '缺陷责任期的具体期限：(.*)个月。\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["缺陷责任期的具体期限"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    factor = int(factor)
                    quexianzerenqi = factor
                    if factor < 24:
                        factors_error["缺陷责任期的具体期限"] = "第三部分_15. 缺陷责任期与保修 15.2缺陷责任期 数值在24以上"
                        addRemarkInDoc(word, document, "缺陷责任期的具体期限", f"第三部分_15. 缺陷责任期与保修 15.2缺陷责任期数值在24以上")
                    else:
                        factors_ok.append("缺陷责任期的具体期限")
                except:
                    factors_error["缺陷责任期的具体期限"] = "第三部分_15. 缺陷责任期与保修 15.2缺陷责任期提取错误"
                    addRemarkInDoc(word, document, "缺陷责任期的具体期限", f"第三部分_15. 缺陷责任期与保修 15.2缺陷责任期提取错误")
            else:
                factors_error["缺陷责任期的具体期限"] = "第三部分_15. 缺陷责任期与保修 15.2缺陷责任期缺陷责任期的具体期限未填写完整"
                addRemarkInDoc(word, document, "缺陷责任期的具体期限", f"第三部分_15. 缺陷责任期与保修 15.2缺陷责任期缺陷责任期的具体期限未填写完整")
        except:
            missObject += "第三部分_15. 缺陷责任期与保修 15.2缺陷责任期要素“缺陷责任期的具体期限”缺失\n"

        try:
            match = '质量保证金采用以下第(.*)种方式'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["质量保证金采用方式"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("质量保证金采用方式")
            else:
                factors_error["质量保证金采用方式"] = "第三部分_15. 缺陷责任期与保修 15.3.1 承包人提供质量保证金的方式 质量保证金采用方式未填写完整"
                addRemarkInDoc(word, document, "质量保证金采用以下第", f"第三部分_15. 缺陷责任期与保修 15.3.1 承包人提供质量保证金的方式质量保证金采用方式未填写完整")
        except:
            missObject += "第三部分_15. 缺陷责任期与保修 15.3.1 承包人提供质量保证金的方式 要素“质量保证金采用方式”缺失\n"
        factors_to_inform['质量保证金采用方式'] = '第三部分_15. 缺陷责任期与保修 15.3.1 承包人提供质量保证金的方式原则上采用第2种'
        addRemarkInDoc(word, document, "质量保证金采用以下第", f"第三部分_15. 缺陷责任期与保修 15.3.1 承包人提供质量保证金的方式原则上采用第2种")

        money_15_3_1_2 = None
        try:
            match = '（2）(.*)%的工程款；\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["%的工程款"] = factor
            if factor != "" and factor != "；" and factor != "。":
                money_15_3_1_2 = factor
                factors_ok.append("%的工程款")
            else:
                factors_error["%的工程款"] = "第三部分_15. 缺陷责任期与保修 15.3.1 承包人提供质量保证金的方式 %的工程款未填写完整"
                addRemarkInDoc(word, document, "%的工程款", f"第三部分_15. 缺陷责任期与保修 15.3.1 承包人提供质量保证金的方式 %的工程款未填写完整")
        except:
            missObject += "第三部分_15. 缺陷责任期与保修 15.3.1 承包人提供质量保证金的方式 要素“%的工程款”缺失\n"
        factors_to_inform['%的工程款'] = '第三部分_15. 缺陷责任期与保修 15.3.1 承包人提供质量保证金的方式 质保金不应少于3%。若少于请要求说明必要性，并由审计部门、财务部门核实'
        addRemarkInDoc(word, document, "%的工程款",
                       f"第三部分_15. 缺陷责任期与保修 15.3.1 承包人提供质量保证金的方式 质保金不应少于3%。若少于请要求说明必要性，并由审计部门、财务部门核实")

        try:
            match = '质量保证金的扣留采取以下第(.*)种方式'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["质量保证金的扣留采取方式"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("质量保证金的扣留采取方式")
            else:
                factors_error["质量保证金的扣留采取方式"] = "第三部分_15. 缺陷责任期与保修 15.3.2 质量保证金的扣留  质量保证金的扣留采取方式未填写完整"
                addRemarkInDoc(word, document, "质量保证金的扣留采取以下第", f"第三部分_15. 缺陷责任期与保修 15.3.2 质量保证金的扣留  质量保证金的扣留采取方式未填写完整")
        except:
            missObject += "第三部分_15. 缺陷责任期与保修 15.3.2 质量保证金的扣留  要素“质量保证金的扣留采取方式”缺失\n"
        factors_to_inform['质量保证金的扣留采取方式'] = '第三部分_15. 缺陷责任期与保修 15.3.2 质量保证金的扣留  原则上采用第2种'
        addRemarkInDoc(word, document, "质量保证金的扣留采取以下第", f"第三部分_15. 缺陷责任期与保修 15.3.2 质量保证金的扣留  原则上采用第2种")

        try:
            match = '质量保证金为工程结算审定金额的(.*)％.*'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["质量保证金为工程结算审定金额的【】％"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != money_15_3_1_2.replace(" ", "").replace("；", "").replace("。", ""):
                    factors_error["质量保证金为工程结算审定金额的【】％"] = "第三部分_15. 缺陷责任期与保修 15.3.2 质量保证金的扣留  应与15.3.1第（2）款的空一致"
                    addRemarkInDoc(word, document, "关于质量保证金的补充约定",
                                   f"第三部分_15. 缺陷责任期与保修 15.3.2 质量保证金的扣留  应与15.3.1第（2）款的空一致")
                else:
                    factors_ok.append("质量保证金为工程结算审定金额的【】％")
            else:
                factors_error["质量保证金为工程结算审定金额的【】％"] = "第三部分_15. 缺陷责任期与保修 15.3.2 质量保证金的扣留  质量保证金为工程结算审定金额的【】％未填写完整"
                addRemarkInDoc(word, document, "关于质量保证金的补充约定",
                               f"第三部分_15. 缺陷责任期与保修 15.3.2 质量保证金的扣留  质量保证金为工程结算审定金额的【】％未填写完整")
        except:
            missObject += "第三部分_15. 缺陷责任期与保修 15.3.2 质量保证金的扣留  要素“质量保证金为工程结算审定金额的【】％”缺失\n"

        try:
            match = '在缺限责任期满后无质量问题(.*)天内由发包人无息'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["在缺限责任期满后无质量问题【】天"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    factor = int(factor)
                    if factor > 28 and factor < 15:
                        factors_error["在缺限责任期满后无质量问题【】天"] = "第三部分_15. 缺陷责任期与保修 15.3.2 质量保证金的扣留  数值在15-28之间选择"
                        addRemarkInDoc(word, document, "在缺限责任期满后无质量问题",
                                       f"第三部分_15. 缺陷责任期与保修 15.3.2 质量保证金的扣留  数值在15-28之间选择")
                    else:
                        factors_ok.append("在缺限责任期满后无质量问题【】天")
                except:
                    factors_error["在缺限责任期满后无质量问题【】天"] = "第三部分_15. 缺陷责任期与保修 15.3.2 质量保证金的扣留  提取错误"
                    addRemarkInDoc(word, document, "在缺限责任期满后无质量问题", f"第三部分_15. 缺陷责任期与保修 15.3.2 质量保证金的扣留  提取错误")
            else:
                factors_error["在缺限责任期满后无质量问题【】天"] = "第三部分_15. 缺陷责任期与保修 15.3.2 质量保证金的扣留  在缺限责任期满后无质量问题【】天未填写完整"
                addRemarkInDoc(word, document, "在缺限责任期满后无质量问题",
                               f"第三部分_15. 缺陷责任期与保修 15.3.2 质量保证金的扣留  在缺限责任期满后无质量问题【】天未填写完整")
        except:
            missObject += "第三部分_15. 缺陷责任期与保修 15.3.2 质量保证金的扣留  要素“在缺限责任期满后无质量问题【】天”缺失\n"
        factors_to_inform['在缺限责任期满后无质量问题'] = '第三部分_15. 缺陷责任期与保修 15.3.2 质量保证金的扣留  请由财务部门核实该支付日期'
        addRemarkInDoc(word, document, "在缺限责任期满后无质量问题", f"第三部分_15. 缺陷责任期与保修 15.3.2 质量保证金的扣留  请由财务部门核实该支付日期")

        try:
            match = '双方同意(.*)扣留防水质保金。\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["双方是否扣留防水质保金"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if factor != '是' and factor != '否':
                    factors_error["双方是否扣留防水质保金"] = "第三部分_15. 缺陷责任期与保修 15.3.2 质量保证金的扣留  填写内容为是、否"
                    addRemarkInDoc(word, document, "扣留防水质保金", f"第三部分_15. 缺陷责任期与保修 15.3.2 质量保证金的扣留  填写内容为是、否")
                else:
                    factors_ok.append("双方是否扣留防水质保金")
            else:
                factors_error["双方是否扣留防水质保金"] = "第三部分_15. 缺陷责任期与保修 15.3.2 质量保证金的扣留  双方是否扣留防水质保金未填写完整"
                addRemarkInDoc(word, document, "扣留防水质保金", f"第三部分_15. 缺陷责任期与保修 15.3.2 质量保证金的扣留  双方是否扣留防水质保金未填写完整")
        except:
            missObject += "第三部分_15. 缺陷责任期与保修 15.3.2 质量保证金的扣留  要素“双方是否扣留防水质保金”缺失\n"

        try:
            match = '承包人收到保修通知并到达工程现场的合理时间：(.*)小时以内。\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["15.4.3 修复通知"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    factor = int(factor)
                    if factor > 48:
                        factors_error["15.4.3 修复通知"] = "第三部分_15. 缺陷责任期与保修 15.4.3 修复通知 数值为小于等于48"
                        addRemarkInDoc(word, document, "15.4.3 修复通知", f"第三部分_15. 缺陷责任期与保修 15.4.3 修复通知 数值为小于等于48")
                    else:
                        factors_ok.append("15.4.3 修复通知")
                except:
                    factors_error["15.4.3 修复通知"] = "第三部分_15. 缺陷责任期与保修 15.4.3 修复通知 提取错误"
                    addRemarkInDoc(word, document, "15.4.3 修复通知", f"第三部分_15. 缺陷责任期与保修 15.4.3 修复通知 提取错误")
            else:
                factors_error["15.4.3 修复通知"] = "第三部分_15. 缺陷责任期与保修 15.4.3 修复通知 15.4.3 修复通知未填写完整"
                addRemarkInDoc(word, document, "15.4.3 修复通知", f"第三部分_15. 缺陷责任期与保修 15.4.3 修复通知 15.4.3 修复通知未填写完整")
        except:
            missObject += "第三部分_15. 缺陷责任期与保修 15.4.3 修复通知 要素“15.4.3 修复通知”缺失\n"

        # 3.16
        try:
            match = '（5）承包人未能按合同约定时限完成相应工作及未能按施工进度计划及时完成合同约定的工作，造成工期延误的，承包人自逾期之日起向发包人支付(.*)元/天的违约金'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["16.2.2承包人违约的责任_5_1"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    factor = int(factor)
                    if factor != content_7_5_2_1:
                        factors_error["16.2.2承包人违约的责任_5_1"] = "第三部分_16. 违约 16.2.2承包人违约的责任 本条第一个空与7.5.2条的第一个空应该一致"
                        addRemarkInDoc(word, document,
                                       "（5）承包人未能按合同约定时限完成相应工作及未能按施工进度计划及时完成合同约定的工作，造成工期延误的，承包人自逾期之日起向发包人支付",
                                       f"第三部分_16. 违约 16.2.2承包人违约的责任 本条第一个空与7.5.2条的第一个空应该一致")
                    else:
                        factors_ok.append("16.2.2承包人违约的责任_5_1")
                except:
                    factors_error["16.2.2承包人违约的责任_5_1"] = "第三部分_16. 违约 16.2.2承包人违约的责任 提取错误"
                    addRemarkInDoc(word, document, "（5）承包人未能按合同约定时限完成相应工作及未能按施工进度计划及时完成合同约定的工作，造成工期延误的，承包人自逾期之日起向发包人支付",
                                   f"第三部分_16. 违约 16.2.2承包人违约的责任 16.2.2承包人违约的责任 提取错误")
            else:
                factors_error["16.2.2承包人违约的责任_5_1"] = "第三部分_16. 违约 16.2.2承包人违约的责任 16.2.2承包人违约的责任_5_1未填写完整"
                addRemarkInDoc(word, document, "（5）承包人未能按合同约定时限完成相应工作及未能按施工进度计划及时完成合同约定的工作，造成工期延误的，承包人自逾期之日起向发包人支付",
                               f"第三部分_16. 违约 16.2.2承包人违约的责任 16.2.2承包人违约的责任_5_1未填写完整")
        except:
            missObject += "第三部分_16. 违约 16.2.2承包人违约的责任要素“16.2.2承包人违约的责任_5_1”缺失\n"

        try:
            match = '违约金累计计算，但总额不超过合同总价款的(.*)%，工期不予顺延；\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["16.2.2承包人违约的责任_5_2"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    factor = int(factor)
                    if factor != content_7_5_2_3:
                        factors_error["16.2.2承包人违约的责任_5_2"] = "第三部分_16. 违约 16.2.2承包人违约的责任要素 本条第二个空与7.5.2条的第三个空应该一致"
                        addRemarkInDoc(word, document,
                                       "（5）承包人未能按合同约定时限完成相应工作及未能按施工进度计划及时完成合同约定的工作，造成工期延误的，承包人自逾期之日起向发包人支付",
                                       f"第三部分_16. 违约 16.2.2承包人违约的责任要素 本条第二个空与7.5.2条的第三个空应该一致")
                    else:
                        factors_ok.append("16.2.2承包人违约的责任_5_2")
                except:
                    factors_error["16.2.2承包人违约的责任_5_2"] = "第三部分_16. 违约 16.2.2承包人违约的责任要素 提取错误"
                    addRemarkInDoc(word, document, "（5）承包人未能按合同约定时限完成相应工作及未能按施工进度计划及时完成合同约定的工作，造成工期延误的，承包人自逾期之日起向发包人支付",
                                   f"第三部分_16. 违约 16.2.2承包人违约的责任要素 提取错误")
            else:
                factors_error["16.2.2承包人违约的责任_5_2"] = "第三部分_16. 违约 16.2.2承包人违约的责任要素 16.2.2承包人违约的责任_5_2未填写完整"
                addRemarkInDoc(word, document, "（5）承包人未能按合同约定时限完成相应工作及未能按施工进度计划及时完成合同约定的工作，造成工期延误的，承包人自逾期之日起向发包人支付",
                               f"第三部分_16. 违约 16.2.2承包人违约的责任要素 16.2.2承包人违约的责任_5_2未填写完整")
        except:
            missObject += "第三部分_16. 违约 16.2.2承包人违约的责任要素 要素“16.2.2承包人违约的责任_5_2”缺失\n"

        # 3.21
        try:
            match = '21.6 开票信息\n单位名称：(.*)；\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["21.6单位名称"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("21.6单位名称")
            else:
                factors_error["21.6单位名称"] = "第三部分_21.6 开票信息 21.6单位名称未填写完整"
                addRemarkInDoc(word, document, "21.6 开票信息", f"第三部分_21.6 开票信息 21.6单位名称未填写完整")
        except:
            missObject += "第三部分_21.6 开票信息 要素“21.6单位名称”缺失\n"

        try:
            match = '21.6 开票信息\n单位名称：.*；\n税号：(.*)；\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["21.6税号"] = factor
            if factor != "" and factor != "；" and factor != "。":
                Unified = UnifiedSocialCreditIdentifier()
                if Unified.check_social_credit_code(factor) == False:
                    factors_error["21.6税号"] = "第三部分_21.6 开票信息 税号错误"
                    addRemarkInDoc(word, document, "21.6 开票信息", f"第三部分_21.6 开票信息 税号错误")
                else:
                    factors_ok.append("21.6税号")
            else:
                factors_error["21.6税号"] = "第三部分_21.6 开票信息 21.6税号未填写完整"
                addRemarkInDoc(word, document, "21.6 开票信息", f"第三部分_21.6 开票信息 21.6税号未填写完整")
        except:
            missObject += "要素“21.6税号”缺失\n"

        try:
            match = '21.6 开票信息\n单位名称：.*；\n税号：.*；\n地址：(.*)；\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["21.6地址"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("21.6地址")
            else:
                factors_error["21.6地址"] = "第三部分_21.6 开票信息 21.6地址未填写完整"
                addRemarkInDoc(word, document, "21.6 开票信息", f"第三部分_21.6 开票信息 21.6地址未填写完整")
        except:
            missObject += "第三部分_21.6 开票信息 要素“21.6地址”缺失\n"

        try:
            match = '21.6 开票信息\n单位名称：.*；\n税号：.*；\n地址：.*；\n电话：(.*)；\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["21.6电话"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if isTelPhoneNumber(factor) == "Error":
                    factors_error["21.6电话"] = "第三部分_21.6 开票信息 电话填写错误"
                    addRemarkInDoc(word, document, "21.6 开票信息", f"第三部分_21.6 开票信息 电话填写错误")
                else:
                    factors_ok.append("21.6电话")
            else:
                factors_error["21.6电话"] = "第三部分_21.6 开票信息 21.6电话未填写完整"
                addRemarkInDoc(word, document, "21.6 开票信息", f"第三部分_21.6 开票信息 21.6电话未填写完整")
        except:
            missObject += "第三部分_21.6 开票信息 要素“21.6电话”缺失\n"

        try:
            match = '21.6 开票信息\n单位名称：.*；\n税号：.*；\n地址：.*；\n电话：.*；\n开户银行：(.*)；\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t",
                                                                                                           "").replace(
                '\t', '')
            factors["21.6开户银行"] = factor
            if factor != "" and factor != "；" and factor != "。":
                if is_bankcard(factor) == False:
                    factors_error["21.6开户银行"] = "第三部分_21.6 开票信息 开户银行填写错误"
                    addRemarkInDoc(word, document, "21.6 开票信息", f"第三部分_21.6 开票信息 开户银行填写错误")
                else:
                    factors_ok.append("21.6开户银行")
            else:
                factors_error["21.6开户银行"] = "第三部分_21.6 开票信息 21.6开户银行未填写完整"
                addRemarkInDoc(word, document, "21.6 开票信息", f"第三部分_21.6 开票信息 21.6开户银行未填写完整")
        except:
            missObject += "第三部分_21.6 开票信息 要素“21.6开户银行”缺失\n"

        try:
            match = '21.6 开票信息\n单位名称：.*；\n税号：.*；\n地址：.*；\n电话：.*；\n开户银行：.*；\n账号：(.*)。\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "")
            factors["21.6账号"] = factor
            if factor != "" and factor != "；" and factor != "。":
                factors_ok.append("21.6账号")
            else:
                factors_error["21.6账号"] = "第三部分_21.6 开票信息 21.6账号未填写完整"
                addRemarkInDoc(word, document, "21.6开户银行", f"第三部分_21.6 开票信息 21.6账号未填写完整")
        except:
            missObject += "第三部分_21.6 开票信息 要素“21.6账号”缺失\n"

    # 附件1-9
    if 1 == True:
        for t in tables:
            top = t.rows[0]
            tops = []
            for cell in top.cells:
                tops.append(
                    cell.text.replace("\n", "").replace("\t", "").replace("\r", ""))

            # 附件1
            if tops[0] == '单位工程名称':
                factors_to_inform["承包人承揽工程项目一览表"] = '附件 承包人承揽工程项目一览表:请工程技术部门对此确认'
                addRemarkInDoc(word, document, "单位工程名称", f"附件 承包人承揽工程项目一览表:请工程技术部门对此确认")
                t = tables[0]
                cell_all = []
                for row in t.rows:
                    cells = []
                    for cell in row.cells:
                        cells.append(
                            cell.text.replace("\n", "").replace("\t", "").replace("\r", ""))
                    cell_all.append(cells)

                if table_ok(cell_all) == False:
                    addRemarkInDoc(word, document, "单位工程名称", f"附件 承包人承揽工程项目一览表:附件 承包人承揽工程项目一览表表格填写错误")

            # 附件2
            if tops[2] == '规格型号' and tops[3] == '单位':
                factors_to_inform["发包人供应材料设备一览表"] = '附件 发包人供应材料设备一览表:请工程技术部门对此确认'
                addRemarkInDoc(word, document, "发包人供应材料设备一览表", f"附件 发包人供应材料设备一览表:请工程技术部门对此确认")
                t = tables[1]
                cell_all = []
                for row in t.rows:
                    cells = []
                    for cell in row.cells:
                        cells.append(
                            cell.text.replace("\n", "").replace("\t", "").replace("\r", ""))
                    cell_all.append(cells)

                if table_ok(cell_all) == False:
                    addRemarkInDoc(word, document, "发包人供应材料设备一览表", f"附件 发包人供应材料设备一览表:附件 发包人供应材料设备一览表表格填写错误")

            # 附件4
            if tops[0] == '文件名称':
                factors_to_inform["主要建设工程文件目录"] = '附件 主要建设工程文件目录:请工程技术部门对此确认'
                addRemarkInDoc(word, document, "主要建设工程文件目录", f"附件 主要建设工程文件目录:请工程技术部门对此确认")
                t = tables[2]
                cell_all = []
                for row in t.rows:
                    cells = []
                    for cell in row.cells:
                        cells.append(
                            cell.text.replace("\n", "").replace("\t", "").replace("\r", ""))
                    cell_all.append(cells)

                if table_ok(cell_all) == False:
                    addRemarkInDoc(word, document, "主要建设工程文件目录", f"附件 主要建设工程文件目录:附件 主要建设工程文件目录表格填写错误")

            # 附件5
            if tops[1] == '机械或设备名称':
                factors_to_inform["承包人用于本工程施工的机械设备表"] = '请工程技术部门对此确认'
                addRemarkInDoc(word, document, "承包人用于本工程施工的机械设备表", f"附件 承包人用于本工程施工的机械设备表:请工程技术部门对此确认")
                t = tables[3]
                cell_all = []
                for row in t.rows:
                    cells = []
                    for cell in row.cells:
                        cells.append(
                            cell.text.replace("\n", "").replace("\t", "").replace("\r", ""))
                    cell_all.append(cells)

                if table_ok(cell_all) == False:
                    addRemarkInDoc(word, document, "承包人用于本工程施工的机械设备表", f"附件 承包人用于本工程施工的机械设备表:附件 承包人用于本工程施工的机械设备表表格填写错误")

            # 附件6
            try:
                if tops[4] == '主要资历、经验及承担过的项目':
                    factors_to_inform["承包人主要施工管理人员表"] = '附件 承包人主要施工管理人员表:请核实与主合同约定内容是否一致，请工程技术部门对此确认'
                    addRemarkInDoc(word, document, "承包人主要施工管理人员表", f"附件 承包人主要施工管理人员表:请核实与主合同约定内容是否一致，请工程技术部门对此确认")
                    t = tables[4]
                    cell_all = []
                    for row in t.rows:
                        cells = []
                        for cell in row.cells[1:]:
                            cells.append(
                                cell.text.replace("\n", "").replace("\t", "").replace("\r", ""))
                        cell_all.append(cells)
                    # print(cell_all)

                    if table_ok(cell_all) == False:
                        addRemarkInDoc(word, document, "承包人主要施工管理人员表", f"附件 承包人主要施工管理人员表:附件 承包人主要施工管理人员表表格填写错误")
            except:
                pass

            # 附件7
            try:
                if tops[4] == '主要资历、经验及承担过的项目':
                    factors_to_inform["分包人主要施工管理人员表"] = '附件 分包人主要施工管理人员表:请核实与主合同约定内容是否一致，请工程技术部门对此确认'
                    addRemarkInDoc(word, document, "分包人主要施工管理人员表", f"附件 分包人主要施工管理人员表:请核实与主合同约定内容是否一致，请工程技术部门对此确认")
                    t = tables[5]
                    cell_all = []
                    for row in t.rows:
                        cells = []
                        for cell in row.cells[1:]:
                            cells.append(
                                cell.text.replace("\n", "").replace("\t", "").replace("\r", ""))
                        cell_all.append(cells)
                    # print(cell_all)

                    if table_ok(cell_all) == False:
                        addRemarkInDoc(word, document, "分包人主要施工管理人员表", f"附件 分包人主要施工管理人员表:附件 分包人主要施工管理人员表表格填写错误")
            except:
                pass

            # 附件11
            if tops[0] == '序号' and tops[1] == '名称':
                factors_to_inform["材料暂估价表"] = '附件 暂估价一览表:请工程技术部门对此确认。'
                addRemarkInDoc(word, document, '材料暂估价表', f"附件 暂估价一览表:请工程技术部门对此确认。")
                t = tables[6]
                cell_all = []
                for row in t.rows:
                    cells = []
                    for cell in row.cells:
                        cells.append(
                            cell.text.replace("\n", "").replace("\t", "").replace("\r", ""))
                    cell_all.append(cells)
                # print(cell_all)

                if table_ok(cell_all) == False:
                    addRemarkInDoc(word, document, "材料暂估价表", f"附件 暂估价一览表:附件 暂估价一览表-1表格填写错误")

            if tops[0] == '序号' and tops[1] == '名称':
                factors_to_inform["工程设备暂估价表"] = '请工程技术部门对此确认。'
                addRemarkInDoc(word, document, '工程设备暂估价表', f"附件 暂估价一览表:请工程技术部门对此确认。")
                t = tables[7]
                cell_all = []
                for row in t.rows:
                    cells = []
                    for cell in row.cells:
                        cells.append(
                            cell.text.replace("\n", "").replace("\t", "").replace("\r", ""))
                    cell_all.append(cells)
                # print(cell_all)

                if table_ok(cell_all) == False:
                    addRemarkInDoc(word, document, "工程设备暂估价表", f"附件 暂估价一览表:附件 暂估价一览表-2表格填写错误")

            if tops[0] == '序号' and tops[1] == '专业工程名称':
                factors_to_inform["专业工程暂估价表"] = '附件 暂估价一览表:请工程技术部门对此确认。'
                addRemarkInDoc(word, document, '专业工程暂估价表', f"附件 暂估价一览表:请工程技术部门对此确认。")
                t = tables[8]
                cell_all = []
                for row in t.rows:
                    cells = []
                    for cell in row.cells:
                        cells.append(
                            cell.text.replace("\n", "").replace("\t", "").replace("\r", ""))
                    cell_all.append(cells)
                # print(cell_all)

                if table_ok(cell_all) == False:
                    addRemarkInDoc(word, document, "专业工程暂估价表", f"附件 暂估价一览表:附件 暂估价一览表-3表格填写错误")

        # 附件3
        try:
            match = '工程质量保修书\n发包人[（]全称[）]：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["工程质量保修书_发包人"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    if factor != fabaoren:
                        factors_error["工程质量保修书_发包人"] = "附件 工程质量保修书:不与合同的发包人名称一致"
                        addRemarkInDoc(word, document, "发包人（全称）：", f"附件 工程质量保修书:不与合同的发包人名称一致")
                    else:
                        factors_ok.append("工程质量保修书_发包人")
                except:
                    factors_error["工程质量保修书_发包人"] = "附件 工程质量保修书:合同的发包人提取失败"
                    addRemarkInDoc(word, document, "发包人（全称）：", f"附件 工程质量保修书:合同的发包人提取失败")
            else:
                factors_error["工程质量保修书_发包人"] = "附件 工程质量保修书:工程质量保修书_发包人未填写完整"
                addRemarkInDoc(word, document, "发包人（全称）：", f"附件 工程质量保修书:工程质量保修书_发包人未填写完整")
        except:
            missObject += "附件 工程质量保修书:要素“工程质量保修书_发包人”缺失\n"

        factors_to_inform["发包人（全称）："] = '附件 工程质量保修书:请工程技术部门对此确认'
        addRemarkInDoc(word, document, "发包人（全称）：", f"附件 工程质量保修书:请工程技术部门对此确认")

        try:
            match = '工程质量保修书\n发包人[（]全称[）]：.*\n承包人[（]全称[）]：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["工程质量保修书_承包人"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    if factor != chengbaoren:
                        factors_error["工程质量保修书_承包人"] = "附件 工程质量保修书:不与合同的承包人名称一致"
                        addRemarkInDoc(word, document, "承包人（全称）：", f"附件 工程质量保修书:不与合同的承包人名称一致")
                    else:
                        factors_ok.append("工程质量保修书_承包人")
                except:
                    factors_error["工程质量保修书_承包人"] = "附件 工程质量保修书:合同的承包人提取失败"
                    addRemarkInDoc(word, document, "承包人（全称）：", f"附件 工程质量保修书:合同的承包人提取失败")
            else:
                factors_error["工程质量保修书_承包人"] = "附件 工程质量保修书:工程质量保修书_承包人未填写完整"
                addRemarkInDoc(word, document, "承包人（全称）：", f"附件 工程质量保修书:工程质量保修书_承包人未填写完整")
        except:
            missObject += "附件 工程质量保修书:要素“工程质量保修书_承包人”缺失\n"

        factors_to_inform["承包人（全称）："] = '附件 工程质量保修书:请工程技术部门对此确认'
        addRemarkInDoc(word, document, "承包人（全称）：", f"附件 工程质量保修书:请工程技术部门对此确认")

        try:
            match = '.*发包人和承包人根据《中华人民共和国建筑法》和《建设工程质量管理条例》，经协商一致就(.*)[（]工程全称[）]签订工程质量保修书。\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["工程质量保修书_工程全称"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    if factor != gongchengmingcheng:
                        factors_error["工程质量保修书_工程全称"] = "附件 工程质量保修书:码云与合同协议书第一条第1款工程名称一致"
                        addRemarkInDoc(word, document, "发包人和承包人根据《中华人民共和国建筑法》和《建设工程质量管理条例》，经协商一致就",
                                       f"附件 工程质量保修书:码云与合同协议书第一条第1款工程名称一致")
                    else:
                        factors_ok.append("工程质量保修书_工程全称")
                except:
                    factors_error["工程质量保修书_工程全称"] = "附件 工程质量保修书:合同协议书第一条第1款工程名称提取失败"
                    addRemarkInDoc(word, document, "发包人和承包人根据《中华人民共和国建筑法》和《建设工程质量管理条例》，经协商一致就",
                                   f"附件 工程质量保修书 合同协议书第一条第1款工程名称提取失败")
            else:
                factors_error["工程质量保修书_工程全称"] = "附件 工程质量保修书:工程质量保修书_工程全称未填写完整"
                addRemarkInDoc(word, document, "发包人和承包人根据《中华人民共和国建筑法》和《建设工程质量管理条例》，经协商一致就", f"附件 工程质量保修书 工程质量保修书_工程全称未填写完整")
        except:
            missObject += "附件  工程质量保修书:要素“工程质量保修书_工程全称”缺失\n"

        try:
            match = '三、缺陷责任期\n工程缺陷责任期为(.*)个月.*'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors["三、缺陷责任期"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    if factor != quexianzerenqi:
                        factors_error["三、缺陷责任期"] = "附件 工程质量保修书:没有与专用条款15.2约定的期限一致"
                        addRemarkInDoc(word, document, "三、缺陷责任期", f"附件 工程质量保修书:没有与专用条款15.2约定的期限一致")
                    else:
                        factors_ok.append("三、缺陷责任期")
                except:
                    factors_error["三、缺陷责任期"] = "缺附件 工程质量保修书:陷责任期提取失败"
                    addRemarkInDoc(word, document, "三、缺陷责任期", f"附件 工程质量保修书:缺陷责任期提取失败")
            else:
                factors_error["三、缺陷责任期"] = "附件 工程质量保修书:三、缺陷责任期未填写完整"
                addRemarkInDoc(word, document, "三、缺陷责任期", f"附件 工程质量保修书:三、缺陷责任期未填写完整")
        except:
            missObject += "附件 工程质量保修书:要素“三、缺陷责任期”缺失\n"

        try:
            match = '工程质量保修书由发包人、承包人在工程竣工验收前共同签署，作为施工合同附件，其有效期限至保修期满。\n*发包人[（]公章[）]：(.*)承包人[（]公章[）]：(.*)\n'
            factor = re.findall(match, text)[0]
            factor = [factor[0].replace(" ", "").replace("；", "").replace("。", "").replace("；", "").replace("。", ""),
                      factor[1].replace(" ", "").replace("；", "").replace("。", "").replace("；", "").replace("。", "")]
            factors["附件 工程质量保修书 结尾公章"] = factor
            if factor != ["", ""]:
                factors_ok.append("附件 工程质量保修书 结尾公章")
            else:
                factors_error["附件 工程质量保修书 结尾公章"] = "附件 工程质量保修书 :附件 工程质量保修书 结尾公章未填写完整"
                addRemarkInDoc(word, document, "工程质量保修书由发包人、承包人在工程竣工验收前共同签署，作为施工合同附件，其有效期限至保修期满。", f"附件 工程质量保修书 结尾公章未填写完整")
        except:
            missObject += "附件 工程质量保修书 :要素“附件 工程质量保修书 结尾公章”缺失\n"

        try:
            match = '工程质量保修书由发包人、承包人在工程竣工验收前共同签署，作为施工合同附件，其有效期限至保修期满。\n.*\n地  址：(.*)地  址：(.*)\n'
            factor = re.findall(match, text)[0]
            factor = [factor[0].replace(" ", "").replace("；", "").replace("。", "").replace("；", "").replace("。", ""),
                      factor[1].replace(" ", "").replace("；", "").replace("。", "").replace("；", "").replace("。", "")]
            factors["附件 工程质量保修书 结尾地址"] = factor
            if factor != ["", ""]:
                try:
                    if factor != dizhi:
                        factors_error["附件 工程质量保修书 结尾地址"] = "附件 工程质量保修书:与合同协议书落款处内容不一致"
                        addRemarkInDoc(word, document, "工程质量保修书由发包人、承包人在工程竣工验收前共同签署，作为施工合同附件，其有效期限至保修期满。",
                                       f"附件 工程质量保修书 与合同协议书落款处内容不一致")
                    else:
                        factors_ok.append("附件 工程质量保修书 结尾地址")
                except:
                    factors_error["附件 工程质量保修书 结尾地址"] = "附件 工程质量保修书:合同协议书落款地址提取错误"
                    addRemarkInDoc(word, document, "工程质量保修书由发包人、承包人在工程竣工验收前共同签署，作为施工合同附件，其有效期限至保修期满。",
                                   f"附件 工程质量保修书:合同协议书落款地址提取错误")
            else:
                factors_error["附件 工程质量保修书 结尾地址"] = "附件 工程质量保修书:附件 工程质量保修书 结尾地址未填写完整"
                addRemarkInDoc(word, document, "工程质量保修书由发包人、承包人在工程竣工验收前共同签署，作为施工合同附件，其有效期限至保修期满。", f"附件 工程质量保修书 结尾地址未填写完整")
        except:
            missObject += "附件 工程质量保修书:要素“附件 工程质量保修书 结尾地址”缺失\n"

        try:
            match = '工程质量保修书由发包人、承包人在工程竣工验收前共同签署，作为施工合同附件，其有效期限至保修期满。.|\n*法定代表人[（]签字[）]：(.*)法定代表人[（]签字[）]：(.*)\n'
            factor = re.findall(match, text)[0]
            factor = [factor[0].replace(" ", "").replace("；", "").replace("。", "").replace("；", "").replace("。", ""),
                      factor[1].replace(" ", "").replace("；", "").replace("。", "").replace("；", "").replace("。", "")]
            factors["附件 工程质量保修书 结尾法定代表人"] = factor
            if factor != ["", ""]:
                try:
                    if factor != fadingdaibiaoren:
                        factors_error["附件 工程质量保修书 结尾法定代表人"] = "附件 工程质量保修书:与合同协议书落款处内容不一致"
                        addRemarkInDoc(word, document, "工程质量保修书由发包人、承包人在工程竣工验收前共同签署，作为施工合同附件，其有效期限至保修期满。",
                                       f"附件 工程质量保修书 :与合同协议书落款处内容不一致")
                    else:
                        factors_ok.append("附件 工程质量保修书 结尾法定代表人")
                except:
                    factors_error["附件 工程质量保修书 结尾法定代表人"] = "附件 工程质量保修书:合同协议书法定代表人提取错误"
                    addRemarkInDoc(word, document, "工程质量保修书由发包人、承包人在工程竣工验收前共同签署，作为施工合同附件，其有效期限至保修期满。",
                                   f"附件 工程质量保修书:合同协议书法定代表人提取错误")
            else:
                factors_error["附件 工程质量保修书 结尾法定代表人"] = "附件 工程质量保修书:附件 工程质量保修书结尾法定代表人未填写完整"
                addRemarkInDoc(word, document, "工程质量保修书由发包人、承包人在工程竣工验收前共同签署，作为施工合同附件，其有效期限至保修期满。", f"附件 工程质量保修书 结尾法定代表人未填写完整")
        except:
            missObject += "附件 工程质量保修书:要素“附件 工程质量保修书结尾法定代表人”缺失\n"

        try:
            match = '工程质量保修书由发包人、承包人在工程竣工验收前共同签署，作为施工合同附件，其有效期限至保修期满。.|\n*委托代理人[（]签字[）]：(.*)委托代理人[（]签字[）]：(.*)\n'
            factor = re.findall(match, text)[0]
            factor = [factor[0].replace(" ", "").replace("；", "").replace("。", "").replace("；", "").replace("。", ""),
                      factor[1].replace(" ", "").replace("；", "").replace("。", "").replace("；", "").replace("。", "")]
            factors[" 附件：工程质量保修书结尾委托代理人"] = factor
            if factor != ["", ""]:
                try:
                    if factor != fadingdaibiaoren:
                        factors_error[" 附件：工程质量保修书结尾委托代理人"] = " 附件：工程质量保修书:与合同协议书落款处内容不一致"
                        addRemarkInDoc(word, document, "工程质量保修书由发包人、承包人在工程竣工验收前共同签署，作为施工合同附件，其有效期限至保修期满。",
                                       f" 附件：工程质量保修书:与合同协议书落款处内容不一致")
                    else:
                        factors_ok.append(" 附件：工程质量保修书结尾委托代理人")
                except:
                    factors_error[" 附件：工程质量保修书结尾委托代理人"] = " 附件：工程质量保修书:合同协议书委托代理人提取错误"
                    addRemarkInDoc(word, document, "工程质量保修书由发包人、承包人在工程竣工验收前共同签署，作为施工合同附件，其有效期限至保修期满。",
                                   f" 附件：工程质量保修书:合同协议书委托代理人提取错误")
            else:
                factors_error[" 附件：工程质量保修书结尾委托代理人"] = " 附件：工程质量保修书: 附件：工程质量保修书结尾委托代理人未填写完整"
                addRemarkInDoc(word, document, "工程质量保修书由发包人、承包人在工程竣工验收前共同签署，作为施工合同附件，其有效期限至保修期满。", f" 附件：工程质量保修书结尾委托代理人未填写完整")
        except:
            missObject += " 附件：工程质量保修书:要素“ 附件：工程质量保修书结尾委托代理人”缺失\n"

        try:
            match = '工程质量保修书由发包人、承包人在工程竣工验收前共同签署，作为施工合同附件，其有效期限至保修期满。.|\n电  话：(.*)电  话：(.*)\n'
            factor = re.findall(match, text)[0]
            factor = [factor[0].replace(" ", "").replace("；", "").replace("。", "").replace("；", "").replace("。", ""),
                      factor[1].replace(" ", "").replace("；", "").replace("。", "").replace("；", "").replace("。", "")]
            factors[" 附件：工程质量保修书结尾电话"] = factor
            if factor != ["", ""]:
                try:
                    if factor != dianhua:
                        factors_error[" 附件：工程质量保修书结尾电话"] = " 附件：工程质量保修书:与合同协议书落款处内容不一致"
                        addRemarkInDoc(word, document, "工程质量保修书由发包人、承包人在工程竣工验收前共同签署，作为施工合同附件，其有效期限至保修期满。",
                                       f" 附件：工程质量保修书:与合同协议书落款处内容不一致")
                    else:
                        factors_ok.append(" 附件：工程质量保修书结尾电话")
                except:
                    factors_error[" 附件：工程质量保修书结尾电话"] = " 附件：工程质量保修书:合同协议书电话提取错误"
                    addRemarkInDoc(word, document, "工程质量保修书由发包人、承包人在工程竣工验收前共同签署，作为施工合同附件，其有效期限至保修期满。",
                                   f" 附件：工程质量保修书合同协议书电话提取错误")
            else:
                factors_error[" 附件：工程质量保修书结尾电话"] = " 附件：工程质量保修书: 附件：工程质量保修书结尾电话未填写完整"
                addRemarkInDoc(word, document, "工程质量保修书由发包人、承包人在工程竣工验收前共同签署，作为施工合同附件，其有效期限至保修期满。", f" 附件：工程质量保修书结尾电话未填写完整")
        except:
            missObject += " 附件：工程质量保修书:要素“ 附件：工程质量保修书结尾电话”缺失\n"

        factors_to_inform[" 附件：工程质量保修书结尾"] = ' 附件：工程质量保修书:请核实下列信息是否与合同协议书内容是否一致'
        addRemarkInDoc(word, document, "工程质量保修书由发包人、承包人在工程竣工验收前共同签署，作为施工合同附件，其有效期限至保修期满。", f"请核实下列信息是否与合同协议书内容是否一致")

        # 附件8
        try:
            match = '履约担保\n(.*)[（]发包人名称[）].*'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors[" 附件：履约担保：发包人名称_1"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    if factor != fabaoren:
                        factors_error[" 附件：履约担保：发包人名称_1"] = " 附件：履约担保:与合同发包人不一致"
                        addRemarkInDoc(word, document, '（发包人名称，以下简称“发包人”）与', f" 附件：履约担保:与合同发包人不一致")
                    else:
                        factors_ok.append(" 附件：履约担保：发包人名称_1")
                except:
                    factors_error[" 附件：履约担保：发包人名称_1"] = " 附件：履约担保:合同发包人提取错误"
                    addRemarkInDoc(word, document, '（发包人名称，以下简称“发包人”）与', f" 附件：履约担保:合同发包人提取错误")
            else:
                factors_error[" 附件：履约担保：发包人名称_1"] = " 附件：履约担保: 附件：履约担保：发包人名称_1未填写完整"
                addRemarkInDoc(word, document, '（发包人名称，以下简称“发包人”）与', f" 附件：履约担保: 附件：履约担保：发包人名称_1未填写完整")
        except:
            missObject += " 附件：履约担保:要素“ 附件：履约担保：发包人名称_1”缺失\n"

        try:
            match = '鉴于(.*)[（]发包人名称，以下简称“发包人”[）].*'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors[" 附件：履约担保：发包人名称0"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    if factor != fabaoren:
                        factors_error[" 附件：履约担保：发包人名称0"] = " 附件：履约担保:与合同发包人不一致"
                        addRemarkInDoc(word, document, '（发包人名称，以下简称“发包人”）与', f" 附件：履约担保:与合同发包人不一致")
                    else:
                        factors_ok.append(" 附件：履约担保：发包人名称0")
                except:
                    factors_error[" 附件：履约担保：发包人名称0"] = " 附件：履约担保:合同发包人提取错误"
                    addRemarkInDoc(word, document, '（发包人名称，以下简称“发包人”）与', f" 附件：履约担保:合同发包人提取错误")
            else:
                factors_error[" 附件：履约担保：发包人名称0"] = " 附件：履约担保: 附件：履约担保：发包人名称0未填写完整"
                addRemarkInDoc(word, document, '（发包人名称，以下简称“发包人”）与', f" 附件：履约担保: 附件：履约担保：发包人名称0未填写完整")
        except:
            missObject += " 附件：履约担保:要素“ 附件：履约担保：发包人名称0”缺失\n"

        try:
            match = '鉴于.*[（]发包人名称，以下简称“发包人”[）]与\n(.*)[（]承包人名称[）][（]以下称“承包人”[）]于.*'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors[" 附件：履约担保：承包人名称0"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    if factor != chengbaoren:
                        factors_error[" 附件：履约担保：承包人名称0"] = " 附件：履约担保:与合同发包人不一致"
                        addRemarkInDoc(word, document, '（承包人名称）（以下称“承包人”）于', f" 附件：履约担保:与合同承包人不一致")
                    else:
                        factors_ok.append(" 附件：履约担保：承包人名称0")
                except:
                    factors_error[" 附件：履约担保：承包人名称0"] = " 附件：履约担保:合同承包人提取错误"
                    addRemarkInDoc(word, document, '（承包人名称）（以下称“承包人”）于', f" 附件：履约担保:合同承包人提取错误")
            else:
                factors_error[" 附件：履约担保：承包人名称0"] = " 附件：履约担保: 附件：履约担保：承包人名称0未填写完整"
                addRemarkInDoc(word, document, '（承包人名称）（以下称“承包人”）于', f" 附件：履约担保: 附件：履约担保：承包人名称0未填写完整")
        except:
            missObject += " 附件：履约担保:要素“ 附件：履约担保：承包人名称0”缺失\n"

        try:
            match = '我方愿意无条件地、不可撤销地就承包人履行与你方签订的合同，向你方提供连带责任担保。\n1. 担保金额人民币(（大写）.*?）).*\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "").replace("\ue5e5", "")
            if factor == '':
                factors_error[" 附件：履约担保：担保金额"] = " 附件：履约担保：担保金额无法提取"
                addRemarkInDoc(word, document, "担保金额人民币（大写）", f" 附件：履约担保：担保金额无法提取")
            else:
                factor = factor.replace(" ", "").replace("；", "").replace("。", "").replace("；", "").replace("。", "")
                factors[" 附件：履约担保：担保金额"] = factor
                pattern = '（大写）(.*)元（¥(.*)）.*'
                try:
                    l = re.findall(pattern, factor)
                    small = float(l[0][1])
                    if check_money1(factor) != False and small == danbaojin:
                        factors_ok.append(" 附件：履约担保：担保金额")
                    else:
                        factors_error[" 附件：履约担保：担保金额"] = " 附件：履约担保：担保金额未填写完整或大小写不一致或与3.7担保金额计算公式结果不一致"
                        addRemarkInDoc(word, document, "担保金额人民币（大写）", f" 附件：履约担保：担保金额未填写完整或大小写不一致或与3.7担保金额计算公式结果不一致")
                except:
                    factors_error[" 附件：履约担保：担保金额"] = " 附件：履约担保：担保金额未填写完整或大小写不一致或与3.7担保金额计算公式结果不一致"
                    addRemarkInDoc(word, document, "担保金额人民币（大写）", f" 附件：履约担保：担保金额未填写完整")
        except:
            missObject += "要素“ 附件：履约担保：担保金额”缺失\n"

        try:
            match = '担 保 人：(.*)[（]盖单位章[）]'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors[" 附件：履约担保：承包人名称"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    if factor != chengbaoren:
                        factors_error[" 附件：履约担保：承包人名称"] = " 附件：履约担保:与合同承包人不一致"
                        addRemarkInDoc(word, document, '（发包人名称，以下简称“发包人”）与', f" 附件：履约担保:与合同承包人不一致")
                    else:
                        factors_ok.append(" 附件：履约担保：承包人名称")
                except:
                    factors_error[" 附件：履约担保：承包人名称"] = " 附件：履约担保:合同承包人提取错误"
                    addRemarkInDoc(word, document, '（发包人名称，以下简称“发包人”）与', f" 附件：履约担保:合同承包人提取错误")
            else:
                factors_error[" 附件：履约担保：承包人名称"] = " 附件：履约担保：承包人名称未填写完整"
                addRemarkInDoc(word, document, '（发包人名称，以下简称“发包人”）与', f" 附件：履约担保：承包人名称未填写完整")
        except:
            missObject += "要素“ 附件：履约担保：承包人名称”缺失\n"

        try:
            match = '法定代表人或其委托代理人：(.*)[（]签字[）]'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors[" 附件：履约担保：结尾签字"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    if factor != chengbaoren_fading:
                        factors_error[" 附件：履约担保：结尾签字"] = " 附件：履约担保:结尾担保方与合同约定不一致"
                        addRemarkInDoc(word, document, '（发包人名称，以下简称“发包人”）与', f" 附件：履约担保:结尾担保方与合同约定不一致")
                    else:
                        factors_ok.append(" 附件：履约担保：结尾签字")
                except:
                    factors_error[" 附件：履约担保：结尾签字"] = " 附件：履约担保:合同担保方法定代表人或其委托代理人提取失败"
                    addRemarkInDoc(word, document, '（发包人名称，以下简称“发包人”）与', f" 附件：履约担保:合同担保方法定代表人或其委托代理人提取失败")
            else:
                factors_error[" 附件：履约担保：结尾签字"] = " 附件：履约担保：结尾签字未填写完整"
                addRemarkInDoc(word, document, '（发包人名称，以下简称“发包人”）与', f" 附件：履约担保：结尾签字未填写完整")
        except:
            missObject += "要素“ 附件：履约担保：结尾签字”缺失\n"

        try:
            match = '地    址：.*\n(.*年.*月.*日)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors[" 附件：履约担保：结尾时间"] = factor
            if get_strtime(factor) != False:
                x = get_strtime(factor)
                year = x.split('-')[0]
                month = x.split('-')[1]
                day = x.split('-')[2]
                x = datetime.date(int(year), int(month), int(day))
                try:
                    if x != sign_date:
                        factors_error[" 附件：履约担保：结尾时间"] = " 附件：履约担保:签订时间与合同时间不一致"
                        addRemarkInDoc(word, document, '（发包人名称，以下简称“发包人”）与', f" 附件：履约担保:签订时间与合同时间不一致")
                    else:
                        factors_ok.append(" 附件：履约担保：结尾时间")
                except:
                    factors_error[" 附件：履约担保：结尾时间"] = " 附件：履约担保:合同时间提取错误"
                    addRemarkInDoc(word, document, '（发包人名称，以下简称“发包人”）与', f" 附件：履约担保:合同时间提取错误")
            else:
                factors_error[" 附件：履约担保：结尾时间"] = " 附件：履约担保：结尾时间未填写完整"
                addRemarkInDoc(word, document, '（发包人名称，以下简称“发包人”）与', f" 附件：履约担保：结尾时间未填写完整")
        except:
            missObject += "要素“ 附件：履约担保：结尾时间”缺失\n"

        # 附件9
        try:
            match = '预付款担保\n(.*)[（]发包人名称[）].*'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors[" 附件：预付款担保 ：发包人名称_1"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    if factor != fabaoren:
                        factors_error[" 附件：预付款担保 ：发包人名称_1"] = " 附件：预付款担保:发包人名称_1与合同发包人不一致"
                        addRemarkInDoc(word, document, '（承包人名称）（以下称“承包人”）与', f" 附件：预付款担保:发包人名称_1与合同发包人不一致")
                    else:
                        factors_ok.append(" 附件：预付款担保 ：发包人名称_1")
                except:
                    factors_error[" 附件：预付款担保 ：发包人名称_1"] = " 附件：预付款担保:合同发包人提取错误"
                    addRemarkInDoc(word, document, '（承包人名称）（以下称“承包人”）与', f" 附件：预付款担保:合同发包人提取错误")
            else:
                factors_error[" 附件：预付款担保 ：发包人名称_1"] = " 附件：预付款担保 ：发包人名称_1未填写完整"
                addRemarkInDoc(word, document, '（承包人名称）（以下称“承包人”）与', f" 附件：预付款担保 ：发包人名称_1未填写完整")
        except:
            missObject += " 附件：预付款担保:要素“ 附件：预付款担保 ：发包人名称_1”缺失\n"

        try:
            match = '根据(.*)[（]承包人名称[）][（]以下称“承包人”[）].*'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors[" 附件：预付款担保 ：承包人"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    if factor != fabaoren:
                        factors_error[" 附件：预付款担保 ：承包人"] = " 附件：预付款担保: 附件：预付款担保开头承包人与合同承包人不一致"
                        addRemarkInDoc(word, document, '（承包人名称）（以下称“承包人”）与', f" 附件：预付款担保: 附件：预付款担保开头承包人与合同承包人不一致")
                    else:
                        factors_ok.append(" 附件：预付款担保 ：承包人")
                except:
                    factors_error[" 附件：预付款担保 ：承包人"] = " 附件：预付款担保:合同承包人提取错误"
                    addRemarkInDoc(word, document, '（承包人名称）（以下称“承包人”）与', f" 附件：预付款担保:合同承包人提取错误")
            else:
                factors_error[" 附件：预付款担保 ：承包人"] = " 附件：预付款担保 ：承包人未填写完整"
                addRemarkInDoc(word, document, '（承包人名称）（以下称“承包人”）与', f" 附件：预付款担保 ：承包人未填写完整")
        except:
            missObject += "要素“ 附件：预付款担保 ：承包人”缺失\n"

        try:
            match = '根据.*[（]承包人名称[）][（]以下称“承包人”[）].*与\n(.*)[（]发包人名称[）][（]以下简称“发包人”[）]\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors[" 附件：预付款担保 ：发包人名称_2"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    if factor != chengbaoren:
                        factors_error[" 附件：预付款担保 ：发包人名称_2"] = " 附件：预付款担保:发包人名称_2与合同承包人不一致"
                        addRemarkInDoc(word, document, '（承包人名称）（以下称“承包人”）与', f" 附件：预付款担保:发包人名称_2与合同承包人不一致")
                    else:
                        factors_ok.append(" 附件：预付款担保 ：发包人名称_2")
                except:
                    factors_error[" 附件：预付款担保 ：发包人名称_2"] = " 附件：预付款担保:合同承包人提取错误"
                    addRemarkInDoc(word, document, '（承包人名称）（以下称“承包人”）与', f"合同承包人提取错误")
            else:
                factors_error[" 附件：预付款担保 ：发包人名称_2"] = " 附件：预付款担保 ：发包人名称_2未填写完整"
                addRemarkInDoc(word, document, '（承包人名称）（以下称“承包人”）与', f" 附件：预付款担保 ：发包人名称_2未填写完整")
        except:
            missObject += "要素“ 附件：预付款担保 ：发包人名称_2”缺失\n"

        try:
            match = '与\n.*[（]发包人名称[）][（]以下简称“发包人”[）]\n于(.*)签订.|\n*附件10.*'
            factor = re.findall(match, text)[1].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors[" 附件：预付款担保 ：签订时间"] = factor
            if get_strtime(factor) != False:
                x = get_strtime(factor)
                year = x.split('-')[0]
                month = x.split('-')[1]
                day = x.split('-')[2]
                x = datetime.date(int(year), int(month), int(day))
                try:
                    if x != sign_date:
                        factors_error[" 附件：预付款担保 ：签订时间"] = " 附件：预付款担保:签订时间与合同时间不一致"
                        addRemarkInDoc(word, document, '（承包人名称）（以下称“承包人”）与', f" 附件：预付款担保:签订时间与合同时间不一致")
                    else:
                        factors_ok.append(" 附件：预付款担保 ：签订时间")
                except:
                    factors_error[" 附件：预付款担保 ：签订时间"] = " 附件：预付款担保:合同时间提取错误"
                    addRemarkInDoc(word, document, '（承包人名称）（以下称“承包人”）与', f" 附件：预付款担保:合同时间提取错误")
            else:
                factors_error[" 附件：预付款担保 ：签订时间"] = " 附件：预付款担保 ：签订时间未填写完整"
                addRemarkInDoc(word, document, '（承包人名称）（以下称“承包人”）与', f" 附件：预付款担保 ：签订时间未填写完整")
        except:
            missObject += "要素“ 附件：预付款担保 ：签订时间”缺失\n"

        try:
            match = '承包人按约定的金额向你方提交一份预付款担保，即有权得到你方支付相等金额的预付款。我方愿意就你方提供给承包人的预付款为承包人提供连带责任担保。\n1. 担保金额人民币(（大写）.*?）).*\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("；", "").replace("。", "").replace("\ue5e5", "")
            if factor == '':
                factors_error[" 附件：预付款担保：担保金额"] = " 附件：预付款担保：担保金额无法提取"
                addRemarkInDoc(word, document, "承包人按约定的金额向你方提交一份预付款担保，即有权得到你方支付相等金额的预付款。我方愿意就你方提供给承包人的预付款为承包人提供连带责任担保。",
                               f" 附件：预付款担保：担保金额无法提取")
            else:
                factor = factor.replace(" ", "").replace("；", "").replace("。", "").replace("；", "").replace("。", "")
                factors[" 附件：预付款担保：担保金额"] = factor
                pattern = '（大写）(.*)元（¥(.*)）.*'
                try:
                    l = re.findall(pattern, factor)
                    small = float(l[0][1])
                    if check_money1(factor) != False and small == danbaojin:
                        factors_ok.append(" 附件：预付款担保：担保金额")
                    else:
                        factors_error[" 附件：预付款担保：担保金额"] = " 附件：预付款担保：担保金额未填写完整或大小写不一致或与3.7担保金额计算公式结果不一致"
                        addRemarkInDoc(word, document,
                                       "承包人按约定的金额向你方提交一份预付款担保，即有权得到你方支付相等金额的预付款。我方愿意就你方提供给承包人的预付款为承包人提供连带责任担保。",
                                       f" 附件：预付款担保：担保金额未填写完整或大小写不一致或与3.7担保金额计算公式结果不一致")
                except:
                    factors_error[" 附件：预付款担保：担保金额"] = " 附件：预付款担保：担保金额未填写完整或大小写不一致或与3.7担保金额计算公式结果不一致"
                    addRemarkInDoc(word, document,
                                   "承包人按约定的金额向你方提交一份预付款担保，即有权得到你方支付相等金额的预付款。我方愿意就你方提供给承包人的预付款为承包人提供连带责任担保。",
                                   f" 附件：预付款担保：担保金额未填写完整")
        except:
            missObject += "要素“ 附件：预付款担保：担保金额”缺失\n"

        try:
            match = '担保人：(.*)[（]盖单位章[）]\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors[" 附件：预付款担保 ：承包人名称"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    if factor != chengbaoren:
                        factors_error[" 附件：预付款担保 ：承包人名称"] = " 附件：预付款担保:结尾担保人与合同承包人不一致"
                        addRemarkInDoc(word, document, '（承包人名称）（以下称“承包人”）与', f" 附件：预付款担保:结尾担保人与合同承包人不一致")
                    else:
                        factors_ok.append(" 附件：预付款担保 ：承包人名称")
                except:
                    factors_error[" 附件：预付款担保 ：承包人名称"] = " 附件：预付款担保:合同承包人提取错误"
                    addRemarkInDoc(word, document, '（承包人名称）（以下称“承包人”）与', f" 附件：预付款担保:合同承包人提取错误")
            else:
                factors_error[" 附件：预付款担保 ：承包人名称"] = " 附件：预付款担保 ：结尾担保人未填写完整"
                addRemarkInDoc(word, document, '（承包人名称）（以下称“承包人”）与', f" 附件：预付款担保 ：结尾担保人未填写完整")
        except:
            missObject += "要素“ 附件：预付款担保 ：承包人名称”缺失\n"

        try:
            match = '法定代表人或其委托代理人：(.*)[（]签字[）]\n'
            factor = re.findall(match, text)[1].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors[" 附件：预付款担保 ：结尾签字"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    if factor != chengbaoren_fading:
                        factors_error[" 附件：预付款担保 ：结尾签字"] = " 附件：预付款担保:结尾担保方法定代表人或委托代理人与合同约定不一致"
                        addRemarkInDoc(word, document, '（承包人名称）（以下称“承包人”）与', f" 附件：预付款担保:结尾担保方法定代表人或委托代理人与合同约定不一致")
                    else:
                        factors_ok.append(" 附件：预付款担保 ：结尾签字")
                except:
                    factors_error[" 附件：预付款担保 ：结尾签字"] = " 附件：预付款担保:合同担保方法定代表人或其委托代理人提取失败"
                    addRemarkInDoc(word, document, '（承包人名称）（以下称“承包人”）与', f" 附件：预付款担保:合同担保方法定代表人或其委托代理人提取失败")
            else:
                factors_error[" 附件：预付款担保 ：结尾签字"] = " 附件：预付款担保 ：结尾签字未填写完整"
                addRemarkInDoc(word, document, '（承包人名称）（以下称“承包人”）与', f" 附件：预付款担保 ：结尾签字未填写完整")
        except:
            missObject += "要素“ 附件：预付款担保 ：结尾签字”缺失\n"

        try:
            match = '地    址：.*\n(.*年.*月.*日)\n'
            factor = re.findall(match, text)[1].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors[" 附件：预付款担保 ：结尾时间"] = factor
            if get_strtime(factor) != False:
                x = get_strtime(factor)
                year = x.split('-')[0]
                month = x.split('-')[1]
                day = x.split('-')[2]
                x = datetime.date(int(year), int(month), int(day))
                try:
                    if x != sign_date:
                        factors_error[" 附件：预付款担保 ：结尾时间"] = " 附件：预付款担保:结尾签订时间与合同时间不一致"
                        addRemarkInDoc(word, document, '（承包人名称）（以下称“承包人”）与', f" 附件：预付款担保:结尾签订时间与合同时间不一致")
                    else:
                        factors_ok.append(" 附件：预付款担保 ：结尾时间")
                except:
                    factors_error[" 附件：预付款担保 ：结尾时间"] = " 附件：预付款担保:合同时间提取错误"
                    addRemarkInDoc(word, document, '（承包人名称）（以下称“承包人”）与', f" 附件：预付款担保:合同时间提取错误")
            else:
                factors_error[" 附件：预付款担保 ：结尾时间"] = " 附件：预付款担保 ：结尾时间未填写完整"
                addRemarkInDoc(word, document, '（承包人名称）（以下称“承包人”）与', f" 附件：预付款担保 ：结尾时间未填写完整")
        except:
            missObject += "要素“ 附件：预付款担保 ：结尾时间”缺失\n"

    # 附件10-14
    if 1 == True:
        # 附件10
        factorx = '（工程名称）《建设工程施工合同》（以下称“主合同”）'

        factors_to_inform[" 附件：支付担保"] = ' 附件：支付担保:若合同协议书或专用条款并未约定发包人的支付担保，则本附件不适用'
        addRemarkInDoc(word, document, factorx, f" 附件：支付担保:若合同协议书或专用条款并未约定发包人的支付担保，则本附件不适用")

        try:
            match = '支付担保\n(.*)[（]承包人[）]：\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors[" 附件：支付担保_承包人"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    if factor != chengbaoren:
                        factors_error[" 附件：支付担保_承包人"] = " 附件：支付担保_承包人与合同承包人不一致"
                        addRemarkInDoc(word, document, factorx, f" 附件：支付担保_承包人与合同承包人不一致")
                    else:
                        factors_ok.append(" 附件：支付担保_承包人")
                except:
                    factors_error[" 附件：支付担保_承包人"] = " 附件：支付担保:合同承包人提取错误"
                    addRemarkInDoc(word, document, factorx, f" 附件：支付担保:合同承包人提取错误")
            else:
                factors_error[" 附件：支付担保_承包人"] = " 附件：支付担保_承包人未填写完整"
                addRemarkInDoc(word, document, factorx, f" 附件：支付担保_承包人未填写完整")
        except:
            missObject += " 附件：支付担保:要素“ 附件：支付担保_承包人”缺失\n"

        try:
            match = '鉴于你方作为承包人已经与(.*)[（]发包人名称[）].*'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors[" 附件：支付担保_发包人"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    if factor != fabaoren:
                        factors_error[" 附件：支付担保_发包人"] = " 附件：支付担保_发包人与合同发包人不一致"
                        addRemarkInDoc(word, document, factorx, f" 附件：支付担保_发包人与合同发包人不一致")
                    else:
                        factors_ok.append(" 附件：支付担保_发包人")
                except:
                    factors_error[" 附件：支付担保_发包人"] = " 附件：支付担保:合同发包人提取错误"
                    addRemarkInDoc(word, document, factorx, f" 附件：支付担保:合同发包人提取错误")
            else:
                factors_error[" 附件：支付担保_发包人"] = " 附件：支付担保_发包人未填写完整"
                addRemarkInDoc(word, document, factorx, f" 附件：支付担保_发包人未填写完整")
        except:
            missObject += "要素“ 附件：支付担保_发包人”缺失\n"

        try:
            match = '鉴于你方作为承包人已经与.*[（]发包人名称[）][（]以下称“发包人”[）]于(.*)签订.*'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors[" 附件：支付担保 ：签订时间"] = factor
            if get_strtime(factor) != False:
                x = get_strtime(factor)
                year = x.split('-')[0]
                month = x.split('-')[1]
                day = x.split('-')[2]
                x = datetime.date(int(year), int(month), int(day))
                try:
                    if x != sign_date:
                        factors_error[" 附件：支付担保 ：签订时间"] = " 附件：支付担保:签订时间与合同时间不一致"
                        addRemarkInDoc(word, document, factorx, f" 附件：支付担保:签订时间与合同时间不一致")
                    else:
                        factors_ok.append(" 附件：支付担保 ：签订时间")
                except:
                    factors_error[" 附件：支付担保 ：签订时间"] = " 附件：支付担保:合同时间提取错误"
                    addRemarkInDoc(word, document, factorx, f" 附件：支付担保:合同时间提取错误")
            else:
                factors_error[" 附件：支付担保 ：签订时间"] = " 附件：支付担保 ：签订时间未填写完整"
                addRemarkInDoc(word, document, factorx, f" 附件：支付担保 ：签订时间未填写完整")
        except:
            missObject += "要素“ 附件：支付担保 ：签订时间”缺失\n"

        factors_to_inform[" 附件：支付担保_一_3"] = '应与双方就发包人支付担保的约定一致。'
        addRemarkInDoc(word, document, '3. 我方保证的金额是主合同约定的工程款的', f"应与双方就发包人支付担保的约定一致。")

        factors_to_inform[" 附件：支付担保_二_2"] = '应与双方就发包人支付担保的约定一致。'
        addRemarkInDoc(word, document, '2. 我方保证的期间为：自本合同生效之日起至主合同约定的工程款支付完毕之日后', f"应与双方就发包人支付担保的约定一致。")

        factors_to_inform[" 附件：支付担保_四_3"] = '应与双方就发包人支付担保的约定一致。'
        addRemarkInDoc(word, document, '四、代偿的安排', f" 附件：支付担保_四_3：应与双方就发包人支付担保的约定一致。")

        factors_to_inform[" 附件：支付担保_五_5"] = '应与双方就发包人支付担保的约定一致。'
        addRemarkInDoc(word, document, '5. 我方解除保证责任后，你方应自我方保证责任解除之日起', f"应与双方就发包人支付担保的约定一致。")

        try:
            match = '八、保函的生效.|\n*担保人：(.*)[（]盖章[）]\n.|\n*附件11.*'
            factor = re.findall(match, text)[2].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors[" 附件：支付担保 ：担保人名称"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    if factor != fabaoren:
                        factors_error[" 附件：支付担保 ：担保人名称"] = " 附件：支付担保:与合同发包人不一致"
                        addRemarkInDoc(word, document, '八、保函的生效', f" 附件：支付担保:与合同发包人不一致")
                    else:
                        factors_ok.append(" 附件：支付担保 ：担保人名称")
                except:
                    factors_error[" 附件：支付担保 ：担保人名称"] = " 附件：支付担保:合同发包人提取错误"
                    addRemarkInDoc(word, document, '八、保函的生效', f" 附件：支付担保:合同发包人提取错误")
            else:
                factors_error[" 附件：支付担保 ：担保人名称"] = " 附件：支付担保 ：担保人名称未填写完整"
                addRemarkInDoc(word, document, '八、保函的生效', f" 附件：支付担保 ：担保人名称未填写完整")
        except:
            missObject += "要素“ 附件：支付担保 ：担保人名称”缺失\n"

        try:
            match = '八、保函的生效.|\n*法定代表人或委托代理人：(.*)[（]签字[）]\n.|\n*附件11.*'
            factor = re.findall(match, text)[2].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors[" 附件：支付担保 ：结尾签字"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    if factor != fabaoren_fading:
                        factors_error[" 附件：支付担保 ：结尾签字"] = " 附件：支付担保:结尾担保方法定代表人或委托代理人与合同约定是否一致"
                        addRemarkInDoc(word, document, '八、保函的生效', f" 附件：支付担保:结尾担保方法定代表人或委托代理人与合同约定是否一致")
                    else:
                        factors_ok.append(" 附件：支付担保 ：结尾签字")
                except:
                    factors_error[" 附件：支付担保 ：结尾签字"] = " 附件：支付担保:合同担保方法定代表人或其委托代理人提取失败"
                    addRemarkInDoc(word, document, '八、保函的生效', f" 附件：支付担保:合同担保方法定代表人或其委托代理人提取失败")
            else:
                factors_error[" 附件：支付担保 ：结尾签字"] = " 附件：支付担保 ：结尾签字未填写完整"
                addRemarkInDoc(word, document, '八、保函的生效', f" 附件：支付担保 ：结尾签字未填写完整")
        except:
            missObject += "要素“ 附件：支付担保 ：结尾签字”缺失\n"

        try:
            match = '地    址：.*\n(.*年.*月.*日)\n'
            factor = re.findall(match, text)[1].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors[" 附件：支付担保 ：结尾时间"] = factor
            if get_strtime(factor) != False:
                x = get_strtime(factor)
                year = x.split('-')[0]
                month = x.split('-')[1]
                day = x.split('-')[2]
                x = datetime.date(int(year), int(month), int(day))
                try:
                    if x != sign_date:
                        factors_error[" 附件：支付担保 ：结尾时间"] = " 附件：支付担保:签订时间与合同时间不一致"
                        addRemarkInDoc(word, document, '八、保函的生效', f" 附件：支付担保:签订时间与合同时间不一致")
                    else:
                        factors_ok.append(" 附件：支付担保 ：结尾时间")
                except:
                    factors_error[" 附件：支付担保 ：结尾时间"] = " 附件：支付担保:合同时间提取错误"
                    addRemarkInDoc(word, document, '八、保函的生效', f" 附件：支付担保:合同时间提取错误")
            else:
                factors_error[" 附件：支付担保 ：结尾时间"] = " 附件：支付担保 ：结尾时间未填写完整"
                addRemarkInDoc(word, document, '八、保函的生效', f" 附件：支付担保 ：结尾时间未填写完整")
        except:
            missObject += "要素“ 附件：支付担保 ：结尾时间”缺失\n"

        # 附件12
        try:
            match = '发包人[（]全称[）]：(.*)[（]以下简称甲方[）]\n承包人.*\n为了加强本工程建设项目的安全、文明施工管理，创建文明施工工地，防止人员伤亡、火灾、治安、及重大经济损失事故的发生，经双方协商签订本协议书。\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors[" 附件：安全、文明施工责任书_发包人"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    if factor != fabaoren:
                        factors_error[" 附件：安全、文明施工责任书_发包人"] = " 附件：安全、文明施工责任书 发包人与合同不一致"
                        addRemarkInDoc(word, document,
                                       "为了加强本工程建设项目的安全、文明施工管理，创建文明施工工地，防止人员伤亡、火灾、治安、及重大经济损失事故的发生，经双方协商签订本协议书。",
                                       f" 附件：安全、文明施工责任书 发包人与合同不一致")
                    else:
                        factors_ok.append(" 附件：安全、文明施工责任书_发包人")
                except:
                    factors_error[" 附件：安全、文明施工责任书_发包人"] = " 附件：安全、文明施工责任书 发包人提取失败"
                    addRemarkInDoc(word, document,
                                   " 附件：安全、文明施工责任书 为了加强本工程建设项目的安全、文明施工管理，创建文明施工工地，防止人员伤亡、火灾、治安、及重大经济损失事故的发生，经双方协商签订本协议书。",
                                   f"发包人提取失败")
            else:
                factors_error[" 附件：安全、文明施工责任书_发包人"] = " 附件：安全、文明施工责任书_发包人未填写完整"
                addRemarkInDoc(word, document, "为了加强本工程建设项目的安全、文明施工管理，创建文明施工工地，防止人员伤亡、火灾、治安、及重大经济损失事故的发生，经双方协商签订本协议书。",
                               f" 附件：安全、文明施工责任书 合同发包人提取失败")
        except:
            missObject += "要素“ 附件：安全、文明施工责任书_发包人”缺失\n"

        try:
            match = '承包人[（]全称[）]：(.*)[（]以下简称乙方[）]\n为了加强本工程建设项目的安全、文明施工管理，创建文明施工工地，防止人员伤亡、火灾、治安、及重大经济损失事故的发生，经双方协商签订本协议书。'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors[" 附件：安全、文明施工责任书_承包人"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    if factor != chengbaoren:
                        factors_error[" 附件：安全、文明施工责任书_承包人"] = " 附件：安全、文明施工责任书 承包人与合同不一致"
                        addRemarkInDoc(word, document,
                                       "为了加强本工程建设项目的安全、文明施工管理，创建文明施工工地，防止人员伤亡、火灾、治安、及重大经济损失事故的发生，经双方协商签订本协议书。",
                                       f" 附件：安全、文明施工责任书 承包人与合同不一致")
                    else:
                        factors_ok.append(" 附件：安全、文明施工责任书_承包人")
                except:
                    factors_error[" 附件：安全、文明施工责任书_承包人"] = " 附件：安全、文明施工责任书 承包人提取失败"
                    addRemarkInDoc(word, document,
                                   " 附件：安全、文明施工责任书 为了加强本工程建设项目的安全、文明施工管理，创建文明施工工地，防止人员伤亡、火灾、治安、及重大经济损失事故的发生，经双方协商签订本协议书。",
                                   f"承包人提取失败")
            else:
                factors_error[" 附件：安全、文明施工责任书_承包人"] = " 附件：安全、文明施工责任书_承包人未填写完整"
                addRemarkInDoc(word, document, "为了加强本工程建设项目的安全、文明施工管理，创建文明施工工地，防止人员伤亡、火灾、治安、及重大经济损失事故的发生，经双方协商签订本协议书。",
                               f" 附件：安全、文明施工责任书 合同承包人提取失败")
        except:
            missObject += "要素“ 附件：安全、文明施工责任书_承包人”缺失\n"

        try:
            match = '第七条：本协议自双方签字盖章后生效。\n甲  方[（]盖公章[）]：(.*)乙  方.*\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors[" 附件：安全、文明施工责任书_结尾签字_甲方"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    if factor != fabaoren:
                        factors_error[" 附件：安全、文明施工责任书_结尾签字_甲方"] = " 附件：安全、文明施工责任书_结尾签字_甲方与合同发包人不一致"
                        addRemarkInDoc(word, document, "第七条：本协议自双方签字盖章后生效。", f" 附件：安全、文明施工责任书_结尾签字_甲方与合同发包人不一致")
                    else:
                        factors_ok.append(" 附件：安全、文明施工责任书_结尾签字_甲方")
                except:
                    factors_error[" 附件：安全、文明施工责任书_结尾签字_甲方"] = " 附件：安全、文明施工责任书 合同发包人提取失败"
                    addRemarkInDoc(word, document, "第七条：本协议自双方签字盖章后生效。", f" 附件：安全、文明施工责任书 合同发包人提取失败")
            else:
                factors_error[" 附件：安全、文明施工责任书_结尾签字_甲方"] = " 附件：安全、文明施工责任书_结尾签字_甲方未填写完整"
                addRemarkInDoc(word, document, "第七条：本协议自双方签字盖章后生效。", f" 附件：安全、文明施工责任书_结尾签字_甲方未填写完整")
        except:
            missObject += "要素“ 附件：安全、文明施工责任书_结尾签字_甲方”缺失\n"

        try:
            match = '第七条：本协议自双方签字盖章后生效。\n甲  方[（]盖公章[）]：.*乙  方[（]盖公章[）]：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors[" 附件：安全、文明施工责任书_结尾签字_乙方"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    if factor != chengbaoren:
                        factors_error[" 附件：安全、文明施工责任书_结尾签字_乙方"] = " 附件：安全、文明施工责任书_结尾签字_乙方与合同承包人不一致"
                        addRemarkInDoc(word, document, "第七条：本协议自双方签字盖章后生效。", f" 附件：安全、文明施工责任书_结尾签字_乙方与合同承包人不一致")
                    else:
                        factors_ok.append(" 附件：安全、文明施工责任书_结尾签字_乙方")
                except:
                    factors_error[" 附件：安全、文明施工责任书_结尾签字_乙方"] = " 附件：安全、文明施工责任书 合同承包人提取失败"
                    addRemarkInDoc(word, document, "第七条：本协议自双方签字盖章后生效。", f" 附件：安全、文明施工责任书 合同承包人提取失败")
            else:
                factors_error[" 附件：安全、文明施工责任书_结尾签字_乙方"] = " 附件：安全、文明施工责任书_结尾签字_乙方未填写完整"
                addRemarkInDoc(word, document, "第七条：本协议自双方签字盖章后生效。", f" 附件：安全、文明施工责任书_结尾签字_乙方未填写完整")
        except:
            missObject += "要素“ 附件：安全、文明施工责任书_结尾签字_乙方”缺失\n"

        try:
            match = '第七条：本协议自双方签字盖章后生效。\n.*\n法定代表人：(.*)法定代表人.*'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors[" 附件：安全、文明施工责任书_结尾签字_甲方法定"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    if factor != fabaoren_fading:
                        factors_error[" 附件：安全、文明施工责任书_结尾签字_甲方法定"] = " 附件：安全、文明施工责任书_结尾签字_甲方法定与合同发包人不一致"
                        addRemarkInDoc(word, document, "第七条：本协议自双方签字盖章后生效。", f" 附件：安全、文明施工责任书_结尾签字_甲方法定与合同发包人不一致")
                    else:
                        factors_ok.append(" 附件：安全、文明施工责任书_结尾签字_甲方法定")
                except:
                    factors_error[" 附件：安全、文明施工责任书_结尾签字_甲方法定"] = " 附件：安全、文明施工责任书 合同发包人提取失败"
                    addRemarkInDoc(word, document, "第七条：本协议自双方签字盖章后生效。", f" 附件：安全、文明施工责任书 合同发包人提取失败")
            else:
                factors_error[" 附件：安全、文明施工责任书_结尾签字_甲方法定"] = " 附件：安全、文明施工责任书_结尾签字_甲方法定未填写完整"
                addRemarkInDoc(word, document, "第七条：本协议自双方签字盖章后生效。", f" 附件：安全、文明施工责任书_结尾签字_甲方法定未填写完整")
        except:
            missObject += "要素“ 附件：安全、文明施工责任书_结尾签字_甲方法定”缺失\n"

        try:
            match = '第七条：本协议自双方签字盖章后生效。\n.*\n法定代表人：.*法定代表人：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors[" 附件：安全、文明施工责任书_结尾签字_甲方法定"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    if factor != fabaoren_fading:
                        factors_error[" 附件：安全、文明施工责任书_结尾签字_甲方法定"] = " 附件：安全、文明施工责任书_结尾签字_甲方法定与合同发包人不一致"
                        addRemarkInDoc(word, document, "第七条：本协议自双方签字盖章后生效。", f" 附件：安全、文明施工责任书_结尾签字_甲方法定与合同发包人不一致")
                    else:
                        factors_ok.append(" 附件：安全、文明施工责任书_结尾签字_甲方法定")
                except:
                    factors_error[" 附件：安全、文明施工责任书_结尾签字_甲方法定"] = " 附件：安全、文明施工责任书 合同发包人提取失败"
                    addRemarkInDoc(word, document, "第七条：本协议自双方签字盖章后生效。", f" 附件：安全、文明施工责任书 合同发包人提取失败")
            else:
                factors_error[" 附件：安全、文明施工责任书_结尾签字_甲方法定"] = " 附件：安全、文明施工责任书_结尾签字_甲方法定未填写完整"
                addRemarkInDoc(word, document, "第七条：本协议自双方签字盖章后生效。", f" 附件：安全、文明施工责任书_结尾签字_甲方法定未填写完整")
        except:
            missObject += "要素“ 附件：安全、文明施工责任书_结尾签字_甲方法定”缺失\n"

        try:
            match = '第七条：本协议自双方签字盖章后生效。\n.*\n.*\n委托代理人：(.*)委托.*'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors[" 附件：安全、文明施工责任书_结尾签字_甲方代理"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    if factor != fabaoren_daili:
                        factors_error[" 附件：安全、文明施工责任书_结尾签字_甲方代理"] = " 附件：安全、文明施工责任书_结尾签字_甲方代理与合同代理人不一致"
                        addRemarkInDoc(word, document, "第七条：本协议自双方签字盖章后生效。", f" 附件：安全、文明施工责任书_结尾签字_甲方代理与合同代理人不一致")
                    else:
                        factors_ok.append(" 附件：安全、文明施工责任书_结尾签字_甲方代理")
                except:
                    factors_error[" 附件：安全、文明施工责任书_结尾签字_甲方代理"] = " 附件：安全、文明施工责任书 合同代理人提取失败"
                    addRemarkInDoc(word, document, "第七条：本协议自双方签字盖章后生效。", f" 附件：安全、文明施工责任书 合同代理人提取失败")
            else:
                factors_error[" 附件：安全、文明施工责任书_结尾签字_甲方代理"] = " 附件：安全、文明施工责任书_结尾签字_甲方代理未填写完整"
                addRemarkInDoc(word, document, "第七条：本协议自双方签字盖章后生效。", f" 附件：安全、文明施工责任书_结尾签字_甲方代理未填写完整")
        except:
            missObject += "要素“ 附件：安全、文明施工责任书_结尾签字_甲方代理”缺失\n"

        try:
            match = '第七条：本协议自双方签字盖章后生效。\n.*\n.*\n委托代理人：.*委托代理人：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors[" 附件：安全、文明施工责任书_结尾签字_乙方代理"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    if factor != chengbaoren_daili:
                        factors_error[" 附件：安全、文明施工责任书_结尾签字_乙方代理"] = " 附件：安全、文明施工责任书_结尾签字_乙方代理与合同代理人不一致"
                        addRemarkInDoc(word, document, "第七条：本协议自双方签字盖章后生效。", f" 附件：安全、文明施工责任书_结尾签字_乙方代理与合同代理人不一致")
                    else:
                        factors_ok.append(" 附件：安全、文明施工责任书_结尾签字_乙方代理")
                except:
                    factors_error[" 附件：安全、文明施工责任书_结尾签字_乙方代理"] = " 附件：安全、文明施工责任书 合同代理人提取失败"
                    addRemarkInDoc(word, document, "第七条：本协议自双方签字盖章后生效。", f" 附件：安全、文明施工责任书 合同代理人提取失败")
            else:
                factors_error[" 附件：安全、文明施工责任书_结尾签字_乙方代理"] = " 附件：安全、文明施工责任书_结尾签字_乙方代理未填写完整"
                addRemarkInDoc(word, document, "第七条：本协议自双方签字盖章后生效。", f" 附件：安全、文明施工责任书_结尾签字_乙方代理未填写完整")
        except:
            missObject += "要素“ 附件：安全、文明施工责任书_结尾签字_乙方代理”缺失\n"

        # 附件13
        try:
            match = '工程名称：(.*)\n项目地址：.*\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors[" 附件：廉政合同_工程名称"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    if factor != gongchengmingcheng:
                        factors_error[" 附件：廉政合同_工程名称"] = " 附件：廉政合同_工程名称与合同不同"
                        addRemarkInDoc(word, document, "为加强工程项目建设中的廉政建设，规范工程建设发包、承包双方的各项活动", f" 附件：廉政合同_工程名称与合同不同")
                    else:
                        factors_ok.append(" 附件：廉政合同_工程名称")
                except:
                    factors_error[" 附件：廉政合同_工程名称"] = " 附件：廉政合同_合同工程名称提取错误"
                    addRemarkInDoc(word, document, "为加强工程项目建设中的廉政建设，规范工程建设发包、承包双方的各项活动", f" 附件：廉政合同_合同工程名称提取错误")
            else:
                factors_error[" 附件：廉政合同_工程名称"] = " 附件：廉政合同_工程名称未填写完整"
                addRemarkInDoc(word, document, "为加强工程项目建设中的廉政建设，规范工程建设发包、承包双方的各项活动", f" 附件：廉政合同_工程名称未填写完整")
        except:
            missObject += "要素“ 附件：廉政合同_工程名称”缺失\n"

        try:
            match = '工程名称：.*\n项目地址：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors[" 附件：廉政合同_工程地址"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    if factor != gongchengdidian:
                        factors_error[" 附件：廉政合同_工程地址"] = " 附件：廉政合同_工程地址与合同不同"
                        addRemarkInDoc(word, document, "为加强工程项目建设中的廉政建设，规范工程建设发包、承包双方的各项活动", f" 附件：廉政合同_工程地址与合同不同")
                    else:
                        factors_ok.append(" 附件：廉政合同_工程地址")
                except:
                    factors_error[" 附件：廉政合同_工程地址"] = "合同工程地址提取错误"
                    addRemarkInDoc(word, document, "为加强工程项目建设中的廉政建设，规范工程建设发包、承包双方的各项活动", f" 附件：廉政合同_合同工程地址提取错误")
            else:
                factors_error[" 附件：廉政合同_工程地址"] = " 附件：廉政合同_工程地址未填写完整"
                addRemarkInDoc(word, document, "为加强工程项目建设中的廉政建设，规范工程建设发包、承包双方的各项活动", f" 附件：廉政合同_工程地址未填写完整")
        except:
            missObject += "要素“ 附件：廉政合同_工程地址”缺失\n"

        try:
            match = '工程名称：.*\n项目地址：.*\n发包单位名称[（]以下称甲方[）]：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors[" 附件：廉政合同_发包单位名称"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    if factor != gongchengdidian:
                        factors_error[" 附件：廉政合同_发包单位名称"] = " 附件：廉政合同_发包单位名称与合同不同"
                        addRemarkInDoc(word, document, "为加强工程项目建设中的廉政建设，规范工程建设发包、承包双方的各项活动", f" 附件：廉政合同_发包单位名称与合同不同")
                    else:
                        factors_ok.append(" 附件：廉政合同_发包单位名称")
                except:
                    factors_error[" 附件：廉政合同_发包单位名称"] = " 附件：廉政合同_合同发包单位名称提取错误"
                    addRemarkInDoc(word, document, "为加强工程项目建设中的廉政建设，规范工程建设发包、承包双方的各项活动", f" 附件：廉政合同_合同发包单位名称提取错误")
            else:
                factors_error[" 附件：廉政合同_发包单位名称"] = " 附件：廉政合同_发包单位名称未填写完整"
                addRemarkInDoc(word, document, "为加强工程项目建设中的廉政建设，规范工程建设发包、承包双方的各项活动", f" 附件：廉政合同_发包单位名称未填写完整")
        except:
            missObject += "要素“ 附件：廉政合同_发包单位名称”缺失\n"

        try:
            match = '工程名称：.*\n项目地址：.*\n发包单位名称[（]以下称甲方[）]：.*\n承包单位名称[（]以下称乙方[）]：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors[" 附件：廉政合同_承包单位名称"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    if factor != gongchengdidian:
                        factors_error[" 附件：廉政合同_承包单位名称"] = " 附件：廉政合同_承包单位名称与合同不同"
                        addRemarkInDoc(word, document, "为加强工程项目建设中的廉政建设，规范工程建设发包、承包双方的各项活动", f" 附件：廉政合同_承包单位名称与合同不同")
                    else:
                        factors_ok.append(" 附件：廉政合同_承包单位名称")
                except:
                    factors_error[" 附件：廉政合同_承包单位名称"] = "合同承包单位名称提取错误"
                    addRemarkInDoc(word, document, "为加强工程项目建设中的廉政建设，规范工程建设发包、承包双方的各项活动", f" 附件：廉政合同_合同承包单位名称提取错误")
            else:
                factors_error[" 附件：廉政合同_承包单位名称"] = " 附件：廉政合同_承包单位名称未填写完整"
                addRemarkInDoc(word, document, "为加强工程项目建设中的廉政建设，规范工程建设发包、承包双方的各项活动", f" 附件：廉政合同_承包单位名称未填写完整")
        except:
            missObject += "要素“ 附件：廉政合同_承包单位名称”缺失\n"

        try:
            match = '第八条 本合同一式十份，甲、乙双方各执五份。\n甲方单位：(.*)[（]盖  章[）].*乙方单位.*'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors[" 附件：廉政合同结尾甲方单位"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    if factor != fabaoren:
                        factors_error[" 附件：廉政合同结尾甲方单位"] = " 附件：廉政合同结尾甲方单位与合同发包人不同"
                        addRemarkInDoc(word, document, "第八条 本合同一式十份，甲、乙双方各执五份。", f" 附件：廉政合同结尾甲方单位与合同发包人不同")
                    else:
                        factors_ok.append(" 附件：廉政合同结尾甲方单位")
                except:
                    factors_error[" 附件：廉政合同结尾甲方单位"] = " 附件：廉政合同_合同发包人提取错误"
                    addRemarkInDoc(word, document, "第八条 本合同一式十份，甲、乙双方各执五份。", f" 附件：廉政合同_合同发包人提取错误")
            else:
                factors_error[" 附件：廉政合同结尾甲方单位"] = " 附件：廉政合同结尾甲方单位未填写完整"
                addRemarkInDoc(word, document, "第八条 本合同一式十份，甲、乙双方各执五份。", f" 附件：廉政合同结尾甲方单位未填写完整")
        except:
            missObject += "要素“ 附件：廉政合同结尾甲方单位”缺失\n"

        try:
            match = '第八条 本合同一式十份，甲、乙双方各执五份。\n甲方单位：.*乙方单位：(.*)[（]盖  章[）]\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors[" 附件：廉政合同结尾乙方单位"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    if factor != chengbaoren:
                        factors_error[" 附件：廉政合同结尾乙方单位"] = " 附件：廉政合同结尾乙方单位与合同承包人不同"
                        addRemarkInDoc(word, document, "第八条 本合同一式十份，甲、乙双方各执五份。", f" 附件：廉政合同结尾乙方单位与合同承包人不同")
                    else:
                        factors_ok.append(" 附件：廉政合同结尾乙方单位")
                except:
                    factors_error[" 附件：廉政合同结尾乙方单位"] = " 附件：廉政合同_合同承包人提取错误"
                    addRemarkInDoc(word, document, "第八条 本合同一式十份，甲、乙双方各执五份。", f" 附件：廉政合同_合同承包人提取错误")
            else:
                factors_error[" 附件：廉政合同结尾乙方单位"] = " 附件：廉政合同结尾乙方单位未填写完整"
                addRemarkInDoc(word, document, "第八条 本合同一式十份，甲、乙双方各执五份。", f" 附件：廉政合同结尾乙方单位未填写完整")
        except:
            missObject += "要素“ 附件：廉政合同结尾乙方单位”缺失\n"

        try:
            match = '第八条 本合同一式十份，甲、乙双方各执五份。.|\n*法定代表人：(.*)法定代表人.*\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors[" 附件：廉政合同结尾甲方法定代表人"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    if factor != fabaoren_fading:
                        factors_error[" 附件：廉政合同结尾甲方法定代表人"] = " 附件：廉政合同结尾甲方法定代表人与合同发包人法定代表人不同"
                        addRemarkInDoc(word, document, "第八条 本合同一式十份，甲、乙双方各执五份。", f" 附件：廉政合同结尾甲方法定代表人与合同发包人法定代表人不同")
                    else:
                        factors_ok.append(" 附件：廉政合同结尾甲方法定代表人")
                except:
                    factors_error[" 附件：廉政合同结尾甲方法定代表人"] = " 附件：廉政合同_合同发包人法定代表人提取错误"
                    addRemarkInDoc(word, document, "第八条 本合同一式十份，甲、乙双方各执五份。", f" 附件：廉政合同_合同发包人法定代表人提取错误")
            else:
                factors_error[" 附件：廉政合同结尾甲方法定代表人"] = " 附件：廉政合同结尾甲方法定代表人未填写完整"
                addRemarkInDoc(word, document, "第八条 本合同一式十份，甲、乙双方各执五份。", f" 附件：廉政合同结尾甲方法定代表人未填写完整")
        except:
            missObject += "要素“ 附件：廉政合同结尾甲方法定代表人”缺失\n"

        try:
            match = '第八条 本合同一式十份，甲、乙双方各执五份。.|\n*法定代表人：.*法定代表人：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors[" 附件：廉政合同结尾乙方法定代表人"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    if factor != chengbaoren_fading:
                        factors_error[" 附件：廉政合同结尾乙方法定代表人"] = " 附件：廉政合同结尾乙方法定代表人与合同承包人法定代表人不同"
                        addRemarkInDoc(word, document, "第八条 本合同一式十份，甲、乙双方各执五份。", f" 附件：廉政合同结尾乙方法定代表人与合同承包人法定代表人不同")
                    else:
                        factors_ok.append(" 附件：廉政合同结尾乙方法定代表人")
                except:
                    factors_error[" 附件：廉政合同结尾乙方法定代表人"] = " 附件：廉政合同_合同承包人法定代表人提取错误"
                    addRemarkInDoc(word, document, "第八条 本合同一式十份，甲、乙双方各执五份。", f" 附件：廉政合同_合同承包人法定代表人提取错误")
            else:
                factors_error[" 附件：廉政合同结尾乙方法定代表人"] = " 附件：廉政合同结尾乙方法定代表人未填写完整"
                addRemarkInDoc(word, document, "第八条 本合同一式十份，甲、乙双方各执五份。", f" 附件：廉政合同结尾乙方法定代表人未填写完整")
        except:
            missObject += "要素“ 附件：廉政合同结尾乙方法定代表人”缺失\n"

        try:
            match = '第八条 本合同一式十份，甲、乙双方各执五份。.|\n(.*年.*月.*日).*年.*月.*日\n*附件14.*'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors[" 附件：廉政合同：甲方结尾时间"] = factor
            if get_strtime(factor) != False:
                x = get_strtime(factor)
                year = x.split('-')[0]
                month = x.split('-')[1]
                day = x.split('-')[2]
                x = datetime.date(int(year), int(month), int(day))
                try:
                    if x != sign_date:
                        factors_error[" 附件：廉政合同：甲方结尾时间"] = "签订时间与合同时间不一致"
                        addRemarkInDoc(word, document, '第八条 本合同一式十份，甲、乙双方各执五份。', f"签订时间与合同时间不一致")
                    else:
                        factors_ok.append(" 附件：廉政合同：甲方结尾时间")
                except:
                    factors_error[" 附件：廉政合同：甲方结尾时间"] = " 附件：廉政合同_合同时间提取错误"
                    addRemarkInDoc(word, document, '第八条 本合同一式十份，甲、乙双方各执五份。', f" 附件：廉政合同_合同时间提取错误")
            else:
                factors_error[" 附件：廉政合同：甲方结尾时间"] = " 附件：廉政合同：甲方结尾时间未填写完整"
                addRemarkInDoc(word, document, '第八条 本合同一式十份，甲、乙双方各执五份。', f" 附件：廉政合同：甲方结尾时间未填写完整")
        except:
            missObject += "要素“ 附件：廉政合同：甲方结尾时间”缺失\n"

        try:
            match = '第八条 本合同一式十份，甲、乙双方各执五份。.|\n.*年.*月.*日(.*年.*月.*日)\n*附件14.*'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors[" 附件：廉政合同：乙方结尾时间"] = factor
            if get_strtime(factor) != False:
                x = get_strtime(factor)
                year = x.split('-')[0]
                month = x.split('-')[1]
                day = x.split('-')[2]
                x = datetime.date(int(year), int(month), int(day))
                try:
                    if x != sign_date:
                        factors_error[" 附件：廉政合同：乙方结尾时间"] = " 附件：廉政合同_签订时间与合同时间不一致"
                        addRemarkInDoc(word, document, '第八条 本合同一式十份，甲、乙双方各执五份。', f" 附件：廉政合同_签订时间与合同时间不一致")
                    else:
                        factors_ok.append(" 附件：廉政合同：乙方结尾时间")
                except:
                    factors_error[" 附件：廉政合同：乙方结尾时间"] = " 附件：廉政合同_合同时间提取错误"
                    addRemarkInDoc(word, document, '第八条 本合同一式十份，甲、乙双方各执五份。', f" 附件：廉政合同_合同时间提取错误")
            else:
                factors_error[" 附件：廉政合同：乙方结尾时间"] = " 附件：廉政合同：乙方结尾时间未填写完整"
                addRemarkInDoc(word, document, '第八条 本合同一式十份，甲、乙双方各执五份。', f" 附件：廉政合同：乙方结尾时间未填写完整")
        except:
            missObject += "要素“ 附件：廉政合同：乙方结尾时间”缺失\n"

        # 附件14
        try:
            match = '为在施工合同的实施过程中创造安全、文明、和谐的施工环境，切实搞好本项目的安全管理工作，本项目业主(.*)[（]以下简称“甲方”[）].*'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors[" 附件：安全生产合同甲方"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    if factor != fabaoren:
                        factors_error[" 附件：安全生产合同甲方"] = " 附件：安全生产合同_甲方与发包人不一致"
                        addRemarkInDoc(word, document, "为在施工合同的实施过程中创造安全、文明、和谐的施工环境，切实搞好本项目的安全管理工作，", f" 附件：安全生产合同_甲方与发包人不一致")
                    else:
                        factors_ok.append(" 附件：安全生产合同甲方")
                except:
                    factors_error[" 附件：安全生产合同甲方"] = " 附件：安全生产合同_合同发包人提取错误"
                    addRemarkInDoc(word, document, "为在施工合同的实施过程中创造安全、文明、和谐的施工环境，切实搞好本项目的安全管理工作，", f" 附件：安全生产合同_合同发包人提取错误")
            else:
                factors_error[" 附件：安全生产合同甲方"] = " 附件：安全生产合同甲方未填写完整"
                addRemarkInDoc(word, document, "为在施工合同的实施过程中创造安全、文明、和谐的施工环境，切实搞好本项目的安全管理工作，", f" 附件：安全生产合同甲方未填写完整")
        except:
            missObject += "要素“ 附件：安全生产合同甲方”缺失\n"

        try:
            match = '为在施工合同的实施过程中创造安全、文明、和谐的施工环境，切实搞好本项目的安全管理工作，本项目业主.*[（]以下简称“甲方”[）]与承包人(.*)[（]以下简称“乙方”[）].*'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors[" 附件：安全生产合同乙方"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    if factor != chengbaoren:
                        factors_error[" 附件：安全生产合同乙方"] = "乙方与承包人不一致"
                        addRemarkInDoc(word, document, "为在施工合同的实施过程中创造安全、文明、和谐的施工环境，切实搞好本项目的安全管理工作，", f" 附件：安全生产合同_乙方与承包人不一致")
                    else:
                        factors_ok.append(" 附件：安全生产合同乙方")
                except:
                    factors_error[" 附件：安全生产合同乙方"] = " 附件：安全生产合同_合同承包人提取错误"
                    addRemarkInDoc(word, document, "为在施工合同的实施过程中创造安全、文明、和谐的施工环境，切实搞好本项目的安全管理工作，", f" 附件：安全生产合同_合同承包人提取错误")
            else:
                factors_error[" 附件：安全生产合同乙方"] = " 附件：安全生产合同乙方未填写完整"
                addRemarkInDoc(word, document, "为在施工合同的实施过程中创造安全、文明、和谐的施工环境，切实搞好本项目的安全管理工作，", f" 附件：安全生产合同乙方未填写完整")
        except:
            missObject += "要素“ 附件：安全生产合同乙方”缺失\n"

        try:
            match = '本合同正本一式二份，副本六分，合同双方各执正本一份，副本三份。双方法定代表或其授权的代理人签署与加盖章后生效，全部工程竣工验收后失效。.|\n*发包人：(.*)[（]盖章[）].*承包人.*'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors[" 附件：安全生产合同结尾发包人"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    if factor != fabaoren:
                        factors_error[" 附件：安全生产合同结尾发包人"] = " 附件：安全生产合同结尾发包人与合同不同"
                        addRemarkInDoc(word, document,
                                       "本合同正本一式二份，副本六分，合同双方各执正本一份，副本三份。双方法定代表或其授权的代理人签署与加盖章后生效，全部工程竣工验收后失效。",
                                       f" 附件：安全生产合同结尾发包人与合同不同")
                    else:
                        factors_ok.append(" 附件：安全生产合同结尾发包人")
                except:
                    factors_error[" 附件：安全生产合同结尾发包人"] = " 附件：安全生产合同_合同发包人提取错误"
                    addRemarkInDoc(word, document,
                                   "本合同正本一式二份，副本六分，合同双方各执正本一份，副本三份。双方法定代表或其授权的代理人签署与加盖章后生效，全部工程竣工验收后失效。",
                                   f" 附件：安全生产合同_合同发包人提取错误")
            else:
                factors_error[" 附件：安全生产合同结尾发包人"] = " 附件：安全生产合同结尾发包人未填写完整"
                addRemarkInDoc(word, document, "本合同正本一式二份，副本六分，合同双方各执正本一份，副本三份。双方法定代表或其授权的代理人签署与加盖章后生效，全部工程竣工验收后失效。",
                               f" 附件：安全生产合同结尾发包人未填写完整")
        except:
            missObject += "要素“ 附件：安全生产合同结尾发包人”缺失\n"

        try:
            match = '本合同正本一式二份，副本六分，合同双方各执正本一份，副本三份。双方法定代表或其授权的代理人签署与加盖章后生效，全部工程竣工验收后失效。.|\n*发包人：.*承包人：(.*)[（]盖章[）].*\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors[" 附件：安全生产合同结尾承包人"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    if factor != chengbaoren:
                        factors_error[" 附件：安全生产合同结尾承包人"] = " 附件：安全生产合同结尾承包人与合同不同"
                        addRemarkInDoc(word, document,
                                       "本合同正本一式二份，副本六分，合同双方各执正本一份，副本三份。双方法定代表或其授权的代理人签署与加盖章后生效，全部工程竣工验收后失效。",
                                       f" 附件：安全生产合同结尾承包人与合同不同")
                    else:
                        factors_ok.append(" 附件：安全生产合同结尾承包人")
                except:
                    factors_error[" 附件：安全生产合同结尾承包人"] = "合同承包人提取错误"
                    addRemarkInDoc(word, document,
                                   "本合同正本一式二份，副本六分，合同双方各执正本一份，副本三份。双方法定代表或其授权的代理人签署与加盖章后生效，全部工程竣工验收后失效。",
                                   f" 附件：安全生产合同_合同承包人提取错误")
            else:
                factors_error[" 附件：安全生产合同结尾承包人"] = " 附件：安全生产合同结尾承包人未填写完整"
                addRemarkInDoc(word, document, "本合同正本一式二份，副本六分，合同双方各执正本一份，副本三份。双方法定代表或其授权的代理人签署与加盖章后生效，全部工程竣工验收后失效。",
                               f" 附件：安全生产合同结尾承包人未填写完整")
        except:
            missObject += "要素“ 附件：安全生产合同结尾承包人”缺失\n"

        try:
            match = '本合同正本一式二份，副本六分，合同双方各执正本一份，副本三份。双方法定代表或其授权的代理人签署与加盖章后生效，全部工程竣工验收后失效。.|\n其授权的代理人：(.*)其授权的代理人.*'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors[" 附件：安全生产合同结尾发包人代理人\代表人"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    if factor != fabaoren_fading and factor != fabaoren_daili:
                        factors_error[" 附件：安全生产合同结尾发包人代理人\代表人"] = " 附件：安全生产合同结尾发包人代理人\代表人与合同不同"
                        addRemarkInDoc(word, document,
                                       "本合同正本一式二份，副本六分，合同双方各执正本一份，副本三份。双方法定代表或其授权的代理人签署与加盖章后生效，全部工程竣工验收后失效。",
                                       f" 附件：安全生产合同结尾发包人代理人\代表人与合同不同")
                    else:
                        factors_ok.append(" 附件：安全生产合同结尾发包人代理人\代表人")
                except:
                    factors_error[" 附件：安全生产合同结尾发包人代理人\代表人"] = " 附件：安全生产合同_合同发包人代理人\代表人提取错误"
                    addRemarkInDoc(word, document,
                                   "本合同正本一式二份，副本六分，合同双方各执正本一份，副本三份。双方法定代表或其授权的代理人签署与加盖章后生效，全部工程竣工验收后失效。",
                                   f" 附件：安全生产合同_合同发包人代理人\代表人提取错误")
            else:
                factors_error[" 附件：安全生产合同结尾发包人代理人\代表人"] = " 附件：安全生产合同结尾发包人代理人\代表人未填写完整"
                addRemarkInDoc(word, document, "本合同正本一式二份，副本六分，合同双方各执正本一份，副本三份。双方法定代表或其授权的代理人签署与加盖章后生效，全部工程竣工验收后失效。",
                               f" 附件：安全生产合同结尾发包人代理人\代表人未填写完整")
        except:
            missObject += "要素“ 附件：安全生产合同结尾发包人代理人\代表人”缺失\n"

        try:
            match = '本合同正本一式二份，副本六分，合同双方各执正本一份，副本三份。双方法定代表或其授权的代理人签署与加盖章后生效，全部工程竣工验收后失效。.|\n其授权的代理人：.*其授权的代理人：(.*)\n'
            factor = re.findall(match, text)[0].replace(" ", "").replace("。", "").replace("；", "").replace("\t", "")
            factors[" 附件：安全生产合同结尾承包人代理人\代表人"] = factor
            if factor != "" and factor != "；" and factor != "。":
                try:
                    if factor != chengbaoren_fading and factor != chengbaoren_daili:
                        factors_error[" 附件：安全生产合同结尾承包人代理人\代表人"] = " 附件：安全生产合同结尾承包人代理人\代表人与合同不同"
                        addRemarkInDoc(word, document,
                                       "本合同正本一式二份，副本六分，合同双方各执正本一份，副本三份。双方法定代表或其授权的代理人签署与加盖章后生效，全部工程竣工验收后失效。",
                                       f" 附件：安全生产合同结尾承包人代理人\代表人与合同不同")
                    else:
                        factors_ok.append(" 附件：安全生产合同结尾承包人代理人\代表人")
                except:
                    factors_error[" 附件：安全生产合同结尾承包人代理人\代表人"] = " 附件：安全生产合同_合同承包人代理人\代表人提取错误"
                    addRemarkInDoc(word, document,
                                   "本合同正本一式二份，副本六分，合同双方各执正本一份，副本三份。双方法定代表或其授权的代理人签署与加盖章后生效，全部工程竣工验收后失效。",
                                   f" 附件：安全生产合同_合同承包人代理人\代表人提取错误")
            else:
                factors_error[" 附件：安全生产合同结尾承包人代理人\代表人"] = " 附件：安全生产合同结尾承包人代理人\代表人未填写完整"
                addRemarkInDoc(word, document, "本合同正本一式二份，副本六分，合同双方各执正本一份，副本三份。双方法定代表或其授权的代理人签署与加盖章后生效，全部工程竣工验收后失效。",
                               f" 附件：安全生产合同结尾承包人代理人\代表人未填写完整")
        except:
            missObject += "要素“ 附件：安全生产合同结尾承包人代理人\代表人”缺失\n"

    try:
        if missObject != "":
            addRemarkInDoc(word, document, "", missObject)
        copy_path = processed_file_sava_dir + "/" + filePath.split("/")[-1]
        filePath = str_insert(copy_path, copy_path.index(".doc"), "(已审查)")
        print(filePath)
        document.SaveAs(filePath)
        document.Close()
    except Exception as ex:
        print(ex)
    # print("中间文件已删除")
    # print(factors, factors_ok, factors_error, factors_to_inform)
    word.Quit()
    return factors, factors_ok, factors_error, factors_to_inform


if __name__ == "__main__":
    from docx import Document

    filePath = r'C:\Users\12259\Desktop\测试文件\总\建筑合同.docx'
    document = Document(filePath)
    paragraghs = document.paragraphs
    tables = document.tables
    text = ""
    for p in paragraghs:
        if p.text != "":
            # 把半角全角符号一律转全角 add by qy
            text = text.replace(':', '：')
            text = text.replace('(', '（')
            text = text.replace(')', '）')
            text = text.replace('\ue5e5', ' ').replace('\u3000', ' ')
            text += p.text + "\n"
    processed_file_sava_dir = r'C:\Users\12259\Desktop'
    filePath0 = r'C:\Users\12259\Desktop\测试文件\总\招标文件.docx'
    factors, factors_ok, factors_error, factors_to_inform = processFunc3(tables, text, filePath,
                                                                         processed_file_sava_dir, filePath0)
    print(factors)
    print(factors_ok)
    print(factors_error)
    print(factors_to_inform)
