# -*- coding:utf-8 -*-
# @ModuleName: DocReader
# @Function:
# @Author: huhonghui
# @email: 1241328737@qq.com
# @Time: 2021/6/18 17:10
import shutil

from ContractCheckResult import ContractCheckResult
from ContractType import ContractType
from docx import Document
from win32com import client as wc
import getFactorsFromContract
from over_all_description import get_over_all_file
import os
import None_standard_contract


class DocReader:
    type_list = {
        "就本合同工程施工及有关事项": ContractType.ConstructionContract,
        "就房屋租赁事宜达成协议如下": ContractType.RentContract,
        "与乙方就具体采购项目签订采购合同后": ContractType.PurchaseAndWarehousingContract,
        "负责物业范围内大厅、广场、道路、楼道、卫生间、楼梯、电": ContractType.PropertyManagementContract,
        "买卖合同": ContractType.BuySellContract,
    }

    def __init__(self, file_path, processed_file_sava_dir, filePath_zhaobiao=""):
        self.file_path = None
        self.processed_file_sava_dir = processed_file_sava_dir  # iu界面中传过来的路径，如"D:/abc"不包含最后的“/”
        self.has_new_file = self.doc_to_docx(file_path)
        self.filePath_zhaobiao = filePath_zhaobiao

    def to_info(self):
        # code here
        try:
            document = Document(self.file_path)
            paragraghs = document.paragraphs

            text = ""
            for p in paragraghs:
                if p.text != "":
                    # 把半角全角符号一律转全角 add by qy
                    text = text.replace(':', '：')
                    text = text.replace('(', '（')
                    text = text.replace(')', '）')
                    text = text.replace('\ue5e5', ' ').replace('\u3000', ' ')
                    text += p.text + "\n"

            contract_type = ContractType.NotSure
            if '非标准' in self.file_path:
                # add by qy
                contract_type = ContractType.NotSure
            else:
                for key in DocReader.type_list.keys():
                    if key in text:
                        contract_type = DocReader.type_list[key]
                        print("判断出该合同属于：", contract_type)  ######################################
                        break
        except Exception as e:
            print(e, "请确定\"" + self.file_path + "\"不是空文档！")  ##################################
            return "docx_blank"
        try:
            if contract_type == ContractType.NotSure:
                # print("合同类型未匹配成功！请检查文档内容是否合乎规范！！")  ######################################
                print("合同为非标准合同，对该合同进行风险性检测和完整性检测！")
                # 新增点，作用为在调用完None_standard_contract.Buy_sell_contract()函数后，关闭word进程
                tables = document.tables  # 把文档的表格s一起传过去 add by qy

                factors, factors_ok, factors_error, factors_to_inform, word = self.checker(contract_type, text, tables)
                self.remove_new_file()
                return ContractCheckResult(contract_type, factors, factors_ok, factors_error, factors_to_inform)
            else:
                tables = document.tables  # 把文档的表格s一起传过去 add by qy
                factors, factors_ok, factors_error, factors_to_inform = self.checker(contract_type, text, tables)
                self.remove_new_file()
                return ContractCheckResult(contract_type, factors, factors_ok, factors_error, factors_to_inform)
        except Exception as ex:
            print(self.file_path, '审查出错', ex)
            factors = {}
            factors_ok = []
            factors_error = {}
            factors_to_inform = {}
        finally:
            return ContractCheckResult(contract_type, factors, factors_ok, factors_error, factors_to_inform)

            # return "type_not_sure"
        '''
        tables = document.tables  # 把文档的表格s一起传过去 add by qy

        factors, factors_ok, factors_error, factors_to_inform = self.checker(contract_type, text, tables)
        return ContractCheckResult(contract_type, factors, factors_ok, factors_error, factors_to_inform)
        '''

        # return ContractCheckResult(contract_type, factors, factors_ok, factors_error, factors_to_inform)

    # 在线接口调用
    def deal_one(self, file_type, standard):
        try:
            document = Document(self.file_path)
            paragraghs = document.paragraphs

            text = ""
            for p in paragraghs:
                if p.text != "":
                    # 把半角全角符号一律转全角 add by qy
                    text = text.replace(':', '：')
                    text = text.replace('(', '（')
                    text = text.replace(')', '）')
                    text = text.replace('\ue5e5', ' ').replace('\u3000', ' ')
                    text += p.text + "\n"
            if standard == 'standard':
                print("合同为标准合同")

                tables = document.tables  # 把文档的表格s一起传过去 add by qy
                if file_type == 0:
                    c_type = ContractType.BuySellContract
                    type_name = '买卖合同'
                    print("买卖合同标准审查")
                    factors, factors_ok, factors_error, factors_to_inform = getFactorsFromContract.buy_sell_contract(
                        text, self.file_path, self.processed_file_sava_dir)

                if file_type == 1:
                    c_type = ContractType.RentContract
                    type_name = '租赁合同'
                    print("租赁合同标准审查")
                    factors, factors_ok, factors_error, factors_to_inform = getFactorsFromContract.rent_contract(
                        text, self.file_path, self.processed_file_sava_dir)


                if file_type == 2:
                    c_type = ContractType.PurchaseAndWarehousingContract
                    type_name = '采购合同'
                    print("采购合同标准审查")
                    factors, factors_ok, factors_error, factors_to_inform = getFactorsFromContract.purchase_and_warehousing_contract(
                        text, tables, self.file_path, self.processed_file_sava_dir)

                if file_type == 3:
                    c_type = ContractType.PropertyManagementContract
                    type_name = '物业管理合同'
                    print("物业管理合同标准审查")
                    factors, factors_ok, factors_error, factors_to_inform = getFactorsFromContract.property_management_contract(
                        text, tables, self.file_path, self.processed_file_sava_dir)

                if file_type == 4:
                    type_name = '建筑合同'
                    print("建筑合同标准审查")
                    c_type = ContractType.ConstructionContract
                    factors, factors_ok, factors_error, factors_to_inform = getFactorsFromContract.construction_contract(
                        text, tables, self.file_path, self.processed_file_sava_dir)

                self.remove_new_file()
            else:
                c_type = ContractType.NotSure
                if file_type == 4:
                    factors = {}
                    factors_ok = []
                    factors_error = {}
                    factors_to_inform = {}
                    return ContractCheckResult(ContractType.NotSure, factors, factors_ok, factors_error,
                                               factors_to_inform)
                if file_type == 0:
                    type_name = '买卖合同'
                    print("买卖合同非标准审查")
                    factors, factors_ok, factors_error, factors_to_inform, _ = None_standard_contract.Buy_Sell_contract(
                        self.file_path, self.processed_file_sava_dir)
                if file_type == 1:
                    type_name = '租赁合同'
                    print("租赁合同非标准审查")
                    factors, factors_ok, factors_error, factors_to_inform, _ = None_standard_contract.lease_contract(
                        self.file_path, self.processed_file_sava_dir)
                if file_type == 2:
                    type_name = '采购合同'
                    print("采购合同非标准审查")
                    factors, factors_ok, factors_error, factors_to_inform, _ = None_standard_contract.purchase_contract(
                        self.file_path, self.processed_file_sava_dir)
                if file_type == 3:
                    type_name = '物业管理合同'
                    print("物业管理合同非标准审查")
                    factors, factors_ok, factors_error, factors_to_inform, _ = None_standard_contract.property_management_contract(
                        self.file_path, self.processed_file_sava_dir)
                self.remove_new_file()
        except Exception as e:
            print(self.file_path, '审查出错', e)
            factors = {}
            factors_ok = []
            factors_error = {}
            factors_to_inform = {}
        finally:
            return ContractCheckResult(c_type, factors, factors_ok, factors_error, factors_to_inform, type_name)

    # 在写了ui界面后，传进来的 file 应该是绝对路径 2021.7.8日修改 by：胡洪辉
    def doc_to_docx(self, file):
        if file[-3:] == "doc":
            flg = False
            if ' ' in file:
                flg = True
                shutil.copy(file, file.replace(' ', ''))
                file = file.replace(' ', '')
            word = wc.Dispatch("Word.Application")
            print(file)
            doc = word.Documents.Open(file)
            doc.SaveAs(file + "x", 12)
            doc.Close()
            if flg:
                os.remove(self.file_path)
            self.file_path = file + "x"
            return True
        else:
            if ' ' in file:
                self.file_path = file.replace(' ', '')
                shutil.copy(file, self.file_path)
                return True
            else:
                self.file_path = file
                return False

    def remove_new_file(self):
        if self.has_new_file:
            os.remove(self.file_path)

    def checker(self, contract_type, text, tables=None):
        if '非标准物业' in self.file_path:
            return None_standard_contract.property_management_contract(self.file_path, self.processed_file_sava_dir)
        elif '非标准买卖' in self.file_path:
            return None_standard_contract.Buy_Sell_contract(self.file_path, self.processed_file_sava_dir)
        elif '非标准租赁' in self.file_path:
            return None_standard_contract.lease_contract(self.file_path, self.processed_file_sava_dir)
        elif '非标准采购' in self.file_path:
            return None_standard_contract.purchase_contract(self.file_path, self.processed_file_sava_dir)
        if contract_type == ContractType.BuySellContract:
            # ++++++++++++修改点+++++++++++++
            return getFactorsFromContract.buy_sell_contract(text, self.file_path, self.processed_file_sava_dir)

        if contract_type == ContractType.RentContract:
            return getFactorsFromContract.rent_contract(text, self.file_path, self.processed_file_sava_dir)

        if contract_type == ContractType.ConstructionContract:
            return getFactorsFromContract.construction_contract(text, tables, self.file_path,
                                                                self.processed_file_sava_dir, self.filePath_zhaobiao)

        if contract_type == ContractType.PurchaseAndWarehousingContract:
            return getFactorsFromContract.purchase_and_warehousing_contract(text, tables, self.file_path,
                                                                            self.processed_file_sava_dir)

        if contract_type == ContractType.PropertyManagementContract:
            return getFactorsFromContract.property_management_contract(text, tables, self.file_path,
                                                                       self.processed_file_sava_dir)


# 下面是测试程序功能的代码，读取一个文件夹中的所有word文件，并分析要素
# 由于增加了ui界面，之后的版本传进DocReader(file)中的file参数需要是绝对路径


if __name__ == '__main__':
    # # 项目根目录的绝对路径
    # prject_path = os.path.abspath(os.path.dirname(__file__))[:-4].replace("\\","/")
    #
    # # 数据目录绝对路径
    # file_dir_path = prject_path + "/烟草合同案卷模板数据/"
    #
    # file_list = os.listdir(file_dir_path)
    # for file in file_list:
    #     if "doc" in file:
    #         print("-----------------------------------\n正在审查文件：" + file)  ######################################
    #         d = DocReader(file_dir_path + file)#最后一个参数为审核文件保存地址
    #         contract_check_result = d.to_info()
    #         if contract_check_result != "docx_blank" and contract_check_result != "type_not_sure":
    #             print("要素",contract_check_result.factors)
    #             print("审核通过", contract_check_result.factors_ok)  ######################################
    #             print("审核错误", contract_check_result.factors_error)  ######################################
    #             print("审核提示", contract_check_result.factors_to_inform)  ######################################
    #         d.remove_new_file()

    # 由于UI界面跑程序，如果有bug会直接闪退，为了方便代码改bug，在这里测试程序，如果有bug会报错提示。add by huhonghui 2021.9.8

    file = "E:/Postgraduate/fuzzymatching/data/合同6.docx"
    export_dir = "C:/Users/Mr.Blonde/Desktop"

    over_all_info = []
    d = DocReader(file, export_dir)
    contract_check_result = d.to_info()

    over_all_info.append([file, contract_check_result])
    get_over_all_file(over_all_info, "C:/Users/Mr.Blonde/Desktop")
