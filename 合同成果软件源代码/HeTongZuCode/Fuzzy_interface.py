#!/usr/bin/env python
# -*- encoding: UTF-8 -*-
"""
@Project: ChinaTobaccoContract_V2
@File: Fuzzy_interface.py
@Author: Mr.Blonde
@Date: 2021/12/13 10:22
"""
import numpy as np
import re
import docx
from fuzzymatching import *
import jieba



# 计算聚集分数
def c_dists(list_):
    len_ = len(list_)
    list_.append(list_[-1])
    list_tmp1 = np.array(list_[0:-1])
    list_tmp = np.array(list_[1:])
    # 重复次数大于2的元素重复次数之和
    list_ = list_[0:-1]
    myset = set(list_)
    times = 0
    s = 2  # 重复次数细粒度,可算作奖励系数
    for item in myset:
        c_ = list_.count(item)
        if c_ >= s:
            times += c_
    dist_ = ((np.sum(list_tmp - list_tmp1)) + 1) / (len_ * (1 + times / s))
    return dist_

def get_subidex(list_):
    min_ = 1000000
    subIdx = []
    len_ = len(list_)
    for i in range(len_):
        for j in range(i+1 , len_+1):
            min_tmp = c_dists(list_[i:j])
            if min_tmp < min_:
                min_ = min_tmp
                # 左，右，重复数
                subIdx = [[i, j]]
            if min_tmp == min_:
                subIdx.append([i, j])
    if len(subIdx) == 1:
        return subIdx[0]
    elif len(subIdx) > 1:
        subIdx_tmp = np.array(subIdx)
        subIdx_tmp = subIdx_tmp[:, 1] - subIdx_tmp[:, 0]
        subIdx_idex = np.argsort(-subIdx_tmp)
        return subIdx[subIdx_idex[0]]
    else:
        return [0, 0]



def get_checkItems(items):
    dict_ = {}
    for i in items.keys():
        dict_tmp = {}
        for j in items[i].keys():
            for k in items[i][j].keys():
                dict_tmp[k] = items[i][j][k]
        dict_[i] = dict_tmp
    return dict_

def find_indexInParagraphs(sbstr,paragraphs3):
    idxs = []
    # paragraphs3=fullToHalf()
    sbstr = fullToHalf(sbstr)
    for i,v in enumerate(paragraphs3):
        v=fullToHalf(v)
        if sbstr in v:
            idxs.append(i)
    return  idxs


# 输入要处理的文件

def fuzzmatching_(check_items1,check_items2,file='./data/合同1.docx'):
    # 加载jieba自定义词库
    import os
    print(os.getcwd())
    jieba.load_userdict(r"C:\WorkSpace\tobacco-master\HeTongZuCode\userdict.txt")
    # 加载文件
    doc = docx.Document(file)
    paragraphs1 = []  # jieba 分词
    paragraphs2 = []  # 段落分词
    paragraphs3 = []  # 用于打批注用的分句;.,。，；\n
    tables = []
    items1 = get_checkItems(check_items1)
    items2 = get_checkItems(check_items2)
    # _________________________________字词变量配置___________________________________________________#
    vote1 = 0  # 如果是0则根据每个item的阈值不采用投票法，如果不是则就是投票候选个数大于vote_num个，根据vote_type定
    vote_num1 = 9  # 投票法阈值
    default_method1 = 'text2vec'  # vote等于0的时候选用的similar算法
    vote_type1 = 'standard'  # standard 几种候选中出现最多的，mean 几种方式分数平均找出最大的vote_num个
    similar_methods1 = ['Simple Ratio', 'Partial Ratio', 'Token Sort Ratio', 'Token Set Ratio', 'text2vec']  # 投票法先不用

    # _________________________________段落长句变量配置___________________________________________________#
    vote2 = 0  # 如果是0则根据每个item的阈值不采用投票法，如果不是则就是投票候选个数大于vote_num个，根据vote_type定
    vote_num2 = 0  # 投票法阈值
    default_method2 = 'text2vec'  # vote等于0的时候选用的similar算法
    vote_type2 = 'standard'  # standard 几种候选中出现最多的，mean 几种方式分数平均找出最大的vote_num个
    similar_methods2 = ['Simple Ratio', 'Partial Ratio', 'Token Sort Ratio', 'Token Set Ratio', 'text2vec']  # 投票法先不用

    for paragraph in doc.paragraphs:
        p = fullToHalf(paragraph.text.replace(" ", ""))
        p = jieba.lcut(p)
        if len(p) == 0:
            continue
        #print(p)
        for i in p:
            paragraphs1.append(i)
    paragraphs1 = sorted(set(paragraphs1), key=paragraphs1.index)

    #print("加入表格之前：###################################\n",paragraphs1)
    # 段落，长句
    for paragraph in doc.paragraphs:
        p = fullToHalf(paragraph.text.replace(' ', ''))
        ps = re.split(r'。|;', p)
        for i in ps:
            if i.replace(" ", '') != '':
                if len(i) > 30:
                    paragraphs2.append(i)  # 大于30才可以
        # 3
        ps = re.split(r';|；|。', paragraph.text)
        for i in ps:
            if i!="":
                paragraphs3.append(i)

    for table in doc.tables:
        tables.append(table)

    #表格加入para1
    _table_list = []
    if tables:
        for i, row in enumerate(table.rows):  # 读每行
            row_content = []
            for cell in row.cells:  # 读一行中的所有单元格
                c = cell.text
                if c not in row_content:
                    row_content.append(c)
            # print(row_content)
            _table_list.append(row_content)
        # print(tables)
        # print(_table_list)
        for i in _table_list:
            for j in i:
                j = fullToHalf(j.replace(" ", "").replace("\n", ""))
                if j != "":
                    paragraphs1.append(j)

    #print("加入表格后：#####################################\n",paragraphs1)

    list_items1 = []
    list_items2 = []
    for key in items1.keys():
        for key1 in items1[key].keys():
            list_items1.append(key1)

    for key in items2.keys():
        for key1 in items2[key].keys():
            list_items2.append(key1)

    # --------------------------字词计算相似度------------------------------------
    if vote1 > 0:
        scores1 = np.zeros((len(similar_methods1), len(list_items1), len(paragraphs1)))
    else:
        scores1 = np.zeros((len(list_items1), len(paragraphs1)))
    if vote1 == 0:
        scores1 = fuzzyMatching(list_items1, paragraphs1,default_method1)
    else:
        if vote_type1 == 'standard':
            for i, method in enumerate(similar_methods1):
                scores1[i] = fuzzyMatching(list_items1, paragraphs1, method)

        else:
            for i, method in enumerate(similar_methods1):
                scores1[i] = fuzzyMatching(list_items1, paragraphs1, method)
            scores1 = np.mean(np.array(scores1), axis=0)
            # print(scores1.shape)
    # --------------------------长句段落计算相似度------------------------------------
    if vote2 > 0:
        scores2 = np.zeros((len(similar_methods2), len(list_items2), len(paragraphs2)))
    else:
        scores2 = np.zeros((len(list_items2), len(paragraphs2)))
    if vote2 == 0:
        scores2 = fuzzyMatching(list_items2, paragraphs2, default_method2)
    else:
        if vote_type2 == 'standard':
            for i, method in enumerate(similar_methods2):
                scores2[i] = fuzzyMatching(list_items2, paragraphs2, method)
        else:
            for i, method in enumerate(similar_methods2):
                scores2[i] = fuzzyMatching(list_items2, paragraphs2, method)
            scores2 = np.mean(np.array(scores2), axis=0)
            # print(scores2.shape)

    rss1 = extract_items(items1, paragraphs1, scores1, vote1, vote_type1, vote_num1)
    rss2 = extract_items(items2, paragraphs2, scores2, vote2, vote_type2, vote_num2)

    # 返回短词字典，长词字典，以及相应提取项，以及要
    #tyh
    # print("\n#######################################短句短词########################################")
    # for i in rss1.keys():
    #     item = rss1[i]
    #     # print(item)
    #     print(f'\n-{i}-------------------------------')
    #     for j in item[0]:
    #         if len(j) > 0:
    #             print(j)
    #
    # print("\n#######################################长句段落########################################")
    # for i in rss2.keys():
    #     item = rss2[i]
    #     # print(item)
    #     print(f'\n-{i}-------------------------------')
    #     for j in item[0]:
    #         print(j)
    #tyh
    # 做完整性检查 大项：{check_lsit:[],check_flg:True or Flase}
    Missing_Dict={}
    Missing_Dict_long={}
    Missing_list=[]
    check_intergrity = {}
    check_intergrity_long = {}
    check_indexList = {}
    #先处理短词短句, 同时计算批注位置
    for i in check_items1.keys():
        # print("i的值：",i)
        match_item = []
        positions = []
        positions_key = {}
        for j in rss1[i][0]:
            if len(j) > 0:
                positions_ = []
                for k in j:
                    match_item.append(k[0])
                    positions_.extend(find_indexInParagraphs(k[1],paragraphs3))
                positions_ = list(set(positions_))
                positions_.sort()
                positions_key[j[0][0]]={"positions":positions_,'score':j[0][2]}
                positions.extend(positions_)
        positions.sort()
        check_lsit = []
        check_flg = True
        for j in check_items1[i].keys():
            # print("j的值：",j)
            check_flg_1 = False
            for k in check_items1[i][j].keys():
                # print("k的值：",k)
                if k in match_item:
                   check_flg_1 =True
            if check_flg_1==False:
                check_flg = False
                check_lsit.append(0)
                Missing_list.append(j)
                Missing_Dict[i]=Missing_list
            else:
                check_lsit.append(1)
        Missing_list = []
        position = get_subidex(positions) #找到打标位置区间
        socre = 0
        c_position = 0
        #print(position, positions, positions_key)
        if len(positions) > 0:
            r_positions = positions[position[0]:position[1]]
            for j in r_positions:
                for k in positions_key.keys():
                    p = positions_key[k]['positions']
                    s = positions_key[k]['score']
                    if j in p:
                        if socre < s:
                            socre = s
                            c_position = j
        flg_str = paragraphs3[c_position]
        if c_position == 0:
            flg_str = ""

        check_intergrity[i] = {"check_flg":check_flg,"flg_str": flg_str,"check_lsit":check_lsit,}
    # print(check_intergrity)
    #doc.Close()
    for i in check_items2.keys():
        # print("i的值：",i)
        match_item = []
        positions = []
        positions_key = {}
        for j in rss2[i][0]:
            if len(j) > 0:
                positions_ = []
                for k in j:
                    match_item.append(k[0])
                    positions_.extend(find_indexInParagraphs(k[1],paragraphs3))
                positions_ = list(set(positions_))
                positions_.sort()
                positions_key[j[0][0]]={"positions":positions_,'score':j[0][2]}
                positions.extend(positions_)
        positions.sort()
        check_lsit = []
        check_flg = True
        for j in check_items2[i].keys():
            # print("j的值：",j)
            check_flg_1 = False
            for k in check_items2[i][j].keys():
                # print("k的值：",k)
                if k in match_item:
                   check_flg_1 =True
            if check_flg_1==False:
                check_flg = False
                check_lsit.append(0)
                Missing_list.append(j)
                Missing_Dict_long[i]=Missing_list
            else:
                check_lsit.append(1)
        Missing_list = []
        position = get_subidex(positions) #找到打标位置区间
        socre = 0
        c_position = 0
        #print(position, positions, positions_key)
        if len(positions) > 0:
            r_positions = positions[position[0]:position[1]]
            for j in r_positions:
                for k in positions_key.keys():
                    p = positions_key[k]['positions']
                    s = positions_key[k]['score']
                    if j in p:
                        if socre < s:
                            socre = s
                            c_position = j
        flg_str = paragraphs3[c_position]
        if c_position == 0:
            flg_str = ""

        check_intergrity_long[i] = {"check_flg":check_flg,"flg_str": flg_str,"check_lsit":check_lsit,}
    # return check_intergrity,doc,check_items1,check_items2
    # print('#############MissingDict##########\n',Missing_Dict)
    # print('#############check_integrity_long##########\n',check_intergrity_long)
    # print('#############Missing_long##########\n',Missing_Dict_long)
    #print()
    return check_intergrity,check_intergrity_long,check_items1,check_items2,Missing_Dict,Missing_Dict_long,paragraphs3
