#!/usr/bin/env python
# -*- encoding: UTF-8 -*-
"""
@Project: ChinaTobaccoContract_V2_new_version_riskAdd
@File: check_Tender.py
@Author: QinYang
@Date: 2022/4/6 19:13
"""
import docx
from win32com import client as wc
import re

# 全角半角转换表
definedConverts = {
    "　": "",
    "“": '"',
    "”": '"',
    "！": "!",
    "￥": "￥",
    "……": "",
    "（": "(",
    "）": ")",
    "——": "",
    "【": "",
    "】": "",
    "；": ";",
    "’": "'",
    "：": ":",
    "|": "",
    "、": "",
    "？": "?",
    "\t": "",
}


# 全角转半角
def fullToHalf(str_):
    for from_, to_ in definedConverts.items():
        # print(from_, to_)
        str_ = str_.replace(from_, to_)
    return str_

def doc_to_docx(file):
    if file[-3:] == "doc":
        word = wc.Dispatch("Word.Application")
        print(file)
        doc = word.Documents.Open(file)
        doc.SaveAs(file + "x", 12)
        doc.Close()
        file_path = file + "x"
        return file_path

    else:
        return file



# # 找出七项招标风险
hetongbiaodi = ['合同标的', '项目名称','合同标的:', '项目名称:','项目:','标的：']
fukuanfangshi = ['付款方式', '转账方式', '项目付款方式','付款方式:', '转账方式:', '项目付款方式:']
hetongjiakuan = ['总价', '总额', '合同价款', '价款', '元人民币', '总费用', '大写：人民币','总价:', '总额:', '合同价款:']
yanshoutiaokuan = ['验收条款', '验收标准', '验收不合格', '验收合格']
weiyuetiaokuan = ['违约条款', '处理违约办法', '处理违约', '违约责任','违约条款:', '处理违约办法', '处理违约', '违约责任:']
zhenyijiejue = ['争议解决条款', '争议解决办法', '协商解决', '解决争议','争议解决条款:', '争议解决办法:', '协商解决', '解决争议']
hetongshengxiao = ['合同生效时间', '生效时间', '起生效', '并生效', '生效条款', '合同生效','合同生效时间:', '生效时间:']

def check_c(paragraphs):
    check_dict = {
        '合同标的':[],
        '付款方式': [],
        '合同价款': [],
        '验收条款': [],
        '违约条款': [],
        '争议解决': [],
        '合同生效': [],
    }
    for paragraph in paragraphs:
        for i in hetongbiaodi:
            if i in paragraph:
                check_dict['合同标的'].append(paragraph)
        for i in fukuanfangshi:
            if i in paragraph:
                check_dict['付款方式'].append(paragraph)
        for i in hetongjiakuan:
            if i in paragraph:
                check_dict['合同价款'].append(paragraph)
        for i in yanshoutiaokuan:
            if i in paragraph:
                check_dict['验收条款'].append(paragraph)
        for i in weiyuetiaokuan:
            if i in paragraph:
                check_dict['违约条款'].append(paragraph)
        for i in zhenyijiejue:
            if i in paragraph:
                check_dict['争议解决'].append(paragraph)
        for i in hetongshengxiao:
            if i in paragraph:
                check_dict['合同生效'].append(paragraph)
    return check_dict


def compare(c_check_dict,t_check_dict):
    c_miss = []
    t_miss = []
    # 合同、招标文件内容是否缺失
    for key in c_check_dict.keys():
        if len(c_check_dict[key])==0:
            c_miss.append(key)
    for key in t_check_dict.keys():
        if len(t_check_dict[key]) == 0:
            t_miss.append(key)
    # 招标文件内容对比
    pair_dict = {}
    for key in c_check_dict.keys():
        list_one = c_check_dict[key]
        list_two = t_check_dict[key]
        pairs = []
        for i in list_one:
            for j in list_two:
                # if i in j or j in i:
                if i == j:
                    pairs.append([i,j])
        pair_dict[key]=pairs
    return c_miss,t_miss,pair_dict

def compare_api(c_file,t_file):
    c_file = doc_to_docx(c_file)
    t_file = doc_to_docx(t_file)

    c_doc = docx.Document(c_file)
    t_doc = docx.Document(t_file)
    c_paragraphs = []
    t_paragraphs = []
    # 分段：
    for paragraph in c_doc.paragraphs:
        p = fullToHalf(paragraph.text.replace(' ', ''))
        ps = re.split(r'。|;', p)
        for i in ps:
            if i.replace(" ", '') != '':
                c_paragraphs.append(i)
    for paragraph in t_doc.paragraphs:
        p = fullToHalf(paragraph.text.replace(' ', ''))
        ps = re.split(r'。|;', p)
        for i in ps:
            if i.replace(" ", '') != '':
                t_paragraphs.append(i)

    c_miss,t_miss,pair_dict = compare(check_c(c_paragraphs),check_c(t_paragraphs))

    print("--------------------------")
    if len(c_miss)==0:
        print("合同的对比要素无缺失")
    else:
        for i in c_miss:
            print(f"合同的对比要素 {i} 缺失")
    print("--------------------------")
    if len(t_miss)==0:
        print("招标文件的对比要素无缺失")
    else:
        for i in t_miss:
            print(f"招标文件的对比要素 {i} 缺失")
    print("--------------------------")
    print("合同、招标文件对比条款：")
    for key in pair_dict.keys():
        flg = 0
        for i in pair_dict[key]:
            if len(i[0])>10 and len(i[1])>10:
                flg = 1
            else:
                if flg==0:
                    flg = 2

        if flg==0:
            print(f"对比要素 {key} 不一致", pair_dict[key])
        elif flg==1:
            print(f"对比要素 {key} 一致", pair_dict[key])
        else:
            print(f"对比要素 {key} 基本一致", pair_dict[key])

# contract_path  tender_path
if __name__ == '__main__':
    compare_api("E:\projects\pycharm\hetong\范例/1\合同文本 (3).doc", "E:\projects\pycharm\hetong\范例/1\（6.8）考试测评服务定点供应商项目-会审后改.docx")
    compare_api("E:\projects\pycharm\hetong\范例/2\合同文本 (3).doc", "E:\projects\pycharm\hetong\范例/2\招标文件客户满意度调查2021----修改版0811.docx")
    compare_api("E:\projects\pycharm\hetong\范例/3\合同文本 (4).doc","E:\projects\pycharm\hetong\范例/3/7.30稿-招标文件-中国烟草总公司四川省公司四川诚至诚烟草投资有限责任公司常年法律顾问项目.docx")